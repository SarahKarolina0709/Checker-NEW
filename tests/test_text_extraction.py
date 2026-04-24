# -*- coding: utf-8 -*-
"""Tests fuer nicegui_app.text_extraction (Datei-zu-Text + Finding-Serialisierung)."""
from __future__ import annotations

import os
import sys
import tempfile
from dataclasses import dataclass, field
from typing import Any, Dict

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from nicegui_app import text_extraction as te


@dataclass
class _Finding:
    severity: str = 'info'
    code: str = ''
    message: str = ''
    category: str = ''
    source_text: str = ''
    target_text: str = ''
    segment_index: int = -1
    meta: Dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# extract_text - Standardfaelle
# ---------------------------------------------------------------------------
class TestExtractText:
    def test_missing_path_returns_marker(self, tmp_path):
        result = te.extract_text(str(tmp_path / 'gibts_nicht.txt'))
        assert result.startswith('[Datei nicht gefunden')

    def test_empty_path(self):
        result = te.extract_text('')
        assert result.startswith('[Datei nicht gefunden')

    def test_txt_file_utf8(self, tmp_path):
        p = tmp_path / 'sample.txt'
        p.write_text('Hallo Welt\nÄÖÜß', encoding='utf-8')
        result = te.extract_text(str(p))
        assert 'Hallo Welt' in result
        assert 'ÄÖÜß' in result

    def test_md_file(self, tmp_path):
        p = tmp_path / 'readme.md'
        p.write_text('# Titel\n\nText.', encoding='utf-8')
        assert '# Titel' in te.extract_text(str(p))

    def test_py_file(self, tmp_path):
        p = tmp_path / 'code.py'
        p.write_text('def foo(): return 1', encoding='utf-8')
        assert 'def foo' in te.extract_text(str(p))

    def test_unknown_extension_returns_empty(self, tmp_path):
        p = tmp_path / 'file.xyz'
        p.write_text('whatever', encoding='utf-8')
        assert te.extract_text(str(p)) == ''

    def test_doc_format_message(self, tmp_path):
        p = tmp_path / 'old.doc'
        p.write_bytes(b'fake')
        result = te.extract_text(str(p))
        assert '.doc' in result and 'externem Konverter' in result

    def test_corrupt_pdf_returns_message(self, tmp_path):
        # Kaputte PDF: Beide Reader scheitern, ohne OCR Fallback-Marker
        p = tmp_path / 'broken.pdf'
        p.write_bytes(b'not a real pdf')
        result = te.extract_text(str(p))
        # Entweder Fehlermeldung oder leerer-OCR-Fallback. Wichtig: kein crash.
        assert isinstance(result, str)


# ---------------------------------------------------------------------------
# extract_text - DOCX (Bug 53: Tabellen muessen mitextrahiert werden)
# ---------------------------------------------------------------------------
class TestExtractDocx:
    def _make_docx(self, tmp_path, paragraphs, tables=None):
        try:
            from docx import Document
        except ImportError:
            pytest.skip('python-docx nicht installiert')
        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        for tbl in (tables or []):
            t = doc.add_table(rows=len(tbl), cols=len(tbl[0]))
            for r_idx, row in enumerate(tbl):
                for c_idx, val in enumerate(row):
                    t.rows[r_idx].cells[c_idx].text = val
        path = tmp_path / 'test.docx'
        doc.save(str(path))
        return str(path)

    def test_docx_paragraphs(self, tmp_path):
        path = self._make_docx(tmp_path, ['Erster Absatz', 'Zweiter Absatz'])
        result = te.extract_text(path)
        assert 'Erster Absatz' in result
        assert 'Zweiter Absatz' in result

    def test_docx_table_cells_extracted(self, tmp_path):
        # Bug 53 regression: Tabellen wurden frueher ignoriert
        path = self._make_docx(
            tmp_path,
            ['Vorwort'],
            tables=[[['Kopf A', 'Kopf B'], ['Wert 100', 'Wert 250']]],
        )
        result = te.extract_text(path)
        assert 'Vorwort' in result
        assert 'Kopf A' in result
        assert 'Wert 100' in result
        assert 'Wert 250' in result

    def test_docx_empty_cells_skipped(self, tmp_path):
        path = self._make_docx(
            tmp_path, [],
            tables=[[['Inhalt', '   ', '']]],
        )
        result = te.extract_text(path)
        assert 'Inhalt' in result


# ---------------------------------------------------------------------------
# Finding <-> Dict
# ---------------------------------------------------------------------------
class TestFindingDict:
    def test_finding_to_dict_basics(self):
        f = _Finding(severity='major', code='X1', message='msg', source_text='s', target_text='t')
        d = te.finding_to_dict(f)
        assert d['severity'] == 'major'
        assert d['code'] == 'X1'
        assert d['source_text'] == 's'

    def test_finding_to_dict_none_text_fields_become_empty_string(self):
        f = _Finding(source_text=None, target_text=None)  # type: ignore[arg-type]
        d = te.finding_to_dict(f)
        assert d['source_text'] == ''
        assert d['target_text'] == ''

    def test_finding_to_dict_meta_none_safe(self):
        @dataclass
        class _F2:
            severity: str = 'info'
            code: str = ''
            message: str = ''
            category: str = ''
            source_text: str = ''
            target_text: str = ''
            segment_index: int = -1
            meta: Any = None
        d = te.finding_to_dict(_F2())
        assert d['meta'] == {}

    def test_dict_to_finding_roundtrip(self):
        d = {'severity': 'critical', 'code': 'C1', 'message': 'm', 'category': 'cat',
             'source_text': 'src', 'target_text': 'tgt', 'segment_index': 5,
             'meta': {'k': 'v'}}
        f = te.dict_to_finding(d, _Finding)
        assert f.severity == 'critical'
        assert f.segment_index == 5
        assert f.meta == {'k': 'v'}

    def test_dict_to_finding_none_severity_becomes_info(self):
        # Bug 52 regression: explizites null in Session-JSON
        d = {'severity': None, 'code': 'X', 'message': 'y'}
        f = te.dict_to_finding(d, _Finding)
        assert f.severity == 'info'

    def test_dict_to_finding_none_segment_index_becomes_minus_one(self):
        d = {'segment_index': None}
        f = te.dict_to_finding(d, _Finding)
        assert f.segment_index == -1

    def test_dict_to_finding_missing_keys_get_defaults(self):
        f = te.dict_to_finding({}, _Finding)
        assert f.severity == 'info'
        assert f.code == ''
        assert f.segment_index == -1
        assert f.meta == {}

    def test_dict_to_finding_meta_none_becomes_empty_dict(self):
        f = te.dict_to_finding({'meta': None}, _Finding)
        assert f.meta == {}
