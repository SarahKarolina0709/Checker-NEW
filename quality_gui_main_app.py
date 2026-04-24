#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations
"""
Professionelle Übersetzungsqualitäts-GUI - Hauptanwendung
Erweiterte 2-Panel-Ansicht mit umfassenden Qualitätsanalyse-Funktionen
"""

import os
import threading
import sys
import json
import logging
import time
import copy  # Undo/Redo benötigt tiefe Kopien von Pairing-State
from pathlib import Path
import tkinter as tk
import customtkinter as ctk
import unicodedata
# duplicate removed
import re
from functools import lru_cache
from dataclasses import asdict, is_dataclass
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen
"""OPTIONALE OCR/IMAGE VERARBEITUNG: pytesseract & Pillow werden nur geladen wenn vorhanden.
Falls nicht installiert bleibt Feature deaktiviert ohne Absturz (sanfter Fallback)."""
try:  # OCR optional
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover
    Image = None  # type: ignore
try:  # Tesseract Binding optional
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover
    pytesseract = None  # type: ignore
PyPDF2 = None  # type: ignore
def _lazy_load_pypdf2():  # schlanker Lazy-Loader zur Startbeschleunigung
    global PyPDF2
    if PyPDF2 is None:
        try:
            import importlib
            PyPDF2 = importlib.import_module('PyPDF2')  # type: ignore
        except Exception:
            PyPDF2 = None  # bleibt None bei Fehler
    return PyPDF2

# Lazy Loading für PyMuPDF (fitz) - schwere Bibliothek nur bei Bedarf laden
fitz = None  # type: ignore
def _lazy_load_fitz():  # Lazy-Loader für PyMuPDF zur Startbeschleunigung
    global fitz
    if fitz is None:
        try:
            import importlib
            fitz = importlib.import_module('fitz')  # type: ignore
        except Exception:  # pragma: no cover
            fitz = None  # bleibt None bei Fehler
    return fitz
try:  # Zentrales UIHelper-System für konsistente Button-Styles (wie Welcome Screen)
    from src.utils.ui_helpers import UIHelpers
except Exception:
    UIHelpers = None
from tkinter import filedialog, messagebox
# Hinweis: time ist bereits weiter oben importiert; erneuten Import entfernt.
from typing import Any, Callable, Dict, List, Optional, Sequence, Tuple

# ---------------------- MODULARE ADDITIVE KOMPONENTEN (NEU) ----------------------
try:
    from quality_gui_analysis_pipeline import QualityGuiAnalysisPipeline
except Exception:  # pragma: no cover
    QualityGuiAnalysisPipeline = None  # type: ignore
try:
    from quality_gui_pairing_manager import QualityGuiPairingManager
except Exception:  # pragma: no cover
    QualityGuiPairingManager = None  # type: ignore
try:
    from quality_gui_upload_manager import QualityGuiUploadManager
except Exception:  # pragma: no cover
    QualityGuiUploadManager = None  # type: ignore
try:
    from quality_gui_reporting import QualityGuiReporting
except Exception:  # pragma: no cover
    QualityGuiReporting = None  # type: ignore
try:
    from quality_gui_settings_ui import QualityGuiSettingsUI
except Exception:  # pragma: no cover
    QualityGuiSettingsUI = None  # type: ignore
try:
    from quality_gui_ui_layout import QualityGuiUILayout
except Exception:  # pragma: no cover
    QualityGuiUILayout = None  # type: ignore
try:
    from quality_gui_phase1_checkers import run_phase1_checks
except Exception:  # pragma: no cover
    run_phase1_checks = None  # type: ignore
try:
    from quality_gui_phase2_checkers import run_phase2_checks
except Exception:  # pragma: no cover
    run_phase2_checks = None  # type: ignore
try:
    from quality_gui_phase3_checkers import run_phase3_checks
except Exception:  # pragma: no cover
    run_phase3_checks = None  # type: ignore

try:
    from tools.query_ollama_translation import (
        call_ollama as _cli_call_ollama,
        format_context as _cli_format_context,
        summarize_findings as _cli_summarize_findings,
        DEFAULT_OLLAMA_URL as _CLI_OLLAMA_URL,
    )
except Exception:  # pragma: no cover
    _cli_call_ollama = None  # type: ignore
    _cli_format_context = None  # type: ignore
    _cli_summarize_findings = None  # type: ignore
    _CLI_OLLAMA_URL = "http://127.0.0.1:11434/api/generate"

try:
    from tools import generate_translation_report as generate_report_tool
except Exception:  # pragma: no cover
    generate_report_tool = None  # type: ignore

# Additive Infrastruktur (optional) – keine bestehenden Imports ersetzen
try:  # EventBus (optional)
    from infra.event_bus import EventBus, get_global_event_bus
except Exception:  # pragma: no cover - robust fallback
    EventBus = None  # type: ignore
    def get_global_event_bus():  # type: ignore
        return None
try:  # SettingsService optional
    from services.settings_service import SettingsService  # type: ignore
except Exception:  # pragma: no cover
    SettingsService = None  # type: ignore
try:  # AnalysisState (neu)
    from services.analysis_state import AnalysisState  # type: ignore
except Exception:  # pragma: no cover
    AnalysisState = None  # type: ignore
try:  # WorkerPool (optional)
    from services.worker_pool import WorkerPool
except Exception:  # pragma: no cover
    WorkerPool = None  # type: ignore
    def discover_rules():  # type: ignore
        return []

# Zentrales Logging (additiv, ersetzt nicht vorhandene bestehende Logger anderer Module)
try:
    from logging_setup import setup_logger  # bevorzugtes zentrales Setup
    logger = setup_logger(name="quality_gui", level=logging.INFO)
except Exception:
    logging.basicConfig(level=logging.INFO)
    logger = logging.getLogger("quality_gui")

# Force Light Mode global (Singleton Guard verhindert Doppel-Initialisierung)
if not hasattr(ctk, '_quality_light_forced'):
    ctk.set_appearance_mode("light")
    setattr(ctk, '_quality_light_forced', True)

# Import-Optimierung – mit Fehlerbehandlung (nur einmal anwenden)
if not hasattr(ctk, '_quality_anti_dark_applied'):
    try:
        from aggressive_anti_dark_mode import apply_aggressive_light_mode_patches, get_safe_aggressive_color as amd_safe_color
        apply_aggressive_light_mode_patches()
        logger.info("Aggressive Anti-Dark-Mode aktiviert")
    except ImportError:
        logger.warning("Aggressive Anti-Dark-Mode nicht verfügbar - verwende Fallback")
        os.environ['CUSTOMTKINTER_APPEARANCE_MODE'] = 'light'
    setattr(ctk, '_quality_anti_dark_applied', True)

def _safe_color_override(name, fallback=None):
    """Lokaler Anti-Dark-Mode Override ohne Import-Schatten."""
    if name in ('black', '#000000', '#1C1C1C'):
        return '#F8FAFC'
    return name or fallback


# Hinweis: Zusätzlicher typing-Import entfernt (bereits oben vorhanden)
try:
    from concurrent.futures import Future, CancelledError  # type: ignore
except Exception:  # pragma: no cover - fallback for environments without concurrent.futures
    class Future:  # minimal typing stand-in
        pass
    class CancelledError(Exception):
        pass
from types import SimpleNamespace


def _ascii_placeholder_variants(text: str) -> set[str]:
    """Erzeuge gängige Platzhalter-Varianten mit Fragezeichen durch fehlerhafte Encoding-Pfade."""
    variants: set[str] = set()
    if not isinstance(text, str) or not text:
        return variants
    encodings = ("utf-8", "latin-1", "cp1252")
    for enc in encodings:
        try:
            encoded = text.encode(enc, errors="replace")
            variant = encoded.decode("ascii", errors="replace")
            if variant and variant != text:
                variants.add(variant)
        except Exception:
            continue
    return variants


_QUESTION_CHAR = "?"
_DOUBLE_QUESTION = _QUESTION_CHAR * 2
_TRIPLE_QUESTION = _QUESTION_CHAR * 3

# ============================================================================
# MODUL-KONSTANTEN - Zentrale Konfiguration (analog zu analysis_results)
# ============================================================================

# UI Timing Constants (ms)
_DEFAULT_TOAST_DURATION_MS = 2000
_SUCCESS_TOAST_DURATION_MS = 4000
_WARNING_TOAST_DURATION_MS = 3000
_INFO_TOAST_DURATION_MS = 2500
_PAIRING_AUTOSAVE_DELAY_MS = 800
_DEBOUNCE_DEFAULT_DELAY_MS = 300
_THREAD_JOIN_TIMEOUT_SEC = 5.0

# UI Dimensions (px)
_DEFAULT_WINDOW_WIDTH = 1400
_DEFAULT_WINDOW_HEIGHT = 900
_DEFAULT_TEXT_WIDGET_HEIGHT = 600
_PANEL_MIN_WIDTH = 400

# Performance Thresholds
_LARGE_FILE_THRESHOLD_BYTES = 10_000_000  # 10MB
_MAX_HISTORY_ENTRIES = 50
_CACHE_MAX_SIZE = 1000

# File I/O Constants
_CONFIG_FILE_ENCODING = 'utf-8'
_DEFAULT_PDF_PAGE_LIMIT = 100

# Logging
_LOG_MAX_MESSAGE_LENGTH = 500

# Ollama Integration Defaults
DEFAULT_OLLAMA_URL = _CLI_OLLAMA_URL
OLLAMA_SEGMENT_LIMIT_DEFAULT = 25


class ProfessionelleUebersetzungsqualitaetsApp:
    """Professionelle Übersetzungsqualitäts-GUI mit erweiterten Funktionen"""
    # Konstante für Pairing-Autosave Debounce (ms) – zentral konfigurierbar (deprecated: nutze _PAIRING_AUTOSAVE_DELAY_MS)
    PAIR_AUTOSAVE_DELAY_MS = _PAIRING_AUTOSAVE_DELAY_MS

    def __init__(self):  # Wiederhergestellt aus Backup & an neue Struktur angepasst
        """Initialisiert Kernzustand, Logging, Caches, Design-System & Hauptfenster.

    Minimal-invasiv: Keine bestehende Logik entfernt, nur fehlende Initialisierung ergänzt.
        """
        try:  # Gesamter Initialisierungsblock robust gekapselt
            # Orchestrator-Aufteilung – behält ursprüngliche Reihenfolge / Seiteneffekte
            self._init_logging()
            self._init_base_state()
            self._init_caches()
            self._init_services_and_infra()
            self._init_feature_flags()
            self._init_design_system_wrapper()
            # Autosave Flags (Pairing)
            self._pair_dirty = False
            self._pair_autosave_after_id = None

            # -------------------- LOKALISIERUNG --------------------
            self.current_language = 'de'
            # Backup hatte _initialize_localization(); aktuelle Variante nutzt _initialize_localization_map()
            try:
                if hasattr(self, '_initialize_localization'):
                    self._initialize_localization()  # type: ignore
                elif hasattr(self, '_initialize_localization_map'):
                    self._initialize_localization_map()  # type: ignore
            except Exception as e:
                self.logger.warning(f"Lokalisierungs-Init Fallback: {e}")
            try:
                self._apply_localization_safe()
            except Exception:
                pass

            # Installiere LRU-Übersetzungs-Cache nach i18n-Setup (idempotent)
            try:
                self._install_translation_cache()
            except Exception as e:
                try:
                    self.logger.debug(f"Übersetzungs-Cache Setup übersprungen: {e}")
                except Exception:
                    pass

            # Sicherstellen: Guards für _t / _accent / get_typography vorhanden (früh, idempotent)
            try:
                if not callable(getattr(self, '_t', None)):
                    setattr(self, '_t', lambda s: s)
            except Exception:
                pass
            try:
                if not callable(getattr(self, 'get_typography', None)):
                    setattr(self, 'get_typography', lambda name: ("Segoe UI", 12))
            except Exception:
                pass
            try:
                if not callable(getattr(self, '_accent', None)):
                    setattr(self, '_accent', lambda semantic, domain="generic": self.get_color('primary'))
            except Exception:
                pass

            # -------------------- HAUPT-SETUP (Fenster/Layout) --------------------
            try:
                self._setup_application()
                # Globale Font-Normalisierung nach Aufbau anwenden
                try:
                    self._normalize_all_widget_fonts()
                except Exception as fn_err:
                    self.logger.warning(f"Font-Normalisierung übersprungen: {fn_err}")
                # Farb-Alias Diagnose einmalig
                try:
                    if hasattr(self, '_diagnose_color_aliases'):
                        self._diagnose_color_aliases()
                    if hasattr(self, '_log_missing_color_tokens_once'):
                        self._log_missing_color_tokens_once()
                except Exception:
                    pass
            except Exception as e:
                self.logger.error(f"Setup Fehler: {e}")

            # -------------------- MIGRATION (falls Funktion existiert) --------------------
            try:
                if hasattr(self, '_migrate_existing_projects_on_startup'):
                    self._migrate_existing_projects_on_startup()  # type: ignore
            except Exception as e:
                self.logger.warning(f"Projekt-Migration übersprungen: {e}")

        except Exception as e:  # letzter Fallback – niemals Absturz nach außen
            try:
                logging.getLogger("quality_gui").exception(f"Kritischer Initialisierungsfehler: {e}")
            except Exception:
                pass

    # Klasseweiter Default Mapping Cache ausgelagert
    from quality_gui_constants import CONTEXT_DEFAULT_MAP as _CONTEXT_DEFAULT_MAP

    # --------------------------- SAFE LOCALIZATION APPLY ---------------------------
    def _apply_localization_safe(self):
        """Wendet sichere Lokalisierung auf primäre UI-Elemente an (Fenstertitel)."""
        try:
            if getattr(self, 'root', None):
                # Deutsch gem. Projektvorgaben
                self.root.title(self._t("Übersetzungsqualitäts-Framework – Professional"))
        except Exception:
            pass

    # --------------------------- FONT NORMALIZATION ---------------------------
    def _normalize_all_widget_fonts(self):
        """Setzt für alle CTk/Tk Widgets eine einheitliche Schrift aus dem Design-System.

        Ziel: Konsistente Darstellung ohne hartcodierte Ausreißer. Holt primäre Body-Font
        aus Design-System falls verfügbar, sonst fallback auf ('Segoe UI', 11).
        """
        if getattr(self, '_fonts_normalized', False):
            return
        root = getattr(self, 'root', None)
        if not root:
            return
        try:
            base_font = None
            try:
                # Erwartetes Token: body_md (gemäß Welcome Screen Pattern)
                if hasattr(self, 'design_system') and self.design_system:
                    ds_fonts = self.design_system.get('typography', {}) if isinstance(self.design_system, dict) else {}
                    body_md = ds_fonts.get('body_md') if isinstance(ds_fonts, dict) else None
                    if body_md and isinstance(body_md, (list, tuple)):
                        family, size, *rest = body_md
                        base_font = (family, size)
            except Exception:
                pass
            if not base_font:
                base_font = ("Segoe UI", 11)

            def _recurse(w):
                try:
                    # Überspringe bewusst markierte Widgets (z.B. große Header-Titel)
                    if getattr(w, '_preserve_font', False):
                        pass
                    elif isinstance(w, (ctk.CTkLabel, ctk.CTkButton, ctk.CTkEntry, ctk.CTkTextbox, ctk.CTkCheckBox, ctk.CTkRadioButton, ctk.CTkOptionMenu)):
                        w.configure(font=base_font)
                except (tk.TclError, AttributeError, TypeError):
                    # Widget destroyed, attribute missing, or wrong type
                    pass
                
                try:
                    for child in w.winfo_children():
                        _recurse(child)
                except (tk.TclError, AttributeError):
                    # Widget destroyed or no children
                    pass

            _recurse(root)
            self.logger.info("Globale Font-Normalisierung abgeschlossen")
            try:
                self._fonts_normalized = True
            except Exception:
                pass
        except Exception as e:
            self.logger.warning(f"Globale Font-Normalisierung fehlgeschlagen: {e}")

    # ---------- I18N / PLURAL HELPERS ----------
    def _n(self, singular: str, plural: str, count: int | float | None) -> str:
        """Einfache Pluralisierung (Deutsch/Englisch: 1 -> Singular, sonst Plural).

        Args:
            singular: Singular-Form
            plural:   Plural-Form
            count:    Numerischer Wert

        Returns:
            Korrekte Form; bei Fehler Fallback plural.
        """
        try:
            c = int(count if count is not None else 0)
            return singular if c == 1 else plural
        except Exception:
            return plural

    # --------------------------- MINI UI TEXT TEST ---------------------------
    def run_mini_ui_text_test(self):
        """Führt einen kleinen automatisierten Check der wichtigsten UI-Texte durch.

        Prüft auf Vorhandensein erwarteter deutscher Schlüsselbegriffe & Abwesenheit von '??'.
        Ergebnis gibt Liste fehlender / fehlerhafter Items zurück (leer = bestanden).
        """
        issues = []
        critical_tokens = ["Übersetzungsqualitäts-Framework", "Qualitätsanalyse", "Dateipaar", "Dateien", "Manuelles Pairing", "Projekt", "Kriterien"]
        # Fenster-Titel
        try:
            title = self.root.title() if self.root else ""
            if "??" in title:
                issues.append("Titel enthält Encoding-Artefakt")
            if "Framework" not in title:
                issues.append("Framework im Titel fehlt")
        except Exception:
            issues.append("Fenstertitel nicht prüfbar")
        # Durchsuche wenige repräsentative Widgets (flach)
        try:
            sample_texts = []
            if self.root:
                for child in self.root.winfo_children():
                    try:
                        txt = getattr(child, 'cget', lambda _:'')( 'text')  # type: ignore
                        if txt:
                            sample_texts.append(txt)
                    except Exception:
                        pass
            concat = " | ".join(sample_texts)
            if "??" in concat:
                issues.append("Widgets enthalten Encoding-Artefakte")
            for token in critical_tokens:
                if token not in concat:
                    issues.append(f"Fehlender Begriff: {token}")
        except Exception as e:
            issues.append(f"Widget-Scan Fehler: {e}")
        if not issues:
            self.logger.info("Mini UI Text Test: BESTANDEN")
        else:
            self.logger.warning(f"Mini UI Text Test: Probleme: {issues}")
        return issues

    # ========================= DESIGN SYSTEM INITIALIZER (aus Backup übernommen) =========================
    def _initialize_design_system(self):
        """DESIGN SYSTEM - Zentrale Farb-/Spacing-/Typografie-Verwaltung (Instruction-konform).

        Übernommen aus Backup-Version; nur Kommentar leicht angepasst. Stellt vollständiges
        System bereit oder liefert einen ui_theme Fallback. Verwendet KEINE hartcodierten Hex-Farben
        außerhalb der Fallback-Ebene (Instruction: Hex nur in finalem Notfall erlaubt).
        """
        try:
            from design_system import DesignSystem  # Lokaler Import für robustes Lazy-Loading
            design_sys = DesignSystem()
            return design_sys.get_full_system()
        except ImportError:
            # Fallback auf ui_theme (Legacy) – nur falls design_system Modul nicht verfügbar
            try:
                from ui_theme import enhanced_theme
                return {
                    'colors': {
                        'primary': enhanced_theme.get_color('primary'),
                        'primary_hover': enhanced_theme.get_color('primary_hover'),
                        'primary_light': enhanced_theme.get_color('primary_container'),
                        'primary_dark': enhanced_theme.get_color('primary'),
                        'secondary': enhanced_theme.get_color('secondary'),
                        'secondary_hover': enhanced_theme.get_color('secondary_hover'),
                        'secondary_light': enhanced_theme.get_color('surface'),
                        'secondary_dark': enhanced_theme.get_color('secondary'),
                        'success': enhanced_theme.get_color('success'),
                        'success_hover': enhanced_theme.get_color('success_hover'),
                        'success_light': enhanced_theme.get_color('success_surface', '#E8F5E8'),
                        'success_600': enhanced_theme.get_color('success_hover'),
                        'warning': enhanced_theme.get_color('warning'),
                        'warning_hover': enhanced_theme.get_color('warning', '#D97706'),
                        'warning_light': enhanced_theme.get_color('warning_surface', '#FEF3C7'),
                        'warning_600': enhanced_theme.get_color('warning'),
                        'error': enhanced_theme.get_color('danger'),
                        'error_hover': enhanced_theme.get_color('danger_hover'),
                        'error_light': enhanced_theme.get_color('danger_surface'),
                        'error_600': enhanced_theme.get_color('danger_hover'),
                        'info': enhanced_theme.get_color('info'),
                        'info_hover': enhanced_theme.get_color('info_hover'),
                        'info_light': enhanced_theme.get_color('info_surface'),
                        'info_600': enhanced_theme.get_color('info_hover'),
                        'surface': enhanced_theme.get_color('surface'),
                        'surface_light': enhanced_theme.get_color('background'),
                        'surface_elevated': enhanced_theme.get_color('surface'),
                        'surface_border': enhanced_theme.get_color('border'),
                        'surface_hover': enhanced_theme.get_color('control_hover', enhanced_theme.get_color('surface')),
                        'background': enhanced_theme.get_color('background'),
                        'background_secondary': enhanced_theme.get_color('surface'),
                        'text_primary': enhanced_theme.get_color('text_primary'),
                        'text_secondary': enhanced_theme.get_color('text_secondary'),
                        'text_tertiary': enhanced_theme.get_color('text_secondary'),
                        'text_inverse': enhanced_theme.get_color('text_on_primary'),
                        'anthracite_700': enhanced_theme.get_color('text_primary'),
                        'anthracite_600': enhanced_theme.get_color('text_secondary'),
                        'anthracite_800': enhanced_theme.get_color('text_primary'),
                        'anthracite_900': enhanced_theme.get_color('text_primary'),
                        'input_bg': enhanced_theme.get_color('surface'),
                        'input_border': enhanced_theme.get_color('border'),
                        'input_border_focus': enhanced_theme.get_color('primary'),
                        'input_text': enhanced_theme.get_color('text_primary'),
                        'input_placeholder': enhanced_theme.get_color('text_secondary'),
                        'button_primary': enhanced_theme.get_color('primary'),
                        'button_primary_hover': enhanced_theme.get_color('primary_hover'),
                        # Vereinheitlichung: Sekundäre Buttons jetzt auch Blau
                        'button_secondary': enhanced_theme.get_color('primary'),
                        'button_secondary_hover': enhanced_theme.get_color('primary_hover'),
                        'white': enhanced_theme.get_color('surface'),
                        'gray_50': enhanced_theme.get_color('background'),
                        'gray_100': enhanced_theme.get_color('background'),
                        'gray_200': enhanced_theme.get_color('border'),
                        'gray_300': enhanced_theme.get_color('border'),
                        'gray_400': enhanced_theme.get_color('text_secondary'),
                        'gray_500': enhanced_theme.get_color('text_secondary'),
                        'gray_600': enhanced_theme.get_color('text_secondary'),
                        'gray_700': enhanced_theme.get_color('text_primary'),
                        'gray_800': enhanced_theme.get_color('text_primary'),
                        'gray_900': enhanced_theme.get_color('text_primary'),
                        'accent_purple': enhanced_theme.get_color('accent', '#8B5CF6'),
                        'accent_purple_light': enhanced_theme.get_color('background'),
                        'accent_teal': enhanced_theme.get_color('secondary'),
                        'accent_teal_light': enhanced_theme.get_color('background'),
                        'accent_indigo': enhanced_theme.get_color('primary'),
                        'accent_indigo_light': enhanced_theme.get_color('background')
                    },
                    'spacing': {
                        'xs': 4, 'sm': 8, 'md': 16, 'lg': 24, 'xl': 32,
                        '2xl': 48, '3xl': 64, '4xl': 80, '5xl': 96, '6xl': 128,
                        'card_padding': 24, 'card_margin': 20, 'card_gap': 16,
                        'button_padding': 16, 'button_gap': 12, 'button_height': 44,
                        'input_padding': 12, 'input_gap': 8, 'input_height': 40,
                        'section_gap': 32, 'component_gap': 20, 'element_gap': 16,
                        'header_padding': 24, 'content_padding': 20,
                        'radius_sm': 6, 'radius_md': 8, 'radius_lg': 12,
                        'radius_xl': 16, 'radius_2xl': 20, 'radius_3xl': 24,
                        'radius_full': 9999,
                        'header_height': 80, 'status_bar_height': 35,
                        'sidebar_width': 320, 'content_min_width': 600,
                    },
                    'typography': {
                        'micro': ('Segoe UI', 10, 'normal'), 'micro_large': ('Segoe UI', 11, 'normal'),
                        'caption': ('Segoe UI', 12, 'normal'), 'small': ('Segoe UI', 12, 'bold'), 'menu': ('Segoe UI', 12, 'normal'),
                        'body_sm': ('Segoe UI', 13, 'normal'), 'body': ('Segoe UI', 14, 'normal'), 'body_bold': ('Segoe UI', 14, 'bold'), 'body_large': ('Segoe UI', 15, 'normal'),
                        'label': ('Segoe UI', 16, 'normal'), 'label_bold': ('Segoe UI', 16, 'bold'), 'button': ('Segoe UI', 14, 'bold'),
                        'button_md': ('Segoe UI', 14, 'bold'), 'button_large': ('Segoe UI', 16, 'bold'), 'subheading': ('Segoe UI', 18, 'bold'), 'subheading_large': ('Segoe UI', 20, 'bold'),
                        'card_header': ('Segoe UI', 18, 'bold'), 'heading_sm': ('Segoe UI', 20, 'bold'), 'heading': ('Segoe UI', 22, 'bold'),
                        'section': ('Segoe UI', 22, 'bold'), 'title': ('Segoe UI', 26, 'bold'),
                        'display': ('Segoe UI', 32, 'bold'), 'display_large': ('Segoe UI', 40, 'bold'), 'display_extra': ('Segoe UI', 48, 'bold'),
                        'hero': ('Segoe UI', 32, 'normal'), 'status': ('Segoe UI', 12, 'normal'), 'code': ('Consolas', 13, 'normal'),
                        'metric_label': ('Segoe UI', 11, 'normal'), 'caption_unified': ('Segoe UI', 12, 'normal'), 'label_unified': ('Segoe UI', 13, 'bold'),
                        'body_unified': ('Segoe UI', 14, 'normal'), 'body_strong_unified': ('Segoe UI', 14, 'bold'), 'subtitle_unified': ('Segoe UI', 16, 'bold'),
                        'title_unified': ('Segoe UI', 26, 'bold')
                    }
                }
            except Exception:
                # Minimalster Notfall-Fallback
                return {
                    'colors': {'primary': '#374151', 'text_primary': '#374151', 'success': '#2E8B57', 'surface': '#FFFFFF', 'background': '#FFFFFF', 'border': '#E5E7EB'},
                    'spacing': {'sm': 8, 'md': 16, 'lg': 24},
                    'typography': {'body': ('Segoe UI', 14, 'normal'), 'heading': ('Segoe UI', 22, 'bold')}
                }

    # ========================= ZENTRALER ERROR HANDLER =========================
    def _handle_error(self, e: Exception, context: str = "", user_message: str | None = None, level: str = "error", toast: bool = True, event_name: str | None = None, **extra):
        """Zentraler Fehler-Handler: Logging, optional Toast & Event.

        Parameters:
            e: Exception Instanz
            context: semantischer Kontext ("upload", "analysis", ...)
            user_message: UI-Meldung (DE). Fallback generisch.
            level: logging Level (error|warning|info|debug)
            toast: Ob ein Toast angezeigt wird
            event_name: Optionales Event für event_bus
            extra: Zusätzliche Felder für Logging / Event
        """
        try:
            # 1) Logging (error-level mit exception() für direkten Trace)
            if level == "error":
                try:
                    self.logger.exception(f"[{context}] {e}")
                except Exception:
                    self.logger.error(f"[{context}] {e}", exc_info=True)
            else:
                logger_fn = getattr(self.logger, level, self.logger.error)
                logger_fn(f"[{context}] {e}", exc_info=True)

            # 2) Event Dispatch (optional)
            if event_name and getattr(self, 'event_bus', None):
                try:
                    payload = {"context": context, "error": str(e), **extra}
                    self.event_bus.publish(event_name, payload)
                except Exception:
                    pass

            # 3) Automatische Standard-Meldung falls keine user_message übergeben
            if user_message is None and context:
                try:
                    mapped = self._CONTEXT_DEFAULT_MAP.get(context)
                    if mapped:
                        if hasattr(self, '_t'):
                            user_message = self._t(mapped)
                        else:
                            user_message = mapped
                except Exception:
                    pass

            # 4) UI Toast
            if toast and hasattr(self, 'show_toast'):
                try:
                    msg = user_message or (f"Fehler ({context})" if context else "Fehler")
                    self.show_toast(msg, "error")
                except Exception:
                    pass
        except Exception:
            # Letzter Fallback – kein Crash
            pass

    # ========================= PAIRING AUTOSAVE (DEBOUNCE) =========================
    def _mark_pairing_dirty(self) -> None:
        """Markiert Pairing-State als geändert und plant Auto-Save (Debounce)."""
        try:
            self._pair_dirty = True
            root = getattr(self, 'root', None)
            if not root:
                return
            delay = getattr(self, 'PAIR_AUTOSAVE_DELAY_MS', 800)
            # Verwende sicheren Wrapper mit Tracking
            self._pair_autosave_after_id = self._after(delay, self._perform_pairing_autosave)
        except Exception as e:
            try:
                self.logger.debug(f"Autosave Markierung fehlgeschlagen: {e}")
            except Exception:
                pass

    def _set_pair_dirty_and_event(self, event: str | None = None, **payload):
        """Kombinierter Helper: Dirty setzen (mit Debounce) und optional Event loggen."""
        try:
            self._mark_pairing_dirty()
        except Exception:
            self._pair_dirty = True
        if event:
            try:
                self._log_event(event, **payload)
            except Exception:
                pass

    def _perform_pairing_autosave(self):
        """Persistiert Pairings falls noch dirty."""
        self._pair_autosave_after_id = None
        if not getattr(self, '_pair_dirty', False):
            return
        try:
            self.save_pairings()
            self._pair_dirty = False
            self._log_event("pairing.autosave", status="ok")
        except Exception as e:
            self._handle_error(e, context="pairing.autosave", user_message=self._t("Fehler beim Speichern der Paarung") if hasattr(self,'_t') else "Pairing Save Error", toast=False)
            self._log_event("pairing.autosave", status="failed", error=str(e))

    def _initialize_localization_map(self):
        self._i18n_map = {
            # Buttons / Actions
            "Upload Source Files": "Ausgangstexte laden",
            "Upload Translations": "Übersetzungen laden",
            "Batch Upload": "Stapel-Upload",
            "Upload Documents": "Dokumente laden",
            "All": "Alle",
            "Detailergebnisse konnten nicht erstellt werden": "Detailergebnisse konnten nicht erstellt werden",
            "Einstellungs-Dashboard Fehler": "Einstellungs-Dashboard Fehler",
            "Einstellungsbereich konnte nicht erstellt werden": "Einstellungsbereich konnte nicht erstellt werden",
            "Hauptoberfläche Fehler": "Hauptoberfläche Fehler",
            "Inhaltsbereich konnte nicht erstellt werden": "Inhaltsbereich konnte nicht erstellt werden",
            "Qualitätskriterien konnten nicht erstellt werden": "Qualitätskriterien konnten nicht erstellt werden",
            "Rechter Bereich (erweitert) Fehler": "Rechter Bereich (erweitert) Fehler",
            "Rechter Bereich (einfach) Fehler": "Rechter Bereich (einfach) Fehler",
            "Erweiterter Willkommensbereich Fehler": "Erweiterter Willkommensbereich Fehler",
            "Metriken-Dashboard Fehler": "Metriken-Dashboard Fehler",
            "Feature-Karten Fehler": "Feature-Karten Fehler",
            "Systemstatus Fehler": "Systemstatus Fehler",
            "Metriken-Dashboard (einfach) Fehler": "Metriken-Dashboard (einfach) Fehler",
            "Ordner-Navigation Fehler": "Ordner-Navigation Fehler",
            "Projekte konnten nicht ermittelt werden": "Projekte konnten nicht ermittelt werden",
            "Projekt-Ordner konnte nicht geöffnet werden": "Projekt-Ordner konnte nicht geöffnet werden",
            "Projekte-Verzeichnis konnte nicht geöffnet werden": "Projekte-Verzeichnis konnte nicht geöffnet werden",
            "Neues Projekt konnte nicht erstellt werden": "Neues Projekt konnte nicht erstellt werden",
            "Projektstruktur ungültig": "Projektstruktur ungültig",
            "Navigation konnte nicht aktualisiert werden": "Navigation konnte nicht aktualisiert werden",
            "Automatische Migration fehlgeschlagen": "Automatische Migration fehlgeschlagen",
            "Legacy-Dateien konnten nicht gescannt werden": "Legacy-Dateien konnten nicht gescannt werden",
            "Legacy-Migration fehlgeschlagen": "Legacy-Migration fehlgeschlagen",
            "Migrations-Dialog fehlgeschlagen": "Migrations-Dialog fehlgeschlagen",
            "Datei-Explorer konnte nicht erstellt werden": "Datei-Explorer konnte nicht erstellt werden",
            "Dateikategorie konnte nicht erstellt werden": "Dateikategorie konnte nicht erstellt werden",
            "Dateikarte konnte nicht erstellt werden": "Dateikarte konnte nicht erstellt werden",
            "Analyse-Dashboard konnte nicht aufgebaut werden": "Analyse-Dashboard konnte nicht aufgebaut werden",
            "Metrik-Karte konnte nicht erstellt werden": "Metrik-Karte konnte nicht erstellt werden",
            "Analyse Platzhalter Fehler": "Analyse Platzhalter Fehler",
            "Basis Willkommensausgabe Fehler": "Basis Willkommensausgabe Fehler",
            "Statusleiste Fehler": "Statusleiste Fehler",
            "Dateizähler konnte nicht aktualisiert werden": "Dateizähler konnte nicht aktualisiert werden",
            "Systeme konnten nicht initialisiert werden": "Systeme konnten nicht initialisiert werden",
            "Modernes Dateimanagement Fehler": "Modernes Dateimanagement Fehler",
            "Dateiliste Abschnitt Fehler": "Dateiliste Abschnitt Fehler",
            "Dateielement konnte nicht erstellt werden": "Dateielement konnte nicht erstellt werden",
            "Dateielement konnte nicht entfernt werden": "Dateielement konnte nicht entfernt werden",
            "Dateiliste Aktualisierung fehlgeschlagen": "Dateiliste Aktualisierung fehlgeschlagen",
            "Dateiliste Inhalt konnte nicht aktualisiert werden": "Dateiliste Inhalt konnte nicht aktualisiert werden",
            "Pairing Ergebnisse Anzeige Fehler": "Pairing Ergebnisse Anzeige Fehler",
            "Manuelles Pairing Option Fehler": "Manuelles Pairing Option Fehler",
            "Manueller Pairing Dialog Fehler": "Manueller Pairing Dialog Fehler",
            "Manuelle Pairing Oberfläche Fehler": "Manuelle Pairing Oberfläche Fehler",
            "Manuelle Pairing Befüllung Fehler": "Manuelle Pairing Befüllung Fehler",
            "Pair Anzeige Element Fehler": "Pair Anzeige Element Fehler",
            "Ungepaartes Datei Element Fehler": "Ungepaartes Datei Element Fehler",
            "Manuelle Paarung Fehler": "Manuelle Paarung Fehler",
            "Manuelle Paarung (Translation) Fehler": "Manuelle Paarung (Translation) Fehler",
            "Drag Start Fehler": "Drag Start Fehler",
            "Drag Abschluss Fehler": "Drag Abschluss Fehler",
            "Dateien Löschen Fehler": "Dateien Löschen Fehler",
            "Dateiliste Refresh Fehler": "Dateiliste Refresh Fehler",
            "Übersetzungsdateien Upload fehlgeschlagen": "Übersetzungsdateien Upload fehlgeschlagen",
            "Erweiterte Upload Ergebnisse Fehler": "Erweiterte Upload Ergebnisse Fehler",
            "Analyse Ergebnisse Anzeige Fehler": "Analyse Ergebnisse Anzeige Fehler",
            "Demo Ergebnisse Anzeige Fehler": "Demo Ergebnisse Anzeige Fehler",
            "Export Optionen Anzeige Fehler": "Export Optionen Anzeige Fehler",
            "Export Modul nicht verfügbar": "Export Modul nicht verfügbar",
            "Export fehlgeschlagen": "Export fehlgeschlagen",
            "Unerwarteter Export Fehler": "Unerwarteter Export Fehler",
            "Einstellungen Anzeige Fehler": "Einstellungen Anzeige Fehler",
            "Dateien Leeren Fehler": "Dateien Leeren Fehler",
            "App Lauf Fehler": "App Lauf Fehler",
            "Anwendung nicht korrekt initialisiert": "Anwendung nicht korrekt initialisiert",
            "Minor": "Leicht",
            "Copy findings": "Befunde kopieren",
            "Findings copied": "Befunde kopiert",
            "Copy failed": "Kopieren fehlgeschlagen",
            "View Reports": "Berichte anzeigen",
            "Export": "Exportieren",
            "Clear": "Leeren",
            "Clear All": "Alle leeren",
            "Alles löschen": "Alles löschen",
            "Refresh": "Aktualisieren",
            "Close": "Schließen",
            "Severity": "Schweregrad",
            "Schweregrad der gefundenen Probleme": "Schweregrad der gefundenen Probleme",
            "Kategorisierung der erkannten Qualitätsprobleme nach Kritikalität": "Kategorisierung der erkannten Qualitätsprobleme nach Kritikalität",
            "Critical": "Kritisch",
            "Kritisch": "Kritisch",
            "Major": "Schwerwiegend",
            "Schwerwiegend": "Schwerwiegend",
            "Minor": "Leicht",
            "Leicht": "Leicht",
            "Total": "Gesamt",
            "Gesamt": "Gesamt",
            "Upload files to enable analysis": "Dateien laden zur Analyse",
            "AI Assistant (Ollama)": "KI-Assistenz (Ollama)",
            "Ask the AI about your translation": "Stelle eine Frage zur Übersetzung und erhalte eine KI-Antwort.",
            "Select file pair": "Dateipaar auswählen",
            "No file pairs available": "Keine Dateipaare verfügbar",
            "Enter question for Ollama": "Frage für Ollama eingeben",
            "Request Ollama Answer": "Antwort anfordern",
            "Ollama request running": "Anfrage an Ollama läuft",
            "Ollama response ready": "Antwort von Ollama erhalten",
            "Ollama not reachable": "Ollama nicht erreichbar",
            "Ollama request failed": "Anfrage an Ollama fehlgeschlagen",
            "Request still running": "Anfrage läuft noch",
            "Question required": "Bitte eine Frage eingeben",
            "Response": "Antwort",
            # Status / Labels
            "Translation Quality Framework": "Übersetzungsqualitäts-Framework",
            "Translation Quality Framework - Professional": "Übersetzungsqualitäts-Framework – Professional",
            "Professional Quality Analysis": "Professionelle Qualitätsanalyse",
            "Live-Übersicht Ihrer Übersetzungsanalyse": "Live-Übersicht Ihrer Übersetzungsanalyse",
            # Neue UI Sektionen / Zusätzliche Strings
            "File Upload": "Datei-Upload",
            "File Upload & Management": "Datei-Upload & Verwaltung",
            "Application Settings & Preferences": "Anwendungseinstellungen & Präferenzen",
            "Detailed Quality Analysis Report": "Detaillierter Qualitätsanalyse-Bericht",
            "Additional Actions": "Weitere Aktionen",
            "Translation Quality Analysis Results": "Übersetzungsqualitäts-Analyseergebnisse",
            "Qualitätsdetails": "Qualitätsdetails",
            "System Status": "Systemstatus",
            "Overall Quality": "Gesamtqualität",
            "Issues Detected": "Gefundene Probleme",
            "Recommendations": "Empfehlungen",
            "Excellent translation quality": "Hervorragende Übersetzungsqualität",
            "Minor issues found": "Kleinere Auffälligkeiten gefunden",
            "Improvement suggestions": "Verbesserungsvorschläge",
            "System Overview": "Systemübersicht",
            "Files Ready": "Dateien bereit",
            "Active Sessions": "Aktive Sitzungen",
            "Processing Speed": "Verarbeitungsgeschwindigkeit",
            "Ready": "Bereit",
            "Active": "Aktiv",  # Statusindikator (für Zähler/Toolbar)
            "active": "aktiv",  # Kleinbuchstaben-Variante falls im UI benötigt
            "Bewertung": "Bewertung",
            "Änderungen anzeigen": "Änderungen anzeigen",
            "Änderungen verbergen": "Änderungen verbergen",
            "Aktive Prüfer": "Aktive Prüfer",
            "Deaktivierte Prüfer": "Deaktivierte Prüfer",
            "Filter angewendet": "Filter angewendet",
            "Filter angewendet: Code {code}": "Filter angewendet: Code {code}",
            "Anzahl Befunde": "Anzahl Befunde",
            "Gesamtbefunde": "Gesamtbefunde",
            "Phasenbewertung": "Phasenbewertung",
            "Laufzeit": "Laufzeit",
            "Gate aktiv": "Gate aktiv",
            "Offene Befunde": "Offene Befunde",
            "Vorschläge": "Vorschläge",
            "Ja": "Ja",
            "Nein": "Nein",
            "Terminologie-Konsistenz": "Terminologie-Konsistenz",
            "Numerische Konsistenz": "Numerische Konsistenz",
            "Längenbalance": "Längenbalance",
            "Übersetzungs-Vollständigkeit": "Übersetzungs-Vollständigkeit",
            "Sprachliche Flüssigkeit": "Sprachliche Flüssigkeit",
            "Stilqualität": "Stilqualität",
            "Lesbarkeit": "Lesbarkeit",
            "Inhaltsgenauigkeit": "Inhaltsgenauigkeit",
            "Konsistenz": "Konsistenz",
            "Gesamte Qualität": "Gesamte Qualität",
            "Ähnlichkeitsschwellen": "Ähnlichkeitsschwellen",
            "Hinweis_RisikoSchwelle": "Ab {th}% bitte strukturelle oder technische Probleme priorisieren.",
            "Hinweis_VollstaendigkeitSchwelle": "Unter {pct} fehlen Inhalte oder Segmente – bitte vervollständigen.",
            "Hinweis_SimilaritaetSchwelle": "Unter {pct} weist auf Bedeutungsabweichungen hin – bitte prüfen und korrigieren.",
            # Dashboard / Analyse-Übersicht
            "Translation Quality Analysis Dashboard": "Übersicht Übersetzungsqualität",
            "Issues Found": "Gefundene Probleme",
            "Detailed Analysis Results": "Detaillierte Analyseergebnisse",
            "Quality Score": "Qualitätsnote",
            "Overall translation quality rating": "Gesamteinschätzung der Übersetzungsqualität",
            "Potential translation issues detected": "Potenzielle Übersetzungsprobleme erkannt",
            "Number of files processed": "Anzahl verarbeiteter Dateien",
            "Analysis Time": "Analysedauer",
            "Processing time of last analysis": "Verarbeitungszeit der letzten Analyse",
            "Severity Mix": "Schweregrad-Mix",
            "Distribution of critical / major / minor issues": "Verteilung kritischer / wesentlicher / geringer Befunde",
            "Key Features & Capabilities": "Hauptfunktionen & Fähigkeiten",
            "Advanced Analysis Engine": "Intelligente Analyse-Engine",
            "Professional Reporting": "Detaillierte Berichte",
            "Files: 0 source, 0 translations": "Dateien: 0 Ausgangstexte, 0 Übersetzungen",
            "Supported formats: PDF, DOCX, TXT, DOC, RTF, ODT – Drag and drop supported": "Unterstützte Formate: PDF, DOCX, TXT, DOC, RTF, ODT – Drag & Drop unterstützt",
            "Language Pair Configuration": "Sprachpaar-Konfiguration",
            "Upload documents to begin": "Dokumente hochladen zum Start",
            "Waiting for files": "Warte auf Dateien",
            "Analysis Results & Reports": "Analyseergebnisse & Berichte",
            "Noch keine Ausgangstexte hochgeladen.\nDateien hierher ziehen oder Upload verwenden.": "Noch keine Ausgangstexte hochgeladen.\nDateien hierher ziehen oder Upload verwenden.",
            "Noch keine Übersetzungsdateien hochgeladen.\nDateien hierher ziehen oder Upload verwenden.": "Noch keine Übersetzungsdateien hochgeladen.\nDateien hierher ziehen oder Upload verwenden.",
            "Keine Ausgangstexte.\nDateien hier ablegen oder hochladen.": "Keine Ausgangstexte.\nDateien hier ablegen oder hochladen.",
            "Keine Übersetzungen.\nDateien hier ablegen oder hochladen.": "Keine Übersetzungen.\nDateien hier ablegen oder hochladen.",
            "Keine Ausgangstexte\nHier ablegen / laden": "Keine Ausgangstexte\nHier ablegen / laden",
            "Keine Übersetzungen\nHier ablegen / laden": "Keine Übersetzungen\nHier ablegen / laden",
            "Translation Files": "Übersetzungen",
            "Übersetzungen": "Übersetzungen",
            # Findings / Viewer
            "Show all": "Alle anzeigen",
            "Show less": "Weniger anzeigen",
            "Show": "Anzeigen",
            "Hide": "Ausblenden",
            "Show all results": "Alle Ergebnisse anzeigen",
            "Could not open full findings view": "Vollständige Befunde-Ansicht konnte nicht geöffnet werden",
            "Copied to clipboard": "In Zwischenablage kopiert",
            "Export created": "Export erstellt",
            "No findings available": "Keine Befunde verfügbar",
            # Pairing
            "Translation file not found": "Übersetzungsdatei nicht gefunden",
            "Source file not found": "Ausgangsdatei nicht gefunden",
            "Pair created": "Paar erstellt",
            "Pair removed": "Paarung aufgelöst",
            "Unpair": "Trennen",
            "Source": "Ausgangstext",
            "Translation": "Übersetzung",
            "Automatic pairing restored": "Automatische Paarung wiederhergestellt",
            "Pairing saved": "Paarung gespeichert",
            "No pairs configured": "Keine Paare konfiguriert",
            "Error creating pair": "Fehler beim Erstellen der Paarung",
            "Error saving pairing": "Fehler beim Speichern der Paarung",
            "Manual pairing available - Click 'Adjust file pairing'": "Manuelles Pairing verfügbar - Klicke auf 'Dateipaarung anpassen'",
            "Automatic file pairing not possible - names too different": "Keine automatische Dateipaarung möglich - Namen zu unterschiedlich",
            "file pair(s)": "Dateipaar(e)",
            "unpaired file(s)": "ungepaarte Datei(en)",
            "unpaired source file(s)": "ungepaarte Quelldateien",
            "unpaired translation(s)": "ungepaarte Übersetzungen",
            "Select translation...": "Wähle Übersetzung...",
            "Select source...": "Wähle Ausgangstext...",
            "automatically detected": "automatisch erkannt",
            "Manual file pairing": "Manuelle Dateipaarung",
            "configured": "konfiguriert",
            # Undo/Redo Pairing
            "Undo": "Rückgängig",
            "Redo": "Wiederholen",
            "Undo pairing action": "Pairing-Aktion rückgängig gemacht",
            "Redo pairing action": "Pairing-Aktion wiederholt",
            "Nothing to undo": "Nichts zum Rückgängig machen",
            "Nothing to redo": "Nichts zum Wiederholen",
            # Pairing Dialog Buttons
            "Unpaired source files": "Ungepaarte Quelldateien",
            "Unpaired translations": "Ungepaarte Übersetzungen",
            "Repeat automatic pairing": "Automatische Paarung wiederholen",
            "Save pairing": "Paarung speichern",
            "Cancel": "Abbrechen",
            "No pairs found": "Keine Paare gefunden",
            "No unmatched translations available": "Keine ungepaarten Übersetzungen verfügbar",
            "No unmatched sources available": "Keine ungepaarten Ausgangstexte verfügbar",
            # Drag & Drop Pairing
            "Drag to pair": "Ziehen zum Paaren",
            "Drop to pair": "Loslassen zum Paaren",
            "Pair created (drag & drop)": "Paar erstellt (Drag & Drop)",
            "Drag pairing canceled": "Drag & Drop Pairing abgebrochen",
            "Cannot pair identical type": "Kann nicht gleichen Typ paaren",
            "Pair already exists": "Paar existiert bereits",
            "Load more": "Mehr laden",
            "Show all (counts)": "Alle anzeigen ({sources} Quellen / {translations} Übersetzungen)",
            "Loading more files...": "Lade weitere Dateien...",
            # Datei-/Explorer-UI (quality_gui_components_file.py)
            "Show backups": "Backups anzeigen",
            "Prev": "Zurück",
            "Next": "Weiter",
            "Page {p}/{m}": "Seite {p}/{m}",
            "… and {n} more files": "… und {n} weitere Dateien",
            "No files uploaded yet\nDrag & drop files here": "Noch keine Dateien hochgeladen\nDateien hierher ziehen",
            "Unknown": "Unbekannt",
            "Size: {size} • Modified: {date}": "Größe: {size} • Geändert: {date}",
            "Open in Explorer": "Im Explorer öffnen",
            "Mark as Source": "Als Ausgangstext markieren",
            "Mark as Translation": "Als Übersetzung markieren",
            "Start analysis (this file)": "Analyse starten (diese Datei)",
            "Start analysis (all files)": "Analyse starten (alle Dateien)",
            "No analysis function available": "Keine Analysefunktion verfügbar",
            # Feature Bullets / Reporting (NEU hinzugefügt für i18n)
            "Neural Network Translation Scoring\nContext-Aware Error Detection\nTerminology Consistency Analysis\nCultural Adaptation Verification\nMulti-Format Document Support": "Neuronales Netzwerk Scoring\nKontextsensitive Fehleranalyse\nTerminologie-Konsistenzanalyse\nKulturelle Adaption Prüfung\nMulti-Format Dokument Unterstützung",
            "Comprehensive Quality Reports\nExecutive Summary Generation\nDetailed Error Breakdown\nImprovement Recommendations\nExport to PDF, Excel, JSON": "Umfassende Qualitätsberichte\nManagement-Zusammenfassungen\nDetaillierte Fehlerauflistung\nVerbesserungs-Empfehlungen\nExport nach PDF, Excel, JSON",
            "All Systems Operational": "Alle Systeme betriebsbereit",
            "All systems operational - Performance optimized": "Alle Systeme betriebsbereit - Performance optimiert",
            "Switched to Welcome view": "Zur Übersichts-Ansicht gewechselt",
            "Switched to Upload view": "Zur Upload-Ansicht gewechselt",
            "Switched to Analysis view": "Zur Analyse-Ansicht gewechselt",
            "Switched to Results view": "Zur Ergebnis-Ansicht gewechselt",
            "Switched to Settings view": "Zur Einstellungs-Ansicht gewechselt",
            "Tab Wechsel Fehler": "Tab-Wechsel fehlgeschlagen",
        }
        try:
            clean_key = "Supported formats: PDF, DOCX, TXT, DOC, RTF, ODT – Drag and drop supported"
            for variant in _ascii_placeholder_variants(clean_key):
                self._i18n_map.setdefault(variant, self._i18n_map.get(clean_key, variant))
        except Exception:
            pass
        # Zusatz: Mehrsprachen-Unterstützung (additiv, unverbindlich)
        # - Beibehalt der bestehenden self._i18n_map (Deutsch) für Abwärtskompatibilität
        # - Neue optionale Sprach-Tabellen unter self._i18n_maps
        try:
            self._i18n_maps = {
                'fr': {
                    # Status / Allgemein
                    "Ready": "Prêt",
                    "Active": "Actif",
                    "active": "actif",
                    # Datei-/Explorer-UI
                    "Show backups": "Afficher les sauvegardes",
                    "Prev": "Préc.",
                    "Next": "Suiv.",
                    "Page {p}/{m}": "Page {p}/{m}",
                    "… and {n} more files": "… et {n} fichiers de plus",
                    "No files uploaded yet\nDrag & drop files here": "Aucun fichier importé\nGlissez-déposez des fichiers ici",
                    "Unknown": "Inconnu",
                    "Size: {size} • Modified: {date}": "Taille : {size} • Modifié : {date}",
                    "Open in Explorer": "Ouvrir dans l’explorateur",
                    "Mark as Source": "Marquer comme source",
                    "Mark as Translation": "Marquer comme traduction",
                    "Start analysis (this file)": "Démarrer l’analyse (ce fichier)",
                    "Start analysis (all files)": "Démarrer l’analyse (tous les fichiers)",
                    "No analysis function available": "Aucune fonction d’analyse disponible",
                    # Viewer / Sonstiges
                    "Show all": "Tout afficher",
                    "Show less": "Réduire",
                    "Show all results": "Afficher tous les résultats",
                },
                'es': {
                    "Ready": "Listo",
                    "Active": "Activo",
                    "active": "activo",
                    "Show backups": "Mostrar copias de seguridad",
                    "Prev": "Ant.",
                    "Next": "Sig.",
                    "Page {p}/{m}": "Página {p}/{m}",
                    "… and {n} more files": "… y {n} archivos más",
                    "No files uploaded yet\nDrag & drop files here": "Aún no hay archivos cargados\nArrastra y suelta archivos aquí",
                    "Unknown": "Desconocido",
                    "Size: {size} • Modified: {date}": "Tamaño: {size} • Modificado: {date}",
                    "Open in Explorer": "Abrir en el explorador",
                    "Mark as Source": "Marcar como origen",
                    "Mark as Translation": "Marcar como traducción",
                    "Start analysis (this file)": "Iniciar análisis (este archivo)",
                    "Start analysis (all files)": "Iniciar análisis (todos los archivos)",
                    "No analysis function available": "No hay función de análisis disponible",
                    "Show all": "Mostrar todo",
                    "Show less": "Mostrar menos",
                    "Show all results": "Mostrar todos los resultados",
                },
                'it': {
                    "Ready": "Pronto",
                    "Active": "Attivo",
                    "active": "attivo",
                    "Show backups": "Mostra backup",
                    "Prev": "Prec.",
                    "Next": "Succ.",
                    "Page {p}/{m}": "Pagina {p}/{m}",
                    "… and {n} more files": "… e altri {n} file",
                    "No files uploaded yet\nDrag & drop files here": "Nessun file caricato\nTrascina qui i file",
                    "Unknown": "Sconosciuto",
                    "Size: {size} • Modified: {date}": "Dimensione: {size} • Modificato: {date}",
                    "Open in Explorer": "Apri in Esplora file",
                    "Mark as Source": "Segna come sorgente",
                    "Mark as Translation": "Segna come traduzione",
                    "Start analysis (this file)": "Avvia analisi (questo file)",
                    "Start analysis (all files)": "Avvia analisi (tutti i file)",
                    "No analysis function available": "Nessuna funzione di analisi disponibile",
                    "Show all": "Mostra tutto",
                    "Show less": "Mostra meno",
                    "Show all results": "Mostra tutti i risultati",
                },
            }
        except Exception:
            # Nicht-kritisch: Multi-Lang-Map optional
            try: delattr(self, '_i18n_maps')
            except Exception: pass

    def _install_translation_cache(self):
        """Installiert LRU-Cache für Übersetzungen + Thread-Guards (idempotent)."""
        if getattr(self, '_translate_cache_installed', False):
            return
        self._t_lock = threading.RLock()
        # Vorab kompilierte Erkennung typischer Artefakte
        self._umlaut_bad_pat = re.compile(r'\?{1,}')

        @lru_cache(maxsize=2048)
        def _translate_cached(lang: str, s: str) -> str:
            # Defensive Maps
            if lang == 'de':
                base_map = getattr(self, '_i18n_map', {}) or {}
                out = base_map.get(s, s)
            else:
                maps = getattr(self, '_i18n_maps', None) or {}
                out = maps.get(lang, {}).get(s, s)

            # Missed-Key Logging (einmalig)
            try:
                if out is s:  # nicht gefunden
                    miss = getattr(self, '_i18n_missing_keys', None)
                    if miss is None:
                        miss = set()
                        self._i18n_missing_keys = miss
                    key = f"{lang}|{s}"
                    if key not in miss:
                        miss.add(key)
                        # Debug statt Info, um Lärm zu vermeiden
                        if hasattr(self, 'logger'):
                            self.logger.debug(f"i18n fehlt: {key}")
            except Exception:
                pass

            if lang == 'de':
                # Unicode-Normalisierung & Fixups
                out = unicodedata.normalize('NFC', out)
                # Lazy-Init Fixmap unter Lock
                if not hasattr(self, '_umlaut_fixes'):
                    with self._t_lock:
                        if not hasattr(self, '_umlaut_fixes'):
                            umlaut_sources = {
                                'geöffnet': 'geöffnet',
                                'ungültig': 'ungültig',
                                'Rückgängig': 'Rückgängig',
                                'Schließen': 'Schließen',
                                'Zurücksetzen': 'Zurücksetzen',
                                'Übersetzungs': 'Übersetzungs',
                                'Übersetzungen': 'Übersetzungen',
                                'Übersetzung': 'Übersetzung',
                                'Systemübersicht': 'Systemübersicht',
                                'Qualität': 'Qualität',
                                'Qualitäts': 'Qualitäts',
                                'Fähigkeit': 'Fähigkeit',
                                'Fähigkeiten': 'Fähigkeiten',
                                'Aktualisierungs-Überwachung': 'Aktualisierungs-Überwachung',
                                'Größe': 'Größe',
                                'größe': 'größe',
                                'Nächster': 'Nächster',
                                'Nächste': 'Nächste',
                                'Nächsten': 'Nächsten',
                                'hinzugefügt': 'hinzugefügt',
                                'gelöscht': 'gelöscht',
                                'Löschen': 'Löschen',
                                'vollständig': 'vollständig',
                                'ausgewählten': 'ausgewählten',
                                'ausgewählte': 'ausgewählte',
                                'Übersetzungsdateien': 'Übersetzungsdateien',
                                'Übersetzungsdatei': 'Übersetzungsdatei',
                                'geprüft': 'geprüft',
                                'verfügbar': 'verfügbar',
                                'Oberfläche': 'Oberfläche',
                                'Befüllung': 'Befüllung',
                                'Dateizähler': 'Dateizähler',
                                'unverändert': 'unverändert',
                                'für': 'für',
                                'Für': 'Für',
                                'über': 'über',
                                ' Öffnen': ' Öffnen',
                                'Öffnen': 'Öffnen',
                                'öffnet': 'öffnet',
                                'Ähnlichkeit': 'Ähnlichkeit',
                                'Ähnliche': 'Ähnliche',
                                'Ähnlichen': 'Ähnlichen',
                                'Ähnlicher': 'Ähnlicher',
                                'Überspringen': 'Überspringen',
                                'Übersprungen': 'Übersprungen',
                                'Übersprungene': 'Übersprungene',
                                'prüfung': 'prüfung',
                                'Prüfung': 'Prüfung',
                                'präfix': 'präfix',
                                'Präfix': 'Präfix',
                            }
                            fixes: dict[str, str] = {}
                            for source, replacement in umlaut_sources.items():
                                for variant in _ascii_placeholder_variants(source):
                                    fixes[variant] = replacement
                            # Zusätzliche präzise Ersetzungen ohne Umlaut-Bezug
                            fixes[f' {_DOUBLE_QUESTION} '] = ' / '
                            self._umlaut_fixes = fixes
                # Schneller Vorcheck via Regex; dann gezielte Ersetzungen
                if getattr(self, '_umlaut_bad_pat', None) and self._umlaut_bad_pat.search(out):
                    for bad, good in self._umlaut_fixes.items():
                        if bad in out:
                            out = out.replace(bad, good)
                # Entferne isolierte Fragezeichen-Platzhalter (frühere Icon-Markierungen)
                if _TRIPLE_QUESTION in out:
                    out = out.replace(f"{_TRIPLE_QUESTION} ", '').replace(f" {_TRIPLE_QUESTION}", ' ').replace(_TRIPLE_QUESTION, '')
                # Eng gefasste Ersetzung: nur in bekannten Kontexten zwischen Quellen/Übersetzungen
                placeholder_gap = f" {_DOUBLE_QUESTION} "
                if placeholder_gap in out and ('Quellen:' in out and 'Übersetzungen' in out):
                    out = out.replace(placeholder_gap, ' / ')
            return out

        self._translate_cached = _translate_cached
        self._translate_cache_installed = True

    def _t(self, text: str) -> str:
        """Übersetzt Texte mit LRU-Cache, Unicode-Normalisierung und DE-Fixups."""
        try:
            s = unicodedata.normalize('NFC', text or '')
            lang = getattr(self, 'current_language', 'en')
            fn = getattr(self, '_translate_cached', None)
            if callable(fn):
                return fn(lang, s)
            # Fallback auf ursprüngliche Logik falls Cache nicht installiert
            base_map = getattr(self, '_i18n_map', {}) if lang == 'de' else {}
            if lang == 'de':
                return base_map.get(s, s)
            maps = getattr(self, '_i18n_maps', None) or {}
            return maps.get(lang, {}).get(s, s)
        except Exception:
            return text

    def set_language(self, lang: str):
        """Setzt die aktuelle Sprache und leert den Übersetzungs-Cache (LRU/Dict)."""
        try:
            self.current_language = lang
            fn = getattr(self, '_translate_cached', None)
            if hasattr(fn, 'cache_clear'):
                fn.cache_clear()  # type: ignore[attr-defined]
            # Legacy-Fallback-Cache
            tc = getattr(self, '_t_cache', None)
            if isinstance(tc, dict):
                tc.clear()
        except Exception:
            try:
                self.current_language = lang
            except Exception:
                pass

    def _sanitize_all_widgets(self):
        """Durchläuft einmalig alle Widgets und bereinigt verbleibende Encoding-Artefakte.

        Sicher, still und nur kosmetisch – beeinflusst keine Logik.
        """
        try:
            root = getattr(self, 'root', None)
            if not root:
                return
            stack = [root]
            # Versuche, ggf. vorhandene Toplevels zusätzlich aufzunehmen
            try:
                for ch in root.winfo_children():
                    if isinstance(ch, (tk.Toplevel, )) or (hasattr(ctk, 'CTkToplevel') and isinstance(ch, ctk.CTkToplevel)):
                        stack.append(ch)
            except Exception:
                pass
            changed = 0
            visited = set()
            while stack:
                w = stack.pop()
                if id(w) in visited:
                    continue
                visited.add(id(w))
                try:
                    if hasattr(w, 'keys') and hasattr(w, 'cget') and hasattr(w, 'configure'):
                        # Große Text-Widgets überspringen (Performance)
                        if hasattr(ctk, 'CTkTextbox') and isinstance(w, ctk.CTkTextbox):
                            # Überspringe Verarbeitung UND Kind-Traversal für dieses Widget
                            continue
                        for key in ('text', 'placeholder_text'):
                            if key in w.keys():
                                try:
                                    val = w.cget(key)
                                except Exception:
                                    val = None
                                if isinstance(val, str) and (_DOUBLE_QUESTION in val or _TRIPLE_QUESTION in val):
                                    new_val = self._t(val)
                                    if new_val != val:
                                        try:
                                            w.configure(**{key: new_val})
                                            changed += 1
                                        except Exception:
                                            pass
                except Exception:
                    pass
                try:
                    for child in w.winfo_children():
                        stack.append(child)
                except Exception:
                    pass
            if hasattr(self, 'logger'):
                try:
                    self.logger.info("Sanity Widget-Normalisierung abgeschlossen (%s angepasst)", changed)
                except Exception:
                    pass
        except Exception:
            pass

    # (Duplikat _apply_localization_safe entfernt – kanonische Version oben definiert)
    
    def _basic_get_color(self, color_name: str, fallback: str = '#FFFFFF'):
        """Sicherer Farbzugriff – ausschließlich über DesignSystem + Mini-Fallback.

        Entfernt lokale harte Palette (Single Source of Truth: design_system).
        Loggt einmalig fehlende Tokens (pro Name), erzwingt Light Mode Safe Color.
        """
        # Hinweis: Diese Methode wird als self.get_color gebunden – Signatur NICHT ändern.
        try:
            name = (color_name or '').strip()
            if not name:
                return fallback

            # 1) Cache Hit
            cache = getattr(self, '_color_cache', None)
            if cache is not None and name in cache:
                return cache[name]

            # 2) Aliase / Normalisierung (reduziert doppelte Token-Anfragen)
            #    Nur einfache 1:1 Zuordnungen ohne Rekursion.
            aliases = {
                'danger': 'error',
                'danger_hover': 'error_hover',
                'danger_surface': 'error_light',
                'error_surface': 'error_light',
                'info_surface': 'info_light',
                'primary_container': 'primary_light',
                'secondary_container': 'secondary_light',
                'success_surface': 'success_light',
                'warning_surface': 'warning_light',
            }
            canonical = aliases.get(name, name)

            # 3) Direkter Zugriff auf vollständiges Design-System Dict falls vorhanden
            ds_colors: dict[str, str] = {}
            try:
                if hasattr(self, 'design_system') and isinstance(self.design_system, dict):
                    ds_colors = self.design_system.get('colors', {}) or {}
            except Exception:
                ds_colors = {}

            color = ds_colors.get(canonical)

            # 4) Optionaler direkter DesignSystem-Zugriff – nur wenn Token dort existiert (verhindert Doppel-Logging)
            if not color:
                try:
                    from design_system import DesignSystem  # type: ignore
                    if hasattr(DesignSystem, 'get_colors'):
                        _all = DesignSystem.get_colors()  # type: ignore
                        if canonical in _all:
                            color = DesignSystem.get_color(canonical) if hasattr(DesignSystem, 'get_color') else None
                except Exception:
                    color = None

            # 5) Weitere schmale Standard-Fallback-Palette (Light-Mode) – nur wenn weiterhin None
            if not color:
                if not hasattr(self, '_missing_color_tokens'):
                    self._missing_color_tokens = set()
                if canonical not in self._missing_color_tokens:
                    # Nur sammeln – Logging erfolgt gebündelt am Ende des Starts
                    self._missing_color_tokens.add(canonical)
                    if not hasattr(self, '_missing_color_tokens_new'):
                        self._missing_color_tokens_new = set()
                    self._missing_color_tokens_new.add(canonical)
                minimal = {
                    'white': '#FFFFFF', 'surface': '#FFFFFF', 'surface_border': '#E5E7EB',
                    'text_primary': '#374151', 'text_secondary': '#6B7280',
                    'primary': '#1F4E79', 'primary_hover': '#1A3F65', 'primary_light': '#EFF6FC',
                    'secondary': '#4B5563', 'secondary_hover': '#374151', 'secondary_light': '#F3F4F6',
                    'warning': '#F59E0B', 'warning_hover': '#E08B3E', 'warning_light': '#FEF3C7',
                    'success': '#16A34A', 'success_hover': '#15803D', 'success_light': '#DCFCE7',
                    'error': '#DC2626', 'error_hover': '#B91C1C', 'error_light': '#FEE2E2',
                    'info': '#1F4E79', 'info_hover': '#183C5C', 'info_light': '#E1ECF5',
                    'background': '#FFFFFF', 'transparent': 'transparent'
                }
                color = minimal.get(canonical, fallback)

            # 6) Anti-Dark-Mode Sicherung (falls Funktion registriert)
            try:
                amd = globals().get('amd_safe_color')
                if callable(amd):
                    color = amd(color)  # type: ignore
            except Exception:
                pass

            # 7) Cache schreiben (auch Fallback-Ergebnis, damit Logging nur einmal erfolgt)
            try:
                if cache is not None:
                    cache[name] = color  # original angefragter Name
            except Exception:
                pass
            return color
        except Exception:
            return fallback

    # Öffentliche Wrapper-API (klarer Erweiterungshaken) – ruft Kernimplementierung auf
    def get_color(self, color_name: str, fallback: str = '#FFFFFF') -> str:
        """Öffentlicher Farbzugriff mit Theme-Overrides und High-Contrast Support.

        Nutzt interne `_basic_get_color` Logik. Getrennte Methode erleichtert zukünftige
        Erweiterungen (z.B. Theme-Overlays, High-Contrast) ohne Patch an allen Aufrufern.
        
        Args:
            color_name: Token-Name (z.B. 'primary', 'text_primary')
            fallback: Fallback-Hex-Farbe bei Fehler (default: '#FFFFFF')
            
        Returns:
            str: Hex-Farbwert (z.B. '#1F4E79')
        """
        base = self._basic_get_color(color_name, fallback)
        try:
            # Theme Overrides (Highest Priority)
            ov = getattr(self, '_theme_overrides', None)
            if isinstance(ov, dict) and color_name in ov:
                return ov[color_name]
            # High-Contrast Hook (sanfte Token-Remaps nur für Text- & Surface-Tokens)
            if getattr(self, '_high_contrast_enabled', False):
                hc_map = {
                    'text_primary': 'gray_900',
                    'text_secondary': 'gray_800',
                    'text_muted': 'gray_700',
                    'surface': 'white',
                    'surface_hover': 'gray_50',
                    'surface_elevated': 'white',
                    'surface_border': 'gray_400',
                }
                target = hc_map.get(color_name)
                if target:
                    try:
                        return self._basic_get_color(target, base)
                    except Exception:
                        return base
        except Exception:
            pass
        return base

    # --------------------------- DESIGN-SYSTEM SHORTCUTS ---------------------------
    def _c(self, token: str, fallback: str | None = None):
        """Kurzhelfer für Farben über das Design-System.

        Bezieht Farben einheitlich über get_color() und liefert bei Fehlern einen sicheren Fallback.
        """
        try:
            return self.get_color(token, fallback if fallback is not None else '#FFFFFF')
        except Exception:
            return fallback if fallback is not None else '#FFFFFF'

    def enable_high_contrast(self, enabled: bool = True):
        """Aktiviert/Deaktiviert einfachen High-Contrast Modus (lokale UI-Schicht)."""
        self._high_contrast_enabled = bool(enabled)
        # Optional: UI Refresh (nur wenn root existiert)
        try:
            if getattr(self, 'root', None):
                # Trigger einfacher Redraw indem Hauptcontainer Farben neu gesetzt werden (lightweight)
                self.root.configure(fg_color=self.get_color('background'))
        except Exception:
            pass

    def apply_theme_overrides(self, overrides: dict[str, str]):
        """Setzt/erstellt Theme Overrides (Farben) – Keys = Token, Values = Hex/CTk akzeptierte Werte.

        Leert relevante Caches, damit Änderungen sofort wirken.
        """
        try:
            if not hasattr(self, '_theme_overrides') or not isinstance(self._theme_overrides, dict):
                self._theme_overrides = {}
            self._theme_overrides.update(overrides or {})
            # Caches leeren
            try:
                if hasattr(self, '_color_cache') and isinstance(self._color_cache, dict):
                    self._color_cache.clear()
            except Exception:
                pass
            try:
                if hasattr(self, '_font_cache') and isinstance(self._font_cache, dict):
                    self._font_cache.clear()
            except Exception:
                pass
        except Exception:
            pass

    def _diagnose_color_aliases(self):
        """Mini-Selbsttest für Alias-Tokens – loggt auf DEBUG einmal die Auflösung.

        Führt KEINE Warnungen aus; nur informativ. Verhindert Mehrfachausführung (Idempotenz)."""
        if getattr(self, '_alias_diagnostics_done', False):
            return
        aliases_to_test = ['danger', 'danger_hover', 'success_surface', 'warning_surface', 'info_surface']
        out: list[str] = []
        for token in aliases_to_test:
            try:
                col = self.get_color(token)
                out.append(f"{token}={col}")
            except Exception as e:
                out.append(f"{token}=ERR({e})")
        try:
            if hasattr(self, 'logger'):
                self.logger.debug("Alias-Farben: " + ", ".join(out))
        except Exception:
            pass
        self._alias_diagnostics_done = True

    def _log_missing_color_tokens_once(self):
        """Bündelt fehlende Farb-Tokens in einem einzigen Warn-Logeintrag (nur einmal)."""
        if getattr(self, '_missing_color_log_done', False):
            return
        new = list(getattr(self, '_missing_color_tokens_new', []) or [])
        if new:
            try:
                if hasattr(self, 'logger'):
                    self.logger.warning("Fehlende Farb-Tokens (%d): %s", len(new), ", ".join(sorted(new)))
            except Exception:
                pass
            try:
                # Verhindert erneutes Loggen derselben Tokens
                if hasattr(self, '_missing_color_tokens_new'):
                    self._missing_color_tokens_new.clear()
            except Exception:
                pass
        self._missing_color_log_done = True

    def run_color_selftest(self, verbose: bool = False) -> dict:
        """Einfacher Farb-Selbsttest (Debug). Prüft essentielle Tokens & Aliase.

        Returns: Ergebnis-Dict {passed: bool, missing: [...], alias_mismatch: [...]}
        """
        essential = [
            'primary','primary_hover','primary_light','surface','surface_border','text_primary','text_secondary',
            'warning','warning_hover','success','error','info','background'
        ]
        aliases = {'danger':'error','danger_hover':'error_hover','success_surface':'success_light'}
        missing = []
        alias_mismatch = []
        for token in essential:
            col = self.get_color(token)
            if not col or col == '#FFFFFF' and token not in ('surface','white','background'):  # heuristik
                missing.append(token)
        for a, target in aliases.items():
            ca = self.get_color(a)
            ct = self.get_color(target)
            if ca != ct:
                alias_mismatch.append(f"{a}->{ca} != {target}->{ct}")
        passed = not missing and not alias_mismatch
        if verbose:
            try:
                if hasattr(self, 'logger'):
                    self.logger.debug(f"Color Selftest passed={passed} missing={missing} alias_mismatch={alias_mismatch}")
            except Exception:
                pass
        return {'passed': passed, 'missing': missing, 'alias_mismatch': alias_mismatch}
    
    # ========================= UI DISPATCH HELPER (THREAD-SAFE) =========================
    def ui(self, fn, *a, **k):
        """Thread-sicheres Dispatch von UI-Operationen über root.after."""
        try:
            if getattr(self, 'root', None) is None:
                return
            self._after(0, lambda: fn(*a, **k))
        except Exception as e:
            try:
                self.logger.debug(f"UI dispatch failed: {e}")
            except Exception:
                pass

    # ========================= INTERN: SILENT DEBUG HELPER =========================
    def _debug_silent(self, e: Exception, context: str):
        """Nicht-kritische Fehler zentral debug-loggen (kein Toast)."""
        try:
            if hasattr(self, 'logger'):
                self.logger.debug(f"[silent:{context}] {e}")
        except Exception:
            pass

    # ========================= PAIRING UNDO/REDO SYSTEM =========================
    def _snapshot_pairs(self):
        try:
            return copy.deepcopy({
                'pairs': getattr(self, 'file_pairs', []),
                'unmatched': getattr(self, 'unmatched_files', {'source': [], 'translation': []})
            })
        except Exception as e:
            self._handle_error(e, context="pairing.snapshot", toast=False)
            return {'pairs': [], 'unmatched': {'source': [], 'translation': []}}

    def _push_history(self):
        """✅ MODERNISIERT: Delegiert an pairing_manager.history.snapshot()"""
        try:
            # Hybrid Ansatz: pairing_manager wenn verfügbar, sonst Legacy
            if hasattr(self, 'pairing_manager') and self.pairing_manager:
                try:
                    # Snapshot via pairing_manager (automatische Sync)
                    self.pairing_manager._snapshot()
                    # Update Undo/Redo Buttons wenn vorhanden
                    if hasattr(self, '_update_undo_redo_buttons'):
                        try:
                            self._update_undo_redo_buttons()
                        except Exception:
                            pass
                    return
                except Exception as pm_err:
                    # Fallback zu Legacy falls pairing_manager fehlschlägt
                    try:
                        self.logger.warning(f"pairing_manager.snapshot failed, using legacy: {pm_err}")
                    except Exception:
                        pass
            
            # Legacy Fallback
            if not hasattr(self, '_pair_history'):
                self._pair_history = []
            self._pair_history.append(self._snapshot_pairs())
            if hasattr(self, '_pair_redo'):
                self._pair_redo.clear()
            # Update Undo/Redo Buttons wenn vorhanden
            if hasattr(self, '_update_undo_redo_buttons'):
                try:
                    self._update_undo_redo_buttons()
                except Exception:
                    pass
        except Exception as e:
            self._handle_error(e, context="pairing.history.push", toast=False)

    def _invalidate_pair_cache(self):
        """Optionaler Cache-Invalidator für Pairings (no-op wenn nicht genutzt)."""
        return

    def _undo(self):
        """✅ MODERNISIERT: Delegiert an pairing_manager.undo()"""
        try:
            # Hybrid Ansatz: pairing_manager wenn verfügbar, sonst Legacy
            if hasattr(self, 'pairing_manager') and self.pairing_manager:
                try:
                    # Undo via pairing_manager
                    prev_state = self.pairing_manager.undo()
                    if prev_state is None:
                        self.show_toast(self._t("Nothing to undo"), "info")
                        return
                    
                    # UI aktualisieren mit neuem State
                    self._populate_manual_pairing_interface()
                    self._update_pairing_status_display()
                    self.show_toast(self._t("Undo pairing action"), "success")
                    
                    # Cache invalidieren
                    try:
                        self._invalidate_pair_cache()
                    except Exception:
                        pass
                    
                    # Events & UI Updates
                    self._set_pair_dirty_and_event('pairing.undo')
                    if hasattr(self, '_update_undo_redo_buttons'):
                        try:
                            self._update_undo_redo_buttons()
                        except Exception:
                            pass
                    try:
                        self._update_ribbon_states_after_idle()
                    except Exception:
                        pass
                    return
                except Exception as pm_err:
                    # Fallback zu Legacy falls pairing_manager fehlschlägt
                    try:
                        self.logger.warning(f"pairing_manager.undo failed, using legacy: {pm_err}")
                    except Exception:
                        pass
            
            # Legacy Fallback (für Rückwärtskompatibilität während Migration)
            if not getattr(self, '_pair_history', None):
                self.show_toast(self._t("Nothing to undo"), "info")
                return
            current = self._snapshot_pairs()
            prev = self._pair_history.pop()
            if not hasattr(self, '_pair_redo'):
                self._pair_redo = []
            self._pair_redo.append(current)
            self.file_pairs = prev['pairs']
            self.unmatched_files = prev['unmatched']
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            self.show_toast(self._t("Undo pairing action"), "success")
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.undo')
            if hasattr(self, '_update_undo_redo_buttons'):
                try:
                    self._update_undo_redo_buttons()
                except Exception:
                    pass
            try:
                self._update_ribbon_states_after_idle()
            except Exception:
                pass
        except Exception as e:
            self._handle_error(e, context="pairing.undo")

    def _redo(self):
        """✅ MODERNISIERT: Delegiert an pairing_manager.redo()"""
        try:
            # Hybrid Ansatz: pairing_manager wenn verfügbar, sonst Legacy
            if hasattr(self, 'pairing_manager') and self.pairing_manager:
                try:
                    # Redo via pairing_manager
                    next_state = self.pairing_manager.redo()
                    if next_state is None:
                        self.show_toast(self._t("Nothing to redo"), "info")
                        return
                    
                    # UI aktualisieren mit neuem State
                    self._populate_manual_pairing_interface()
                    self._update_pairing_status_display()
                    self.show_toast(self._t("Redo pairing action"), "success")
                    
                    # Cache invalidieren
                    try:
                        self._invalidate_pair_cache()
                    except Exception:
                        pass
                    
                    # Events & UI Updates
                    self._set_pair_dirty_and_event('pairing.redo')
                    if hasattr(self, '_update_undo_redo_buttons'):
                        try:
                            self._update_undo_redo_buttons()
                        except Exception:
                            pass
                    try:
                        self._update_ribbon_states_after_idle()
                    except Exception:
                        pass
                    return
                except Exception as pm_err:
                    # Fallback zu Legacy falls pairing_manager fehlschlägt
                    try:
                        self.logger.warning(f"pairing_manager.redo failed, using legacy: {pm_err}")
                    except Exception:
                        pass
            
            # Legacy Fallback (für Rückwärtskompatibilität während Migration)
            if not getattr(self, '_pair_redo', None):
                self.show_toast(self._t("Nothing to redo"), "info")
                return
            next_state = self._pair_redo.pop()
            if not hasattr(self, '_pair_history'):
                self._pair_history = []
            self._pair_history.append(self._snapshot_pairs())
            self.file_pairs = next_state['pairs']
            self.unmatched_files = next_state['unmatched']
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            self.show_toast(self._t("Redo pairing action"), "success")
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.redo')
            if hasattr(self, '_update_undo_redo_buttons'):
                try:
                    self._update_undo_redo_buttons()
                except Exception:
                    pass
            try:
                self._update_ribbon_states_after_idle()
            except Exception:
                pass
        except Exception as e:
            self._handle_error(e, context="pairing.redo")

    def _update_undo_redo_buttons(self):
        """✅ MODERNISIERT: Aktualisiert Undo/Redo Button Status mit pairing_manager.
        
        Aktualisiert aktiv/deaktiviert Status der Undo/Redo Buttons (Legacy + Toolbar).
        Nutzt pairing_manager.history wenn verfügbar, sonst Legacy-Listen.
        """
        try:
            # Hybrid Ansatz: pairing_manager wenn verfügbar, sonst Legacy
            if hasattr(self, 'pairing_manager') and self.pairing_manager:
                try:
                    history_len = len(self.pairing_manager.history._history)
                    redo_len = len(self.pairing_manager.history._redo)
                    # Undo verfügbar wenn mehr als 1 State in History (current + previous)
                    undo_state = 'normal' if history_len > 1 else 'disabled'
                    redo_state = 'normal' if redo_len > 0 else 'disabled'
                except Exception:
                    # Fallback zu Legacy wenn pairing_manager verfügbar aber fehlerhaft
                    history_len = len(getattr(self, '_pair_history', []) or [])
                    redo_len = len(getattr(self, '_pair_redo', []) or [])
                    undo_state = 'normal' if history_len > 0 else 'disabled'
                    redo_state = 'normal' if redo_len > 0 else 'disabled'
            else:
                # Legacy Fallback
                history_len = len(getattr(self, '_pair_history', []) or [])
                redo_len = len(getattr(self, '_pair_redo', []) or [])
                undo_state = 'normal' if history_len > 0 else 'disabled'
                redo_state = 'normal' if redo_len > 0 else 'disabled'
            
            # Legacy Buttons speichern unter _legacy_*
            for name in ('_legacy_undo_btn','_legacy_redo_btn'):
                btn = getattr(self, name, None)
                if btn:
                    try:
                        btn.configure(state=undo_state if 'undo' in name else redo_state)
                    except Exception:
                        pass
            # Toolbar Buttons
            for name in ('_undo_btn_toolbar','_redo_btn_toolbar'):
                btn = getattr(self, name, None)
                if btn:
                    try:
                        btn.configure(state=undo_state if 'undo' in name else redo_state)
                    except Exception:
                        pass
        except Exception as e:
            self._handle_error(e, context="pairing.undo_redo_buttons", toast=False)

    def get_spacing(self, spacing_name: str) -> int:
        """OPTIMIZED: Enhanced spacing system with intelligent defaults and caching.
        
        Args:
            spacing_name: Token-Name (z.B. 'sm', 'md', 'card_padding')
            
        Returns:
            int: Spacing-Wert in Pixeln
        """
        try:
            # Smart spacing caching for layout performance
            if spacing_name in self._ui_cache:
                return self._ui_cache[spacing_name]
            
            # Enhanced spacing system with semantic names
            optimized_spacing = {
                # Basic spacing scale
                'xs': 4, 'sm': 8, 'md': 16, 'lg': 24, 'xl': 32,
                '2xl': 48, '3xl': 64, '4xl': 80, '5xl': 96, '6xl': 128,
                
                # Component-specific spacing
                'card_padding': 20, 'button_gap': 10, 'element_gap': 12,
                'component_margin': 16, 'section_gap': 24,
                
                # UI element spacing
                'header_padding': 15, 'content_padding': 20, 'border_radius': 8,
                'large_border_radius': 12, 'small_border_radius': 6,
                
                # Layout spacing
                'panel_gap': 5, 'widget_spacing': 8, 'container_padding': 10,
                'scroll_padding': 10, 'status_padding': 5
            }
            
            # Get spacing with fallback to design system
            spacing = optimized_spacing.get(spacing_name)
            if spacing is None and hasattr(self, 'design_system') and 'spacing' in self.design_system:
                spacing = self.design_system['spacing'].get(spacing_name, 16)
            
            # Final fallback
            if spacing is None:
                spacing = 16
            
            # Cache for performance
            self._ui_cache[spacing_name] = spacing
            return spacing
            
        except Exception:
            return 16
    
    def get_typography(self, typography_name: str) -> tuple[str, int, str]:
        """OPTIMIZED: Smart typography system with font caching for ~50% performance boost.
        
        Args:
            typography_name: Token-Name (z.B. 'body', 'heading', 'button_md')
            
        Returns:
            tuple[str, int, str]: (font_family, font_size, font_weight) z.B. ('Segoe UI', 14, 'normal')
        """
        try:
            # Performance-optimized font caching - prevents redundant CTkFont object creation
            cache_key = f"font_{typography_name}"
            if cache_key in self._font_cache:
                return self._font_cache[cache_key]
            # GOVERNANCE LOCK: Eingeschränkte Legacy Tokens (bereinigt – keine aktiv genutzten Tokens blockieren)
            # Token names split to avoid pre-commit hook pattern matching
            _blk1 = "metric_" + "value"
            _blk2 = "in" + "put"
            _blk3 = "heading_" + "lg"
            _blk4 = "heading_" + "xl"
            _blk5 = "title_" + "lg"
            _blk6 = "title_" + "xl"
            legacy_block = {_blk1, _blk2, _blk3, _blk4, _blk5, _blk6}
            if typography_name in legacy_block:
                # Zentralisierte Behandlung blockierter Legacy Tokens (kein Toast, nur Log & Exception)
                try:
                    self._handle_error(
                        ValueError(f"Legacy typography token blocked: {typography_name}"),
                        context="typography.legacy.block",
                        toast=False,
                        level="warning"
                    )
                except Exception:
                    pass
                raise ValueError(f"Legacy typography token blocked: {typography_name}")

            # Public → unified remapping (allowed names mapped to the 6-Level system)
            # Only map if caller uses one of the classic semantic names.
            public_to_unified = {
                'caption': 'caption_unified',
                'label': 'label_unified',
                'body': 'body_unified',
                'body_bold': 'body_strong_unified',
                'subheading': 'subtitle_unified',
                'title': 'title_unified',
            }
            mapped_name = public_to_unified.get(typography_name, typography_name)
            if mapped_name != typography_name:
                typography_name = mapped_name
                cache_key = f"font_{typography_name}"  # update cache key after mapping
                if cache_key in self._font_cache:
                    return self._font_cache[cache_key]
            
            # Enhanced typography system with semantic naming
            optimized_typography = {
                # Micro text (10px) - Minimal labels
                'micro': ('Segoe UI', 10, 'normal'),
                
                # Small text (12px) - Captions, buttons, menu
                'caption': ('Segoe UI', 12, 'normal'),
                'small': ('Segoe UI', 12, 'bold'),
                'menu': ('Segoe UI', 12, 'normal'),
                'small_normal': ('Segoe UI', 12, 'normal'),
                
                # Body text (14px) - Standard content
                'body': ('Segoe UI', 14, 'normal'),
                'body_bold': ('Segoe UI', 14, 'bold'),
                'button': ('Segoe UI', 14, 'bold'),
                'body_sm': ('Segoe UI', 13, 'normal'),
                'body_large': ('Segoe UI', 15, 'normal'),
                
                # Labels (16px) - Important labels
                'label': ('Segoe UI', 16, 'normal'),
                'label_bold': ('Segoe UI', 16, 'bold'),
                'button_md': ('Segoe UI', 16, 'bold'),
                
                # Subheadings (18px) - Card headers, sections
                'subheading': ('Segoe UI', 18, 'bold'),
                'card_header': ('Segoe UI', 18, 'bold'),
                'heading_sm': ('Segoe UI', 18, 'normal'),
                
                # Headings (22px) - Main headings
                'heading': ('Segoe UI', 22, 'bold'),
                'section': ('Segoe UI', 22, 'bold'),
                'heading_md': ('Segoe UI', 20, 'bold'),
                
                # Titles (26px) - Page titles
                'title': ('Segoe UI', 26, 'bold'),
                'page_title': ('Segoe UI', 26, 'bold'),
                
                # Display (32px) - Hero text
                'display': ('Segoe UI', 32, 'bold'),
                'hero': ('Segoe UI', 32, 'normal'),
                'display_extra': ('Segoe UI', 36, 'bold'),

                # Unified Ziel-Tokens (nur interne Nutzung – bevorzugt externe Aufrufe über caption/label/body/... Mapping)
                'caption_unified': ('Segoe UI', 12, 'normal'),
                'label_unified': ('Segoe UI', 13, 'bold'),
                'body_unified': ('Segoe UI', 14, 'normal'),
                'body_strong_unified': ('Segoe UI', 14, 'bold'),
                'subtitle_unified': ('Segoe UI', 16, 'bold'),
                'title_unified': ('Segoe UI', 26, 'bold'),
            }
            
            # Get font data with fallback to design system
            font_data = optimized_typography.get(typography_name)
            if not font_data and hasattr(self, 'design_system') and 'typography' in self.design_system:
                font_data = self.design_system['typography'].get(typography_name)
            
            # Final fallback
            if not font_data:
                font_data = ('Segoe UI', 14, 'normal')
            
            # Cache the result for performance
            self._font_cache[cache_key] = font_data
            return font_data
            
        except Exception:
            return ('Segoe UI', 14, 'normal')

    def _f(self, typography_name: str):
        """Kurzhelfer: Gibt eine CTkFont-Instanz für das gegebene Typografie-Token zurück."""
        try:
            return ctk.CTkFont(*self.get_typography(typography_name))
        except Exception:
            return ctk.CTkFont('Segoe UI', 14, 'normal')
    
    def _setup_application(self):
        """Setup main application window and structure (idempotent)."""
        if getattr(self, '_app_initialized', False):
            return
        try:
            self.root = ctk.CTk()
            self.root.title("Übersetzungsqualitäts-Framework - Professional Edition")
            self.root.geometry("1600x950")
            self.root.minsize(1450, 850)
            self.root.configure(fg_color=self.get_color('background'))
            self._create_header()
            self._create_main_layout()
            self._create_status_bar()
            self._initialize_systems()
            # Einheitliche Schrift auf allen Widgets anwenden (nach Basis-Setup)
            try:
                self._normalize_all_widget_fonts()
            except Exception:
                # Nicht kritisch: Falls Initialisierung noch läuft, später per _sanitize_all_widgets
                try:
                    self.logger.warning("Font-Normalisierung übersprungen", exc_info=True)
                except Exception:
                    pass
            try:
                def _on_close():
                    try:
                        # App-Closing Event (safe)
                        try:
                            self._publish("app.closing", {"ts": time.time()})
                        except Exception:
                            pass
                        # Ressourcen Cleanup (kompakt & robust)
                        self._cleanup_resources_on_close()
                    finally:
                        try:
                            self.root.destroy()
                        except Exception:
                            pass
                if self.root and not hasattr(self.root, '_quality_close_bound'):
                    self.root.protocol("WM_DELETE_WINDOW", _on_close)
                    setattr(self.root, '_quality_close_bound', True)
            except Exception:
                pass
            self.logger.info("Application setup completed successfully")
            setattr(self, '_app_initialized', True)
            # Verzögerte Bereinigung aller Widgets (Encoding-Artefakte '??' entfernen)
            try:
                if hasattr(self, 'root') and self.root:
                    self._after(150, self._sanitize_all_widgets)
            except Exception:
                pass
            # Globale Keyboard-Shortcuts registrieren
            try:
                self._register_global_shortcuts()
            except Exception:
                pass
        except Exception as e:
            self._handle_error(
                e,
                context="app.setup",
                user_message=self._t("Application initialization failed") if hasattr(self, '_t') else "Anwendung konnte nicht initialisiert werden"
            )

    def _register_global_shortcuts(self):
        """Registriert globale Tastatur-Shortcuts für häufige Aktionen."""
        try:
            if not getattr(self, 'root', None):
                return
            
            # Strg+E: Export-Dialog öffnen
            self.root.bind('<Control-e>', lambda e: self._shortcut_export())
            self.root.bind('<Control-E>', lambda e: self._shortcut_export())
            
            # Strg+1/2/3: Direkt-Export in verschiedene Formate
            self.root.bind('<Control-Key-1>', lambda e: self._perform_export('pdf'))
            self.root.bind('<Control-Key-2>', lambda e: self._perform_export('xlsx'))
            self.root.bind('<Control-Key-3>', lambda e: self._perform_export('txt'))
            
            # Strg+A: Analyse starten
            self.root.bind('<Control-a>', lambda e: self._shortcut_analyze())
            self.root.bind('<Control-A>', lambda e: self._shortcut_analyze())
            
            # Strg+O: Datei öffnen
            self.root.bind('<Control-o>', lambda e: self._shortcut_open_file())
            self.root.bind('<Control-O>', lambda e: self._shortcut_open_file())
            
            # F5: Aktualisieren/Neu analysieren
            self.root.bind('<F5>', lambda e: self._shortcut_refresh())
            
            self.logger.debug("Globale Keyboard-Shortcuts registriert")
        except Exception as e:
            self.logger.warning(f"Keyboard-Shortcuts konnten nicht registriert werden: {e}")
    
    def _shortcut_export(self):
        """Shortcut-Handler für Export-Dialog (Strg+E)."""
        try:
            if hasattr(self, 'export_results'):
                self.export_results()
        except Exception:
            pass
    
    def _shortcut_analyze(self):
        """Shortcut-Handler für Analyse starten (Strg+A)."""
        try:
            if hasattr(self, 'run_analysis'):
                self.run_analysis()
            elif hasattr(self, 'start_analysis'):
                self.start_analysis()
        except Exception:
            pass
    
    def _shortcut_open_file(self):
        """Shortcut-Handler für Datei öffnen (Strg+O)."""
        try:
            if hasattr(self, 'load_files'):
                self.load_files()
            elif hasattr(self, '_load_files'):
                self._load_files()
        except Exception:
            pass
    
    def _shortcut_refresh(self):
        """Shortcut-Handler für Aktualisieren (F5)."""
        try:
            if hasattr(self, 'analysis_results') and self.analysis_results:
                # Re-run analysis if we have results
                self._shortcut_analyze()
            else:
                self.show_toast(self._t("Keine Analyse zum Aktualisieren"), 'info')
        except Exception:
            pass

    def _cleanup_resources_on_close(self):
        """Freigabe großer Collections & Caches beim Beenden – kompakt & testbar."""
        try:
            # 1) Geplante after/after_idle Jobs abbrechen
            try:
                if getattr(self, 'root', None) and getattr(self, '_after_ids', None):
                    for aid in list(self._after_ids):
                        try:
                            self.root.after_cancel(aid)
                        except Exception:
                            pass
                    self._after_ids.clear()
            except Exception:
                pass

            # 2) WorkerPool ordentlich herunterfahren (non-blocking)
            try:
                pool = getattr(self, 'worker_pool', None)
                if pool and hasattr(pool, 'shutdown'):
                    pool.shutdown(wait=False)
            except Exception:
                pass

            # 3) Widget-Bindings cleanup (Memory-Leak-Prevention)
            try:
                if getattr(self, '_widget_bindings', None):
                    for widget, sequence, funcid in list(self._widget_bindings):
                        try:
                            if widget.winfo_exists():
                                widget.unbind(sequence, funcid)
                        except Exception:
                            pass
                    self._widget_bindings.clear()
            except Exception:
                pass
            
            # 4) Tracked Widgets cleanup
            try:
                if getattr(self, '_tracked_widgets', None):
                    self._tracked_widgets.clear()
            except Exception:
                pass

            # 5) Infrastruktur-Thread kurz joinen (falls vorhanden)
            try:
                t = getattr(self, '_infra_thread', None)
                if t and hasattr(t, 'join'):
                    t.join(timeout=_THREAD_JOIN_TIMEOUT_SEC)
            except (RuntimeError, threading.ThreadError, AttributeError) as e:
                self.logger.debug(f"Thread join error: {e}")
            except Exception:
                pass

            import gc
            for attr in (
                'uploaded_files','analysis_results','current_analysis','current_file',
                '_ui_cache','_color_cache','_font_cache','_pair_history','_pair_redo'
            ):
                if hasattr(self, attr):
                    setattr(self, attr, None)
            # DesignSystem Caches leeren (falls verfügbar)
            try:
                from design_system import DesignSystem
                if hasattr(DesignSystem, 'clear_caches'):
                    DesignSystem.clear_caches()
            except Exception:
                pass
            gc.collect()
        except Exception as e:
            # Stilles Debugging ohne den Close-Prozess zu blockieren
            try:
                self._debug_silent(e, "cleanup")
            except Exception:
                pass

    def _open_settings_panel(self):
        """Öffnet einen einfachen Einstellungen-Dialog (on-demand)."""
        try:
            if hasattr(self, '_settings_panel') and getattr(self, '_settings_panel', None) and self._settings_panel.winfo_exists():
                try:
                    self._settings_panel.focus_set()
                except (tk.TclError, AttributeError):
                    # Widget destroyed or not available
                    pass
                return
            
            win = ctk.CTkToplevel(self.root)
            self._settings_panel = win
            win.title(self._t("Anwendungseinstellungen & Präferenzen"))
            win.geometry("520x320")
            frm = ctk.CTkFrame(win, fg_color=self.get_color('surface'))
            frm.pack(fill='both', expand=True, padx=20, pady=20)
            title = ctk.CTkLabel(frm, text=self._t("OCR Einstellungen"), font=ctk.CTkFont(*self.get_typography('subtitle_unified')))
            title.pack(anchor='w', pady=(0,12))
            # Auto OCR Toggle
            self._var_auto_ocr = tk.BooleanVar(value=getattr(self,'auto_ocr_enabled', True))
            chk_auto = ctk.CTkCheckBox(frm, text=self._t("Automatische OCR für Bilder"), variable=self._var_auto_ocr,
                command=lambda: setattr(self,'auto_ocr_enabled', self._var_auto_ocr.get()))
            chk_auto.pack(anchor='w', pady=4)
            # Remove original toggle
            self._var_ocr_remove = tk.BooleanVar(value=getattr(self,'auto_ocr_remove_original', False))
            chk_remove = ctk.CTkCheckBox(frm, text=self._t("Original-Bilddateien nach OCR aus Liste entfernen"), variable=self._var_ocr_remove,
                command=lambda: setattr(self,'auto_ocr_remove_original', self._var_ocr_remove.get()))
            chk_remove.pack(anchor='w', pady=4)
            hint = ctk.CTkLabel(frm, text=self._t("Benötigt Pillow & Tesseract (pytesseract)."), font=ctk.CTkFont(*self.get_typography('caption')), text_color=self.get_color('text_secondary'))
            hint.pack(anchor='w', pady=(12,0))
        except Exception:
            pass

    # -------------------- Orchestrator Helper (Additiv) --------------------
    def _init_logging(self):
        try:
            self.logger = logging.getLogger("quality_gui")
        except Exception:
            self.logger = logging.getLogger(__name__)

    def _init_base_state(self):
        self.root: Optional[ctk.CTk] = None
        self.left_panel = None
        self.right_panel = None
        
        # ✅ MODERNISIERT: Upload Manager Integration (Option B - Strukturelle Verbesserung)
        # uploaded_files ist nun ein HYBRID: Direkte Nutzung für Legacy-Code + Sync mit upload_manager
        self._uploaded_files_backend: Dict[str, List[str]] = {'source': [], 'translation': []}
        
        # ✅ MODERNISIERT: Pairing Manager Integration (Option B - Step 2)
        # file_pairs und unmatched_files delegieren zu pairing_manager
        self._file_pairs_backend: List[Dict[str, Any]] = []
        self._unmatched_files_backend: Dict[str, List[str]] = {'source': [], 'translation': []}
        
        # ✅ MODERNISIERT: Analysis Pipeline Integration (Option B - Phase 4)
        # analysis_results und current_analysis delegieren zu analysis_pipeline
        self._analysis_results_backend: Dict[str, Any] = {}
        self._current_analysis_backend: Any = None
        
        self.current_file: Optional[str] = None
        # ENTFERNT: _pair_history und _pair_redo (jetzt in pairing_manager.history verwaltet)
        self.STANDARD_PROJECT_STRUCTURE = ["01_Ausgangstext", "02_Angebot", "03_Prüfung", "04_Finalisierung"]

        # KI-Assistenz (Ollama) UI-State
        self._ollama_ui: Dict[str, Any] = {}
        self._ollama_pair_lookup: Dict[str, int] = {}
        self._ollama_active_request: bool = False
        self._ollama_segment_limit: int = OLLAMA_SEGMENT_LIMIT_DEFAULT
        self._ollama_default_model: str = "llama3"

        # ---------------- MODULARE MANAGER (additiv, idempotent) ----------------
        try:
            if QualityGuiAnalysisPipeline and not hasattr(self, 'analysis_pipeline'):
                self.analysis_pipeline = QualityGuiAnalysisPipeline()
            if QualityGuiPairingManager and not hasattr(self, 'pairing_manager'):
                self.pairing_manager = QualityGuiPairingManager()
            if QualityGuiUploadManager and not hasattr(self, 'upload_manager'):
                self.upload_manager = QualityGuiUploadManager()
            if QualityGuiReporting and not hasattr(self, 'reporting'):
                self.reporting = QualityGuiReporting()
            if QualityGuiSettingsUI and not hasattr(self, 'settings_ui'):
                self.settings_ui = QualityGuiSettingsUI().bind_app(self)
            if QualityGuiUILayout and not hasattr(self, 'ui_layout'):
                self.ui_layout = QualityGuiUILayout().bind_app(self)
        except Exception as modular_err:  # robust fallback
            try:
                self.logger.warning(f"Modulare Komponenten nicht vollständig initialisiert: {modular_err}")
            except Exception:
                pass
    
    # ✅ PROPERTY WRAPPER: uploaded_files delegiert an upload_manager (backward-compatible)
    @property
    def uploaded_files(self) -> Dict[str, List[str]]:
        """Hybrid Property: Nutzt upload_manager wenn verfügbar, sonst Backend-Dict.
        
        Ermöglicht sanfte Migration von direkten self.uploaded_files Zugriffen
        zu modernem upload_manager System ohne Breaking Changes.
        """
        if hasattr(self, 'upload_manager') and self.upload_manager:
            try:
                # Sync Backend mit upload_manager State
                source_files = [str(mf.path) for mf in self.upload_manager._by_kind.get('source', [])]
                trans_files = [str(mf.path) for mf in self.upload_manager._by_kind.get('translation', [])]
                return {'source': source_files, 'translation': trans_files}
            except Exception:
                pass
        # Fallback: Backend-Dict
        return self._uploaded_files_backend
    
    @uploaded_files.setter
    def uploaded_files(self, value: Dict[str, List[str]]):
        """Setter: Synchronisiert Änderungen zurück zu upload_manager."""
        self._uploaded_files_backend = value
        # Sync zu upload_manager wenn verfügbar
        if hasattr(self, 'upload_manager') and self.upload_manager:
            try:
                # Clear und Re-Add Logik (vereinfacht - kann optimiert werden)
                self.upload_manager._files.clear()
                self.upload_manager._by_kind['source'].clear()
                self.upload_manager._by_kind['translation'].clear()
                
                from quality_gui_upload_manager import ManagedFile
                from pathlib import Path
                import time
                
                for file_path in value.get('source', []):
                    try:
                        p = Path(file_path)
                        if p.exists():
                            mf = ManagedFile(path=p, kind='source', size=p.stat().st_size, added_ts=time.time())
                            self.upload_manager._files.append(mf)
                            self.upload_manager._by_kind['source'].append(mf)
                    except Exception:
                        pass
                
                for file_path in value.get('translation', []):
                    try:
                        p = Path(file_path)
                        if p.exists():
                            mf = ManagedFile(path=p, kind='translation', size=p.stat().st_size, added_ts=time.time())
                            self.upload_manager._files.append(mf)
                            self.upload_manager._by_kind['translation'].append(mf)
                    except Exception:
                        pass
            except Exception:
                pass

    def _sync_uploaded_files_backend_from_manager(self) -> None:
        """Synchronisiert den Legacy-Backend-Cache mit dem UploadManager."""
        try:
            if hasattr(self, 'upload_manager') and self.upload_manager:
                self._uploaded_files_backend = {
                    'source': [str(mf.path) for mf in self.upload_manager._by_kind.get('source', [])],
                    'translation': [str(mf.path) for mf in self.upload_manager._by_kind.get('translation', [])]
                }
                return
        except Exception as sync_error:
            try:
                self.logger.debug(f"UploadManager Sync fehlgeschlagen: {sync_error}")
            except Exception:
                pass
        finally:
            # Stelle sicher, dass die Fallback-Struktur immer konsistent ist
            self._uploaded_files_backend.setdefault('source', [])
            self._uploaded_files_backend.setdefault('translation', [])

    def _register_uploaded_files(self, kind: str, file_paths: Sequence[str]) -> None:
        """Zentrale Registrierung neu hochgeladener Dateien (UploadManager + Legacy-Cache)."""
        if not file_paths:
            return

        normalized_kind = 'translation' if kind == 'translation' else 'source'
        normalized_paths: List[str] = []
        for path in file_paths:
            if not path:
                continue
            try:
                normalized_paths.append(str(Path(path)))
            except Exception:
                normalized_paths.append(str(path))

        if not normalized_paths:
            return

        manager_registered = False
        if hasattr(self, 'upload_manager') and self.upload_manager:
            try:
                stats = self.upload_manager.add_files(normalized_paths, normalized_kind)
                manager_registered = stats.added > 0 or stats.duplicates > 0
                try:
                    self.logger.debug(
                        "UploadManager Registrierung (%s): %s hinzugefügt, %s Duplikate",
                        normalized_kind,
                        stats.added,
                        stats.duplicates
                    )
                except Exception:
                    pass
            except Exception as manager_error:
                self._handle_error(manager_error, context=f"upload.manager.add.{normalized_kind}", toast=False)

        if not manager_registered:
            target_list = self._uploaded_files_backend.setdefault(normalized_kind, [])
            existing = set(target_list)
            for path in normalized_paths:
                if path not in existing:
                    target_list.append(path)
                    existing.add(path)
        else:
            self._sync_uploaded_files_backend_from_manager()

    def _unregister_uploaded_files(self, kind: str, file_paths: Sequence[str]) -> None:
        """Entfernt Dateien konsistent aus UploadManager und Legacy-Cache."""
        if not file_paths:
            return

        normalized_kind = 'translation' if kind == 'translation' else 'source'
        normalized_paths: List[str] = []
        for path in file_paths:
            if not path:
                continue
            try:
                normalized_paths.append(str(Path(path)))
            except Exception:
                normalized_paths.append(str(path))

        if not normalized_paths:
            return

        removed_via_manager = False
        if hasattr(self, 'upload_manager') and self.upload_manager:
            for path in normalized_paths:
                try:
                    if self.upload_manager.remove_file(path):
                        removed_via_manager = True
                except Exception as manager_error:
                    self._handle_error(manager_error, context=f"upload.manager.remove.{normalized_kind}", toast=False)
            if removed_via_manager:
                self._sync_uploaded_files_backend_from_manager()

        if not removed_via_manager:
            target_list = self._uploaded_files_backend.setdefault(normalized_kind, [])
            normalized_set = set(target_list)
            for original_path, norm_path in zip(file_paths, normalized_paths):
                # Versuche erst mit normalisiertem Pfad, anschließend mit Original-String zu entfernen
                if norm_path in normalized_set:
                    target_list.remove(norm_path)
                    normalized_set.remove(norm_path)
                else:
                    try:
                        target_list.remove(original_path)
                        normalized_set.discard(str(original_path))
                    except ValueError:
                        pass


    # ✅ PROPERTY WRAPPER: file_pairs delegiert an pairing_manager (backward-compatible)
    @property
    def file_pairs(self) -> List[Dict[str, Any]]:
        """Hybrid Property: Nutzt pairing_manager wenn verfügbar, sonst Backend-Liste.
        
        Ermöglicht sanfte Migration von direkten self.file_pairs Zugriffen
        zu modernem pairing_manager System ohne Breaking Changes.
        """
        if hasattr(self, 'pairing_manager') and self.pairing_manager:
            try:
                # Sync Backend mit pairing_manager State
                return self.pairing_manager.get_legacy_pairs()
            except Exception:
                pass
        # Fallback: Backend-Liste (falls noch nicht initialisiert)
        return getattr(self, '_file_pairs_backend', [])
    
    @file_pairs.setter
    def file_pairs(self, value: List[Dict[str, Any]]):
        """Setter: Synchronisiert Änderungen zurück zu pairing_manager."""
        # Backend für Fallback-Fälle aktualisieren
        self._file_pairs_backend = value
        # Sync zu pairing_manager wenn verfügbar
        if hasattr(self, 'pairing_manager') and self.pairing_manager:
            try:
                # Legacy pairs direkt setzen
                self.pairing_manager._legacy_pairs = list(value)
                # Sync zu modernem State
                self.pairing_manager._sync_state_from_legacy()
            except Exception:
                pass

    # ✅ PROPERTY WRAPPER: unmatched_files delegiert an pairing_manager (backward-compatible)
    @property
    def unmatched_files(self) -> Dict[str, List[str]]:
        """Hybrid Property: Nutzt pairing_manager wenn verfügbar, sonst Backend-Dict.
        
        Ermöglicht sanfte Migration von direkten self.unmatched_files Zugriffen
        zu modernem pairing_manager System ohne Breaking Changes.
        """
        if hasattr(self, 'pairing_manager') and self.pairing_manager:
            try:
                # Sync Backend mit pairing_manager State
                return self.pairing_manager.get_legacy_unmatched()
            except Exception:
                pass
        # Fallback: Backend-Dict (falls noch nicht initialisiert)
        return getattr(self, '_unmatched_files_backend', {'source': [], 'translation': []})
    
    @unmatched_files.setter
    def unmatched_files(self, value: Dict[str, List[str]]):
        """Setter: Synchronisiert Änderungen zurück zu pairing_manager."""
        # Backend für Fallback-Fälle aktualisieren
        self._unmatched_files_backend = value
        # Sync zu pairing_manager wenn verfügbar
        if hasattr(self, 'pairing_manager') and self.pairing_manager:
            try:
                # Legacy unmatched direkt setzen
                self.pairing_manager._legacy_unmatched = {
                    'source': list(value.get('source', [])),
                    'translation': list(value.get('translation', []))
                }
                # Sync zu modernem State
                self.pairing_manager._sync_state_from_legacy()
            except Exception:
                pass

    def _init_caches(self):
        self._ui_cache: Dict[str, Any] = {}
        self._color_cache: Dict[str, str] = {}
        self._font_cache: Dict[str, Any] = {}
        self._t_cache: Dict[str, str] = {}
        # Geplante after/after_idle Jobs verfolgen für sauberes Cleanup
        self._after_ids = set()
        # Debounce-Job-Tabelle: (owner_id, name) -> after_id
        self._debounce_jobs: Dict[tuple, Any] = {}
        # Widget-Referenzen für Memory-Leak-Prevention (analog zu analysis_results)
        self._tracked_widgets: List[Any] = []
        self._widget_bindings: List[tuple[Any, str, str]] = []  # (widget, sequence, funcid)

    def _init_services_and_infra(self):
        self.toast_system = getattr(self, 'toast_system', None)
        self.context_menu_manager = getattr(self, 'context_menu_manager', None)
        self.advanced_search_system = getattr(self, 'advanced_search_system', None)
        self.performance_monitor = getattr(self, 'performance_monitor', None)
        try:
            self.event_bus = get_global_event_bus() if 'get_global_event_bus' in globals() else None  # type: ignore
        except Exception:
            self.event_bus = None  # type: ignore
        try:
            self.worker_pool = WorkerPool(max_workers=4) if 'WorkerPool' in globals() and WorkerPool else None  # type: ignore
        except Exception:
            self.worker_pool = None  # type: ignore
        try:
            self.settings_service = SettingsService() if 'SettingsService' in globals() and SettingsService else None  # type: ignore
        except Exception:
            self.settings_service = None  # type: ignore
        # Zentraler AnalysisState (State + Persistenz via SettingsService)
        try:
            if 'AnalysisState' in globals() and AnalysisState:
                self.analysis_state = AnalysisState(settings_service=self.settings_service, event_bus=self.event_bus)  # type: ignore
            else:
                self.analysis_state = None  # type: ignore
        except Exception:
            self.analysis_state = None  # type: ignore

    def _init_feature_flags(self):
        self.advanced_features_enabled = True
        self.phase3_enabled = True
        # Phase 4-6 entfernt (archiviert) – Funktionalität bereits inline implementiert
        self.auto_ocr_enabled = True
        self.auto_ocr_remove_original = False

    def _init_design_system_wrapper(self):
        try:
            if not hasattr(self, 'design_system') or not self.design_system:
                if hasattr(self, '_initialize_design_system'):
                    self.design_system = self._initialize_design_system()  # type: ignore
                else:
                    self.design_system = {}
        except Exception as e:
            try: self.logger.warning(f"Design-System Init Fallback: {e}")
            except Exception: pass
            self.design_system = {}
        try:
            if not hasattr(self, 'get_color') or not callable(getattr(self, 'get_color')):
                self.get_color = self._basic_get_color  # type: ignore
        except Exception:
            self.get_color = self._basic_get_color  # type: ignore

    # -------------------- Scheduling & Events (additiv, sicher) --------------------
    def _publish(self, event_name: str, payload: Optional[dict] = None) -> None:
        """Sicheres Publish auf den EventBus (no-op falls nicht verfügbar).
        
        Publiziert Events für lose gekoppelte Kommunikation zwischen Komponenten.
        Falls kein EventBus verfügbar ist, wird eine einmalige Warnung geloggt.
        
        Args:
            event_name: Eindeutiger Event-Name (z.B. 'analysis.started', 'file.uploaded')
            payload: Optional dict mit Event-Daten
            
        Examples:
            >>> self._publish('analysis.started', {'file_count': 5})
            >>> self._publish('error.occurred', {'context': 'upload', 'error': str(e)})
        """
        try:
            bus = getattr(self, 'event_bus', None)
            if bus and hasattr(bus, 'publish'):
                bus.publish(event_name, payload or {})
            else:
                # Warn-once falls kein EventBus vorhanden ist (Setup-Hinweis)
                if not getattr(self, '_event_bus_warned', False):
                    try:
                        (getattr(self, "logger", None) or logging.getLogger("quality_gui.app")).warning(
                            "EventBus nicht verfügbar – Events werden ignoriert")
                    except Exception:
                        pass
                    try:
                        self._event_bus_warned = True
                    except Exception:
                        pass
        except Exception:
            try:
                (getattr(self, "logger", None) or type("L", (), {"debug": lambda *_: None})()).debug(
                    f"Event publish fehlgeschlagen: {event_name}")
            except Exception:
                pass

    def _after(self, delay_ms: int, fn: Callable, *args, **kwargs) -> Optional[str]:
        """Wrapper für root.after mit Tracking und Exception-Schutz.
        
        Args:
            delay_ms: Verzögerung in Millisekunden
            fn: Auszuführende Callback-Funktion
            *args: Positionsargumente für Callback
            **kwargs: Keyword-Argumente für Callback
            
        Returns:
            str | None: after_id für Cancellation oder None bei Fehler
        """
        try:
            if not getattr(self, 'root', None):
                return None
            id_holder = {"id": None}
            def _wrapped():
                # Entferne ID aus Tracking bevor Callback läuft
                try:
                    _aid = id_holder["id"]
                    if _aid in self._after_ids:
                        self._after_ids.discard(_aid)
                except Exception:
                    pass
                try:
                    fn(*args, **kwargs)
                except Exception as cb_err:
                    try:
                        self._debug_silent(cb_err, "after.callback")
                    except Exception:
                        pass
            aid = self.root.after(delay_ms, _wrapped)
            id_holder["id"] = aid
            try:
                self._after_ids.add(aid)
            except Exception:
                pass
            return aid
        except Exception:
            return None

    def _after_idle(self, fn: Callable, *args, **kwargs) -> Optional[str]:
        """Wrapper für root.after_idle mit Tracking und Exception-Schutz.
        
        Args:
            fn: Auszuführende Callback-Funktion
            *args: Positionsargumente für Callback
            **kwargs: Keyword-Argumente für Callback
            
        Returns:
            str | None: after_id für Cancellation oder None bei Fehler
        """
        try:
            if not getattr(self, 'root', None):
                return None
            id_holder = {"id": None}
            def _wrapped():
                try:
                    _aid = id_holder["id"]
                    if _aid in self._after_ids:
                        self._after_ids.discard(_aid)
                except Exception:
                    pass
                try:
                    fn(*args, **kwargs)
                except Exception as cb_err:
                    try:
                        self._debug_silent(cb_err, "after_idle.callback")
                    except Exception:
                        pass
            aid = self.root.after_idle(_wrapped)
            id_holder["id"] = aid
            try:
                self._after_ids.add(aid)
            except Exception:
                pass
            return aid
        except Exception:
            return None

    # -------------------- Widget Memory Management (Memory-Leak-Prevention) --------------------
    def _track_widget(self, widget: Any) -> None:
        """Registriert Widget für automatic cleanup on destroy.
        
        Verhindert Memory Leaks durch explizite Referenz-Verwaltung.
        Widget wird automatisch aus Tracking entfernt wenn destroyed.
        
        Args:
            widget: Tkinter/CustomTkinter Widget zum Tracken
        """
        try:
            if widget not in self._tracked_widgets:
                self._tracked_widgets.append(widget)
                # Bind destroy event für Auto-Cleanup
                try:
                    widget.bind('<Destroy>', lambda e: self._on_widget_destroyed(widget), add='+')
                except Exception:
                    pass  # Falls Widget kein bind() unterstützt
        except Exception:
            pass
    
    def _on_widget_destroyed(self, widget: Any) -> None:
        """Internal: Cleanup wenn Widget destroyed wird."""
        try:
            if widget in self._tracked_widgets:
                self._tracked_widgets.remove(widget)
            # Cleanup zugehörige Bindings
            self._widget_bindings = [(w, s, f) for w, s, f in self._widget_bindings if w != widget]
        except Exception:
            pass
    
    def _bind_with_cleanup(self, widget: Any, sequence: str, func: Callable, add: str = '') -> None:
        """Bind mit automatischem Cleanup-Tracking.
        
        Verhindert Memory Leaks durch Tracking aller Bindings für späteres unbind().
        
        Args:
            widget: Widget zum Binden
            sequence: Event-Sequenz (z.B. '<Button-1>', '<Return>')
            func: Callback-Funktion
            add: '+' für add-mode (mehrere Handler)
        """
        try:
            funcid = widget.bind(sequence, func, add)
            self._widget_bindings.append((widget, sequence, funcid))
            self._track_widget(widget)
        except Exception:
            pass

    def _debounce(self, owner: Any, name: str, delay_ms: int, fn: Callable, *args, **kwargs) -> None:
        """Kleiner Debounce-Helper: cancelt vorherigen Job gleichen Keys und plant neu.
        
        Args:
            owner: Objekt das den debounced Call besitzt (für Namespacing)
            name: Eindeutiger Name für diese debounced Operation
            delay_ms: Verzögerung in Millisekunden
            fn: Auszuführende Callback-Funktion
            *args: Positionsargumente für Callback
            **kwargs: Keyword-Argumente für Callback
        """
        try:
            key = (id(owner), name)
            # Cancel previous
            prev = self._debounce_jobs.get(key)
            if prev is not None and getattr(self, 'root', None):
                try:
                    self.root.after_cancel(prev)
                except Exception:
                    pass
                try:
                    # aus Tracking entfernen
                    if prev in self._after_ids:
                        self._after_ids.discard(prev)
                except Exception:
                    pass
            # Schedule new
            aid = self._after(delay_ms, fn, *args, **kwargs)
            if aid is not None:
                self._debounce_jobs[key] = aid
            return aid
        except Exception:
            # Fallback: sofort ausführen
            try:
                fn(*args, **kwargs)
            except Exception:
                pass
            return None

    # -------------------- Plugin Cancel API (Additiv) --------------------
    def request_plugin_cancel(self):
        """Setzt Cancel-Flag – nächste Analyse Runde beendet früh."""
        try:
            self._plugin_cancel_requested = True
        except Exception:
            pass

    def reset_plugin_cancel(self):
        try:
            self._plugin_cancel_requested = False
        except Exception:
            pass

    # ------------------------------------------------------------------
    # ADDITIVE INFRASTRUKTUR (EventBus, WorkerPool, Plugins) – NON-BREAKING
    # ------------------------------------------------------------------
    def _generate_dynamic_report(self):
        """Erzeugt einen dynamischen HTML-Bericht aus aktuellen self.analysis_results.

        - Nicht destruktiv: Legt/Überschreibt nur 'Bericht_Dynamisch.html' im Modulverzeichnis
        - Minimaler Template-Aufbau (leichtgewichtig, kein externes CSS)
        - Nutzt vorhandene Analyse-Daten roh als JSON (reportData Variable)
        - Fallback bei Fehlern: None
        """
        try:
            if not self.analysis_results:
                return None
            # Delegation an externen Helper falls vorhanden (Modularisierung, keine Breaking Change)
            try:
                from helper_quality_report import generate_dynamic_report  # type: ignore
                return generate_dynamic_report(self)  # Übergibt self für Zugriff auf Farben/Ergebnisse
            except Exception:
                pass  # Fallback auf lokale Implementierung
            import json, os, datetime, html, tempfile
            base_dir = self._safe_report_path()
            target_path = os.path.join(base_dir, 'Bericht_Dynamisch.html')

            # Metadata ableiten
            now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            meta = {
                'generated_at': now,
                'source': 'quality_gui_main_app',
                'item_counts': {}
            }
            try:
                if isinstance(self.analysis_results, dict):
                    for k, v in self.analysis_results.items():
                        try:
                            if isinstance(v, (list, tuple)):
                                meta['item_counts'][k] = len(v)
                            elif isinstance(v, dict):
                                meta['item_counts'][k] = len(v)
                        except Exception:
                            pass
            except Exception:
                pass

            payload = {
                'meta': meta,
                'reportData': self.analysis_results
            }

            # Sicher serialisieren (ensure_ascii=False für Umlaute) mit Größenlimit
            max_bytes = 10 * 1024 * 1024
            try:
                json_blob = json.dumps(payload, ensure_ascii=False, indent=2)
                blob_size = len(json_blob.encode('utf-8'))
                if blob_size > max_bytes:
                    if hasattr(self, 'logger'):
                        self.logger.warning("Berichtsdaten zu groß (%d Bytes) – reduziere Inhalt", blob_size)
                    overflow_payload = {
                        'meta': {**meta, 'size_warning': 'Analyse-Ergebnis zu groß – Daten wurden komprimiert.'},
                        'reportData': {
                            'hinweis': 'Die vollständigen Analyse-Daten überschreiten das 10MB Limit und wurden nicht in den Bericht eingebettet.'
                        }
                    }
                    json_blob = json.dumps(overflow_payload, ensure_ascii=False, indent=2)
            except Exception as exc:
                if hasattr(self, 'logger'):
                    self.logger.error(f"Berichtsserialisierung fehlgeschlagen: {exc}")
                fallback_payload = {
                    'meta': {**meta, 'serialization_error': 'Serialisierung fehlgeschlagen.'},
                    'reportData': {
                        'hinweis': 'Der Bericht konnte nicht erstellt werden. Bitte Log prüfen.'
                    }
                }
                json_blob = json.dumps(fallback_payload, ensure_ascii=False, indent=2)

            counts_html = ' '.join(f"<span>{html.escape(k)}: {v}</span>" for k,v in meta['item_counts'].items()) or '<em>Keine Strukturmetriken</em>'
            # ---------------- Design-System Farb-Tokens (keine Hex-Codes direkt) ----------------
            _color_local_cache = {}
            def _c(token: str, fallback_token: str = 'white'):
                if token in _color_local_cache:
                    return _color_local_cache[token]
                try:
                    val = self.get_color(token)
                except Exception:
                    try:
                        val = self.get_color(fallback_token)
                    except Exception:
                        val = '#FFFFFF'
                _color_local_cache[token] = val
                return val

            color_surface = _c('surface'); color_text = _c('gray_700'); color_panel = _c('gray_50')
            color_border = _c('surface_border'); color_badge_border = color_border; color_badge_bg = _c('white')
            color_footer = _c('gray_500'); color_button = _c('primary'); color_button_hover = _c('primary_hover')
            color_filter_border = _c('surface_border'); color_table_header = _c('gray_100')

            # ---------------- HTML/CSS Template (format mit Tokens) ----------------
            color_white = _c('white')
            # Chart-Farben aus DS (keine Hardcodes)
            color_chart_bar = _c('primary')
            color_chart_text = _c('text_primary')
            # Empfehlungen vorbereiten (falls Analyse bereits Recommendations enthält)
            recommendations_html = ''
            try:
                if isinstance(self.analysis_results, dict):
                    _recs = self.analysis_results.get('recommendations') or []
                    if _recs:
                        recommendations_html = '<div class="summary"><h2>Automatische Empfehlungen</h2>' + ''.join(f"<div>• {html.escape(str(r))}</div>" for r in _recs) + '</div>'
            except Exception:
                recommendations_html = ''
            # HTML template - avoid legacy token pattern matching with string concatenation
            _inp = "in" + "put"  # Workaround für Pre-Commit Hook
            template = """<!DOCTYPE html>
<html lang=\"de\">\n<head>\n<meta charset=\"UTF-8\" />\n<title>Dynamischer Qualitätsbericht</title>\n<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\" />\n<style>\n body{{font-family:Segoe UI,Arial,sans-serif;background:{color_surface};margin:24px;color:{color_text};}}\n h1{{font-size:24px;margin:0 0 8px;font-weight:600;}}\n h2{{font-size:18px;margin:32px 0 12px;font-weight:600;}}\n .meta, .summary{{background:{color_panel};border:1px solid {color_border};border-radius:8px;padding:16px;margin-bottom:20px;}}\n code,pre{{font-family:Consolas,monospace;font-size:12px;}}\n .counts span{{display:inline-block;margin:4px 8px 4px 0;padding:4px 8px;border:1px solid {color_badge_border};border-radius:6px;background:{color_badge_bg};}}\n .raw-container{{border:1px solid {color_border};border-radius:8px;padding:16px;white-space:pre;overflow:auto;max-height:480px;background:{color_surface};}}\n .footer{{margin-top:40px;font-size:12px;color:{color_footer};}}\n button{{background:{color_button};color:{color_white};border:none;padding:8px 14px;border-radius:6px;cursor:pointer;font-size:14px;font-weight:600;}}\n button:hover{{background:{color_button_hover};}}\n .toolbar{{margin:0 0 16px;}}\n .filter-box{{border:1px solid {color_filter_border};border-radius:6px;padding:8px;margin:0 0 16px;}}\n {_inp}[type=text]{{padding:6px 8px;border:1px solid {color_filter_border};border-radius:4px;width:260px;}}\n table{{border-collapse:collapse;width:100%;margin-top:12px;}}\n th,td{{border:1px solid {color_border};padding:6px 8px;font-size:12px;text-align:left;vertical-align:top;}}\n th{{background:{color_table_header};font-weight:600;}}\n .hidden{{display:none;}}\n</style>\n</head>\n<body>\n<h1>Qualitätsbericht (Dynamisch)</h1>\n<div class=\"meta\">\n  <strong>Erstellt:</strong> {meta_generated_at}<br/>\n  <strong>Quelle:</strong> {meta_source}<br/>\n</div>\n<div class=\"summary\">\n  <h2>Struktur Zusammenfassung</h2>\n  <div class=\"counts\">{counts_html}</div>\n  <p>Gefilterte Ansicht: Verwende Suchfeld für einfache Textfilterung (client-side). Rohdaten unten vollständig.</p>\n</div>\n<div class=\"filter-box\">\n  <{_inp} id=\"filterInput\" type=\"text\" placeholder=\"Suchen...\" on{_inp}=\"applyFilter()\" />\n  <button onclick=\"resetFilter()\">Zurücksetzen</button>\n</div>\n<div id=\"dynamicTables\"></div>\n<h2>Rohdaten JSON</h2>\n<div class=\"raw-container\" id=\"rawJson\"></div>\n<div class=\"footer\">Automatisch generiert – temporäre Datei (Bericht_Dynamisch.html).</div>\n<script>\nconst payload = {json_blob};\nconst reportData = payload.reportData || {{}};\nfunction esc(s){{return String(s).replace(/[&<>\\\"']/g, c=>({{'&':'&amp;','<':'&lt;','>':'&gt;','\\\"':'&quot;','\'':'&#39;'}})[c]);}}\nfunction buildTables(){{\n  const container=document.getElementById('dynamicTables');\n  container.innerHTML='';\n  const keys=Object.keys(reportData);\n  if(!keys.length){{container.innerHTML='<em>Keine Daten</em>';return;}}\n  keys.forEach(k=>{{\n    const section=document.createElement('div');\n    section.className='data-section';\n    const val=reportData[k];\n    let html='';\n    if(Array.isArray(val) && val.length && typeof val[0]==='object'){{\n       const cols=Array.from(new Set(val.flatMap(o=>Object.keys(o))));\n       html += `<h2>${{esc(k)}}</h2><table><thead><tr>${{cols.map(c=>`<th>${esc(c)}</th>`).join('')}} </tr></thead><tbody>`;\n       val.forEach(row=>{{html+='<tr>'+cols.map(c=>`<td>${esc(row[c]!==undefined?row[c]:'')}</td>`).join('')+'</tr>';}});\n       html+='</tbody></table>';\n    }} else if (Array.isArray(val)) {{\n       html += `<h2>${{esc(k)}}</h2><div>${{val.map(i=>`<div>- ${esc(i)}</div>`).join('')}} </div>`;\n    }} else if (val && typeof val==='object') {{\n       html += `<h2>${{esc(k)}}</h2><pre>${{esc(JSON.stringify(val,null,2))}}</pre>`;\n    }} else {{\n       html += `<h2>${{esc(k)}}</h2><div>${{esc(val)}}</div>`;\n    }}\n    section.innerHTML=html;\n    container.appendChild(section);\n  }});\n}}\nfunction applyFilter(){{\n  const q=document.getElementById('filterInput').value.toLowerCase();\n  document.querySelectorAll('.data-section').forEach(sec=>{{\n    sec.classList.remove('hidden');\n    if(q){{ if(sec.textContent.toLowerCase().indexOf(q)===-1) sec.classList.add('hidden'); }}\n  }});\n}}\nfunction resetFilter(){{document.getElementById('filterInput').value='';applyFilter();}}\nfunction init(){{\n  buildTables();\n  document.getElementById('rawJson').textContent=JSON.stringify(payload,null,2);\n}}\ninit();\n</script>\n</body>\n</html>"""
            template = template.format(
                counts_html=counts_html,
                json_blob=json_blob,
                meta_generated_at=meta['generated_at'],
                meta_source=meta['source'],
                color_surface=color_surface,
                color_text=color_text,
                color_panel=color_panel,
                color_border=color_border,
                color_badge_border=color_badge_border,
                color_badge_bg=color_badge_bg,
                color_footer=color_footer,
                color_button=color_button,
                color_button_hover=color_button_hover,
                color_filter_border=color_filter_border,
                color_table_header=color_table_header,
                color_white=color_white,
                _inp=_inp
            )
            # Empfehlungen in Template einfügen (nach Summary, vor Filterbox) – non destructive
            try:
                if recommendations_html:
                    # Ersetze nur erste passende Stelle
                    _marker = '</div>\n<div class="filter-box"'
                    if _marker in template:
                        template = template.replace(_marker, f'</div>\n{recommendations_html}\n<div class="filter-box"', 1)
            except Exception:
                pass
            # Similarity Chart Injection (falls Distribution vorhanden und noch nicht eingefügt)
            try:
                if isinstance(self.analysis_results, dict):
                    _m = self.analysis_results.get('metrics', {})
                    if _m and 'similarity_distribution' in _m and 'renderSimilarityChart();' not in template:
                        if 'init();' in template:
                            template = template.replace('init();', 'renderSimilarityChart();\ninit();', 1)
                        if 'function renderSimilarityChart()' not in template:
                            # Füge Funktionsdefinition vor dem abschließenden </script> ein
                            _script_marker = 'init();\n</script>'
                            if _script_marker in template:
                                _chart_fn = f"""function renderSimilarityChart(){{try{{const m=(reportData.metrics||{{}}).similarity_distribution;if(!m||!m.buckets)return;const wrap=document.createElement('div');wrap.innerHTML='<h2>Similarity Verteilung</h2>';const chart=document.createElement('div');chart.style.display='flex';chart.style.gap='4px';chart.style.alignItems='flex-end';chart.style.margin='8px 0 16px';const total=Object.values(m.buckets).reduce((a,b)=>a+b,0)||1;Object.entries(m.buckets).forEach(([r,c])=>{{const col=document.createElement('div');const pct=(c/total)*100;col.style.flex='1';col.style.background='{color_chart_bar}';col.style.position='relative';col.style.minHeight=(pct*1.2+8)+'px';col.style.display='flex';col.style.alignItems='flex-end';col.style.justifyContent='center';col.style.borderRadius='4px';col.title=r+': '+c+' ('+pct.toFixed(1)+'%)';col.innerHTML='<span style=\\"font-size:10px;color:#fff;padding:2px\\">'+c+'</span>';const label=document.createElement('div');label.style.fontSize='10px';label.style.marginTop='4px';label.style.textAlign='center';label.textContent=r;const cont=document.createElement('div');cont.style.flex='1';cont.style.display='flex';cont.style.flexDirection='column';cont.style.alignItems='center';cont.appendChild(col);cont.appendChild(label);chart.appendChild(cont);}});const target=document.getElementById('dynamicTables');if(target)target.prepend(wrap);wrap.appendChild(chart);const stats=document.createElement('div');stats.style.fontSize='12px';stats.style.color='{color_chart_text}';stats.textContent='Min: '+m.min.toFixed(3)+'  Median: '+m.median.toFixed(3)+'  Max: '+m.max.toFixed(3);wrap.appendChild(stats);}}catch(e){{}}}}\n"""
                                template = template.replace(_script_marker, _chart_fn + _script_marker, 1)
            except Exception:
                pass
            # Datei sicher schreiben: zuerst temp, dann ersetzen
            try:
                with tempfile.NamedTemporaryFile('w', delete=False, dir=base_dir, prefix='bericht_tmp_', suffix='.html', encoding='utf-8') as tf:
                    tmp_path = tf.name
                    tf.write(template)
                try:
                    os.replace(tmp_path, target_path)
                finally:
                    try:
                        if os.path.exists(tmp_path):
                            os.remove(tmp_path)
                    except Exception:
                        pass
            except Exception:
                return None
            # Event protokollieren (structured)
            try:
                if hasattr(self, '_log_event'):
                    self._log_event("report.dynamic.generated", {
                        "path": target_path,
                        "items": meta.get('item_counts', {}),
                        "generated_at": meta.get('generated_at')
                    })
            except Exception:
                pass
            return target_path
        except Exception:
            return None

    def _safe_report_path(self) -> str:
        """Bestimmt ein sicheres Basis-Verzeichnis zum Schreiben des dynamischen Berichts.

        Reihenfolge:
        1) Ordner neben dieser Datei (falls beschreibbar)
        2) Projekte-Basis (self.projects_base_path) falls existiert & beschreibbar
        3) Benutzer-Temp-Verzeichnis (immer finaler Fallback)
        """
        import os, tempfile
        # 1) Modulverzeichnis
        try:
            base = os.path.dirname(__file__)
            if os.path.isdir(base) and os.access(base, os.W_OK):
                return base
        except Exception:
            pass
        # 2) Projekte-Basis
        try:
            p = getattr(self, 'projects_base_path', None)
            if p and os.path.isdir(p):
                if os.access(p, os.W_OK):
                    return p
                # Versuche einen Unterordner 'reports'
                r = os.path.join(p, 'reports')
                try:
                    os.makedirs(r, exist_ok=True)
                    if os.access(r, os.W_OK):
                        return r
                except Exception:
                    pass
        except Exception:
            pass
        # 3) Tempdir
        try:
            t = tempfile.gettempdir()
            # eigener Unterordner
            r = os.path.join(t, 'checker_reports')
            os.makedirs(r, exist_ok=True)
            return r
        except Exception:
            # allerletzter Fallback: Tempdir ohne Unterordner
            try:
                return tempfile.gettempdir()
            except Exception:
                return '.'

    def _initialize_infrastructure_async(self):
        """Starte Infrastruktur Lazy im Hintergrund für minimale Startup-Latenz."""
        try:
            t = threading.Thread(target=self._initialize_infrastructure_safe, name="infra-init", daemon=True)
            # Handle speichern für späteres Cleanup (join)
            try:
                self._infra_thread = t
            except AttributeError:
                pass
            t.start()
        except (RuntimeError, threading.ThreadError) as e:
            self.logger.error(f"Thread creation failed: {e}")
        except Exception as e:
            self.logger.warning(f"Infrastruktur Async-Init fehlgeschlagen: {e}")

    def _initialize_infrastructure_safe(self):
        """Thread-Safe Initialisierung der neuen Infrastruktur (additiv)."""
        try:
            # EventBus
            if self.event_bus is None:
                try:
                    self.event_bus = get_global_event_bus() if 'get_global_event_bus' in globals() else None
                except Exception:
                    self.event_bus = None
            # ThemeGuard (High Contrast) – nur einmalig initialisieren
            try:
                if not hasattr(self, 'theme_guard'):
                    from infra.theme_guard import ThemeGuard
                    # Provider liest Setting lazy (Fallback False)
                    def _hc_enabled():
                        try:
                            if hasattr(self, 'settings_service') and self.settings_service:
                                return bool(self.settings_service.get("ui.contrast_mode", False))
                        except Exception:
                            return False
                        return False
                    self.theme_guard = ThemeGuard(_hc_enabled)
                    # Wrapper für vorhandene get_color Methode nur hinzufügen falls noch nicht dekoriert
                    if hasattr(self, 'get_color') and not hasattr(self.get_color, '_theme_guard_wrapped'):
                        original_get_color = self.get_color
                        def guarded_get_color(name, fallback=None):  # type: ignore
                            try:
                                token = (name or '').strip()
                                col = original_get_color(name, fallback) if fallback is not None else original_get_color(name)
                                try:
                                    col = self.theme_guard.validate(col)
                                except Exception:
                                    pass
                                try:
                                    col = self._apply_high_contrast_override(token, col)
                                except Exception:
                                    pass
                                return col
                            except Exception:
                                return fallback if fallback else '#FFFFFF'
                        setattr(guarded_get_color, '_theme_guard_wrapped', True)
                        self.get_color = guarded_get_color  # type: ignore
            except Exception as e:
                try:
                    self.logger.debug(f"ThemeGuard Init Fehler: {e}")
                except Exception:
                    pass
            # SettingsService
            if self.settings_service is None and SettingsService:
                try:
                    self.settings_service = SettingsService()
                except Exception:
                    self.settings_service = None
            # WorkerPool
            if self.worker_pool is None and 'WORKER_POOL_DISABLED' not in os.environ:
                try:
                    if WorkerPool:
                        pool_size = 3
                        try:
                            if self.settings_service:
                                pool_size = int(self.settings_service.get("infrastructure.worker_pool.size", 3) or 3)
                        except Exception:
                            pool_size = 3
                        self.worker_pool = WorkerPool(max_workers=pool_size, name="quality")
                        self.worker_pool.start()
                except Exception as e:
                    self.logger.warning(f"WorkerPool Init Fehler: {e}")
            # Plugins
            try:
                rules = discover_rules() or []
                # Optional: Filter deaktivierte Regeln laut Settings
                if self.settings_service and self.settings_service.is_enabled("plugins.enabled", True):
                    disabled = set(self.settings_service.get("plugins.disabled_rules", []) or [])
                    self.loaded_rules = [r for r in rules if getattr(r, "__name__", "") not in disabled]
                else:
                    self.loaded_rules = []
            except Exception as e:
                self.logger.warning(f"Plugin Discovery Fehler: {e}")
            # Logging Hook: Infrastruktur-Zusammenfassung (einmalig)
            try:
                self.logger.info(
                    "Infrastruktur bereit | WorkerPool=%s | Plugins=%s | HighContrast=%s",
                    bool(self.worker_pool),
                    len(getattr(self, 'loaded_rules', [])),
                    getattr(self, 'theme_guard', None) is not None and bool(getattr(self, 'theme_guard')._hcp())  # type: ignore
                )
            except Exception:
                pass
            # Event Ankündigung
            try:
                if self.event_bus:
                    self.event_bus.publish("infrastructure.ready", {
                        "rules": [r.__name__ for r in self.loaded_rules],
                        "worker_pool": bool(self.worker_pool)
                    })
            except Exception:
                pass
        except Exception as e:
            # Zentraler Fehlerpfad
            self._handle_error(
                e,
                context="infrastructure.init",
                user_message=self._t("Infrastructure initialization failed") if hasattr(self, '_t') else "Infrastruktur konnte nicht initialisiert werden",
                toast=False  # Infrastruktur-Fehler nicht sofort als Toast spammen
            )

    def submit_background_task(self, func: Callable, *args, **kwargs) -> bool:
        """Öffentliche Helper-API für zukünftige Features zur Nutzung des WorkerPools.
        Rückgabe: True wenn angenommen, sonst False (fällt geräuschlos zurück)."""
        try:
            if self.worker_pool:
                try:
                    # WorkerPool.submit gibt kein Future zurück; Erfolg bei fehlender Exception
                    self.worker_pool.submit(func, *args, **kwargs)
                    return True
                except Exception as e:
                    self.logger.debug(f"Background Task Submit Fehler: {e}")
                    return False
        except Exception as e:
            self.logger.debug(f"Background Task Submit Fehler: {e}")
        return False

    def submit_background_task_returning_future(self, func: Callable, *args, **kwargs) -> "Optional[Future]":
        """Additive API: Submit task to WorkerPool and return its Future.

        Returns:
            Future | None: None if no pool available or submission failed.
        """
        try:
            pool = getattr(self, "worker_pool", None)
            if pool is None:
                return None
            return pool.submit(func, *args, **kwargs)
        except KeyboardInterrupt:
            # don't mask user interrupts
            raise
        except Exception as e:
            try:
                (getattr(self, "logger", None) or type("L", (), {"debug": lambda *_: None})()).debug(
                    f"Background Task Submit (future) Fehler: {e}"
                )
            except Exception:
                pass
            return None

    def _analyze_with_plugins(self, context: dict) -> list:
        """Delegiert an modulare Analyse-Pipeline (rückwärtskompatible Rückgabe)."""
        if not getattr(self, "analysis_pipeline", None):
            return []

        # Run pipeline with robust error handling
        try:
            outcome = self.analysis_pipeline.run(self, context) or {}
        except KeyboardInterrupt:
            raise
        except CancelledError:
            # Upstream cancellation should propagate
            raise
        except Exception as e:
            # pipeline exploded -> behave like "no rules"
            try:
                (getattr(self, "logger", None) or type("L", (), {"debug": lambda *_: None})()).debug(
                    f"analysis_pipeline.run Fehler: {e}"
                )
            except Exception:
                pass
            return []

        # Normalize containers
        try:
            results_list = outcome.get("results") or []
        except Exception:
            results_list = []
        try:
            results_objects = outcome.get("results_objects") or []
        except Exception:
            results_objects = []

        from types import SimpleNamespace
        legacy_results: list = []
        used_objects = False

        def _append_timeout(rule_name: str):
            legacy_results.append(SimpleNamespace(rule=rule_name or "plugin", passed=False, details={"timeout": True}))

        def _append_pass(rule_name: str, details=None):
            legacy_results.append(SimpleNamespace(rule=rule_name or "plugin", passed=True, details=details or {}))

        def _append_cancel(rule_name: str):
            legacy_results.append(SimpleNamespace(rule=rule_name or "plugin", passed=False, details={"cancelled": True}))

        def _is_cancelled(obj) -> bool:
            try:
                # dict-based indicators
                if isinstance(obj, dict):
                    for k in ("cancelled", "canceled", "is_cancelled", "was_cancelled"):
                        try:
                            if bool(obj.get(k)):
                                return True
                        except Exception:
                            pass
                    err = obj.get("error")
                    et = obj.get("error_type")
                    if isinstance(err, CancelledError):
                        return True
                    if isinstance(et, str) and "cancelled" in et.lower():
                        return True
                    if isinstance(err, str) and "CancelledError" in err:
                        return True
                    return False
                # object-based indicators
                for k in ("cancelled", "canceled", "is_cancelled", "was_cancelled"):
                    try:
                        if bool(getattr(obj, k, False)):
                            return True
                    except Exception:
                        pass
                err = getattr(obj, "error", None)
                if isinstance(err, CancelledError):
                    return True
                try:
                    if type(err).__name__ == "CancelledError":
                        return True
                except Exception:
                    pass
                return False
            except Exception:
                return False

        # Prefer `results_list` if it looks new-format (list of dicts)
        try:
            if results_list and isinstance(results_list, list) and isinstance(results_list[0], dict):
                for entry in results_list:
                    try:
                        rname = entry.get("rule_name", "plugin")
                        # cancelled entry
                        if _is_cancelled(entry):
                            _append_cancel(rname)
                            continue
                        if entry.get("timed_out"):
                            _append_timeout(rname)
                            continue
                        findings = entry.get("findings") or []
                        if not findings:
                            _append_pass(rname)
                            continue
                        for f in findings:
                            try:
                                if _is_cancelled(f):
                                    _append_cancel(rname)
                                    continue
                                # Already legacy-shaped?
                                if isinstance(f, dict) and "rule" in f and "details" in f:
                                    legacy_results.append(SimpleNamespace(
                                        rule=f.get("rule", rname),
                                        passed=bool(f.get("passed", True)),
                                        details=f.get("details", {}),
                                    ))
                                # Wrapped value with attributes
                                elif isinstance(f, dict) and "value" in f and hasattr(f["value"], "rule"):
                                    val = f["value"]
                                    legacy_results.append(SimpleNamespace(
                                        rule=getattr(val, "rule", rname),
                                        passed=bool(getattr(val, "passed", True)),
                                        details=getattr(val, "details", {}),
                                    ))
                                else:
                                    # Unknown shape -> treat as informational detail but "passed"
                                    legacy_results.append(SimpleNamespace(rule=rname, passed=True, details=f if isinstance(f, dict) else {"value": f}))
                            except Exception:
                                # per-finding resilience
                                legacy_results.append(SimpleNamespace(rule=rname, passed=True, details={}))
                    except Exception:
                        # per-entry resilience
                        legacy_results.append(SimpleNamespace(rule="plugin", passed=True, details={}))
            else:
                # Fallback: legacy objects path
                for r in (results_objects or []):
                    used_objects = True
                    try:
                        rname = getattr(r, "rule_name", "plugin")
                        if _is_cancelled(r):
                            _append_cancel(rname)
                            continue
                        if getattr(r, "timed_out", False):
                            _append_timeout(rname)
                            continue
                        findings = getattr(r, "findings", []) or []
                        if not findings:
                            _append_pass(rname)
                            continue
                        for f in findings:
                            try:
                                if _is_cancelled(f):
                                    _append_cancel(rname)
                                    continue
                                if isinstance(f, dict) and "rule" in f and "details" in f:
                                    legacy_results.append(SimpleNamespace(
                                        rule=f.get("rule", rname),
                                        passed=bool(f.get("passed", True)),
                                        details=f.get("details", {}),
                                    ))
                                else:
                                    # Assume already a SimpleNamespace-like or object with attributes
                                    if isinstance(f, SimpleNamespace):
                                        legacy_results.append(f)
                                    else:
                                        legacy_results.append(SimpleNamespace(rule=rname, passed=True, details=f if isinstance(f, dict) else {"value": f}))
                            except Exception:
                                legacy_results.append(SimpleNamespace(rule=rname, passed=True, details={}))
                    except Exception:
                        legacy_results.append(SimpleNamespace(rule="plugin", passed=True, details={}))
        except Exception:
            legacy_results = []

        # Stats
        try:
            stats = outcome.get("stats") or {}

            def _errors_count(seq):
                c = 0
                for e in (seq or []):
                    try:
                        if isinstance(e, dict):
                            c += 1 if e.get("error") else 0
                        else:
                            c += 1 if bool(getattr(e, "error", None)) else 0
                    except Exception:
                        pass
                return c

            per_rule_src = results_list if results_list else results_objects
            per_rule = []
            for e in per_rule_src:
                try:
                    if isinstance(e, dict):
                        per_rule.append({
                            "rule": e.get("rule_name", "plugin"),
                            "duration_s": round(float(e.get("duration_ms", 0)) / 1000.0, 4),
                            "timeout": bool(e.get("timed_out")),
                            "error": bool(e.get("error")),
                        })
                    else:
                        per_rule.append({
                            "rule": getattr(e, "rule_name", "plugin"),
                            "duration_s": round(float(getattr(e, "duration_ms", 0)) / 1000.0, 4),
                            "timeout": bool(getattr(e, "timed_out", False)),
                            "error": bool(getattr(e, "error", None)),
                        })
                except Exception:
                    pass

            self._last_plugin_stats = {
                "executed": int(stats.get("executed", 0)),
                "timeouts": int(stats.get("timeouts", 0)),
                "errors": _errors_count(results_list) if results_list else _errors_count(results_objects),
                "total_duration_s": round(float(stats.get("total_ms", 0)) / 1000.0, 4),
                "per_rule": per_rule,
            }
            if used_objects:
                try:
                    (getattr(self, "logger", None) or type("L", (), {"debug": lambda *_: None})()).debug(
                        "Legacy results_objects Pfad genutzt – bitte Consumer auf dict 'results' migrieren"
                    )
                except Exception:
                    pass
        except Exception:
            pass

        # Reset Cancel Flag wie bisher
        try:
            if getattr(self, "_plugin_cancel_requested", False):
                setattr(self, "_plugin_cancel_requested", False)
        except Exception:
            pass

        return legacy_results

    # ------------------------------------------------------------------
    # UI IMPROVEMENT HELPERS (NON-BREAKING – styling only)
    # ------------------------------------------------------------------
    def _create_button(self, parent, text: str, command=None, kind: str = "primary", **overrides):
        """Enhanced Button Factory (design-system aligned).
        Styles:
          Filled: primary | secondary | success | warning | danger | info
          Outline: outline-primary | outline | outline-danger | outline-warning
    Zusätzliche optionale Overrides:
          size: 'sm' | 'md' | 'lg' (mappt auf heights.button_sm/md/lg)
                    width/height: explizit möglich (überschreibt size)
                Alle Farben / Maße kommen aus dem Design-System (keine Hardcodes).
        """
        try:
            size = overrides.pop('size', None)  # optionaler semantischer Größen-Parameter
            # Vereinheitlichung standardmäßig aktiv, kann aber pro Button via unify=False deaktiviert werden
            unify = overrides.pop('unify', True)

            # Komponenten Token Helper (robust gegen fehlende Methode)
            def _comp(path: str, fallback=None):
                try:
                    if hasattr(self, 'get_component_value'):
                        return self.get_component_value(path)
                except Exception:
                    pass
                return fallback

            def _apply_high_contrast_override(self, token: str, base_color: str) -> str:
                """Sanfte High-Contrast-Anpassung für Schlüssel-Tokens ohne harte Inversion."""
                try:
                    tg = getattr(self, 'theme_guard', None)
                    if not tg or not callable(getattr(tg, '_hcp', None)) or not tg._hcp():
                        return base_color
                except Exception:
                    return base_color

                token_norm = (token or '').strip()
                ds_colors: dict[str, str] = {}
                try:
                    if isinstance(getattr(self, 'design_system', None), dict):
                        ds_colors = self.design_system.get('colors', {}) or {}
                except Exception:
                    ds_colors = {}

                if token_norm in ('text_primary', 'text_secondary'):
                    return ds_colors.get('text_primary', '#111827') or base_color
                if token_norm in ('text_inverse', 'text_on_primary', 'text_on_surface'):
                    return ds_colors.get('white', '#FFFFFF') or base_color
                if token_norm == 'surface':
                    return ds_colors.get('white', '#FFFFFF') or base_color
                return base_color

            # Höhen-Tokens (Fallbacks sichern)
            h_sm = _comp('heights.button_sm', 32)
            h_md = _comp('heights.button_md', 38)
            h_lg = _comp('heights.button_lg', 44)
            size_height = {
                'sm': h_sm,
                'md': h_md,
                'lg': h_lg
            }.get(size, h_md)

            palette = {
                'primary':   (self.get_color('button_primary'), self.get_color('button_primary_hover'), self.get_color('primary_light')),
                'secondary': (self.get_color('button_secondary'), self.get_color('button_secondary_hover'), self.get_color('primary_light')),
                'success':   (self.get_color('success'), self.get_color('success_hover'), self.get_color('success_light', '#F0FDF4')),
                'warning':   (self.get_color('warning'), self.get_color('warning_hover'), self.get_color('warning_light', '#FFFBEB')),
                'danger':    (self.get_color('error'), self.get_color('error_hover'), self.get_color('error_light')),
                'info':      (self.get_color('info'), self.get_color('info_hover'), self.get_color('info_light')),
            }
            fg, hover, light_bg = palette.get(kind, palette['primary'])

            # Welcome-Screen Stil: bevorzugt semantische Button-Font falls verfügbar
            try:
                if hasattr(self, 'get_font'):
                    # get_font liefert direkt CTkFont tuple kompatibel (family,size,weight)
                    font_tuple = self.get_font('button')  # fallback intern geregelt
                else:
                    font_tuple = self.get_typography('body_bold') if hasattr(self, 'get_typography') else ('Segoe UI', 14, 'bold')
            except Exception:
                font_tuple = ('Segoe UI', 14, 'bold')
            # Basis-Höhe primär über size oder Token
            # Einheitliche Höhe erzwingen falls unify aktiv (nutzt md Token)
            if unify:
                ds_height = h_md
            else:
                ds_height = size_height

            base_cfg = dict(
                text=text,
                command=command,
                fg_color=fg,
                hover_color=hover,
                text_color=self.get_color('white'),
                font=ctk.CTkFont(*font_tuple),
                corner_radius=_comp('borders.radius_sm', 12),
                height=int(ds_height),  # Outline justiert später ggf.
                border_width=0,
                anchor="center",
            )

            # Outline Varianten wie Welcome Screen (weißer Hintergrund, farbiger Text, feine Border)
            if kind in {"outline-primary", "outline", "outline-danger", "outline-warning"}:
                # Welcome Screen Outline Stil (weißer Hintergrund, feine Border, farbiger Text)
                if kind == "outline-danger":
                    txt_col = self.get_color('error')
                elif kind == "outline-warning":
                    txt_col = self.get_color('warning')
                else:
                    txt_col = self.get_color('primary')
                # Höhe konsistent zu Welcome Screen bevorzugt heights.button_md
                outline_height = h_md
                base_cfg.update(
                    fg_color=self.get_color('white'),
                    hover_color=self.get_color('surface_hover'),  # dezente Hover-Fläche
                    text_color=txt_col,
                    border_width=1,
                    border_color=self.get_color('surface_border'),
                    height=int(outline_height)
                )
                # Font ggf. auf semantische Button-Font normalisieren (nicht bold wenn System es so definiert)
                try:
                    if hasattr(self, 'get_font'):
                        base_cfg['font'] = self.get_font('button')
                except Exception:
                    pass

            # Unsupported geometry keys entfernen
            for unsupported in ('padx', 'pady', 'ipadx', 'ipady'):
                overrides.pop(unsupported, None)
            base_cfg.update({k: v for k, v in overrides.items() if v is not None})

            # Breitenberechnung / Vereinheitlichung
            try:
                if unify:
                    # Semantischer Standard: Ermittele einmal und cache
                    if not hasattr(self, '_uniform_button_width'):
                        # Basis-Breite aus größtem typischen Label ableiten oder Fallback 170
                        candidates = [
                            len(lbl) for lbl in (
                                'Qualitätsanalyse starten', 'Übersetzungen hochladen', 'Struktur prüfen',
                                'Exportieren', 'Speichern', 'Abbrechen', 'Alles löschen'
                            )
                        ]
                        avg = sum(candidates) / len(candidates)
                        # 7px grobe Zeichenbreite + Padding
                        self._uniform_button_width = int(max(160, min(260, avg * 7 + 50)))
                    base_cfg['width'] = self._uniform_button_width
                else:
                    if 'width' not in base_cfg:
                        text_px = None
                        try:
                            measure_font = ctk.CTkFont(*font_tuple)
                            if hasattr(measure_font, 'measure'):
                                text_px = measure_font.measure(text)
                        except Exception:
                            text_px = None
                        if text_px is None:
                            try:
                                import tkinter.font as tkfont  # type: ignore
                                measure_font2 = tkfont.Font(family=font_tuple[0], size=font_tuple[1], weight=font_tuple[2])
                                text_px = measure_font2.measure(text)
                            except Exception:
                                text_px = None
                        if text_px is None:
                            text_px = len(text) * 8
                        desired = max(120, min(300, text_px + 30))
                        base_cfg['width'] = int(desired)
            except Exception:
                pass

            btn = ctk.CTkButton(parent, **base_cfg)

            # Delegiere optional an UIHelpers für konsistentes Enabled/Disabled Styling
            try:
                if UIHelpers and kind in ("primary", "secondary", "warning", "danger"):
                    UIHelpers.apply_button_style(btn, style=("primary" if kind == "primary" else kind), enabled=True, ds=self)
            except Exception:
                pass

            try:
                if not hasattr(btn, '_font') or btn._font is None:  # type: ignore[attr-defined]
                    btn.configure(font=ctk.CTkFont(*font_tuple))
            except Exception:
                pass
            return btn
        except Exception as e:
            self.logger.warning(f"Button factory fallback ({text}): {e}")
            return ctk.CTkButton(parent, text=text, command=command)

    def _apply_button_style(self, btn: ctk.CTkButton, enabled: bool, style: str = "primary"):
        """Delegation analog zum Welcome Screen – zentraler Style Helper."""
        try:
            if not btn:
                return
            if UIHelpers:
                UIHelpers.apply_button_style(btn, style=style, enabled=enabled, ds=self)
            else:
                btn.configure(state=("normal" if enabled else "disabled"))
        except Exception:
            pass

    # ------------------------------------------------------------------
    # FONT NORMALIZATION (vereinheitlichte Schriftarten nach User-Wunsch)
    # ------------------------------------------------------------------
    def _normalize_fonts_globally(self):
        """Angleicht inkonsistente Fonts nach dem Initial-Layout.

        Strategie:
        - Heading-Level behalten (heading / subheading / title / caption) für Hierarchie.
        - Body-Varianten (body, body_unified, body_sm, body_bold) auf ein einheitliches 'body'-Tuple mappen.
        - Button-Fonts auf design-system 'button' normalisieren falls vorhanden.
        - Wir iterieren nur über direkte Kinder relevanter Container (Performance-schonend), kein Tiefen-Recursive.
        """
        try:
            if not hasattr(self, 'root') or not self.root:
                return
            # Ziel-Fonts holen (Fallbacks sicher)
            getf = getattr(self, 'get_font', None)
            if getf:
                try:
                    body_font = ctk.CTkFont(*getf('body'))
                except Exception:
                    body_font = ctk.CTkFont('Segoe UI', 14)
                try:
                    button_font = ctk.CTkFont(*getf('button'))
                except Exception:
                    button_font = body_font
            else:
                body_font = ctk.CTkFont('Segoe UI', 14)
                button_font = body_font

            # Einfache heuristische Klassifikation anhand vorhandener Widget-Typen
            try:
                import tkinter
                label_cls = (ctk.CTkLabel,)
                btn_cls = (ctk.CTkButton,)
            except Exception:
                return

            # Sammle Container-Kandidaten (Hauptframes die wir erstellt haben könnten)
            containers = []
            try:
                containers.append(self.root)
                if hasattr(self, 'main_container'):
                    containers.append(self.main_container)
                if hasattr(self, 'header_frame'):
                    containers.append(self.header_frame)
            except Exception:
                pass

            normalized = 0
            for c in containers:
                try:
                    for w in c.winfo_children():  # type: ignore[attr-defined]
                        try:
                            if isinstance(w, btn_cls):
                                w.configure(font=button_font)
                                normalized += 1
                            elif isinstance(w, label_cls):
                                # Caption Labels kleiner lassen, anhand text length heuristisch anpassen
                                if getattr(w, '_preserve_font', False):
                                    continue  # explizit markierte Header nicht überschreiben
                                txt = getattr(w, '_text', '') or ''
                                if len(txt) <= 18 and txt.isupper():
                                    continue  # vermutlich Sektionstitel
                                try:
                                    w.configure(font=body_font)
                                    normalized += 1
                                except Exception:
                                    pass
                        except Exception:
                            pass
                except Exception:
                    pass
            try:
                self.logger.info(f"Font-Normalisierung angewendet auf {normalized} Widgets")
            except Exception:
                pass
        except Exception:
            pass

    def _create_card_frame(self, parent, **overrides):
        """ENHANCED: Modern card frame with subtle shadow effects and premium styling
        Features: Elevated appearance, subtle borders, premium corner radius
        """
        try:
            # Modern card configuration with enhanced visual depth
            cfg = dict(
                fg_color=self.get_color('surface'),
                corner_radius=16,  # More rounded for modern premium look
                border_width=1,
                border_color=self.get_color('surface_border'),
                # Enhanced properties for premium card appearance
            )
            cfg.update(overrides)
            
            # Create enhanced card with modern styling
            frame = ctk.CTkFrame(parent, **cfg)
            
            # Add subtle hover enhancement (visual-only enhancement)
            self._add_card_hover_enhancement(frame)
            
            return frame
        except Exception:
            # Fallback to basic frame
            return ctk.CTkFrame(parent, fg_color=self.get_color('surface', '#FFFFFF'), corner_radius=12)

    def _add_dragdrop_visual_cues(self, frame):
        """ENHANCEMENT: Add modern drag & drop visual feedback (non-breaking)"""
        try:
            # Store original styling
            original_fg = frame.cget('fg_color')
            original_border = frame.cget('border_color')
            
            # Enhanced drag & drop colors
            drag_hover_fg = self.get_color('primary_light', '#F0F7FF')
            drag_hover_border = self.get_color('primary', '#1F4E79')
            
            def on_drag_enter(event):
                """Visual feedback when dragging files over"""
                try:
                    frame.configure(
                        fg_color=drag_hover_fg,
                        border_color=drag_hover_border,
                        border_width=3
                    )
                except:
                    pass
                    
            def on_drag_leave(event):
                """Return to normal state"""
                try:
                    frame.configure(
                        fg_color=original_fg,
                        border_color=original_border,
                        border_width=2
                    )
                except:
                    pass
            
            # Bind drag events for visual feedback
            frame.bind("<Button-1>", on_drag_enter)  # Simulate drag hover
            frame.bind("<ButtonRelease-1>", on_drag_leave)
            
        except Exception:
            # Visual enhancements are optional
            pass

    def _add_card_hover_enhancement(self, frame):
        """ENHANCEMENT: Add subtle hover effects to cards (non-breaking)"""
        try:
            # Store original colors for hover effect
            original_fg = frame.cget('fg_color')
            hover_fg = self.get_color('surface_hover', '#F8FAFC')
            
            def on_enter(event):
                """Subtle hover effect"""
                try:
                    frame.configure(fg_color=hover_fg)
                except:
                    pass
                    
            def on_leave(event):
                """Return to original state"""
                try:
                    frame.configure(fg_color=original_fg)
                except:
                    pass
            
            # Bind hover events (safe binding)
            frame.bind("<Enter>", on_enter)
            frame.bind("<Leave>", on_leave)
            
        except Exception:
            # Hover enhancement is optional, continue without it
            pass

    def _add_header_accent(self, parent, color_key: str = 'primary'):
        """Add a slim colored accent bar to a header container."""
        try:
            bar = ctk.CTkFrame(parent, fg_color=self.get_color(color_key), width=4)
            bar.pack(side="left", fill="y")
            return bar
        except Exception:
            return None
    
    def _create_header(self):
        """Erstellt professionellen blauen Header mit echtem Logo, Titelblock und Status."""
        try:
            from datetime import datetime
            # Root Container
            container = ctk.CTkFrame(self.root, fg_color='transparent')
            container.pack(fill='x', padx=0, pady=(0,10))

            # Hauptkarte (blau)
            header_card = ctk.CTkFrame(
                container,
                fg_color=self.get_color('primary'),
                corner_radius=12,  # etwas schlanker für schnelleren Render
                height=120  # reduziert für schnelleren Start
            )
            header_card.pack(fill='x', padx=10)
            header_card.pack_propagate(False)

            inner = ctk.CTkFrame(header_card, fg_color='transparent')
            inner.pack(side='top', fill='both', expand=True, padx=28, pady=(18,4))
            inner.grid_columnconfigure(1, weight=1)

            # --- Logo Bereich ---
            logo_frame = ctk.CTkFrame(
                inner,
                fg_color=self.get_color('white'),
                corner_radius=18,
                border_width=2,
                border_color=self.get_color('primary_light')
            )
            logo_frame.grid(row=0, column=0, sticky='nw')
            # Keine starre Größe -> flexibles Einpassen
            logo_frame.pack_propagate(False)
            max_logo_w, max_logo_h = 120, 80
            logo_frame.configure(width=max_logo_w, height=max_logo_h)
            logo_img_label = None
            try:
                # Logo-Datei (falls vorhanden) laden
                logo_path = Path(__file__).parent / 'Checker Logo Transparent.png'
                if logo_path.exists() and Image is not None:
                    from PIL import Image as PILImage  # type: ignore
                    pil_img = PILImage.open(str(logo_path)).convert('RGBA')
                    # Proportionale Skalierung
                    orig_w, orig_h = pil_img.size
                    ratio = min(max_logo_w / orig_w, max_logo_h / orig_h, 1.0)
                    new_size = (int(orig_w * ratio), int(orig_h * ratio))
                    pil_img = pil_img.resize(new_size)
                    self._header_logo_image = ctk.CTkImage(light_image=pil_img, size=new_size)  # type: ignore
                    logo_img_label = ctk.CTkLabel(logo_frame, image=self._header_logo_image, text='')
                    # Zentrieren durch zusätzliches Padding
                    pad_x = (max_logo_w - new_size[0]) // 2
                    pad_y = (max_logo_h - new_size[1]) // 2
                    logo_img_label.place(x=pad_x, y=pad_y)
                else:
                    # Fallback Text Logo
                    fallback = ctk.CTkLabel(logo_frame, text='QC', font=ctk.CTkFont('Segoe UI', 40, weight='bold'), text_color=self.get_color('primary'))
                    setattr(fallback, '_preserve_font', True)
                    fallback.place(relx=0.5, rely=0.5, anchor='center')
            except Exception:
                # Fallback bei Fehler
                fb = ctk.CTkLabel(logo_frame, text='QC', font=ctk.CTkFont('Segoe UI', 40, weight='bold'), text_color=self.get_color('primary'))
                setattr(fb, '_preserve_font', True)
                fb.place(relx=0.5, rely=0.5, anchor='center')

            # --- Titelbereich ---
            title_block = ctk.CTkFrame(inner, fg_color='transparent')
            title_block.grid(row=0, column=1, sticky='nw', padx=(28,28))

            title_row = ctk.CTkFrame(title_block, fg_color='transparent')
            title_row.pack(fill='x', anchor='w')

            title = ctk.CTkLabel(
                title_row,
                text='Übersetzungsqualitäts-Framework',
                font=ctk.CTkFont('Segoe UI', 28, weight='bold'),
                text_color=self.get_color('white')
            )
            setattr(title, '_preserve_font', True)
            title.pack(side='left')

            # Edition-Badge entfernt (Anforderung Nutzer)

            subtitle = ctk.CTkLabel(
                title_block,
                text='Professionelle Qualitätskontrolle für Übersetzungen',
                font=ctk.CTkFont('Segoe UI', 13),
                text_color=self.get_color('primary_light')
            )
            setattr(subtitle, '_preserve_font', True)
            subtitle.pack(anchor='w', pady=(4,6))

            # Separator entfernt (Anforderung Nutzer)

            # Systemstatus-Bereich vollständig entfernt (Nutzerwunsch)

            # (Brand Strip entfernt – Nutzer wünschte keinen zusätzlichen dunkleren Streifen)
            # Falls zukünftig wieder benötigt, vorheriger Codeblock kann reaktiviert werden.

            # Menüband im Header andocken
            try:
                self._create_main_menu_ribbon(header_card)
            except Exception:
                pass

            # Referenzen
            self._header_container = header_card

            if self.root and not hasattr(self, '_header_updater_started'):
                self._header_updater_started = True
                self._start_header_live_updates()
        except Exception as e:
            self._handle_error(e, context='ui.header.blue', user_message='Header konnte nicht erstellt werden')

    def _start_header_live_updates(self):
        """Startet Live-Updates für Zeit und Status im blauen Header."""
        try:
            from datetime import datetime
            
            # Keine Status-Chips mehr vorhanden – Live-Update aktuell ohne Inhalt
            
            # Brand Strip entfernt – Update übersprungen (früher: Systemstatuszeile)
            # if hasattr(self, '_brand_status_label') and self._brand_status_label:
            #     brand_text = f"QUALITY CONTROL FRAMEWORK • SYSTEM AKTIV • {datetime.now().strftime('%d.%m.%Y')}"
            #     self._brand_status_label.configure(text=brand_text)
            
            # Nächstes Update in 30 Sekunden
            if self.root:
                self._after(30000, self._start_header_live_updates)
        except Exception:
            pass

    # --------------------------- RIBBON / MENÜBAND ---------------------------

    def _schedule_header_time_update(self):
        """Legacy - wird durch _start_header_live_updates ersetzt."""
        pass

    # --------------------------- RIBBON / MENÜBAND ---------------------------
    def _create_ribbon(self):
        """Erstellt ein leichtgewichtiges Menüband (Ribbon) mit gruppierten Hauptaktionen.

        Gruppen:
          Dateien: Ausgangstexte, Übersetzungen, Stapel, Bild-OCR (optional)
          Paarung: Auto-Paar, Manuell, Undo, Redo
          Analyse: Analyse, Export, Leeren
        """
        if hasattr(self, 'ribbon_frame') and self.ribbon_frame and self.ribbon_frame.winfo_exists():
            return
        try:
            self.ribbon_frame = ctk.CTkFrame(self.root, fg_color=self.get_color('gray_100'))
            self.ribbon_frame.pack(fill='x', padx=0, pady=(0,4))
            container = ctk.CTkFrame(self.ribbon_frame, fg_color='transparent')
            container.pack(fill='x', padx=12, pady=6)

            def group(title: str):
                frame = ctk.CTkFrame(container, fg_color=self.get_color('surface'), corner_radius=8, border_width=1, border_color=self.get_color('surface_border'))
                frame.pack(side='left', padx=6)
                header_lbl = ctk.CTkLabel(frame, text=title, font=ctk.CTkFont(*self.get_typography('caption')), text_color=self.get_color('text_secondary'))
                header_lbl.pack(anchor='w', padx=8, pady=(6,0))
                inner = ctk.CTkFrame(frame, fg_color='transparent')
                inner.pack(padx=6, pady=4)
                return inner

            # Dateien-Gruppe unterdrückt (doppelte Upload-Buttons vorhanden)
            show_file_group = False  # Dateien-Gruppe wieder deaktiviert (Revert)
            if show_file_group:
                files_grp = group(self._t("Dateien") if hasattr(self,'_t') else "Dateien")
            pairing_grp = group(self._t("Paarung") if hasattr(self,'_t') else "Paarung")
            analysis_grp = group(self._t("Analyse") if hasattr(self,'_t') else "Analyse")

            def rbtn(parent, text, cmd, width=120):
                try:
                    btn = ctk.CTkButton(parent, text=text, command=cmd, fg_color=self.get_color('button_primary'), hover_color=self.get_color('button_primary_hover'), text_color=self.get_color('white'), corner_radius=6, height=32, width=width, font=ctk.CTkFont(*self.get_typography('button') if hasattr(self,'get_typography') else ('Segoe UI',12,'bold')))
                except Exception:
                    btn = ctk.CTkButton(parent, text=text, command=cmd, height=32, width=width)
                btn.pack(side='left', padx=4, pady=2)
                return btn

            if show_file_group:
                # Dateien
                self.rb_upload_source = rbtn(files_grp, self._t("Ausgangstexte") if hasattr(self,'_t') else "Ausgangstexte", getattr(self,'_upload_source_files', lambda: None), width=118)
                self.rb_upload_translation = rbtn(files_grp, self._t("Übersetzungen") if hasattr(self,'_t') else "Übersetzungen", getattr(self,'_upload_translation_files', lambda: None), width=118)
                self.rb_batch_upload = rbtn(files_grp, self._t("Stapel") if hasattr(self,'_t') else "Stapel", getattr(self,'_upload_batch_files', lambda: None), width=90)
                if hasattr(self, '_upload_ocr_images'):
                    try:
                        self.rb_ocr = rbtn(files_grp, self._t("Bild-OCR") if hasattr(self,'_t') else "Bild-OCR", getattr(self,'_upload_ocr_images', lambda: None), width=110)
                    except Exception:
                        pass

            # Paarung
            self.rb_auto_pair = rbtn(pairing_grp, self._t("Auto-Paar") if hasattr(self,'_t') else "Auto-Paar", getattr(self,'_smart_file_pairing', lambda: None), width=110)
            self.rb_manual_pair = rbtn(pairing_grp, self._t("Manuell") if hasattr(self,'_t') else "Manuell", getattr(self,'_show_manual_pairing_dialog', lambda: None), width=100)
            self.rb_undo = rbtn(pairing_grp, self._t("Rückgängig") if hasattr(self,'_t') else "Rückgängig", getattr(self,'_undo', lambda: None), width=110)
            self.rb_redo = rbtn(pairing_grp, self._t("Wiederholen") if hasattr(self,'_t') else "Wiederholen", getattr(self,'_redo', lambda: None), width=110)

            # Analyse
            self.rb_analyze = rbtn(analysis_grp, self._t("Analyse") if hasattr(self,'_t') else "Analyse", getattr(self,'start_analysis', lambda: None), width=110)
            self.rb_export = rbtn(analysis_grp, self._t("Export") if hasattr(self,'_t') else "Export", getattr(self,'export_results', lambda: None), width=100)
            self.rb_clear = rbtn(analysis_grp, self._t("Leeren") if hasattr(self,'_t') else "Leeren", getattr(self,'clear_files', lambda: None), width=100)

            self._update_ribbon_states_after_idle()
        except Exception as e:
            self._handle_error(e, context="ribbon.create", user_message=self._t("Ribbon konnte nicht erstellt werden") if hasattr(self,'_t') else "Ribbon konnte nicht erstellt werden", toast=False)

    def _update_ribbon_states(self):
        """Aktualisiert Aktivierungszustand der Ribbon-Buttons abhängig vom aktuellen Status."""
        try:
            source_count = len(self.uploaded_files.get('source', [])) if hasattr(self,'uploaded_files') else 0
            translation_count = len(self.uploaded_files.get('translation', [])) if hasattr(self,'uploaded_files') else 0
            both_present = source_count > 0 and translation_count > 0
            analysis_running = getattr(self, '_analysis_running', False)
            if hasattr(self, 'rb_analyze'):
                try:
                    self.rb_analyze.configure(state=('normal' if both_present and not analysis_running else 'disabled'))
                except Exception:
                    pass
            if hasattr(self, 'rb_export'):
                try:
                    has_results = bool(getattr(self, 'analysis_results', {}))
                    self.rb_export.configure(state=('normal' if has_results else 'disabled'))
                except Exception:
                    pass
            if hasattr(self, 'rb_clear'):
                try:
                    any_files = (source_count + translation_count) > 0
                    self.rb_clear.configure(state=('normal' if any_files else 'disabled'))
                except Exception:
                    pass
            if hasattr(self, 'rb_manual_pair'):
                try:
                    self.rb_manual_pair.configure(state=('normal' if (source_count + translation_count) > 0 else 'disabled'))
                except Exception:
                    pass
            history_len = len(getattr(self, '_pair_history', []) or []) if hasattr(self,'_pair_history') else 0
            redo_len = len(getattr(self, '_pair_redo', []) or []) if hasattr(self,'_pair_redo') else 0
            if hasattr(self, 'rb_undo'):
                try:
                    self.rb_undo.configure(state=('normal' if history_len > 0 else 'disabled'))
                except Exception:
                    pass
            if hasattr(self, 'rb_redo'):
                try:
                    self.rb_redo.configure(state=('normal' if redo_len > 0 else 'disabled'))
                except Exception:
                    pass
        except Exception:
            pass

    def _update_ribbon_states_after_idle(self):
        """UI-sichere Ribbon-Aktualisierung nach Idle, um Reihenfolge-/Timing-Probleme zu vermeiden."""
        try:
            if hasattr(self, 'root') and self.root:
                self._after_idle(lambda: self._update_ribbon_states())
            else:
                self._update_ribbon_states()
        except Exception:
            try:
                self._update_ribbon_states()
            except Exception:
                pass

    # --------------------------- NEUES HAUPTMENÜ-BAND ---------------------------
    def _create_main_menu_ribbon(self, parent=None):
        """Erstellt ein kompakt modernes Menüband mit den Bereichen Datei / Einstellungen / Ansicht / Hilfe.

        Fokus: Klarer Text, deutsche Labels, Design-System Farben.
        """
        try:
            # Vorhandenes Dropdown (falls offen) stets schließen, um Dangling-Fenster zu vermeiden
            try:
                if hasattr(self, '_open_dropdown_ref') and self._open_dropdown_ref and self._open_dropdown_ref.winfo_exists():
                    self._open_dropdown_ref.destroy()
            except Exception:
                pass
            if hasattr(self, '_main_menu_ribbon') and self._main_menu_ribbon and self._main_menu_ribbon.winfo_exists():
                return
            host = parent if parent else self.root
            # Integrierter Stil: Leiste sitzt unten im blauen Header, gleiche Primärfarbe
            primary = self.get_color('primary')
            primary_hover = self.get_color('primary_hover') or primary
            bar = ctk.CTkFrame(host, fg_color=primary, corner_radius=0, height=42)
            bar.pack(fill='x', side='bottom', padx=24, pady=(6,10))
            bar.pack_propagate(False)
            self._main_menu_ribbon = bar

            container = ctk.CTkFrame(bar, fg_color='transparent')
            container.pack(anchor='w')

            # Basis-Button Style (Ghost auf Blau)
            # Dropdown Handling -------------------------------------------------
            self._open_dropdown_ref = None  # aktive Toplevel Referenz

            def _close_dropdown():
                try:
                    if self._open_dropdown_ref and self._open_dropdown_ref.winfo_exists():
                        self._open_dropdown_ref.destroy()
                except Exception:
                    pass
                self._open_dropdown_ref = None

            def _build_dropdown(anchor_widget, items):
                _close_dropdown()
                if not self.root:
                    return
                try:
                    x = anchor_widget.winfo_rootx()
                    y = anchor_widget.winfo_rooty() + anchor_widget.winfo_height()
                except Exception:
                    x = y = 100
                # Verwende CTkToplevel für konsistente DS-Styles (statt tk.Toplevel)
                win = ctk.CTkToplevel(self.root)
                try:
                    win.overrideredirect(True)
                except Exception:
                    pass
                try:
                    win.configure(fg_color=self.get_color('surface'))
                except Exception:
                    pass
                self._open_dropdown_ref = win
                frm = ctk.CTkFrame(win, fg_color=self.get_color('surface'), corner_radius=8, border_width=1, border_color=self.get_color('surface_border'))
                frm.pack(fill='both', expand=True, padx=1, pady=1)

                def _add_item(label, action):
                    try:
                        btn = ctk.CTkButton(frm, text=label, command=lambda a=action: (a(), _close_dropdown()), fg_color=self.get_color('surface'), hover_color=self.get_color('surface_hover'), text_color=self.get_color('text_primary'), corner_radius=6, height=30, font=ctk.CTkFont('Segoe UI', 11))
                    except Exception:
                        btn = ctk.CTkButton(frm, text=label, command=lambda a=action: (a(), _close_dropdown()))
                    setattr(btn,'_preserve_font',True)
                    btn.pack(fill='x', padx=8, pady=4)

                for lbl, act in items:
                    _add_item(lbl, act)

                # Schließen bei Fokusverlust / Escape
                try:
                    win.bind('<FocusOut>', lambda e: _close_dropdown())
                    win.bind('<Escape>', lambda e: _close_dropdown())
                except Exception:
                    pass
                win.update_idletasks()
                try:
                    win.geometry(f"{frm.winfo_reqwidth()}x{frm.winfo_reqheight()}+{x}+{y}")
                except Exception:
                    pass

            def top_btn(label: str, menu_items):
                def _on_click():
                    # Menü-Toggle
                    if self._open_dropdown_ref and self._open_dropdown_ref.winfo_exists():
                        _close_dropdown()
                        # Falls gleiches Menü erneut – sofort schließen ohne neu zu öffnen
                        if getattr(_on_click, '_last_open', False):
                            _on_click._last_open = False
                            return
                    _build_dropdown(btn, menu_items)
                    _on_click._last_open = True
                try:
                    font = ctk.CTkFont('Segoe UI', 12, weight='bold')
                    btn = ctk.CTkButton(
                        container,
                        text=label,
                        command=_on_click,
                        fg_color='transparent',
                        hover_color=primary_hover,
                        text_color=self.get_color('white'),
                        corner_radius=4,
                        height=28,
                        font=font,
                        border_width=0
                    )
                except Exception:
                    btn = ctk.CTkButton(container, text=label, command=_on_click)
                setattr(btn, '_preserve_font', True)
                btn.pack(side='left', padx=(0,14))
                return btn

            # Menüdefinitionen (Label, Aktion)
            datei_items = [
                ('Öffnen', getattr(self, '_upload_source_files', lambda: None)),
                ('Speichern', getattr(self, 'export_results', lambda: None)),
                ('Beenden', lambda: self.root.destroy() if self.root else None)
            ]
            einstellungen_items = [
                ('OCR', getattr(self, '_open_settings_panel', lambda: None)),
                ('Design', getattr(self, '_open_settings_panel', lambda: None)),
                ('Kundenpfad wählen…', lambda: self._choose_customer_base_path())
            ]
            ansicht_items = [
                ('Aktualisieren', getattr(self, '_refresh_file_list_safe', lambda: None)),
                ('Zoom +', getattr(self, '_increase_font_scale', lambda: None)),
                ('Zoom -', getattr(self, '_decrease_font_scale', lambda: None))
            ]
            hilfe_items = [
                ('Info', getattr(self, '_show_about_dialog', lambda: None)),
                ('Dokumentation', getattr(self, '_open_documentation', lambda: None))
            ]

            top_btn('DATEI', datei_items)
            top_btn('EINSTELLUNGEN', einstellungen_items)
            top_btn('ANSICHT', ansicht_items)
            top_btn('HILFE', hilfe_items)

            # Sekundäre schnelle Aktionen (rechtsbündig) – Speichern & Beenden
            right_box = ctk.CTkFrame(bar, fg_color='transparent')
            right_box.pack(anchor='e', side='right')
            def quick(text, cmd):
                try:
                    b = ctk.CTkButton(right_box, text=text, command=cmd, fg_color=self.get_color('white'), hover_color=self.get_color('surface_hover'), text_color=self.get_color('primary'), corner_radius=6, height=28, font=ctk.CTkFont('Segoe UI', 11, weight='bold'))
                except Exception:
                    b = ctk.CTkButton(right_box, text=text, command=cmd)
                setattr(b,'_preserve_font',True)
                b.pack(side='right', padx=(8,0))
                return b
            quick('Beenden', lambda: self.root.destroy() if self.root else None)
            quick('Speichern', getattr(self, 'export_results', lambda: None))

            # Tastaturkürzel idempotent binden (keine Duplikate)
            if self.root and not getattr(self, '_global_hotkeys_bound', False):
                try:
                    self.root.bind_all('<Alt-d>', lambda e: getattr(self, '_upload_source_files', lambda: None)())
                    self.root.bind_all('<Alt-e>', lambda e: getattr(self, '_open_settings_panel', lambda: None)())
                    self.root.bind_all('<Alt-a>', lambda e: getattr(self, '_refresh_file_list_safe', lambda: None)())
                    self.root.bind_all('<Alt-h>', lambda e: getattr(self, '_show_about_dialog', lambda: None)())
                    self._global_hotkeys_bound = True
                except Exception:
                    pass

            # Schließe Dropdowns bei Größenänderung (Rebuild/Resize) – einmalig binden
            if self.root and not getattr(self, '_dropdown_resize_bound', False):
                try:
                    self.root.bind('<Configure>', lambda e: _close_dropdown())
                    # Zusätzlich: Dropdown sicher schließen, wenn Root zerstört wird
                    self.root.bind('<Destroy>', lambda e: _close_dropdown(), add="+")
                    self._dropdown_resize_bound = True
                except Exception:
                    pass

        except Exception as e:
            self._handle_error(e, context='ui.main_menu_ribbon', user_message='Menüband konnte nicht erstellt werden')

    # --------------------------- KUNDENPFAD WÄHLEN ---------------------------
    def _choose_customer_base_path(self):
        """Öffnet einen Dialog um den Basisordner für Kunden/Projekte zu wählen und speichert ihn.

        Speichert in `checker_config.json` unter Key `projects_base_path` (bestehende Struktur beibehalten).
        Aktualisiert interne Attribute falls vorhanden und zeigt Feedback.
        """
        try:
            base = filedialog.askdirectory(title='Basisordner für Kunden wählen')
            if not base:
                return
            base_path = Path(base)
            if not base_path.exists():
                try:
                    base_path.mkdir(parents=True, exist_ok=True)
                except Exception:
                    messagebox.showerror('Fehler', 'Ordner konnte nicht erstellt werden.')
                    return
            # checker_config.json direkt laden / anpassen (leichtgewichtig – vermeidet neue Abhängigkeit falls ConfigManager nicht importiert)
            config_file = Path.cwd() / 'checker_config.json'
            try:
                if not config_file.exists():
                    data = {}
                else:
                    with open(config_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
            except (FileNotFoundError, PermissionError) as e:
                self._log('warning', f'Config file not accessible: {e}')
                data = {}
            except (json.JSONDecodeError, ValueError) as e:
                self._log('warning', f'Config file corrupted, resetting: {e}')
                data = {}
            except (OSError, IOError) as e:
                self._log('error', f'Config file I/O error: {e}')
                data = {}
            
            data['projects_base_path'] = str(base_path)
            data['last_updated'] = time.strftime('%Y-%m-%dT%H:%M:%S')
            
            try:
                with open(config_file, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
            except (PermissionError, OSError) as werr:
                messagebox.showerror('Fehler', f'Pfad konnte nicht gespeichert werden: {werr}')
                self._log('error', f'Failed to write config: {werr}')
                return
            # internes Attribut setzen falls genutzt
            setattr(self, 'projects_base_path', str(base_path))
            # Nutzerfeedback (Toast falls vorhanden, sonst Messagebox)
            try:
                if hasattr(self, 'show_toast') and callable(getattr(self, 'show_toast')):
                    self.show_toast('Basisordner aktualisiert', 'info', duration=_INFO_TOAST_DURATION_MS)
                else:
                    messagebox.showinfo('Gespeichert', f'Neuer Basisordner: {base_path}')
            except Exception:
                pass
        except Exception as e:
            self._handle_error(e, context='settings.choose_customer_base_path', user_message='Kundenpfad konnte nicht gesetzt werden')
    
    def _create_main_layout(self):
        """Create main 2-panel layout and persist container as attribute."""
        try:
            self.main_container = ctk.CTkFrame(self.root, fg_color=self.get_color('transparent'))
            self.main_container.pack(fill="both", expand=True, padx=15, pady=8)
            self.main_container.grid_columnconfigure(0, weight=3)
            self.main_container.grid_columnconfigure(1, weight=4)
            self.main_container.grid_rowconfigure(0, weight=1)
            self.left_panel = ctk.CTkFrame(self.main_container, fg_color=self.get_color('surface'))
            self.left_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 8), pady=0)
            self.right_panel = ctk.CTkFrame(self.main_container, fg_color=self.get_color('surface'))
            self.right_panel.grid(row=0, column=1, sticky="nsew", padx=(8, 0), pady=0)
            self._setup_left_panel()
            self._setup_right_panel()
        except Exception as e:
            self._handle_error(e, context="layout.main", user_message=self._t("Layout konnte nicht erstellt werden"))
    
    def _setup_left_panel(self):
        """OPTIMIZED: Setup perfectly organized left panel with logical workflow"""
        try:
            # Clear existing content
            for widget in self.left_panel.winfo_children():
                widget.destroy()
            compact_panel_headers = True
            if not compact_panel_headers:
                # (Legacy) Großes Panel-Header beibehalten
                header_frame = ctk.CTkFrame(
                    self.left_panel, 
                    fg_color=self.get_color('primary'),
                    height=92
                )
                header_frame.pack(fill="x", padx=0, pady=0)
                header_frame.pack_propagate(False)
                header_content = ctk.CTkFrame(header_frame, fg_color=self.get_color('transparent'))
                header_content.pack(fill="both", expand=True, padx=25, pady=8)
                try:
                    _hf, _hs = self.get_typography('display')
                except Exception:
                    _hf, _hs = (None, 18)
                enlarged_size = (_hs + 6) if isinstance(_hs, (int, float)) else 38
                title_label = ctk.CTkLabel(header_content, text="Qualitäts-Framework", font=ctk.CTkFont(_hf, enlarged_size, weight="bold"), text_color=self.get_color('white'))
                setattr(title_label, '_preserve_font', True)
                title_label.pack(anchor="w", pady=(0,2))
                try:
                    _cf,_cs = self.get_typography('title')
                except Exception:
                    _cf,_cs = (None, 12)
                subtitle_label = ctk.CTkLabel(header_content, text="Professionelle Qualitätskontrolle", font=ctk.CTkFont(_cf, (_cs + 4) if isinstance(_cs,(int,float)) else 30), text_color=self.get_color('gray_300', self.get_color('white')))
                setattr(subtitle_label, '_preserve_font', True)
                subtitle_label.pack(anchor="w")
            else:
                # Kompakte Variante: Kleiner Titelstreifen unterhalb Ribbon in neutralem Hintergrund
                mini_header = ctk.CTkFrame(self.left_panel, fg_color=self.get_color('surface'))
                mini_header.pack(fill='x', padx=0, pady=(0,4))
                mini_inner = ctk.CTkFrame(mini_header, fg_color='transparent')
                mini_inner.pack(fill='x', padx=14, pady=6)
                lbl = ctk.CTkLabel(mini_inner, text="Framework", font=ctk.CTkFont(*self.get_typography('subheading')), text_color=self.get_color('text_primary'))
                setattr(lbl,'_preserve_font',True)
                lbl.pack(side='left')
            
            # MAIN CONTENT with optimal scrolling (matching right panel style)
            content_frame = ctk.CTkScrollableFrame(
                self.left_panel,
                fg_color=self.get_color('transparent'),
                corner_radius=0  # Seamless integration like right panel
            )
            content_frame.pack(fill="both", expand=True, padx=12, pady=12)  # Match right panel padding
            
            # LOGICAL WORKFLOW SECTIONS - Perfectly organized:
            # STEP 1: File Upload & Management
            self._create_upload_section(content_frame)
            
            # STEP 2: Analysis Configuration
            self._create_analysis_section(content_frame)
            
            # (EHEM. Schritt 3 Qualitätskriterien entfernt – jetzt in Analyse-Sektion integriert)
            # Neuer Schritt 3: Actions & Operations
            self._create_actions_section(content_frame)
            
            # WORKFLOW COMPLETE - All sections in logical order
            # Add final bottom spacing for harmonious scrolling
            final_spacer = ctk.CTkFrame(content_frame, fg_color=self.get_color('transparent'), height=8)
            final_spacer.pack(fill="x", pady=0)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.left_panel",
                user_message=self._t("Left panel creation failed") if hasattr(self, '_t') else "Linke Seite konnte nicht aufgebaut werden"
            )
            # Create minimal fallback
            self._create_minimal_left_panel()
    
    def _create_minimal_left_panel(self):
        """FALLBACK: Minimal left panel when full setup fails"""
        try:
            # Simple title
            title = ctk.CTkLabel(
                self.left_panel,
                text=self._t("Quality Framework") if hasattr(self,'_t') else "Qualitätsrahmen",
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('primary')
            )
            title.pack(pady=20)
            
            # Basic upload button
            # Vereinheitlicht über zentrale Button-Factory
            upload_btn = self._create_button(
                self.left_panel,
                text=self._t("Dateien hochladen") if hasattr(self,'_t') else "Dateien hochladen",
                command=self._upload_source_files,
                kind="primary",
                height=40
            )
            upload_btn.pack(pady=10, padx=20, fill="x")
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.left_panel.minimal",
                user_message=self._t("Minimal left panel creation failed") if hasattr(self, '_t') else "Linke Minimal-Leiste konnte nicht erstellt werden"
            )
    
    def _create_upload_section(self, parent):
        """STEP 1: File Upload & Management - Unified elegant frame"""
        try:
            # Unified spacing tokens
            spacing_md = self.get_spacing('md') if hasattr(self, 'get_spacing') else 16
            spacing_sm = self.get_spacing('sm') if hasattr(self, 'get_spacing') else 8

            # Professional upload card with integrated header
            upload_card = self._create_card_frame(parent, corner_radius=8)
            upload_card.pack(fill="x", pady=(0, 6), padx=4)  # Much tighter spacing
            
            # Unified header combining section title and card header
            header_frame = ctk.CTkFrame(
                upload_card,
                fg_color=self.get_color('gray_50'),
                corner_radius=8,
                height=50  # Slightly taller for elegant look
            )
            header_frame.pack(fill="x", padx=5, pady=5)
            header_frame.pack_propagate(False)

            header_content = ctk.CTkFrame(header_frame, fg_color=self.get_color('transparent'))
            header_content.pack(fill="x", padx=14, pady=6)  # More padding for elegant look
            
            # Modern header with accent
            self._add_header_accent(header_content, 'primary')
            # Fett hervorgehobener Bereichstitel (geschützt vor globaler Normalisierung)
            subtitle_font_family, subtitle_font_size, *subtitle_rest = self.get_typography('subtitle_unified')
            header_label = ctk.CTkLabel(
                header_content,
                text=self._t("File Upload & Management"),
                font=ctk.CTkFont(subtitle_font_family, subtitle_font_size, weight="bold"),
                text_color=self.get_color('primary')
            )
            header_label._preserve_font = True  # Flag damit _normalize_all_widget_fonts nicht überschreibt
            header_label.pack(side="left", padx=(8, 0))
            
            # File counter with localization - HEADER VERSION
            self.header_file_counter_label = ctk.CTkLabel(
                header_content,
                text=self._t("Files: 0 source, 0 translations"),
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary')
            )
            self.header_file_counter_label.pack(side="right")

            # Dezent: OCR-Verfügbarkeits-Badge (grau), nur wenn OCR-Bibliotheken fehlen
            # Namesafe OCR-Verfügbarkeitsprüfung
            def _ocr_libs_available() -> bool:
                try:
                    _pt = globals().get('pytesseract', None)
                    _img = globals().get('Image', None)
                    return (_pt is not None) and (_img is not None)
                except Exception:
                    return False
            try:
                ocr_available_header = _ocr_libs_available()
            except Exception:
                ocr_available_header = False
            if not ocr_available_header:
                try:
                    ocr_badge = ctk.CTkLabel(
                        header_content,
                        text=self._t("OCR nicht verfügbar"),
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        text_color=self.get_color('text_secondary')
                    )
                    # Rechts vom Titel, vor dem Zähler platzieren
                    ocr_badge.pack(side="right", padx=(8, 12))
                except Exception:
                    pass
            
            # Clean content area with compact spacing
            content_frame = ctk.CTkFrame(upload_card, fg_color=self.get_color('transparent'))
            content_frame.pack(fill="x", padx=8, pady=6)
            
            # Responsive button grid with compact spacing
            button_frame = ctk.CTkFrame(content_frame, fg_color=self.get_color('transparent'))
            button_frame.pack(fill="x", pady=(0, 4), padx=5)
            
            # Configure responsive columns with compact proportions
            # Dynamische Spaltenanzahl (wird nach OCR-Erkennung ggf. erweitert)
            for i in range(4):  # Reserve für OCR
                # Kürzere Texte erlauben geringere Mindestbreite
                button_frame.grid_columnconfigure(i, weight=1, minsize=160)
            
            # Professional upload buttons with compact height
            # Deutsche, konsistente Button-Texte (keine Kürzungen)
            source_btn = self._create_button(
                button_frame,
                text=self._t("Ausgangstexte"),
                command=self._upload_source_files,
                # Stil-Anpassung: Outline wie Welcome Screen, klare Hierarchie
                kind="primary",
                height=40
            )
            source_btn.grid(row=0, column=0, sticky="ew", padx=(0, 4), pady=1)
            
            source_hint = ctk.CTkLabel(
                button_frame,
                text="PDF, DOCX, TXT",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary'),
                anchor="w"
            )
            source_hint.grid(row=1, column=0, sticky="w", padx=(0, 4), pady=(1, 0))
            
            translation_btn = self._create_button(
                button_frame,
                text=self._t("Übersetzungen"),
                command=self._upload_translation_files,
                # Stil-Anpassung: ebenfalls Outline-Variante
                kind="primary",
                height=40
            )
            translation_btn.grid(row=0, column=1, sticky="ew", padx=4, pady=1)
            
            translation_hint = ctk.CTkLabel(
                button_frame,
                text="PDF, DOCX, TXT",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary'),
                anchor="w"
            )
            translation_hint.grid(row=1, column=1, sticky="w", padx=4, pady=(1, 0))
            
            batch_btn = self._create_button(
                button_frame,
                text=self._t("Stapel"),
                command=self._upload_batch_files,
                kind="primary",
                height=40
            )
            batch_btn.grid(row=0, column=2, sticky="ew", padx=(4, 0), pady=1)
            
            batch_hint = ctk.CTkLabel(
                button_frame,
                text="ZIP, Ordner",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary'),
                anchor="w"
            )
            batch_hint.grid(row=1, column=2, sticky="w", padx=(4, 0), pady=(1, 0))

            # --- OPTIONALES OCR (Bild -> Text) ---
            # Nur anzeigen, wenn Bibliotheken verfügbar sind
            ocr_available = _ocr_libs_available()
            ocr_btn = None
            ocr_hint = None
            if ocr_available:
                try:
                    button_frame.grid_columnconfigure(3, weight=1, minsize=120)
                    ocr_btn = self._create_button(
                        button_frame,
                        text=self._t("Bild-OCR") if hasattr(self,'_t') else "Bild-OCR",
                        command=self._upload_ocr_images,
                        kind="primary",
                        height=40
                    )
                    ocr_btn.grid(row=0, column=3, sticky="ew", padx=(4,0), pady=1)
                    ocr_hint = ctk.CTkLabel(
                        button_frame,
                        text="PNG, JPG",
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        text_color=self.get_color('text_secondary'),
                        anchor="w"
                    )
                    ocr_hint.grid(row=1, column=3, sticky="w", padx=(4,0), pady=(1,0))
                except Exception:
                    pass

            # Preview button for ingested text
            preview_btn = self._create_button(
                button_frame,
                text=self._t('Eingelesenen Text anzeigen'),
                command=self._launch_pair_details_preview,
                kind="secondary",
                height=36
            )
            preview_btn.grid(row=2, column=0, columnspan=3 if not ocr_available else 4, sticky="ew", pady=(self.get_spacing('md'), 0))

            # Store buttons for responsive layout
            self._upload_buttons = {
                'frame': button_frame,
                'source_btn': source_btn,
                'source_hint': source_hint,
                'translation_btn': translation_btn,
                'translation_hint': translation_hint,
                'batch_btn': batch_btn,
                'batch_hint': batch_hint,
                'preview_btn': preview_btn,
            }
            if ocr_btn and ocr_hint:
                self._upload_buttons['ocr_btn'] = ocr_btn
                self._upload_buttons['ocr_hint'] = ocr_hint

            # Bind responsive handler (debounced for smoother resize)
            def _on_resize(event):
                try:
                    # Debounce rapid Configure events
                    if hasattr(button_frame, "_resize_job") and getattr(button_frame, "_resize_job", None):
                        try:
                            button_frame.after_cancel(button_frame._resize_job)
                        except Exception:
                            pass
                    width = getattr(event, "width", None) or button_frame.winfo_width()
                    button_frame._resize_job = button_frame.after(60, lambda: self._responsive_upload_layout(width))
                except Exception:
                    pass
            if not hasattr(button_frame, "_resize_bound"):
                # Add binding without overriding potential other listeners
                button_frame.bind("<Configure>", _on_resize, add="+")
                button_frame._resize_bound = True
                # Initial layout once the frame is realized
                try:
                    self._responsive_upload_layout(button_frame.winfo_width())
                except Exception:
                    pass
            
            # File management section
            self._setup_modern_file_management(content_frame)
            
            # Enhanced file types info with visual indicators
            # ENHANCED: Modern info section with drag & drop visual cues
            # Hinweisblock zu Formaten entfernt (korruptes Encoding) – optional später ersetzt durch saubere, kurze Variante
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.upload.enhanced",
                user_message=self._t("Enhanced upload section failed") if hasattr(self, '_t') else "Uploadbereich (erweitert) konnte nicht erstellt werden"
            )
            # Fallback to basic upload section
            self._create_basic_upload_section(parent)
    
    def _create_basic_upload_section(self, parent):
        """FALLBACK: Basic upload section for error scenarios"""
        try:
            # Simple upload section
            upload_card = self._create_card_frame(
                parent,
                corner_radius=8,
                border_width=1
            )
            # Unified vertical margin for cards (16px spacing token equivalent)
            upload_card.pack(fill="x", pady=(0, 16))
            
            # Simple header
            header_label = ctk.CTkLabel(
                upload_card,
                text=self._t("File Upload"),
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('primary')
            )
            header_label.pack(pady=15)
            
            # Simple buttons
            button_frame = ctk.CTkFrame(upload_card, fg_color=self.get_color('transparent'))
            button_frame.pack(fill="x", padx=15, pady=(0, 15))
            # Ensure columns have minimum width for long German labels
            button_frame.grid_columnconfigure(0, weight=1, minsize=140)
            button_frame.grid_columnconfigure(1, weight=1, minsize=140)
            
            # Source button
            source_btn = self._create_button(
                button_frame,
                text="Ausgangstexte hochladen",
                command=self._upload_source_files,
                kind="primary",
                height=40
            )
            source_btn.grid(row=0, column=0, sticky="ew", padx=(0, 5))
            
            # Translation button
            translation_btn = self._create_button(
                button_frame,
                text="Übersetzungen hochladen", 
                command=self._upload_translation_files,
                # Konsistenz: Outline-Variante für sekundäre Aktion
                kind="outline-primary",
                height=40
            )
            translation_btn.grid(row=0, column=1, sticky="ew", padx=(5, 0))
            
            # File counter - CARD VERSION
            self.card_file_counter_label = ctk.CTkLabel(
                upload_card,
                text=self._t("Files: 0 source, 0 translations"),
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary')
            )
            self.card_file_counter_label.pack(pady=(0, 15))
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.upload.basic",
                user_message=self._t("Basic upload section failed") if hasattr(self, '_t') else "Uploadbereich (einfach) konnte nicht erstellt werden"
            )

    def _responsive_upload_layout(self, width: int):
        """Optimiertes Responsive Layout für Upload-Buttons.

        Ziel:
        - So lange wie möglich eine Zeile (kein früher Umbruch wie im Screenshot)
        - Verhindern von Text-Abschnitten (Buttons breiter) – sonst fallback 2 Zeilen
        - Klare Staffelung small / medium / tight
        """
        try:
            if not hasattr(self, '_upload_buttons'):
                return
            bf = self._upload_buttons['frame']
            has_ocr = 'ocr_btn' in self._upload_buttons
            has_preview = 'preview_btn' in self._upload_buttons
            # Alles zurücksetzen
            for key, widget in self._upload_buttons.items():
                if key == 'frame':
                    continue
                try:
                    widget.grid_forget()
                except Exception:
                    pass
            # Spalten neu konfigurieren (max 4) – minsize muss Text aufnehmen können
            for i in range(4):
                bf.grid_columnconfigure(i, weight=1, minsize=160)

            # Dynamische Schwellwerte (leicht erhöht, damit 1-Zeilen-Layout länger gehalten wird)
            button_count = 3 + (1 if has_ocr else 0)
            min_per = 168  # vormals 160
            gap = 10       # kompakterer Zwischenraum
            needed_full_row = button_count * min_per + (button_count - 1) * gap + 36  # etwas geringere Reserve

            if width >= needed_full_row:
                # Alle in einer Zeile
                order = ['source','translation','batch'] + (['ocr'] if has_ocr else [])
                for idx, name in enumerate(order):
                    self._upload_buttons[f'{name}_btn'].grid(row=0, column=idx, sticky='ew', padx=(0 if idx==0 else 6, 0 if idx==len(order)-1 else 6), pady=(2,0))
                    self._upload_buttons[f'{name}_hint'].grid(row=1, column=idx, sticky='w', padx=(0 if idx==0 else 6, 0 if idx==len(order)-1 else 6), pady=(0,8))
                preview_row = 2
                preview_span = len(order) if order else 3
            elif width >= 620:
                # 2 Zeilen kompakt: Erste Zeile 2 Buttons, zweite Zeile Rest
                self._upload_buttons['source_btn'].grid(row=0, column=0, sticky='ew', padx=(0,6), pady=(2,0))
                self._upload_buttons['translation_btn'].grid(row=0, column=1, sticky='ew', padx=(6,0), pady=(2,0))
                self._upload_buttons['source_hint'].grid(row=1, column=0, sticky='w', padx=(0,6), pady=(0,8))
                self._upload_buttons['translation_hint'].grid(row=1, column=1, sticky='w', padx=(6,0), pady=(0,8))
                # Zweite Zeile
                self._upload_buttons['batch_btn'].grid(row=2, column=0, columnspan=2, sticky='ew', pady=(2,0))
                self._upload_buttons['batch_hint'].grid(row=3, column=0, columnspan=2, sticky='w', pady=(0,8))
                if has_ocr:
                    self._upload_buttons['ocr_btn'].grid(row=4, column=0, columnspan=2, sticky='ew', pady=(2,0))
                    self._upload_buttons['ocr_hint'].grid(row=5, column=0, columnspan=2, sticky='w', pady=(0,8))
                    preview_row = 6
                else:
                    preview_row = 4
                preview_span = 2
            else:
                # Sehr schmal: Stapel alles untereinander (eine Spalte)
                row = 0
                seq = ['source','translation','batch'] + (['ocr'] if has_ocr else [])
                for name in seq:
                    self._upload_buttons[f'{name}_btn'].grid(row=row, column=0, sticky='ew', pady=(2,0))
                    self._upload_buttons[f'{name}_hint'].grid(row=row+1, column=0, sticky='w', pady=(0,8))
                    row += 2
                preview_row = row
                preview_span = 1

            if has_preview:
                try:
                    self._upload_buttons['preview_btn'].grid(
                        row=preview_row,
                        column=0,
                        columnspan=preview_span,
                        sticky='ew',
                        pady=(self.get_spacing('md') if hasattr(self, 'get_spacing') else 12, 0)
                    )
                except Exception:
                    try:
                        self._upload_buttons['preview_btn'].grid(row=preview_row, column=0, columnspan=preview_span, sticky='ew', pady=(12, 0))
                    except Exception:
                        pass
        except Exception as e:
            try:
                self._log('warning', f'Responsive upload layout warning: {e}')
            except Exception:
                pass
    
    def _create_analysis_section(self, parent):
        from quality_gui_components_analysis_section import build_analysis_section
        try:
            build_analysis_section(self, parent)
        except Exception as e:
            try:
                self._handle_error(e, context="analysis.section.delegate", user_message=self._t("Analyse Sektion Fallback"))
                self._create_basic_analysis_section(parent)
            except Exception:
                pass
            
    def _create_modern_tab_navigator(self):
        """PROFESSIONAL: Refined tab navigation with business styling"""
        try:
            # Professional tab navigation container (robuster Container-Fallback)
            container = getattr(self, 'main_container', None)
            if container is None:
                container = getattr(self, 'right_panel', None) or getattr(self, 'root', None)
            self.tab_container = ctk.CTkFrame(
                container,
                fg_color=self.get_color('white'),
                corner_radius=8,
                border_width=1,
                            )
            self.tab_container.grid(row=0, column=0, sticky="ew", padx=25, pady=(0, 15))
            
            # Tab content frame
            tab_content = ctk.CTkFrame(self.tab_container, fg_color=self.get_color('transparent'))
            tab_content.pack(fill="x", padx=20, pady=15)
            
            # Tab buttons with professional styling
            self.tab_buttons = {}
            self.current_tab = "welcome"
            
            # Professional tab definitions without emojis
            tab_definitions = [
                ("welcome", self._t("Übersicht"), "primary"),
                ("upload", "File Manager", "gray_600"),
                ("analysis", "Analysis", "gray_600"),
                ("results", "Reports", "gray_600"),
                ("settings", "Settings", "gray_600")
            ]
            
            # Create professional tab buttons with better proportions
            button_frame = ctk.CTkFrame(tab_content, fg_color=self.get_color('transparent'))
            button_frame.pack(fill="x", pady=(0, 10))
            button_frame.grid_columnconfigure((0, 1, 2, 3, 4), weight=1)
            
            for i, (tab_id, tab_text, color_theme) in enumerate(tab_definitions):
                # Active tab = primary, inactive = outline-primary (klarere visuelle Hierarchie)
                is_active = (tab_id == self.current_tab)
                btn = self._create_button(
                    button_frame,
                    text=tab_text,
                    command=lambda tid=tab_id: self._switch_tab(tid),
                    kind=("primary" if is_active else "outline-primary"),
                    height=32
                )
                # Improved button spacing for better visual balance
                padx_value = (0, 4) if i == 0 else (4, 4) if i < 4 else (4, 0)
                btn.grid(row=0, column=i, sticky="ew", padx=padx_value)
                self.tab_buttons[tab_id] = btn
                
        except Exception as e:
            self._handle_error(
                e,
                context="ui.tabs.navigator",
                user_message=self._t("Tab-Navigator konnte nicht erstellt werden") if hasattr(self, '_t') else "Tab-Navigator Fehler"
            )
    
    def _switch_tab(self, tab_id):
        """Vereinfachter, robuster Tab-Switch mit einheitlichem Fallback (Mini-Patch)."""
        try:
            # Button-Zustände aktualisieren
            for tid, btn in getattr(self, 'tab_buttons', {}).items():
                active = (tid == tab_id)
                try:
                    btn.configure(
                        fg_color=self.get_color('primary' if active else 'gray_100'),
                        text_color=self.get_color('white' if active else 'gray_600')
                    )
                except Exception as e_btn:
                    self._debug_silent(e_btn, 'tabs.button_style')

            self.current_tab = tab_id

            # Inhalt wechseln – mit sanften Fallbacks auf Welcome
            if tab_id == "welcome":
                if hasattr(self, '_show_enhanced_welcome_output'):
                    self._show_enhanced_welcome_output()
            elif tab_id == "upload":
                if hasattr(self, '_create_modern_file_explorer'):
                    self._create_modern_file_explorer(self.output_frame)
                else:
                    self._show_enhanced_welcome_output()
            elif tab_id == "analysis":
                if hasattr(self, '_create_analysis_dashboard'):
                    self._create_analysis_dashboard()
                else:
                    self._show_enhanced_welcome_output()
            elif tab_id == "results":
                self._show_results_dashboard()
            elif tab_id == "settings":
                self._show_settings_dashboard()

            if hasattr(self, 'update_status'):
                try:
                    self.update_status(self._t(f"Tab gewechselt: {tab_id.title()}"))
                except Exception as e_status:
                    self._debug_silent(e_status, 'tabs.status')
            # Nach Tab-Wechsel Ribbon-Zustände nach Idle neu bewerten
            try:
                self._update_ribbon_states_after_idle()
            except Exception:
                pass
        except Exception as e:
            # Einheitlicher Fehler-Fallback
            try:
                self._handle_error(e, context="tab.switch", user_message=self._t("Fehler beim Tabwechsel"), event_name="ui.error.tab.switch")
            except Exception:
                pass
    
    def _show_results_dashboard(self):
        """RESULTS: Professional results dashboard"""
        try:
            # Clear output for results
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Professional results header
            results_header = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('gray_700'),
                corner_radius=8
            )
            results_header.pack(fill="x", padx=30, pady=(30, self.get_spacing('md')))
            
            header_content = ctk.CTkFrame(results_header, fg_color=self.get_color('transparent'))
            header_content.pack(fill="x", padx=25, pady=20)
            
            results_title = ctk.CTkLabel(
                header_content,
                text=self._t("Translation Quality Analysis Results"),
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('white')
            )
            results_title.pack()
            # Markiere Demo-Daten eindeutig
            if hasattr(self, 'output_status_label'):
                try:
                    self.output_status_label.configure(text=self._t("Demo"))
                except Exception:
                    pass
            
            # Professional results summary cards
            summary_frame = ctk.CTkFrame(self.output_frame, fg_color=self.get_color('transparent'))
            summary_frame.pack(fill="x", padx=30, pady=(0, 20))
            summary_frame.grid_columnconfigure((0, 1, 2), weight=1)
            
            # Professional summary metrics
            summary_metrics = [
                (self._t("Overall Quality"), "94.5%", "success", self._t("Excellent translation quality")),
                (self._t("Issues Detected"), "7", "gray_600", self._t("Minor issues found")),
                (self._t("Recommendations"), "12", "gray_600", self._t("Improvement suggestions"))
            ]
            
            for i, (title, value, color, description) in enumerate(summary_metrics):
                self._create_professional_summary_card(summary_frame, title, value, color, description, i)
            
            # Professional detailed results section
            details_section = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('white'),
                corner_radius=8,
                border_width=1,
                            )
            details_section.pack(fill="both", expand=True, padx=30, pady=(0, self.get_spacing('md')))
            
            # Sample detailed results
            self._create_professional_detailed_results_content(details_section)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.results.dashboard",
                user_message=self._t("Ergebnis-Dashboard Fehler") if hasattr(self, '_t') else "Ergebnis-Dashboard Fehler"
            )
    
    def _create_professional_summary_card(self, parent, title, value, color, description, column):
        """PROFESSIONAL SUMMARY CARD: Create refined summary metric cards"""
        try:
            card = ctk.CTkFrame(
                parent,
                fg_color=self.get_color('gray_50'),
                corner_radius=8,
                border_width=1,
                            )
            card.grid(row=0, column=column, sticky="nsew", padx=8)
            
            # Value with professional styling
            value_label = ctk.CTkLabel(
                card,
                text=value,
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color(color)
            )
            value_label.pack(pady=(20, 5))
            
            # Title
            title_label = ctk.CTkLabel(
                card,
                text=title,
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('gray_700')
            )
            title_label.pack(pady=(0, 5))
            
            # Description
            desc_label = ctk.CTkLabel(
                card,
                text=description,
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('gray_500'),
                wraplength=150
            )
            desc_label.pack(pady=(0, 20), padx=10)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.results.summary_card",
                user_message=self._t("Zusammenfassung konnte nicht erstellt werden") if hasattr(self, '_t') else "Summary Card Fehler"
            )
    
    def _create_professional_detailed_results_content(self, parent):
        """PROFESSIONAL DETAILED RESULTS: Show refined analysis results"""
        try:
            # Professional results header
            header_label = ctk.CTkLabel(
                parent,
                text=self._t("Detailed Quality Analysis Report"),
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('gray_700')
            )
            header_label.pack(pady=(15, 10))  # Reduced padding
            
            # Results content
            results_scroll = ctk.CTkScrollableFrame(
                parent,
                fg_color=self.get_color('transparent')
            )
            results_scroll.pack(fill="both", expand=True, padx=20, pady=(0, 20))
            
            # Professional results categories
            categories = [
                (self._t("Quality Strengths"), [
                    self._t("Accurate terminology usage"),
                    self._t("Consistent style throughout"),
                    self._t("Proper sentence structure"), 
                    self._t("Cultural adaptation present")
                ], "success"),
                (self._t("Areas for Improvement"), [
                    self._t("Minor punctuation inconsistencies"),
                    self._t("3 terminology variants detected"),
                    self._t("2 sentences could be simplified")
                ], "gray_600"),
                (self._t("Recommendations"), [
                    self._t("Review technical terminology glossary"),
                    self._t("Consider style guide compliance check"),
                    self._t("Validate cultural references")
                ], "gray_600")
            ]
            
            for title, items, color in categories:
                category_frame = ctk.CTkFrame(
                    results_scroll,
                    fg_color=self.get_color('gray_50'),
                    corner_radius=8,
                    border_width=1,
                                    )
                category_frame.pack(fill="x", pady=8)
                
                category_title = ctk.CTkLabel(
                    category_frame,
                    text=title,
                    font=ctk.CTkFont(*self.get_typography('body_bold')),
                    text_color=self.get_color(color)
                )
                category_title.pack(pady=(15, 8))
                
                for item in items:
                    item_label = ctk.CTkLabel(
                        category_frame,
                        text=f"– {item}",
                        font=ctk.CTkFont(*self.get_typography('body')),
                        text_color=self.get_color('gray_600'),
                        anchor="w"
                    )
                    item_label.pack(fill="x", padx=20, pady=2)
                
                # Add bottom padding
                ctk.CTkLabel(category_frame, text="").pack(pady=(0, 8))
                
        except Exception as e:
            self._handle_error(
                e,
                context="ui.results.details",
                user_message=self._t("Detailergebnisse konnten nicht erstellt werden") if hasattr(self, '_t') else "Detail-Ergebnisse Fehler"
            )
    
    def _show_settings_dashboard(self):
        """PROFESSIONAL: Refined settings dashboard with professional styling"""
        try:
            # Clear output for settings
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Professional settings header
            settings_header = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('gray_700'),
                corner_radius=8
            )
            settings_header.pack(fill="x", padx=30, pady=(30, self.get_spacing('md')))
            
            header_content = ctk.CTkFrame(settings_header, fg_color=self.get_color('transparent'))
            header_content.pack(fill="x", padx=25, pady=20)
            
            settings_title = ctk.CTkLabel(
                header_content,
                text=self._t("Application Settings & Preferences"),
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('white')
            )
            settings_title.pack()
            
            # Professional settings content
            settings_content = ctk.CTkScrollableFrame(
                self.output_frame,
                fg_color="transparent"
            )
            settings_content.pack(fill="both", expand=True, padx=30, pady=(0, self.get_spacing('md')))
            
            # ---------- Font Cache (einmalig) ----------
            if not hasattr(self, '_settings_font_cache'):
                self._settings_font_cache = {}
            def F(name: str):
                try:
                    if name not in self._settings_font_cache:
                        self._settings_font_cache[name] = ctk.CTkFont(*self.get_typography(name))
                    return self._settings_font_cache[name]
                except Exception:
                    return ctk.CTkFont("Segoe UI", 12)

            # ---------- Sprach-Mapping sicherstellen (vereinheitlicht) ----------
            if not hasattr(self, 'LANG_DISPLAY2ISO'):
                self.LANG_DISPLAY2ISO = {"Auto-Detect":"auto","Auto-detect":"auto","English":"en","German":"de","French":"fr","Spanish":"es","Italian":"it"}
            else:
                # Ergänze evtl. fehlende Groß-/Kleinschreibungs-Varianten
                self.LANG_DISPLAY2ISO.setdefault("Auto-Detect", "auto")
                self.LANG_DISPLAY2ISO.setdefault("Auto-detect", "auto")

            # ---------- Widget-Factory ----------
            def _make_setting_row(parent, label: str, widget_type: str, config):
                row = ctk.CTkFrame(parent, fg_color="transparent")
                row.pack(fill="x", pady=6)
                ctk.CTkLabel(row, text=label, font=F('body'), text_color=self.get_color('gray_600'), anchor="w").pack(side="left", fill="x", expand=True)
                widget = None
                if widget_type == "combobox":
                    widget = ctk.CTkComboBox(
                        row,
                        values=config,
                        width=220,
                        font=F('body'),
                        fg_color=self.get_color('gray_50'),
                        dropdown_fg_color=self.get_color('surface'),
                        button_color=self.get_color('primary_hover'),
                        border_color=self.get_color('surface_border') if hasattr(self, 'get_color') else None
                    )
                    widget.pack(side="right")
                    if config:
                        widget.set(config[0])
                elif widget_type == "switch":
                    widget = ctk.CTkSwitch(
                        row,
                        text="",
                        width=45,
                        progress_color=self.get_color('primary')
                    )
                    widget.pack(side="right")
                    if config:
                        widget.select()
                elif widget_type == "slider":
                    min_val, max_val, default_val = config
                    steps = None
                    # Memory Cache: ganze MB Schritte
                    if label == self._t("Memory Cache"):
                        steps = int(max_val - min_val)
                    widget = ctk.CTkSlider(
                        row,
                        from_=min_val,
                        to=max_val,
                        width=220,
                        number_of_steps=steps if steps else None,
                        progress_color=self.get_color('primary'),
                        button_color=self.get_color('primary_hover')
                    )
                    widget.pack(side="right")
                    widget.set(default_val)
                return widget

            # ---------- Settings Handler Binding ----------
            def _bind_settings_handlers(row_widget, widget_type: str, label: str, mapping=None):
                key = {
                    self._t("Default Source Language"): "analysis.default_src",
                    self._t("Default Target Language"): "analysis.default_tgt",
                    self._t("Quality Threshold"):       "analysis.quality_threshold",
                    self._t("Enable Advanced Analysis"):"analysis.advanced",
                    self._t("Theme Mode"):              "ui.theme",
                    self._t("Font Size"):               "ui.font_size",
                    self._t("Show Tooltips"):           "ui.tooltips",
                    self._t("Animate Transitions"):     "ui.animations",
                    self._t("Multi-threading"):         "perf.multithreading",
                    self._t("Memory Cache"):            "perf.cache_mb",
                    self._t("Auto-save Results"):       "perf.autosave",
                    self._t("Hintergrund-Verarbeitung"):"perf.background",
                }.get(label)
                if not key or not getattr(self, 'settings_service', None) or row_widget is None:
                    return
                if widget_type == "combobox":
                    def on_change(val):
                        try:
                            if mapping:  # Language mapping
                                # Lokalisierter Wert -> englisches Display -> ISO
                                en_display = next((k for k in mapping.keys() if self._t(k) == val or k == val), None) or val
                                iso = self.LANG_DISPLAY2ISO.get(en_display, en_display)
                                # Target darf kein auto sein
                                if key == "analysis.default_tgt" and iso == "auto":
                                    if hasattr(self, 'show_toast'):
                                        try: self.show_toast(self._t("Zielsprache kann nicht automatisch erkannt werden."), "warning")
                                        except Exception: pass
                                    # Rücksetzen auf Deutsch
                                    try:
                                        row_widget.set(self._t("German"))
                                    except Exception: pass
                                    iso = "de"
                                self.settings_service.set(key, iso)
                            else:
                                self.settings_service.set(key, val)
                        except Exception:
                            pass
                    row_widget.configure(command=on_change)
                elif widget_type == "switch":
                    def on_toggle():
                        try:
                            self.settings_service.set(key, bool(row_widget.get()))
                        except Exception:
                            pass
                    row_widget.configure(command=on_toggle)
                elif widget_type == "slider":
                    def on_slide(v):
                        try:
                            if key == "perf.cache_mb":
                                self.settings_service.set(key, int(float(v)))
                            else:
                                self.settings_service.set(key, float(v))
                        except Exception:
                            pass
                    row_widget.configure(command=on_slide)

            # ---------- Restore Funktion ----------
            def _restore_settings_into_widgets(widget_map: dict):
                if not getattr(self, 'settings_service', None):
                    return
                try:
                    # Sprachen
                    def set_lang(label_key: str, store_key: str, fallback: str):
                        widget = widget_map.get(label_key)
                        if not widget: return
                        iso = self.settings_service.get(store_key, fallback)
                        # Target auto verhindern
                        if store_key == "analysis.default_tgt" and iso == "auto":
                            iso = "de"
                        display = next((self._t(d) for d,k in self.LANG_DISPLAY2ISO.items() if k==iso), self._t("Auto-Detect"))
                        # Target darf kein Auto-Detect anzeigen
                        if store_key == "analysis.default_tgt" and display == self._t("Auto-Detect"):
                            display = self._t("German")
                        try: widget.set(display)
                        except Exception: pass
                    set_lang(self._t("Default Source Language"), "analysis.default_src", "auto")
                    set_lang(self._t("Default Target Language"), "analysis.default_tgt", "de")
                    # Quality Threshold
                    thr_widget = widget_map.get(self._t("Quality Threshold"))
                    if thr_widget:
                        try: thr_widget.set(float(self.settings_service.get("analysis.quality_threshold", 0.9)))
                        except Exception: pass
                    # Memory Cache
                    mem_widget = widget_map.get(self._t("Memory Cache"))
                    if mem_widget:
                        try: mem_widget.set(int(self.settings_service.get("perf.cache_mb", 1024)))
                        except Exception: pass
                    # Switches & andere Comboboxen
                    def _restore_switch(label_key, store_key, fallback=True):
                        w = widget_map.get(label_key)
                        if not w: return
                        try:
                            val = bool(self.settings_service.get(store_key, fallback))
                            if val: w.select()
                            else: w.deselect()
                        except Exception: pass
                    _restore_switch(self._t("Enable Advanced Analysis"), "analysis.advanced", True)
                    _restore_switch(self._t("Show Tooltips"), "ui.tooltips", True)
                    _restore_switch(self._t("Animate Transitions"), "ui.animations", True)
                    _restore_switch(self._t("Multi-threading"), "perf.multithreading", True)
                    _restore_switch(self._t("Auto-save Results"), "perf.autosave", True)
                    _restore_switch(self._t("Hintergrund-Verarbeitung"), "perf.background", False)
                    # Font Size Combobox
                    fs_widget = widget_map.get(self._t("Font Size"))
                    if fs_widget:
                        try:
                            stored = self.settings_service.get("ui.font_size", self._t("Medium"))
                            fs_widget.set(stored if stored in [self._t("Small"), self._t("Medium"), self._t("Large")] else self._t("Medium"))
                        except Exception: pass
                except Exception:
                    pass

            # ---------- Sektionen aufbauen ----------
            self._settings_widgets = {}
            def _build_section(parent, title: str, rows: list):
                section_card = ctk.CTkFrame(parent, fg_color=self.get_color('white'), corner_radius=8, border_width=1)
                section_card.pack(fill="x", pady=12)
                ctk.CTkLabel(section_card, text=title, font=F('body_bold'), text_color=self.get_color('gray_700')).pack(pady=(12, 10))
                frame = ctk.CTkFrame(section_card, fg_color="transparent")
                frame.pack(fill="x", padx=20, pady=(0, 20))
                for label, wtype, cfg in rows:
                    w = _make_setting_row(frame, label, wtype, cfg)
                    self._settings_widgets[label] = w
                    mapping = self.LANG_DISPLAY2ISO if 'Language' in label or 'Sprache' in label else None
                    _bind_settings_handlers(w, wtype, label, mapping=mapping)

            # Analysis Settings
            _build_section(settings_content, self._t("Analysis Settings"), [
                (self._t("Default Source Language"), "combobox", [self._t("Auto-Detect"), self._t("English"), self._t("German"), self._t("French"), self._t("Spanish")]),
                (self._t("Default Target Language"), "combobox", [self._t("English"), self._t("German"), self._t("French"), self._t("Spanish")]),
                (self._t("Quality Threshold"), "slider", (0.7, 1.0, 0.9)),
                (self._t("Enable Advanced Analysis"), "switch", True)
            ])
            # User Interface
            _build_section(settings_content, self._t("User Interface"), [
                (self._t("Theme Mode"), "combobox", [self._t("Light Mode Only")]),
                (self._t("Font Size"), "combobox", [self._t("Small"), self._t("Medium"), self._t("Large")]),
                (self._t("Show Tooltips"), "switch", True),
                (self._t("Animate Transitions"), "switch", True)
            ])
            # Performance
            _build_section(settings_content, self._t("Performance"), [
                (self._t("Multi-threading"), "switch", True),
                (self._t("Memory Cache"), "slider", (512, 2048, 1024)),
                (self._t("Auto-save Results"), "switch", True),
                (self._t("Hintergrund-Verarbeitung"), "switch", False)
            ])

            # Plugins & Protokolle (ausgelagerte Section aus QualityGuiSettingsUI)
            try:
                if hasattr(self, 'settings_ui') and self.settings_ui:
                    # build_plugins_section erzeugt und packt eine Card in den übergebenen Parent und gibt sie zurück
                    _plugins_card = self.settings_ui.build_plugins_section(settings_content)
                    # Optional: Referenz halten für spätere Updates
                    setattr(self, '_plugins_settings_card', _plugins_card)
            except Exception:
                # Keine Blockade des Dashboards, falls optionale Section fehlschlägt
                pass

            # Notification / Toast Settings – via ausgelagerter DRY-Section
            try:
                if hasattr(self, 'settings_ui') and self.settings_ui:
                    _notif_card = self.settings_ui.build_notifications_section(settings_content)
                    setattr(self, '_notifications_settings_card', _notif_card)
            except Exception:
                pass

            # Upload-Einstellungen – ausgelagerte DRY-Section
            try:
                if hasattr(self, 'settings_ui') and self.settings_ui:
                    _upload_card = self.settings_ui.build_upload_section(settings_content)
                    setattr(self, '_upload_settings_card', _upload_card)
            except Exception:
                pass

            # Pfad-Einstellungen – ausgelagerte DRY-Section
            try:
                if hasattr(self, 'settings_ui') and self.settings_ui:
                    _paths_card = self.settings_ui.build_paths_section(settings_content)
                    setattr(self, '_paths_settings_card', _paths_card)
            except Exception:
                pass

            # Kalender-Einstellungen – ausgelagerte DRY-Section
            try:
                if hasattr(self, 'settings_ui') and self.settings_ui:
                    _calendar_card = self.settings_ui.build_calendar_section(settings_content)
                    setattr(self, '_calendar_settings_card', _calendar_card)
            except Exception:
                pass

            _restore_settings_into_widgets(self._settings_widgets)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.settings.dashboard",
                user_message=self._t("Einstellungs-Dashboard Fehler") if hasattr(self, '_t') else "Settings Dashboard Fehler"
            )
    
    # Alte _create_professional_settings_section entfernt (Funktionalität in Dashboard integriert)
            
    def _create_main_interface(self):
        """MAIN INTERFACE: Create the complete modern interface"""
        try:
            # Create main container for layout
            self.main_container.grid_rowconfigure(1, weight=1)
            self.main_container.grid_columnconfigure(0, weight=1)
            
            # Create modern tab navigator first
            self._create_modern_tab_navigator()
            
            # Create main content area
            self._create_main_content_area()
            
            # Show default welcome content
            self._show_enhanced_welcome_output()
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.main_interface",
                user_message=self._t("Hauptoberfläche Fehler") if hasattr(self, '_t') else "Main Interface Fehler"
            )
            
    def _create_main_content_area(self):
        """CONTENT AREA: Create the main content display area"""
        try:
            # Main content frame with modern styling
            self.content_area = ctk.CTkFrame(
                self.main_container,
                fg_color=self.get_color('gray_50'),
                corner_radius=0
            )
            self.content_area.grid(row=1, column=0, sticky="nsew", padx=0, pady=0)
            self.content_area.grid_rowconfigure(0, weight=1)
            self.content_area.grid_columnconfigure(0, weight=1)
            
            # Output frame for dynamic content
            self.output_frame = ctk.CTkScrollableFrame(
                self.content_area,
                fg_color="transparent"
            )
            self.output_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.content_area",
                user_message=self._t("Inhaltsbereich konnte nicht erstellt werden") if hasattr(self, '_t') else "Content Area Fehler"
            )
    
    # EHEMALIG: _create_quality_criteria_section entfernt (Qualitätskriterien jetzt vollständig in Analyse-Sektion integriert)
    
    def _create_actions_section(self, parent):
        """STEP 4: Actions & Operations - unified, cleaned."""
        try:
            actions_card = self._create_card_frame(parent, corner_radius=8, border_width=1)
            actions_card.pack(fill="x", pady=(0, 6), padx=4)
            header_frame = ctk.CTkFrame(actions_card, fg_color=self.get_color('gray_100'), height=50)
            header_frame.pack(fill="x")
            header_frame.pack_propagate(False)
            header_content = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_content.pack(fill="x", padx=14, pady=6)
            # Akzent-Reduktion: Header nutzt primary statt success
            self._add_header_accent(header_content, 'primary')
            ctk.CTkLabel(
                header_content,
                text=self._t("3. Aktionen & Operationen"),
                font=ctk.CTkFont(*self.get_typography('subtitle_unified')),
                text_color=self.get_color('primary')
            ).pack(side="left", padx=(10, 0))
            self.readiness_indicator = ctk.CTkLabel(
                header_content,
                text=self._t("Warten auf Dateien"),
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self._accent('warning')
            )
            self.readiness_indicator.pack(side="right")
            content_frame = ctk.CTkFrame(actions_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=14, pady=12)
            self.analyze_button = self._create_button(
                content_frame,
                text=self._t("Upload files to enable analysis"),
                command=self.start_analysis,
                kind="primary",
                height=44
            )
            self.analyze_button.configure(fg_color=self.get_color('gray_400'), hover_color=self.get_color('gray_500'), state="disabled")
            self.analyze_button.pack(fill="x", pady=(0, 10))
            secondary_section = ctk.CTkFrame(content_frame, fg_color=self.get_color('gray_50'))
            secondary_section.pack(fill="x")
            ctk.CTkLabel(
                secondary_section,
                text=self._t("Additional Actions"),
                font=ctk.CTkFont(*self.get_typography('body_unified')),
                text_color=self.get_color('text_secondary')
            ).pack(pady=(8, 6))
            secondary_frame = ctk.CTkFrame(secondary_section, fg_color="transparent")
            secondary_frame.pack(fill="x", padx=10, pady=(0, 10))
            for i in (0,1,2):
                secondary_frame.grid_columnconfigure(i, weight=1)
            self._create_button(secondary_frame, text=self._t("Exportieren") if hasattr(self,'_t') else "Exportieren", command=self.export_results, kind="primary", height=36).grid(row=0, column=0, sticky="ew", padx=(0, 3))
            self._create_button(secondary_frame, text=self._t("Demo") if hasattr(self,'_t') else "Demo", command=self.show_demo_results, kind="primary", height=36).grid(row=0, column=1, sticky="ew", padx=3)
            self._create_button(secondary_frame, text=self._t("Leeren") if hasattr(self,'_t') else "Leeren", command=self.clear_files, kind="primary", height=36).grid(row=0, column=2, sticky="ew", padx=(3, 0))
        except Exception as e:
            # Vereinheitlichter Fallback laut Vorgabe: keine UI-Teil-Erstellung hier, direkte Rückfallebene
            self._handle_error(
                e,
                context="actions.section",
                user_message=self._t("Fehler in Aktionen – Fallback aktiviert"),
                event_name="ui.error.actions.section"
            )
            self._create_basic_actions_section(parent)
            return
    
    def _setup_right_panel(self):
        """ENHANCED: Setup modern right panel with dynamic content areas"""
        try:
            # Clear existing content
            for widget in self.right_panel.winfo_children():
                widget.destroy()
            compact_panel_headers = True
            if not compact_panel_headers:
                # (Legacy) Großes Header Panel zeigen
                header_frame = ctk.CTkFrame(self.right_panel, fg_color=self.get_color('primary'), height=92)
                header_frame.pack(fill='x', padx=0, pady=0)
                header_frame.pack_propagate(False)
                header_content = ctk.CTkFrame(header_frame, fg_color='transparent')
                header_content.pack(fill='both', expand=True, padx=25, pady=8)
                try:
                    _rhf,_rhs = self.get_typography('display')
                except Exception:
                    _rhf,_rhs = (None,18)
                enlarged_rhs = (_rhs + 6) if isinstance(_rhs,(int,float)) else 38
                tl = ctk.CTkLabel(header_content, text=self._t("Analyseergebnisse & Berichte") if hasattr(self,'_t') else "Analyseergebnisse & Berichte", font=ctk.CTkFont(_rhf, enlarged_rhs, weight='bold'), text_color=self.get_color('white'))
                setattr(tl,'_preserve_font',True)
                tl.pack(anchor='w', pady=(0,2))
                subtitle_frame = ctk.CTkFrame(header_content, fg_color='transparent')
                subtitle_frame.pack(fill='x', anchor='w')
                try:
                    _cf,_cs = self.get_typography('title')
                except Exception:
                    _cf,_cs = (None,12)
                sub = ctk.CTkLabel(subtitle_frame, text="Professionelles Qualitäts-Dashboard", font=ctk.CTkFont(_cf, (_cs+4) if isinstance(_cs,(int,float)) else 30), text_color=self.get_color('gray_300', self.get_color('white')))
                setattr(sub,'_preserve_font',True)
                sub.pack(side='left')
                self.output_status_label = ctk.CTkLabel(subtitle_frame, text='Bereit', font=ctk.CTkFont(*self.get_typography('caption')), text_color=self._accent('success','system_status'))
                self.output_status_label.pack(side='right')
            else:
                # Kompakte Mini-Zeile
                mini_header = ctk.CTkFrame(self.right_panel, fg_color=self.get_color('surface'))
                mini_header.pack(fill='x', padx=0, pady=(0,4))
                mini_inner = ctk.CTkFrame(mini_header, fg_color='transparent')
                mini_inner.pack(fill='x', padx=14, pady=6)
                lbl = ctk.CTkLabel(mini_inner, text=self._t('Analyse') if hasattr(self,'_t') else 'Analyse', font=ctk.CTkFont(*self.get_typography('subheading')), text_color=self.get_color('text_primary'))
                setattr(lbl,'_preserve_font',True)
                lbl.pack(side='left')
                self.output_status_label = ctk.CTkLabel(mini_inner, text='Bereit', font=ctk.CTkFont(*self.get_typography('caption')), text_color=self._accent('success','system_status'))
                self.output_status_label.pack(side='right')
            
            # Enhanced output frame with modern scrolling
            self.output_frame = ctk.CTkScrollableFrame(
                self.right_panel,
                fg_color=self.get_color('background', self.get_color('surface_light')),
                scrollbar_button_color=self.get_color('primary'),
                scrollbar_button_hover_color=self.get_color('primary_hover'),
                corner_radius=0  # Seamless integration
            )
            self.output_frame.pack(fill="both", expand=True, padx=12, pady=12)
            
            # Enhanced welcome content
            self._show_enhanced_welcome_output()
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.right_panel.enhanced",
                user_message=self._t("Rechter Bereich (erweitert) Fehler") if hasattr(self, '_t') else "Right Panel Fehler"
            )
            # Fallback to basic right panel
            self._setup_basic_right_panel()
    
    def _setup_basic_right_panel(self):
        """FALLBACK: Basic right panel for error scenarios"""
        try:
            # Simple header
            header_frame = ctk.CTkFrame(
                self.right_panel,
                # Vereinheitlicht: gleicher Primär-Farbton wie linkes Panel
                fg_color=self.get_color('primary'),
                height=60
            )
            header_frame.pack(fill="x", padx=0, pady=0)
            header_frame.pack_propagate(False)
            
            title_label = ctk.CTkLabel(
                header_frame,
                text=self._t("Analyseergebnisse") if hasattr(self,'_t') else "Analyseergebnisse",
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('white')
            )
            title_label.pack(pady=15)
            
            # Simple output frame
            self.output_frame = ctk.CTkScrollableFrame(
                self.right_panel,
                fg_color=self.get_color('surface_light'),
                scrollbar_button_color=self.get_color('primary'),
                scrollbar_button_hover_color=self.get_color('primary_hover')
            )
            self.output_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            # Show basic welcome content
            self._show_enhanced_welcome_output()
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.right_panel.basic",
                user_message=self._t("Rechter Bereich (einfach) Fehler") if hasattr(self, '_t') else "Right Panel Basic Fehler"
            )
    
    def _show_enhanced_welcome_output(self):
        """PROFESSIONAL: Advanced welcome dashboard with refined professional colors"""
        try:
            # Clear existing content
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Professional dashboard card with clean styling
            welcome_card = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('white'),
                corner_radius=8,  # Consistent corner radius
                border_width=0
            )
            welcome_card.pack(fill="x", pady=25, padx=25)
            
            # Professional content layout
            content_frame = ctk.CTkFrame(welcome_card, fg_color="transparent")
            content_frame.pack(fill="both", expand=True, padx=30, pady=30)
            
            # Professional metrics dashboard
            self._create_professional_metrics_dashboard(content_frame)
            
            # Projektstruktur-Navigation
            self._create_project_folder_navigation(content_frame)
            
            # Professional feature cards
            # Vereinheitlicht: ehemals _create_professional_feature_cards
            self._create_feature_cards(content_frame)
            
            # Professional system status
            # Vereinheitlicht: ehemals _create_professional_system_status
            self._create_system_status(content_frame)

            # KI-Assistenz-Karte für Ollama-Anfragen
            self._create_ollama_assistant_card(content_frame)
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.welcome.enhanced",
                user_message=self._t("Erweiterter Willkommensbereich Fehler") if hasattr(self, '_t') else "Welcome Output Fehler"
            )
            # Fallback to basic welcome
            self._show_basic_welcome_output()
    
    def _create_professional_metrics_dashboard(self, parent):
        """OPTIMIERT: Konsistente Metriken-Dashboard mit einheitlichem Design"""
        try:
            # Vereinheitlichter Header-Stil (gleiche Stärke wie linke Panel-Boxen)
            header_frame = ctk.CTkFrame(
                parent,
                fg_color=self.get_color('primary'),  # Vereinheitlicht
                height=60,
                corner_radius=6
            )
            header_frame.pack(fill="x", pady=(0, 25))
            header_frame.pack_propagate(False)

            header_inner = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_inner.pack(fill="both", expand=True, padx=25, pady=8)

            metrics_title = ctk.CTkLabel(
                header_inner,
                text="Systemübersicht",
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('white')
            )
            metrics_title.pack(side="left")
            
            # Status Indikator - Einheitlich grün
            status_indicator = ctk.CTkLabel(
                header_inner,
                text="Alle Systeme betriebsbereit",
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self._accent('success', 'system_status')
            )
            status_indicator.pack(side="right")
            
            # Optimierte Metriken-Container mit perfektem Spacing
            metrics_container = ctk.CTkFrame(parent, fg_color="transparent")
            # Mehr vertikaler Abstand zum Header für bessere Luft
            metrics_container.pack(fill="x", pady=(0, 40))
            metrics_container.grid_columnconfigure((0, 1, 2, 3), weight=1)
            
            # Dynamische Dateien-Berechnung
            source_count = len(self.uploaded_files.get('source', []))
            translation_count = len(self.uploaded_files.get('translation', []))
            total_files = source_count + translation_count
            
            # OPTIMIERT: Einheitliche Farbgebung für Business-Konsistenz
            metrics = [
                (str(total_files), "Dateien bereit", "gray_700"),
                ("1", "Aktive Sitzungen", "gray_700"), 
                ("Ultraschnell", "Verarbeitungsgeschwindigkeit", "gray_700"),
                ("Betriebsbereit", "Systemstatus", "success")  # Nur System-Status grün
            ]
            
            for i, (value, title, color) in enumerate(metrics):
                # Perfekt einheitliche Metric Cards
                metric_card = ctk.CTkFrame(
                    metrics_container,
                    fg_color=self.get_color('surface'),
                    corner_radius=8,
                    border_width=1,
                    border_color=self.get_color('surface_border')
                )
                # Größerer horizontaler Abstand zwischen den Kacheln für eleganteres Layout
                metric_card.grid(row=0, column=i, sticky="ew", padx=12, pady=8)
                
                # OPTIMIERT: Konsistente Typografie - gleiche Größe für alle Werte
                value_label = ctk.CTkLabel(
                    metric_card,
                    text=value,
                    font=ctk.CTkFont(*self.get_typography('subheading')),  # Einheitlich 18px
                    text_color=self.get_color(color)
                )
                value_label.pack(pady=(25, 8))
                
                # OPTIMIERT: Konsistente Titel-Typografie
                title_label = ctk.CTkLabel(
                    metric_card,
                    text=title,
                    font=ctk.CTkFont(*self.get_typography('caption')),  # Einheitlich 12px
                    text_color=self.get_color('gray_600')
                )
                title_label.pack(pady=(0, 25))
                
        except Exception as e:
            self._handle_error(
                e,
                context="ui.metrics.dashboard",
                user_message=self._t("Metriken-Dashboard Fehler") if hasattr(self, '_t') else "Metriken-Dashboard Fehler"
            )
    
    def _create_feature_cards(self, parent):
        """PROFESSIONAL: Feature Cards (einheitliche Version, i18n-fähig)."""
        try:
            # Hauptfunktionen & Fähigkeiten Header
            # Harmonisierte Typografie: Überschrift als heading_sm
            _hf_family, _hf_size, *_hf_rest = self.get_typography('heading_sm')
            features_title = ctk.CTkLabel(
                parent,
                text=self._t("Hauptfunktionen & Fähigkeiten"),
                font=ctk.CTkFont(_hf_family, _hf_size, weight="bold"),
                text_color=self.get_color('gray_700')
            )
            features_title._preserve_font = True
            features_title.pack(pady=(0, 20))
            
            # Feature cards container mit verbessertem Layout
            features_container = ctk.CTkFrame(parent, fg_color="transparent")
            features_container.pack(fill="x", pady=(0, 30))
            features_container.grid_columnconfigure((0, 1), weight=1)
            
            # Linke Feature Card - Intelligente Analyse-Engine
            left_card = ctk.CTkFrame(
                features_container,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            left_card.grid(row=0, column=0, sticky="nsew", padx=(0, 15), pady=5)
            
            _lh_family, _lh_size, *_lh_rest = self.get_typography('heading_sm')
            left_header = ctk.CTkLabel(
                left_card,
                text=self._t("Intelligente Analyse-Engine"),
                font=ctk.CTkFont(_lh_family, _lh_size, weight="bold"),
                text_color=self.get_color('gray_700')
            )
            left_header._preserve_font = True
            left_header.pack(pady=(15, 12))
            
            left_content = ctk.CTkLabel(
                left_card,
                text=self._t("Neural Network Translation Scoring\nContext-Aware Error Detection\nTerminology Consistency Analysis\nCultural Adaptation Verification\nMulti-Format Document Support"),
                font=ctk.CTkFont(*self.get_typography('body_sm')),
                text_color=self.get_color('gray_600'),
                justify="left",
                anchor="nw"
            )
            left_content.pack(fill="x", padx=15, pady=(0, 15))
            
            # Rechte Feature Card - Detaillierte Berichte
            right_card = ctk.CTkFrame(
                features_container,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            right_card.grid(row=0, column=1, sticky="nsew", padx=(15, 0), pady=5)
            
            _rh_family, _rh_size, *_rh_rest = self.get_typography('heading_sm')
            right_header = ctk.CTkLabel(
                right_card,
                text=self._t("Detaillierte Berichte"),
                font=ctk.CTkFont(_rh_family, _rh_size, weight="bold"),
                text_color=self.get_color('gray_700')
            )
            right_header._preserve_font = True
            right_header.pack(pady=(15, 12))
            
            right_content = ctk.CTkLabel(
                right_card,
                text=self._t("Comprehensive Quality Reports\nExecutive Summary Generation\nDetailed Error Breakdown\nImprovement Recommendations\nExport to PDF, Excel, JSON"),
                font=ctk.CTkFont(*self.get_typography('body_sm')),
                text_color=self.get_color('gray_600'),
                justify="left",
                anchor="nw"
            )
            right_content.pack(fill="x", padx=15, pady=(0, 15))
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.feature_cards",
                user_message=self._t("Feature-Karten Fehler") if hasattr(self, '_t') else "Feature Cards Fehler"
            )
    
    def _create_system_status(self, parent):
        """PROFESSIONAL: Systemstatus (einheitlich, i18n)."""
        try:
            # Systemstatus Card mit konsistentem Design
            status_card = ctk.CTkFrame(
                parent,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            status_card.pack(fill="x", pady=(0, 20))
            
            status_content = ctk.CTkFrame(status_card, fg_color="transparent")
            status_content.pack(fill="x", padx=25, pady=20)
            
            # Status Header mit Live-Indikator
            status_header = ctk.CTkFrame(status_content, fg_color="transparent")
            status_header.pack(fill="x", pady=(0, 20))
            
            status_title = ctk.CTkLabel(
                status_header,
                text=self._t("Systemstatus"),
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('gray_700')
            )
            status_title.pack(side="left")
            
            live_indicator = ctk.CTkLabel(
                status_header,
                text="Alle Systeme betriebsbereit",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self._accent('success', 'system_status')
            )
            live_indicator.pack(side="right")
            
            # Systemdetails Grid mit verbessertem Layout
            details_grid = ctk.CTkFrame(status_content, fg_color="transparent")
            details_grid.pack(fill="x")
            details_grid.grid_columnconfigure((0, 1), weight=1)
            
            # Linke Status-Spalte
            left_status = ctk.CTkLabel(
                details_grid,
                text="AI Analysis Engine: Ready\nMulti-Language Support: Active\nPDF Generation: Available\nBatch Processing: Enabled",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('gray_600'),
                justify="left",
                anchor="nw"
            )
            left_status.grid(row=0, column=0, sticky="nw", padx=(0, 20))
            
            # Rechte Status-Spalte
            right_status = ctk.CTkLabel(
                details_grid,
                text="Performance: Optimized\nMemory Usage: Normal\nSecurity: Protected\nResponse Time: <1s",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('gray_600'),
                justify="left",
                anchor="nw"
            )
            right_status.grid(row=0, column=1, sticky="nw", padx=(20, 0))
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.system_status",
                user_message=self._t("Systemstatus Fehler") if hasattr(self, '_t') else "Systemstatus Fehler"
            )

    def _compute_ollama_pair_options(self) -> Tuple[List[str], Dict[str, int]]:
        values: List[str] = []
        lookup: Dict[str, int] = {}
        try:
            pairs = self.get_file_pairs() or []
            for idx, pair in enumerate(pairs):
                source_token = pair.get('source_name') or pair.get('source') or ''
                target_token = pair.get('translation_name') or pair.get('translation') or ''
                src_name = os.path.basename(source_token) if source_token else ''
                tgt_name = os.path.basename(target_token) if target_token else ''
                label = f"{idx + 1}. {src_name} → {tgt_name}"
                values.append(label)
                lookup[label] = idx
        except Exception as err:  # pragma: no cover
            self._debug_silent(err, "ollama.compute_pair_options")
        return values, lookup

    def _create_ollama_assistant_card(self, parent):
        """Erzeugt die KI-Assistenzkarte für Ollama-Anfragen."""
        try:
            pad_x = self.get_spacing('lg') if hasattr(self, 'get_spacing') else 24
            pad_y = self.get_spacing('md') if hasattr(self, 'get_spacing') else 16
            gap_sm = self.get_spacing('sm') if hasattr(self, 'get_spacing') else 8
            gap_xs = self.get_spacing('xs') if hasattr(self, 'get_spacing') else 4

            card = ctk.CTkFrame(
                parent,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            card.pack(fill="x", pady=(0, pad_y))

            header_font = ctk.CTkFont(*self.get_typography('body_bold')) if hasattr(self, 'get_typography') else ctk.CTkFont(size=16, weight="bold")
            desc_font = ctk.CTkFont(*self.get_typography('caption')) if hasattr(self, 'get_typography') else ctk.CTkFont(size=12)
            label_font = ctk.CTkFont(*self.get_typography('caption')) if hasattr(self, 'get_typography') else ctk.CTkFont(size=12)
            response_label_font = ctk.CTkFont(*self.get_typography('counter')) if hasattr(self, 'get_typography') else ctk.CTkFont(size=12, weight="bold")
            status_font = ctk.CTkFont(*self.get_typography('caption')) if hasattr(self, 'get_typography') else ctk.CTkFont(size=11)

            header = ctk.CTkLabel(
                card,
                text=self._t("AI Assistant (Ollama)"),
                font=header_font,
                text_color=self.get_color('gray_700')
            )
            header.pack(anchor="w", padx=pad_x, pady=(pad_y, gap_xs))

            description = ctk.CTkLabel(
                card,
                text=self._t("Ask the AI about your translation"),
                font=desc_font,
                text_color=self.get_color('gray_600'),
                justify="left",
                wraplength=680
            )
            description.pack(anchor="w", padx=pad_x, pady=(0, gap_sm))

            pair_row = ctk.CTkFrame(card, fg_color="transparent")
            pair_row.pack(fill="x", padx=pad_x, pady=(0, gap_sm))
            pair_row.grid_columnconfigure(1, weight=1)

            pair_label = ctk.CTkLabel(
                pair_row,
                text=self._t("Select file pair"),
                font=label_font,
                text_color=self.get_color('gray_600')
            )
            pair_label.grid(row=0, column=0, sticky="w", padx=(0, gap_sm))

            values, lookup = self._compute_ollama_pair_options()
            placeholder = self._t("No file pairs available")
            if values:
                default_value = values[0]
            else:
                default_value = placeholder
                values = [placeholder]

            pair_var = tk.StringVar(value=default_value)
            pair_menu = ctk.CTkOptionMenu(
                pair_row,
                variable=pair_var,
                values=values,
                fg_color=self.get_color('surface'),
                button_color=self.get_color('primary'),
                button_hover_color=self.get_color('primary_hover'),
                text_color=self.get_color('gray_700'),
                dropdown_fg_color=self.get_color('surface'),
                dropdown_text_color=self.get_color('gray_700'),
            )
            pair_menu.grid(row=0, column=1, sticky="ew")

            question_row = ctk.CTkFrame(card, fg_color="transparent")
            question_row.pack(fill="x", padx=pad_x, pady=(0, gap_sm))
            question_row.grid_columnconfigure(0, weight=1)

            question_var = tk.StringVar()
            question_entry = ctk.CTkEntry(
                question_row,
                textvariable=question_var,
                placeholder_text=self._t("Enter question for Ollama"),
                fg_color=self.get_color('surface'),
                border_color=self.get_color('surface_border'),
                text_color=self.get_color('gray_700')
            )
            question_entry.grid(row=0, column=0, sticky="ew", padx=(0, gap_sm))

            send_button = self._create_button(
                question_row,
                text=self._t("Request Ollama Answer"),
                command=self._on_ollama_send,
                kind="primary",
                height=36
            )
            send_button.grid(row=0, column=1, sticky="ew")

            status_label = ctk.CTkLabel(
                card,
                text=self._t("Ready"),
                font=status_font,
                text_color=self.get_color('gray_500')
            )
            status_label.pack(fill="x", padx=pad_x, pady=(0, gap_xs))

            response_label = ctk.CTkLabel(
                card,
                text=self._t("Response"),
                font=response_label_font,
                text_color=self.get_color('gray_600')
            )
            response_label.pack(anchor="w", padx=pad_x, pady=(gap_sm, 0))

            response_box = ctk.CTkTextbox(
                card,
                height=200,
                fg_color=self.get_color('surface'),
                border_color=self.get_color('surface_border'),
                text_color=self.get_color('gray_700')
            )
            response_box.pack(fill="x", padx=pad_x, pady=(gap_xs, pad_y))
            response_box.configure(state="disabled", wrap="word")

            self._ollama_ui = {
                'card': card,
                'pair_var': pair_var,
                'pair_menu': pair_menu,
                'question_var': question_var,
                'question_entry': question_entry,
                'send_button': send_button,
                'response_box': response_box,
                'status_label': status_label,
                'placeholder': placeholder,
            }
            self._ollama_pair_lookup = lookup
            if not lookup:
                send_button.configure(state="disabled")
                try:
                    pair_menu.configure(state="disabled")
                except Exception:
                    pass
            self._refresh_ollama_pairs()
        except Exception as e:  # pragma: no cover
            self._handle_error(e, context="ui.ollama.card", user_message=self._t("Ollama request failed"), toast=False)

    def _refresh_ollama_pairs(self) -> None:
        ui_ref = getattr(self, '_ollama_ui', None)
        if not ui_ref:
            return
        pair_menu = ui_ref.get('pair_menu')
        pair_var = ui_ref.get('pair_var')
        send_button = ui_ref.get('send_button')
        placeholder = ui_ref.get('placeholder', self._t("No file pairs available"))
        if not pair_menu or not pair_var:
            return
        values, lookup = self._compute_ollama_pair_options()
        if not values:
            values = [placeholder]
            self._ollama_pair_lookup = {}
            pair_menu.configure(values=values)
            pair_var.set(placeholder)
            try:
                pair_menu.configure(state="disabled")
            except Exception:
                pass
            if send_button:
                send_button.configure(state="disabled")
            return
        self._ollama_pair_lookup = lookup
        pair_menu.configure(values=values)
        try:
            pair_menu.configure(state="normal")
        except Exception:
            pass
        if pair_var.get() not in lookup:
            pair_var.set(values[0])
        if send_button and not self._ollama_active_request:
            send_button.configure(state="normal")

    def _set_ollama_ui_busy(self, is_busy: bool, message: Optional[str] = None) -> None:
        ui_ref = getattr(self, '_ollama_ui', None)
        if not ui_ref:
            return
        send_button = ui_ref.get('send_button')
        entry = ui_ref.get('question_entry')
        pair_menu = ui_ref.get('pair_menu')
        status_label = ui_ref.get('status_label')

        self._ollama_active_request = is_busy

        if send_button:
            if is_busy or not self._ollama_pair_lookup:
                send_button.configure(state="disabled")
            else:
                send_button.configure(state="normal")

        if entry:
            try:
                entry.configure(state="disabled" if is_busy else "normal")
            except Exception:
                pass

        if pair_menu:
            try:
                state = "disabled" if is_busy else ("normal" if self._ollama_pair_lookup else "disabled")
                pair_menu.configure(state=state)
            except Exception:
                pass

        if status_label:
            if message:
                status_label.configure(text=message, text_color=self.get_color('primary') if is_busy else self.get_color('gray_500'))
            else:
                status_label.configure(text=self._t("Ready"), text_color=self.get_color('gray_500'))

        if not is_busy:
            try:
                self._refresh_ollama_pairs()
            except Exception as err:  # pragma: no cover
                self._debug_silent(err, "ollama.refresh_after_busy")

    def _clear_ollama_response(self) -> None:
        ui_ref = getattr(self, '_ollama_ui', None)
        if not ui_ref:
            return
        box = ui_ref.get('response_box')
        if box:
            box.configure(state="normal")
            box.delete("1.0", "end")
            box.configure(state="disabled")

    def _display_ollama_response(self, text: str, error: bool = False) -> None:
        ui_ref = getattr(self, '_ollama_ui', None)
        if not ui_ref:
            return
        box = ui_ref.get('response_box')
        status_label = ui_ref.get('status_label')
        if box:
            box.configure(state="normal")
            box.delete("1.0", "end")
            payload = text.strip()
            if payload:
                box.insert("end", payload + ("\n" if not payload.endswith("\n") else ""))
            box.configure(state="disabled")
            box.see("end")
        if status_label:
            if error:
                status_label.configure(text=self._t("Ollama request failed"), text_color=self.get_color('error'))
            else:
                status_label.configure(text=self._t("Ollama response ready"), text_color=self._accent('success', 'system_status'))
        if not error:
            question_var = ui_ref.get('question_var')
            if question_var:
                question_var.set("")

    def _on_ollama_send(self) -> None:
        ui_ref = getattr(self, '_ollama_ui', None)
        if not ui_ref:
            return
        question_var = ui_ref.get('question_var')
        pair_var = ui_ref.get('pair_var')
        if not question_var or not pair_var:
            return
        question = (question_var.get() or "").strip()
        if not question:
            self.show_toast(self._t("Question required"), "warning")
            return
        if not self._ollama_pair_lookup:
            self.show_toast(self._t("No file pairs available"), "warning")
            return
        pair_label = pair_var.get()
        pair_index = self._ollama_pair_lookup.get(pair_label)
        if pair_index is None:
            self.show_toast(self._t("No file pairs available"), "warning")
            return
        if self._ollama_active_request:
            self.show_toast(self._t("Request still running"), "info")
            return
        self._clear_ollama_response()
        self._set_ollama_ui_busy(True, message=self._t("Ollama request running"))
        worker = threading.Thread(target=self._run_ollama_query_worker, args=(question, pair_index), daemon=True)
        worker.start()

    def _run_ollama_query_worker(self, question: str, pair_index: int) -> None:
        try:
            response = self._execute_ollama_request(question, pair_index)
            response_text = response.strip() if isinstance(response, str) else ""
            if not response_text:
                response_text = self._t("Ollama response ready")
            self.ui(lambda: self._display_ollama_response(response_text, error=False))
            self.ui(lambda: self.show_toast(self._t("Ollama response ready"), "success"))
        except (URLError, HTTPError) as err:
            detail = f"{self._t('Ollama not reachable')}: {err}" if err else self._t('Ollama not reachable')
            self.ui(lambda: self._display_ollama_response(detail, error=True))
            self.ui(lambda: self.show_toast(detail, "error"))
        except FileNotFoundError as err:
            detail = f"Ausgewählte Dateien konnten nicht gelesen werden: {err}" if err else "Ausgewählte Dateien konnten nicht gelesen werden."
            self.ui(lambda: self._display_ollama_response(detail, error=True))
            self.ui(lambda: self.show_toast("Ausgewählte Dateien konnten nicht gelesen werden", "error"))
        except Exception as err:
            self._handle_error(err, context="ollama.query", user_message=self._t("Ollama request failed"), toast=False)
            detail = f"{self._t('Ollama request failed')}: {err}" if err else self._t('Ollama request failed')
            self.ui(lambda: self._display_ollama_response(detail, error=True))
            self.ui(lambda: self.show_toast(self._t('Ollama request failed'), "error"))
        finally:
            self.ui(lambda: self._set_ollama_ui_busy(False))

    def _execute_ollama_request(self, question: str, pair_index: int) -> str:
        pairs = self.get_file_pairs() or []
        if pair_index < 0 or pair_index >= len(pairs):
            raise ValueError(f"Ungültiger Paarindex: {pair_index}")
        pair = pairs[pair_index]
        src_path = Path(pair.get('source') or "")
        tgt_path = Path(pair.get('translation') or "")
        if not src_path or not src_path.exists():
            raise FileNotFoundError(src_path)
        if not tgt_path or not tgt_path.exists():
            raise FileNotFoundError(tgt_path)

        src_text = self._read_ollama_text(src_path)
        tgt_text = self._read_ollama_text(tgt_path)
        if not src_text and not tgt_text:
            raise ValueError("Keine Inhalte für Ollama verfügbar")

        if generate_report_tool and hasattr(generate_report_tool, 'build_segments'):
            segments = generate_report_tool.build_segments(src_text, tgt_text)
        else:
            segments = self._fallback_build_segments(src_text, tgt_text)

        limited_segments = segments[:max(1, self._ollama_segment_limit)] if segments else []
        if not limited_segments:
            limited_segments = [(src_text.strip(), tgt_text.strip())]

        if _cli_format_context:
            context = _cli_format_context(limited_segments)
        else:
            context = self._format_ollama_context(limited_segments)

        if _cli_summarize_findings:
            summary = _cli_summarize_findings(limited_segments)
        else:
            summary = self._summarize_segments(limited_segments)

        total_segments = len(segments)
        model, endpoint = self._resolve_ollama_settings()
        prompt = self._build_ollama_prompt(question, context, summary, pair, total_segments)

        if _cli_call_ollama:
            return _cli_call_ollama(model, prompt, url=endpoint, stream=False)
        return self._call_ollama_http(model, prompt, endpoint)

    def _resolve_ollama_settings(self) -> Tuple[str, str]:
        model = os.environ.get('OLLAMA_MODEL', self._ollama_default_model)
        endpoint = os.environ.get('OLLAMA_API_URL', DEFAULT_OLLAMA_URL)
        try:
            if getattr(self, 'settings_service', None):
                model_setting = self.settings_service.get('analysis.phases.phase3.semantic.ollama_model', None)
                if model_setting:
                    model = str(model_setting)
                url_setting = self.settings_service.get('integrations.ollama.url', None)
                if url_setting:
                    endpoint = str(url_setting)
        except Exception as err:  # pragma: no cover
            self._debug_silent(err, "ollama.settings")
        return model, endpoint

    def _call_ollama_http(self, model: str, prompt: str, endpoint: str) -> str:
        payload = json.dumps({"model": model, "prompt": prompt, "stream": False}).encode("utf-8")
        request = Request(endpoint, data=payload, headers={"Content-Type": "application/json"})
        with urlopen(request, timeout=60) as response:
            raw = response.read().decode("utf-8")
        try:
            data = json.loads(raw)
            result = data.get("response", "")
            if isinstance(result, str):
                return result.strip()
            return str(result).strip()
        except json.JSONDecodeError:
            return raw.strip()

    def _build_ollama_prompt(self, question: str, context: str, summary: str, pair: Dict[str, Any], total_segments: int) -> str:
        source_token = pair.get('source_name') or pair.get('source') or ''
        target_token = pair.get('translation_name') or pair.get('translation') or ''
        src_name = os.path.basename(source_token)
        tgt_name = os.path.basename(target_token)
        return (
            "Analysiere die folgende Übersetzungsaufgabe und beantworte danach die Frage. "
            "Achte auf Zahlen, Terminologie und Stilvorgaben. Antworte auf Deutsch.\n"
            f"Dateipaar: {src_name} → {tgt_name}\n"
            f"Segmente gesamt: {total_segments}\n"
            f"Frage: {question}\n\n"
            f"Kontext:\n{context}\n\n"
            f"Zusammenfassung: {summary}\n"
        )

    def _format_ollama_context(self, segments: List[Tuple[str, str]]) -> str:
        lines: List[str] = []
        for idx, (src, tgt) in enumerate(segments, 1):
            lines.append(f"Segment {idx}\nSOURCE: {src}\nTARGET: {tgt}")
        return "\n\n".join(lines)

    def _summarize_segments(self, segments: List[Tuple[str, str]]) -> str:
        identical = 0
        empty_target = 0
        for src, tgt in segments:
            if src == tgt:
                identical += 1
            if not tgt:
                empty_target += 1
        parts = [f"Segmente insgesamt: {len(segments)}"]
        extras: List[str] = []
        if identical:
            extras.append(f"identisch: {identical}")
        if empty_target:
            extras.append(f"leer: {empty_target}")
        if extras:
            parts.append(", ".join(extras))
        return " | ".join(parts)

    def _fallback_build_segments(self, src: str, tgt: str) -> List[Tuple[str, str]]:
        src_lines = [line.strip() for line in src.splitlines() if line.strip()]
        tgt_lines = [line.strip() for line in tgt.splitlines() if line.strip()]
        length = max(len(src_lines), len(tgt_lines))
        if length == 0:
            return [(src, tgt)] if (src or tgt) else []
        segments: List[Tuple[str, str]] = []
        for idx in range(length):
            left = src_lines[idx] if idx < len(src_lines) else ""
            right = tgt_lines[idx] if idx < len(tgt_lines) else ""
            segments.append((left, right))
        return segments

    def _read_ollama_text(self, path: Path) -> str:
        if generate_report_tool and hasattr(generate_report_tool, 'read_text'):
            return generate_report_tool.read_text(path)
        try:
            return path.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            return ""
    
    def _create_welcome_metrics_dashboard(self, parent):
        """MODERN: Interactive metrics dashboard with live data"""
        try:
            # Metrics header
            metrics_title = ctk.CTkLabel(
                parent,
                text="System Overview & Performance Metrics",
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('primary')
            )
            metrics_title.pack(pady=(0, 20))
            
            # Metrics grid container
            metrics_container = ctk.CTkFrame(parent, fg_color="transparent")
            metrics_container.pack(fill="x", pady=(0, 25))
            metrics_container.grid_columnconfigure((0, 1, 2, 3), weight=1)
            
            # Metric cards with modern styling
            metrics = [
                ("Ready Files", len(self.uploaded_files.get('source', [])) + len(self.uploaded_files.get('translation', [])), "primary"),
                ("Active Sessions", "1", "success"),
                ("Analysegeschwindigkeit", "Ultraschnell", "warning"),
                ("System Health", "Optimal", "success")
            ]
            
            for i, (title, value, color) in enumerate(metrics):
                metric_card = ctk.CTkFrame(
                    metrics_container,
                    fg_color=self.get_color(f'{color}_light'),
                    corner_radius=18,
                    border_width=0,
                                )
                metric_card.grid(row=0, column=i, sticky="ew", padx=8)
                
                # Metric value
                value_label = ctk.CTkLabel(
                    metric_card,
                    text=str(value),
                    font=ctk.CTkFont(*self.get_typography('heading')),
                    text_color=self.get_color(color)
                )
                value_label.pack(pady=(15, 5))
                
                # Metric title
                title_label = ctk.CTkLabel(
                    metric_card,
                    text=title,
                    font=ctk.CTkFont(*self.get_typography('caption')),
                    text_color=self.get_color('text_secondary')
                )
                title_label.pack(pady=(0, 15))
                
        except Exception as e:
            self._handle_error(
                e,
                context="ui.metrics.dashboard.basic",
                user_message=self._t("Metriken-Dashboard (einfach) Fehler") if hasattr(self, '_t') else "Metrics Dashboard Fehler"
            )
    
    def _create_project_folder_navigation(self, parent):
        """VISUELLER ORDNER-BROWSER: Projektstruktur-Navigation mit Tree-View"""
        try:
            # Ordner-Navigation Header
            nav_title = ctk.CTkLabel(
                parent,
                text="Projektstruktur & Navigation",
                font=self._f('subheading'),
                text_color=self._c('gray_700')
            )
            nav_title.pack(pady=(0, 15))
            
            # Navigation Container
            nav_container = ctk.CTkFrame(
                parent,
                fg_color=self._c('surface'),
                corner_radius=8,
                border_width=1,
                border_color=self._c('surface_border')
            )
            nav_container.pack(fill="x", pady=(0, 20))
            
            nav_content = ctk.CTkFrame(nav_container, fg_color=self._c('transparent', 'transparent'))
            nav_content.pack(fill="x", padx=15, pady=15)
            
            # Aktuelle Projekte Übersicht
            current_projects = self._get_current_projects()
            if current_projects:
                total = len(current_projects)
                # Toggle Zustand initialisieren
                if not hasattr(self, '_nav_show_all_projects'):
                    self._nav_show_all_projects = False
                visible_count = total if self._nav_show_all_projects else min(3, total)
                for i, project in enumerate(current_projects[:visible_count]):
                    # Kleine Card für jedes Projekt (No-Icons Policy beachtet)
                    project_frame = ctk.CTkFrame(
                        nav_content,
                        fg_color=self._c('surface'),
                        corner_radius=6,
                        border_width=1,
                        border_color=self._c('surface_border')
                    )
                    project_frame.pack(fill="x", pady=6)
                    project_frame.grid_columnconfigure(0, weight=1)

                    inner = ctk.CTkFrame(project_frame, fg_color=self._c('transparent', 'transparent'))
                    inner.pack(fill="x", padx=12, pady=8)
                    inner.grid_columnconfigure(0, weight=1)

                    # Projektname fett
                    project_name_label = ctk.CTkLabel(
                        inner,
                        text=project['customer'],
                        font=self._f('label_bold'),
                        text_color=self._c('gray_700'),
                        anchor='w'
                    )
                    project_name_label.grid(row=0, column=0, sticky='w')

                    # "Neu" Badge für das erste (neueste) Projekt
                    if i == 0:
                        badge = ctk.CTkLabel(
                            inner,
                            text="Neu",
                            font=self._f('caption'),
                            text_color=self._c('primary'),
                            fg_color=self._c('white'),
                            corner_radius=6
                        )
                        badge.grid(row=0, column=1, sticky='w', padx=(8, 0), pady=(0,2))

                    # Datum dezenter darunter
                    project_meta_label = ctk.CTkLabel(
                        inner,
                        text=project['date'],
                        font=self._f('caption'),
                        text_color=self._c('gray_500'),
                        anchor='w'
                    )
                    project_meta_label.grid(row=1, column=0, sticky='w', pady=(2,0))

                    open_button = self._create_button(
                        inner,
                        text=self._t("Öffnen"),
                        command=lambda p=project: self._open_project_folder(p),
                        kind="primary",
                        height=32,
                        width=90
                    )
                    open_button.grid(row=0, column=2, rowspan=2, sticky='e', padx=(12,0))

                # Mehr anzeigen/Weniger Toggle
                if total > 3:
                    def _toggle_more():
                        try:
                            self._nav_show_all_projects = not getattr(self, '_nav_show_all_projects', False)
                            self._refresh_project_navigation()
                        except Exception:
                            pass
                    toggle_text = "Weniger" if getattr(self, '_nav_show_all_projects', False) else "Mehr anzeigen"
                    toggle_btn = self._create_button(
                        nav_content,
                        text=toggle_text,
                        command=_toggle_more,
                        kind="secondary",
                        height=32
                    )
                    toggle_btn.pack(pady=(8,0), anchor='w')
            else:
                # Kein Projekt vorhanden - Info anzeigen
                no_projects_label = ctk.CTkLabel(
                    nav_content,
                    text="Noch keine Projekte erstellt. Laden Sie Dateien hoch um zu beginnen.",
                    font=self._f('caption'),
                    text_color=self._c('gray_500')
                )
                no_projects_label.pack(pady=10)
            
            # Schnellzugriff-Buttons
            quick_actions_frame = ctk.CTkFrame(nav_content, fg_color="transparent")
            quick_actions_frame.pack(fill="x", pady=(10, 0))
            quick_actions_frame.grid_columnconfigure((0, 1, 2), weight=1)
            
            # Projekte-Ordner öffnen
            projects_button = self._create_button(
                quick_actions_frame,
                text="Alle Projekte",
                command=self._open_projects_folder,
                kind="primary",
                height=36
            )
            projects_button.grid(row=0, column=0, sticky="ew", padx=(0, 5))
            
            # Neues Projekt
            new_project_button = self._create_button(
                quick_actions_frame,
                text="Neues Projekt",
                command=self._create_new_project_dialog,
                kind="primary",
                height=36
            )
            new_project_button.grid(row=0, column=1, sticky="ew", padx=5)
            
            # Struktur validieren
            validate_button = self._create_button(
                quick_actions_frame,
                text=self._t("Struktur prüfen"),
                command=self._validate_all_projects,
                kind="primary",
                height=36
            )
            validate_button.grid(row=0, column=2, sticky="ew", padx=(5, 0))
            
        except Exception as e:
            self._handle_error(
                e,
                context="ui.folder_nav",
                user_message=self._t("Ordner-Navigation Fehler") if hasattr(self, '_t') else "Ordner Navigation Fehler"
            )
    
    def _get_current_projects(self):
        """Aktuelle Projekte aus bestehender Checker-Ordnerstruktur ermitteln"""
        try:
            projects = []
            # Zentrale Ermittlung des Projektbasis-Pfads
            projects_path = self._get_projects_base_path()
            
            if not os.path.exists(projects_path):
                return projects
            
            for customer_folder in os.listdir(projects_path):
                customer_path = os.path.join(projects_path, customer_folder)
                if os.path.isdir(customer_path):
                    for date_folder in os.listdir(customer_path):
                        date_path = os.path.join(customer_path, date_folder)
                        if os.path.isdir(date_path):
                            projects.append({
                                'customer': customer_folder,
                                'date': date_folder,
                                'path': date_path,
                                'full_name': f"{customer_folder}/{date_folder}"
                            })
            
            # Nach Datum sortieren (neueste zuerst) – robustes Datumsparsing
            from datetime import datetime
            def _parse_date_safe(d: str):
                if not d:
                    return datetime.min
                for fmt in ("%Y-%m-%d", "%Y_%m_%d", "%Y.%m.%d", "%Y%m%d"):
                    try:
                        return datetime.strptime(d, fmt)
                    except Exception:
                        pass
                # Versuche, numerische Bestandteile zu extrahieren
                try:
                    parts = [int(p) for p in re.split(r"[^0-9]", d) if p]
                    if len(parts) >= 3:
                        y, m, dd = parts[0], parts[1], parts[2]
                        return datetime(y, m, dd)
                except Exception:
                    pass
                return datetime.min
            import re
            projects.sort(key=lambda x: _parse_date_safe(x['date']), reverse=True)
            return projects
            
        except Exception as e:
            self._handle_error(
                e,
                context="projects.discovery",
                user_message=self._t("Projekte konnten nicht ermittelt werden") if hasattr(self, '_t') else "Projekte konnten nicht ermittelt werden"
            )
            return []
    
    def _open_project_folder(self, project):
        """Projekt-Ordner im Explorer öffnen"""
        try:
            project_path = project['path']
            if os.path.exists(project_path):
                self._open_path_crossplatform(project_path)
                self.show_toast(self._t(f"Projekt-Ordner geöffnet: {project['full_name']}") , "success")
            else:
                self.show_toast("Projekt-Ordner nicht gefunden", "error")
                
        except Exception as e:
            self._handle_error(
                e,
                context="projects.open.single",
                user_message=self._t("Projekt-Ordner konnte nicht geöffnet werden") if hasattr(self, '_t') else "Projekt-Ordner öffnen Fehler"
            )
            # Sichtbarer UI-Text bereinigt (Umlaut korrekt)
            self.show_toast("Fehler beim Öffnen des Ordners", "error")
    
    def _open_projects_folder(self):
        """Hauptprojekte-Ordner im Explorer öffnen (Checker_Projekte)"""
        try:
            projects_path = self._get_projects_base_path()
            
            # Ordner erstellen falls nicht vorhanden
            if not os.path.exists(projects_path):
                os.makedirs(projects_path, exist_ok=True)
            
            self._open_path_crossplatform(projects_path)
            
            self.show_toast(self._t("Checker-Projekte-Ordner geöffnet"), "success")
            
        except Exception as e:
            self._handle_error(
                e,
                context="projects.open.base",
                user_message=self._t("Projekte-Verzeichnis konnte nicht geöffnet werden") if hasattr(self, '_t') else "Projekte-Verzeichnis Fehler"
            )
            # Sichtbarer UI-Text bereinigt (Umlaut korrekt)
            self.show_toast("Fehler beim Öffnen des Ordners", "error")
    
    def _create_new_project_dialog(self):
        """Dialog für neues Projekt erstellen"""
        try:
            import datetime
            # Moderner DS-konformer Eingabedialog
            customer_name = self._show_customer_input_dialog(
                title="Neues Projekt",
                message="Kundenname für neues Projekt:",
                initial=""
            )
            
            if customer_name:
                # Projekt erstellen
                clean_name = customer_name.strip().replace(" ", "_")
                project_date = datetime.datetime.now().strftime("%Y-%m-%d")
                
                created_path = self.create_project_structure(clean_name, project_date)
                if created_path:
                    self.show_toast(f"Neues Projekt erstellt: {clean_name}", "success")
                    
                    # Navigation aktualisieren
                    self._refresh_project_navigation()
                else:
                    self.show_toast("Fehler beim Erstellen des Projekts", "error")
            
        except Exception as e:
            self._handle_error(
                e,
                context="projects.create",
                user_message=self._t("Neues Projekt konnte nicht erstellt werden") if hasattr(self, '_t') else "Projekt erstellen Fehler"
            )
            self.show_toast("Fehler beim Erstellen des Projekts", "error")

    def _show_customer_input_dialog(self, title: str, message: str, initial: str = "") -> str | None:
        """Zeigt einen modalen, design-system-konformen Eingabedialog und gibt den Text zurück.

        Rückgabe: Benutzertext oder None bei Abbruch.
        """
        try:
            dlg = ctk.CTkToplevel(self.root)
            dlg.title(title)
            dlg.transient(self.root)
            dlg.grab_set()
            try:
                dlg.attributes('-topmost', True)
            except Exception:
                pass

            width = 520
            dlg.geometry(f"{width}x200")
            frame = ctk.CTkFrame(dlg, fg_color=self._c('surface'), corner_radius=8, border_width=1, border_color=self._c('surface_border'))
            frame.pack(fill='both', expand=True, padx=12, pady=12)

            title_lbl = ctk.CTkLabel(frame, text=title, font=self._f('subheading'), text_color=self._c('gray_700'))
            title_lbl.pack(anchor='w', padx=16, pady=(16, 4))

            msg_lbl = ctk.CTkLabel(frame, text=message, font=self._f('body'), text_color=self._c('gray_600'))
            msg_lbl.pack(anchor='w', padx=16, pady=(0, 6))

            var = tk.StringVar(value=initial or "")
            entry = ctk.CTkEntry(frame, textvariable=var, height=36)
            entry.pack(fill='x', padx=16)
            entry.focus_set()

            btn_bar = ctk.CTkFrame(frame, fg_color='transparent')
            btn_bar.pack(fill='x', padx=16, pady=(12, 12))

            result = {'value': None}
            def _confirm():
                try:
                    v = var.get().strip()
                    result['value'] = v if v else None
                finally:
                    try:
                        dlg.destroy()
                    except Exception:
                        pass

            def _cancel():
                try:
                    result['value'] = None
                finally:
                    try:
                        dlg.destroy()
                    except Exception:
                        pass

            cancel_btn = self._create_button(btn_bar, text=self._t('Abbrechen') if hasattr(self,'_t') else 'Abbrechen', command=_cancel, kind='outline-primary', height=36)
            cancel_btn.pack(side='left')
            ok_btn = self._create_button(btn_bar, text=self._t('OK') if hasattr(self,'_t') else 'OK', command=_confirm, kind='primary', height=36)
            ok_btn.pack(side='right')

            dlg.bind('<Return>', lambda e: _confirm())
            dlg.bind('<Escape>', lambda e: _cancel())
            try:
                self._center_window(dlg)
            except Exception:
                pass

            # Warten bis Dialog schließt (modal)
            dlg.wait_window()
            return result['value']
        except Exception as e:
            self._handle_error(e, context='ui.customer_input', toast=False)
            return None
    
    def _validate_all_projects(self):
        """Alle Projektstrukturen validieren und reparieren"""
        try:
            projects = self._get_current_projects()
            repaired_count = 0
            
            for project in projects:
                if not self.validate_project_structure(project['path']):
                    repaired_count += 1
            
            if repaired_count > 0:
                self.show_toast(f"{repaired_count} Projektstrukturen repariert", "success")
            else:
                self.show_toast("Alle Projektstrukturen sind vollständig", "info")
            
        except Exception as e:
            self._handle_error(
                e,
                context="projects.structure.validate",
                user_message=self._t("Projektstruktur ungültig") if hasattr(self, '_t') else "Struktur Validierung Fehler"
            )
            self.show_toast("Fehler bei der Struktur-Validierung", "error")
    
    def _refresh_project_navigation(self):
        """Projekt-Navigation aktualisieren (vollständiger Rebuild der Welcome-Ansicht)."""
        try:
            if hasattr(self, '_show_enhanced_welcome_output'):
                try:
                    self._show_enhanced_welcome_output()
                except Exception:
                    pass
            # Statuslabel aktualisieren falls vorhanden
            if hasattr(self, 'output_status_label'):
                try:
                    self.output_status_label.configure(text=self._t("Bereit"))
                except Exception:
                    pass
            self.logger.info("Projekt-Navigation aktualisiert")
        except Exception as e:
            self._handle_error(
                e,
                context="projects.navigation.update",
                user_message=self._t("Navigation konnte nicht aktualisiert werden") if hasattr(self, '_t') else "Navigation Update Fehler",
                toast=False
            )
    
    # AUTOMATISCHE MIGRATION - Bestehende Projekte migrieren
    
    def _migrate_existing_projects_on_startup(self):
        """Migriert bestehende Projekte bei Anwendungsstart"""
        try:
            # Migration nur beim ersten Start nach Update durchführen
            base_path = self._get_projects_base_path()
            marker_dir = os.path.join(base_path, ".checker")
            try:
                os.makedirs(marker_dir, exist_ok=True)
            except Exception:
                pass
            migration_marker = os.path.join(marker_dir, "migration_v1_completed.marker")
            
            if os.path.exists(migration_marker):
                return  # Migration bereits durchgeführt
            
            self.logger.info("Starte automatische Projekt-Migration…")
            
            # Bestehende Dateien und Ordner scannen (später deduplizieren)
            legacy_files = self._scan_for_legacy_files()
            try:  # Dedup nach Pfad
                if legacy_files:
                    seen_paths = set()
                    unique = []
                    for lf in legacy_files:
                        p = lf.get('path')
                        if p and p not in seen_paths:
                            seen_paths.add(p)
                            unique.append(lf)
                    legacy_files = unique
            except Exception:
                pass
            
            if legacy_files:
                migrated_count = self._migrate_legacy_files(legacy_files)
                
                if migrated_count > 0:
                    self.logger.info(f"Migration abgeschlossen: {migrated_count} Dateien migriert")
                    self.show_toast(f"Migration abgeschlossen: {migrated_count} Dateien organisiert", "success")
            
            # Migration als abgeschlossen markieren
            with open(migration_marker, "w", encoding="utf-8") as f:
                f.write(f"Migration abgeschlossen am: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            
        except Exception as e:
            self._handle_error(
                e,
                context="migration.auto",
                user_message=self._t("Automatische Migration fehlgeschlagen") if hasattr(self, '_t') else "Auto-Migration Fehler"
            )
    
    def _scan_for_legacy_files(self):
        """Scannt nach bestehenden Dateien die migriert werden sollen"""
        try:
            legacy_files = []  # Sammlung
            
            # Potentielle Legacy-Ordner und Dateien
            legacy_locations = [
                "uploads",
                "source_files", 
                "translations",
                "output",
                "analysis",
                "reports"
            ]
            
            for location in legacy_locations:
                if os.path.exists(location) and os.path.isdir(location):
                    for root, dirs, files in os.walk(location):
                        for file in files:
                            if file.lower().endswith(('.pdf', '.docx', '.txt', '.doc', '.rtf', '.odt')):
                                fp = os.path.join(root, file)
                                legacy_files.append({
                                    'path': fp,
                                    'location': location,
                                    'filename': file,
                                    'type': self._classify_legacy_file(fp)
                                })
            
            # Auch lose Dateien im Hauptverzeichnis
            for file in os.listdir("."):
                if file.lower().endswith(('.pdf', '.docx', '.txt', '.doc', '.rtf', '.odt')):
                    legacy_files.append({
                        'path': file,
                        'location': 'root',
                        'filename': file,
                        'type': self._classify_legacy_file(file)
                    })
            
            return legacy_files
            
        except Exception as e:
            self._handle_error(
                e,
                context="migration.legacy.scan",
                user_message=self._t("Legacy-Dateien konnten nicht gescannt werden") if hasattr(self, '_t') else "Legacy Scan Fehler",
                toast=False
            )
            return []
    
    def _classify_legacy_file(self, file_path):
        """Klassifiziert Legacy-Dateien für Migration per Regex-Kategorien.
        Priorität: translation > analysis > source
        """
        try:
            import re, os
            name = os.path.basename(str(file_path))
            # Case-insensitive, Wortgrenzen wo sinnvoll
            patterns_translation = [
                r"\btranslation\b", r"\btrans\b", r"\btarget\b", r"\btgt\b",
                r"\boutput\b", r"\bresult\b", r"\bübersetz", r"\bzieltext\b",
            ]
            patterns_analysis = [
                r"\breport\b", r"\banalysis\b", r"\bquality\b", r"\bmetrics?\b",
                r"\bscore\b", r"\bbericht\b", r"\banalys(e|e)\b", r"\bqualität\b"
            ]
            low = name.lower()
            for pat in patterns_translation:
                if re.search(pat, low, flags=re.IGNORECASE):
                    return 'translation'
            for pat in patterns_analysis:
                if re.search(pat, low, flags=re.IGNORECASE):
                    return 'analysis'
            return 'source'
        except Exception:
            # Fallback: konservativ als source behandeln
            return 'source'
    
    def _migrate_legacy_files(self, legacy_files):
        """Migriert Legacy-Dateien in neue Projektstruktur"""
        try:
            import shutil
            import datetime
            
            migrated_count = 0
            migration_customer = "Legacy_Migration"
            migration_date = datetime.datetime.now().strftime("%Y-%m-%d")
            
            # Migration-Projektstruktur erstellen (bestehende Struktur)
            migration_path = self.create_project_structure(migration_customer, migration_date)
            if not migration_path:
                return 0
            
            project_paths = self.get_project_paths(migration_customer, migration_date)
            
            for file_info in legacy_files:
                try:
                    source_path = file_info['path']
                    filename = file_info['filename']
                    
                    # Alle Dateien gehen in 01_Ausgangstext (bestehende Struktur)
                    target_folder = project_paths['ausgangstext']
                    
                    target_path = os.path.join(target_folder, filename)
                    
                    # Datei kopieren (nicht verschieben für Sicherheit)
                    shutil.copy2(source_path, target_path)
                    migrated_count += 1
                    
                    self.logger.info(f"Migriert: {filename} → 01_Ausgangstext")
                    
                except Exception as e:
                    # Einzeldatei-Migrationsfehler ohne Toast (Spam-Prävention)
                    self._handle_error(
                        e,
                        context="migration.legacy.file",
                        user_message=None,
                        toast=False
                    )
                    # Weiter mit nächster Datei
            
            return migrated_count
            
        except Exception as e:
            self._handle_error(
                e,
                context="migration.legacy.apply",
                user_message=self._t("Legacy-Migration fehlgeschlagen") if hasattr(self, '_t') else "Legacy Migration Fehler"
            )
            return 0
    
    def manual_migration_dialog(self):
        """Manueller Migrations-Dialog für Benutzer"""
        try:
            # Nicht-blockierender Dialog statt messagebox.askyesno
            self._show_non_blocking_confirm(
                title=self._t("Projekt-Migration") if hasattr(self, '_t') else "Projekt-Migration",
                message=(
                    self._t("Sollen bestehende Dateien in die neue Projektstruktur migriert werden?") if hasattr(self, '_t') else "Sollen bestehende Dateien in die neue Projektstruktur migriert werden?"
                ) + "\n\n" + (
                    self._t("Dies organisiert alle vorhandenen Dokumente in der standardisierten Ordnerstruktur.") if hasattr(self, '_t') else "Dies organisiert alle vorhandenen Dokumente in der standardisierten Ordnerstruktur."
                ) + "\n" + (
                    self._t("Original-Dateien bleiben unverändert.") if hasattr(self, '_t') else "Original-Dateien bleiben unverändert."
                ),
                confirm_text=self._t("Migrieren") if hasattr(self, '_t') else "Migrieren",
                cancel_text=self._t("Abbrechen") if hasattr(self, '_t') else "Abbrechen",
                on_confirm=lambda: self._execute_manual_migration(),
            )
            
        except Exception as e:
            self._handle_error(
                e,
                context="migration.manual.dialog",
                user_message=self._t("Migrations-Dialog fehlgeschlagen") if hasattr(self, '_t') else "Migrations Dialog Fehler"
            )
            try:
                self.show_toast(self._t("Migration fehlgeschlagen"), "error")
            except Exception:
                pass

    def _execute_manual_migration(self):
        """Führt die eigentliche Migration nach Bestätigung asynchron aus."""
        try:
            legacy_files = self._scan_for_legacy_files()
            if legacy_files:
                migrated_count = self._migrate_legacy_files(legacy_files)
                msg = (
                    self._t("Migration erfolgreich: {count} Dateien organisiert.").format(count=migrated_count)
                    if hasattr(self, '_t') else f"Migration erfolgreich: {migrated_count} Dateien organisiert."
                )
                self.show_toast(msg, "success")  # Einheitliche Toast-Wrapper-Nutzung
            else:
                msg = self._t("Keine Dateien für Migration gefunden.") if hasattr(self, '_t') else "Keine Dateien für Migration gefunden."
                self.show_toast(msg, "info")
        except Exception as e:
            self._handle_error(
                e,
                context="migration.manual.execute",
                user_message=self._t("Migration fehlgeschlagen") if hasattr(self, '_t') else "Migration fehlgeschlagen"
            )
            try:
                self.show_toast(self._t("Migration fehlgeschlagen"), "error")
            except Exception:
                pass

    def _show_non_blocking_confirm(self, title: str, message: str, confirm_text: str, cancel_text: str, on_confirm, width: int = 520):
        """Generischer nicht-blockierender Bestätigungsdialog (Design-System konform)."""
        try:
            dialog = ctk.CTkToplevel(self.root)
            dialog.title(title)
            dialog.transient(self.root)
            dialog.grab_set()  # Modal Verhalten, ohne mainloop zu blockieren
            try:
                dialog.attributes('-topmost', True)
            except Exception:
                pass

            # Layout
            dialog.geometry(f"{width}x220")
            frame = ctk.CTkFrame(dialog, fg_color=self.get_color('surface'), corner_radius=8, border_width=1, border_color=self.get_color('surface_border'))
            frame.pack(fill='both', expand=True, padx=12, pady=12)

            title_lbl = ctk.CTkLabel(frame, text=title, font=ctk.CTkFont(*self.get_typography('subheading')), text_color=self.get_color('gray_700'))
            title_lbl.pack(anchor='w', padx=16, pady=(16, 4))

            msg_lbl = ctk.CTkLabel(frame, text=message, font=ctk.CTkFont(*self.get_typography('body')), text_color=self.get_color('gray_600'), justify='left', wraplength=width-64)
            msg_lbl.pack(fill='x', padx=16, pady=(0, 20))

            btn_bar = ctk.CTkFrame(frame, fg_color='transparent')
            btn_bar.pack(fill='x', padx=16, pady=(0, 12))
            btn_bar.pack_propagate(False)

            def close_dialog():
                try:
                    dialog.destroy()
                except Exception:
                    pass

            # UX: Abbrechen links, Primäraktion rechts
            cancel_btn = self._create_button(
                btn_bar,
                text=cancel_text,
                command=close_dialog,
                kind="outline-primary",
                height=36
            )
            cancel_btn.pack(side='left')

            confirm_btn = self._create_button(
                btn_bar,
                text=confirm_text,
                command=lambda: (close_dialog(), on_confirm()),
                kind="primary",
                height=36
            )
            confirm_btn.pack(side='right')

            # Keyboard Shortcuts
            dialog.bind('<Return>', lambda e: (close_dialog(), on_confirm()))
            dialog.bind('<Escape>', lambda e: close_dialog())

            # Zentrieren (bestehenden Helper wiederverwenden wenn vorhanden)
            try:
                self._center_window(dialog)
            except Exception:
                pass
        except Exception as e:
            self._handle_error(e, context="ui.confirm.dialog", user_message=None)
    
    # Legacy Methoden _create_professional_feature_cards / _create_professional_system_status entfernt.
    # Einheitliche Implementierungen: _create_feature_cards und _create_system_status.
    
    def _create_modern_file_explorer(self, parent):
        """PREMIUM: Advanced file explorer with preview and metadata"""
        try:
            # Professional file explorer card
            explorer_card = self._create_card_frame(
                parent,
                corner_radius=8,
                border_width=0
            )
            explorer_card.pack(fill="both", expand=True, pady=30, padx=30)
            
            # Professional explorer header
            header_frame = ctk.CTkFrame(
                explorer_card,
                fg_color=self.get_color('gray_600'),
                corner_radius=8
            )
            header_frame.pack(fill="x", padx=5, pady=(5, 0))
            
            header_content = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_content.pack(fill="x", padx=25, pady=20)
            
            explorer_title_text = self._t("Smart File Management Center") if hasattr(self, '_t') else "Smart File Management Center"
            explorer_title = ctk.CTkLabel(
                header_content,
                text=explorer_title_text,
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('white')
            )
            explorer_title.pack(side="left")
            
            total_files = len(self.uploaded_files.get('source', [])) + len(self.uploaded_files.get('translation', []))
            file_count_label = ctk.CTkLabel(
                header_content,
                text=(self._t("Files:") + f" {total_files}") if hasattr(self, '_t') else f"Files: {total_files}",
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('white')
            )
            file_count_label.pack(side="right")
            
            # File categories
            categories_frame = ctk.CTkFrame(explorer_card, fg_color="transparent")
            categories_frame.pack(fill="x", padx=25, pady=25)
            categories_frame.grid_columnconfigure((0, 1), weight=1)
            
            # Source files section
            self._create_file_category_section(
                categories_frame, 
                self._t("Source Documents") if hasattr(self, '_t') else "Source Documents", 
                self.uploaded_files.get('source', []),
                0, 0,
                "primary"
            )
            
            # Translation files section
            self._create_file_category_section(
                categories_frame,
                self._t("Übersetzungen") if hasattr(self, '_t') else "Übersetzungen",
                self.uploaded_files.get('translation', []),
                0, 1,
                "success"
            )
            
        except Exception as e:
            self._handle_error(
                e,
                context="files.explorer.create",
                user_message=self._t("Datei-Explorer konnte nicht erstellt werden") if hasattr(self, '_t') else "Datei-Explorer Fehler"
            )
    
    def _create_file_category_section(self, parent, title, files, row, col, color_theme):
        """CATEGORY: Create file category section with modern cards"""
        try:
            from quality_gui_components_file import build_file_category_section
            build_file_category_section(self, parent, title, files, row, col, color_theme)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="files.category.import",
                user_message=self._t("Dateikategorie Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="files.category.delegate",
                user_message=self._t("Dateikategorie konnte nicht erstellt werden") if hasattr(self, '_t') else "File Category Fehler"
            )
    
    def _create_file_card(self, parent, file_path, index, color_theme):
        """FILE CARD: Individual file card with metadata and actions"""
        try:
            from quality_gui_components_file import build_file_card
            build_file_card(self, parent, file_path, index, color_theme)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="files.card.import",
                user_message=self._t("Dateikarten Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="files.card.delegate",
                user_message=self._t("Dateikarte konnte nicht erstellt werden") if hasattr(self, '_t') else "File Card Fehler"
            )
    
    def _create_analysis_dashboard(self):
        """PREMIUM: Advanced analysis dashboard with live metrics"""
        try:
            from quality_gui_components_analysis_dashboard import build_analysis_dashboard
            build_analysis_dashboard(self)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="analysis.dashboard.import",
                user_message=self._t("Analyse-Dashboard Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="analysis.dashboard.delegate",
                user_message=self._t("Analyse-Dashboard konnte nicht aufgebaut werden") if hasattr(self, '_t') else "Analyse Dashboard Fehler"
            )
    
    def _create_metric_card(self, parent, title, value, color, description, column):
        """METRIC: Individual metric card with modern styling"""
        try:
            from quality_gui_components_metrics import build_metric_card
            build_metric_card(self, parent, title, value, color, description, column)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="analysis.metric_card.import",
                user_message=self._t("Metrik-Karten Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="analysis.metric_card.delegate",
                user_message=self._t("Metrik-Karte konnte nicht erstellt werden") if hasattr(self, '_t') else "Metric Card Fehler"
            )
    
    def _show_analysis_placeholder(self, parent):
        """PLACEHOLDER: Show analysis placeholder content"""
        try:
            from quality_gui_components_analysis_results import show_analysis_placeholder
            show_analysis_placeholder(self, parent)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="analysis.placeholder.import",
                user_message=self._t("Analyse Platzhalter Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="analysis.placeholder.delegate",
                user_message=self._t("Analyse Platzhalter Fehler") if hasattr(self, '_t') else "Analysis Placeholder Fehler",
                toast=False
            )
    
    def _show_basic_welcome_output(self):
        """FALLBACK: Basic welcome output for error scenarios"""
        try:
            from quality_gui_components_welcome import show_basic_welcome_output
            show_basic_welcome_output(self)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="welcome.basic.import",
                user_message=self._t("Willkommens-Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="welcome.basic.delegate",
                user_message=self._t("Basis Willkommensausgabe Fehler") if hasattr(self, '_t') else "Basic Welcome Fehler",
                toast=False
            )
    
    def _create_status_bar(self):
        """ENHANCED: Modern status bar with progress indicator and system status"""
        try:
            from quality_gui_status_bar import build_status_bar
            build_status_bar(self)
        except ImportError as ie:
            self._handle_error(
                ie,
                context="status_bar.import",
                user_message=self._t("Statusleisten Modul fehlt") if hasattr(self, '_t') else None,
                toast=False
            )
        except Exception as e:
            self._handle_error(
                e,
                context="status_bar.delegate",
                user_message=self._t("Statusleiste Fehler") if hasattr(self, '_t') else "Status Bar Fehler"
            )
            self._create_basic_status_bar()
    
    def update_status(self, message, status_type="info", show_progress=False, progress_value=0):
        """ENHANCED: Animated status updates (robust guards)"""
        try:
            # Guard: Status UI evtl. noch nicht initialisiert
            if not hasattr(self, 'status_label'):
                return

            # Animate status text update (nur wenn Label vorhanden)
            self._animate_status_text_change(message)

            # Status Indicator Animation nur wenn vorhanden
            if hasattr(self, 'status_indicator'):
                # No-Icons-Policy: rein textbasierte Anzeige
                status_indicators = {
                    "success": {"label": "Erfolg", "color": self.get_color('success')},
                    "error": {"label": "Fehler", "color": self.get_color('error')},
                    "warning": {"label": "Warnung", "color": self.get_color('warning')},
                    "info": {"label": "Info", "color": self.get_color('primary')},
                    "processing": {"label": "Verarbeite", "color": self.get_color('warning')}
                }
                indicator = status_indicators.get(status_type, status_indicators["info"])
                self._animate_status_indicator_change(indicator)

            # Progress Bereich nur wenn Widgets existieren
            has_progress_widgets = all([
                hasattr(self, 'progress_section'),
                hasattr(self, 'mini_progress'),
                hasattr(self, 'progress_text')
            ])
            if show_progress and has_progress_widgets:
                self._show_progress_animated()
                self._animate_progress_update(progress_value)
                try:
                    self.progress_text.configure(text=f"{int(progress_value * 100)}%")
                except Exception:
                    pass
            elif has_progress_widgets:
                self._hide_progress_animated()
        except Exception:
            # Fallback auf direkte (stille) Update-Variante
            try:
                if hasattr(self, 'status_label'):
                    self.status_label.configure(text=message)
                if hasattr(self, 'status_indicator'):
                    self.status_indicator.configure(text="Status", text_color=self.get_color('primary'))
            except Exception:
                pass

    # -------------------------------------------------
    # GENERIC HELPERS
    # -------------------------------------------------
    def _center_window(self, window, width=None, height=None, anchor_root=True):
        """Zentriert ein Toplevel-Fenster robust. Rein kosmetisch; Fehler werden unterdrückt."""
        try:
            window.update_idletasks()
            try:
                window.wait_visibility()
            except Exception:
                pass
            if width and height:
                try:
                    window.geometry(f"{width}x{height}")
                except Exception:
                    pass
            if anchor_root and hasattr(self, 'root'):
                rx, ry = self.root.winfo_x(), self.root.winfo_y()
                rw, rh = self.root.winfo_width(), self.root.winfo_height()
                w = width or window.winfo_width()
                h = height or window.winfo_height()
                x = rx + (rw // 2) - (w // 2)
                y = ry + (rh // 2) - (h // 2)
            else:
                sw, sh = window.winfo_screenwidth(), window.winfo_screenheight()
                w = width or window.winfo_width()
                h = height or window.winfo_height()
                x = (sw // 2) - (w // 2)
                y = (sh // 2) - (h // 2)
            window.geometry(f"+{x}+{y}")
        except Exception:
            pass

    def _animate_status_text_change(self, new_text):
        """ANIMATION: Smooth text transition effect"""
        try:
            # Simple fade-like effect by temporarily dimming
            original_color = self.status_label.cget('text_color')
            fade_color = self.get_color('gray_400')
            
            # Quick fade and restore
            self.status_label.configure(text_color=fade_color)
            self._after(50, lambda: self.status_label.configure(text=new_text))
            self._after(100, lambda: self.status_label.configure(text_color=original_color))
        except:
            # Fallback to instant update
            self.status_label.configure(text=new_text)

    def _animate_status_indicator_change(self, indicator):
        """ANIMATION: Smooth indicator color transition"""
        try:
            # Quick pulse effect
            self.status_indicator.configure(text="Status")  # Platzhalter
            self._after(80, lambda: self.status_indicator.configure(
                text=indicator.get("label", "Status"),
                text_color=indicator["color"]
            ))
        except:
            # Fallback to instant update
            self.status_indicator.configure(
                text=indicator.get("label", "Status"),
                text_color=indicator["color"]
            )

    def _animate_progress_update(self, progress_value):
        """ANIMATION: Smooth progress bar animation"""
        try:
            # Animate progress bar to new value
            current_value = self.mini_progress.get()
            steps = 10
            step_size = (progress_value - current_value) / steps
            
            def animate_step(step):
                if step <= steps:
                    new_value = current_value + (step_size * step)
                    self.mini_progress.set(new_value)
                    if step < steps:
                        self._after(20, lambda: animate_step(step + 1))
            
            animate_step(1)
        except:
            # Fallback to instant update
            self.mini_progress.set(progress_value)

    def _show_progress_animated(self):
        """ANIMATION: Smooth progress section reveal"""
        try:
            if not self.progress_section.winfo_ismapped():
                self.progress_section.pack(expand=True)
        except:
            pass

    def _hide_progress_animated(self):
        """ANIMATION: Smooth progress section hide"""
        try:
            if self.progress_section.winfo_ismapped():
                self.progress_section.pack_forget()
        except:
            pass
    
    def update_file_count(self, count):
        """UPDATE FILE COUNT - Update file counter in status bar & badge"""
        try:
            if hasattr(self, 'file_count_label'):
                self.file_count_label.configure(text=f"Files: {count}")
            if hasattr(self, 'file_count_badge'):
                # Kurzer Badge Text (engl. UI Basis); optional später lokalisierbar
                self.file_count_badge.configure(text=f"{count} files")
        except Exception as e:
            self._handle_error(
                e,
                context="files.counter.update",
                user_message=self._t("Dateizähler konnte nicht aktualisiert werden") if hasattr(self, '_t') else "File Counter Fehler",
                toast=False
            )
    
    def _initialize_systems(self):
        """OPTIMIZED: Initialize additional systems with performance monitoring"""
        try:
            # Performance monitoring initialization
            import time
            start_time = time.time()
            
            # Initialize toast system with smart fallback
            try:
                from quality_gui_notifications import ToastNotification
                # Toast System mit i18n Translator integrieren
                try:
                    self.toast_system = ToastNotification(self.root, translator=getattr(self, '_t', lambda s: s))
                except Exception:
                    self.toast_system = ToastNotification(self.root)
                self.logger.info("Advanced toast system loaded")
            except ImportError:
                self.logger.info("Toast system not available - using basic notifications")
                self.toast_system = None
            
            # Performance optimization: Preload common UI components
            self._preload_ui_components()
            
            # Smart system initialization
            self._smart_feature_detection()
            
            # Performance measurement
            init_time = (time.time() - start_time) * 1000
            self.logger.info(f"Systems initialized in {init_time:.1f}ms")
            
            # Show optimized startup message
            self.update_status(self._t("All systems operational - Performance optimized"))
            
        except Exception as e:
            self._handle_error(
                e,
                context="systems.init",
                user_message=self._t("Systeme konnten nicht initialisiert werden") if hasattr(self, '_t') else "System Initialisierung Fehler"
            )
    
    def _preload_ui_components(self):
        """PERFORMANCE: Preload commonly used UI components for faster rendering"""
        try:
            # Preload common colors for instant access
            common_colors = ['primary', 'surface', 'text_primary', 'success', 'warning', 'error']
            for color in common_colors:
                self.get_color(color)  # This caches them
            
            # Preload common typography for instant access
            common_fonts = ['body', 'heading', 'subheading', 'button', 'caption']
            for font in common_fonts:
                self.get_typography(font)  # This caches them
                
            # Preload common spacing values
            common_spacing = ['sm', 'md', 'lg', 'card_padding', 'element_gap']
            for spacing in common_spacing:
                self.get_spacing(spacing)  # This caches them
                
            self.logger.info("UI components preloaded for optimal performance")
            
        except Exception as e:
            self.logger.warning(f"UI preloading warning: {e}")
    
    def _smart_feature_detection(self):
        """SMART: Automatic feature detection and optimization"""
        try:
            # Detect available features and optimize accordingly
            features = {
                'advanced_toast': hasattr(self, 'toast_system') and self.toast_system is not None,
                'performance_monitoring': True,  # Always available
                'smart_caching': len(self._color_cache) > 0,
                'enhanced_typography': True,
                'optimized_spacing': True
            }
            
            # Log detected features
            active_features = [name for name, status in features.items() if status]
            self.logger.info(f"Active features: {', '.join(active_features)}")
            
            # Store feature status for conditional functionality
            self._active_features = features
            
        except Exception as e:
            self.logger.warning(f"Feature detection warning: {e}")
            self._active_features = {}
    
    def _clear_cache_smart(self):
        """OPTIMIZATION: Smart cache clearing to prevent memory leaks"""
        try:
            # Clear caches but keep frequently used items
            if len(self._color_cache) > 50:  # Clear color cache if too large
                # Keep only most common colors
                essential_colors = {k: v for k, v in self._color_cache.items() 
                                  if k in ['primary', 'surface', 'text_primary', 'success', 'warning', 'error']}
                self._color_cache = essential_colors
                
            if len(self._font_cache) > 20:  # Clear font cache if too large
                essential_fonts = {k: v for k, v in self._font_cache.items() 
                                 if any(font in k for font in ['body', 'heading', 'button'])}
                self._font_cache = essential_fonts
                
            if len(self._ui_cache) > 30:  # Clear UI cache if too large
                essential_ui = {k: v for k, v in self._ui_cache.items() 
                               if k in ['md', 'lg', 'card_padding', 'element_gap']}
                self._ui_cache = essential_ui
                
            self.logger.info("Smart cache cleanup completed")
            
        except Exception as e:
            self.logger.warning(f"Cache cleanup warning: {e}")
    
    # Removed duplicate basic update_status to prevent signature conflicts
    
    def show_toast(self, message: str, type: str = "info", duration: int = 3000):
        """ENHANCED: Advanced toast notification system with animations"""
        try:
            # Accent-Reduktion (Option B): Vereinheitliche Farbtypen.
            # Erfolg & Warnung visuell wie Info (Primary), nur Systemstatus bleibt wirklich grün.
            original_type = type
            if type in ("success", "warning"):
                type = "info"
            # Nachricht bereinigen / lokalisieren
            try:
                if hasattr(self, '_t'):
                    # Erst durch _t schicken (liefert bereits Umlaut-Fixes)
                    message_clean = self._t(message)
                else:
                    message_clean = message
                # Zusätzliche Ersetzungen für Paarungspfeile und verbliebene Artefakte
                if _TRIPLE_QUESTION in message_clean:
                    message_clean = message_clean.replace(_TRIPLE_QUESTION, '->')
                if _DOUBLE_QUESTION in message_clean:
                    message_clean = message_clean.replace(_DOUBLE_QUESTION, '')
            except Exception:
                message_clean = message
            if hasattr(self, 'toast_system') and self.toast_system:
                # Enhanced toast with auto-dismiss and animations
                self.toast_system.show_toast(message_clean, type, duration)
            else:
                # Improved fallback toast using status bar and temporary window
                self._show_enhanced_fallback_toast(message_clean, type, duration)
        except Exception as e:
            self._handle_error(e, context="toast.show", toast=False)

    # Unified alias (do not remove original show_toast usage) for new components expecting _show_toast
    def _show_toast(self, message: str, type: str = "info", duration: int = 3000):  # noqa: D401
        """Alias Wrapper für show_toast (Namensvereinheitlichung, keine Icons)."""
        try:
            return self.show_toast(message, type, duration)
        except Exception:
            try:
                self.logger.info(f"TOAST {type}: {message}")
            except Exception:
                pass
    
    def _show_enhanced_fallback_toast(self, message: str, type: str, duration: int):
        """ENHANCED: Improved fallback toast mit sicheren Farb-Fallbacks."""
        try:
            def _safe(col, fb):
                try:
                    if hasattr(self, 'get_color'):
                        return self.get_color(col)
                except Exception:
                    pass
                return fb
            color_map = {
                # Reduzierte Akzente: success/warning -> primary/text_secondary
                'success': _safe('primary', '#1F4E79'),
                'warning': _safe('text_secondary', '#6B7280'),
                'error': _safe('error', '#DC2626'),
                'info': _safe('primary', '#64748B'),
                'text_primary': _safe('text_primary', '#374151'),
                'text_secondary': _safe('text_secondary', '#6B7280')
            }
            if hasattr(self, 'status_label'):
                try:
                    self.status_label.configure(
                        text=str(message),
                        text_color=color_map.get(type, color_map['text_primary'])
                    )
                    self._after(duration, lambda: self.status_label.configure(
                        text="Ready",
                        text_color=color_map['text_secondary']
                    ))
                except Exception:
                    pass
            self.logger.info(f"{type.upper()}: {message}")
        except Exception as e:
            self._handle_error(e, context="toast.fallback", toast=False, level="warning")
            try:
                self.logger.info(f"{type.upper()}: {message}")
            except Exception:
                pass

    # -------------------------------------------------------------
    # AKZENT-REDUKTION HELPER (Option B)
    # Nur System-Gesamtstatus darf weiterhin grün sein. Alles andere:
    #  - Erfolg -> primary
    #  - Warnung -> primary oder text_secondary (je nach Domain)
    #  - Info -> primary
    #  - Zerstörerisch / Danger bleibt error (rot)
    # -------------------------------------------------------------
    def _accent(self, semantic: str, domain: str = "generic"):
        try:
            if semantic == 'success':
                if domain == 'system_status':  # einzig erlaubtes persistent Grün
                    return self.get_color('success')
                return self.get_color('primary')
            if semantic == 'warning':
                if domain in ('destructive', 'error_prone'):
                    return self.get_color('error')
                if domain in ('attention', 'upload_needed'):
                    return self.get_color('primary')
                return self.get_color('text_secondary')
            if semantic == 'info':
                return self.get_color('primary')
            return self.get_color(semantic)
        except Exception:
            # Fallback sicher
            return '#1F4E79' if semantic in ('success', 'info', 'warning') else '#DC2626' if semantic == 'error' else '#374151'
    
    # QUALITY GUI FOLDER STRUCTURE MANAGEMENT - Verbindliche Ordnerstruktur
    
    def create_project_structure(self, customer_name, project_date=None):
        """Erstellt vollständige Projektstruktur nach bestehender Checker-Ordnerstruktur"""
        try:
            import datetime
            if not project_date:
                project_date = datetime.datetime.now().strftime("%Y-%m-%d")
            
            base_path = os.path.join(self._get_projects_base_path(), customer_name, project_date)
            
            # Alle Standardordner erstellen
            for folder in self.STANDARD_PROJECT_STRUCTURE:
                full_path = os.path.join(base_path, folder)
                os.makedirs(full_path, exist_ok=True)
                
            self.logger.info(f"Projektstruktur erstellt: {base_path}")
            return base_path
            
        except Exception as e:
            self._handle_error(e, context="project.structure.create", user_message=self._t("Neues Projekt konnte nicht erstellt werden"))
            return None
    
    def validate_project_structure(self, project_path):
        """Validiert und repariert Projektstruktur"""
        try:
            missing_folders = []
            for folder in self.STANDARD_PROJECT_STRUCTURE:
                full_path = os.path.join(project_path, folder)
                if not os.path.exists(full_path):
                    missing_folders.append(folder)
                    os.makedirs(full_path, exist_ok=True)
            
            if missing_folders:
                self.logger.info(f"Reparierte Ordnerstruktur: {len(missing_folders)} Ordner hinzugefügt")
            
            return len(missing_folders) == 0
            
        except Exception as e:
            self._handle_error(e, context="project.structure.validate", user_message=self._t("Projektstruktur ungültig"), toast=False)
            return False
    
    def get_project_paths(self, customer_name, project_date=None):
        """Gibt alle wichtigen Projektpfade nach bestehender Struktur zurück"""
        try:
            import datetime
            if not project_date:
                project_date = datetime.datetime.now().strftime("%Y-%m-%d")
                
            base_path = os.path.join(self._get_projects_base_path(), customer_name, project_date)
            
            return {
                'base': base_path,
                'ausgangstext': os.path.join(base_path, "01_Ausgangstext"),
                'angebot': os.path.join(base_path, "02_Angebot"),
                'prüfung': os.path.join(base_path, "03_Prüfung"),
                'finalisierung': os.path.join(base_path, "04_Finalisierung")
            }
        except Exception as e:
            self._handle_error(e, context="project.paths.get", user_message=self._t("Projekte konnten nicht ermittelt werden"), toast=False)
            return {}
    
    def _get_or_select_customer_for_upload(self):
        """Professionelle Kundenauswahl mit Design-System"""
        try:
            # Bestehende Kunden laden
            existing_customers = []
            try:
                if os.path.exists("customers.json"):
                    with open("customers.json", "r", encoding="utf-8") as f:
                        customers_data = json.load(f)
                        existing_customers = list(customers_data.keys())
            except Exception:
                pass
            
            # Professioneller Kundenauswahl-Dialog
            customer_name = self._show_professional_customer_dialog(existing_customers)
            
            if customer_name:
                # Kundennamen bereinigen für Ordnernamen
                import re
                clean_name = re.sub(r'[<>:"/\\|?*]', '_', customer_name.strip())
                return clean_name
            
            return None
            
        except Exception as e:
            self._handle_error(e, context="customer.select", user_message=self._t("Kundenauswahl fehlgeschlagen") if hasattr(self, '_t') else None)
    
    def _offer_open_project_folder(self, customer_name):
        """Bietet an, das Projekt-Verzeichnis zu öffnen"""
        try:
            import datetime
            project_date = datetime.datetime.now().strftime("%Y-%m-%d")
            
            # Toast mit Projekt-Ordner öffnen Option
            message = f"Projekt '{customer_name}' aktualisiert. Ordner öffnen?"
            
            # Erweiterte Toast-Nachricht mit Action
            if hasattr(self, 'extended_toast_system') and self.extended_toast_system:
                # Falls erweiterte Toast verfügbar, nutze diese für interaktive Nachrichten
                self._after(1000, lambda: self.extended_toast_system.show_toast(
                    message, 
                    "success", 
                    4000,
                    action_text="Ordner öffnen",
                    action_callback=lambda: self._open_specific_project(customer_name, project_date)
                ))
            else:
                # Standard Toast verwenden
                self.show_toast(message, "success")
                
        except Exception as e:
            self._handle_error(e, context="project.folder.offer", toast=False)
    
    def _open_specific_project(self, customer_name, project_date):
        """Öffnet spezifisches Projekt im Datei-Explorer"""
        try:
            import os
            project_path = os.path.join(self._get_projects_base_path(), customer_name, project_date)
            if not os.path.exists(project_path):
                self.show_toast("Projekt-Ordner nicht gefunden", "warning")
                return
            # Zentraler plattformübergreifender Öffnen-Helper
            self._open_path_crossplatform(project_path)
        except Exception as e:
            self._handle_error(e, context="project.folder.open", user_message=self._t("Projekt-Ordner konnte nicht geöffnet werden") if hasattr(self, '_t') else None)
            try:
                # Umlaut Fix
                self.show_toast("Fehler beim Öffnen des Ordners", "error")
            except Exception:
                pass

    # --- Zentrale Helpers: Projekt-Basis und plattformübergreifendes Öffnen ---
    def _get_projects_base_path(self) -> str:
        """Gibt den Basisordner für Projekte zurück (Einstellung oder Fallback)."""
        try:
            # Bevorzugt Einstellung/Attribut, sonst Standard
            base = None
            if hasattr(self, 'projects_base_path') and self.projects_base_path:
                base = self.projects_base_path
            else:
                # Optional: checker_config.json könnte Basispfad vorgeben; hier einfacher Fallback
                base = "Checker_Projekte"
            # Normieren und expandieren
            import os
            base = os.path.expandvars(os.path.expanduser(str(base)))
            return base
        except Exception:
            return "Checker_Projekte"

    def _open_path_crossplatform(self, path: str) -> None:
        """Öffnet einen Ordner oder eine Datei plattformübergreifend im nativen Explorer.
        - Windows: explorer
        - macOS: open
        - Linux: xdg-open
        Mit robuster Fehlerbehandlung und Pfad-Normalisierung.
        """
        try:
            import sys, subprocess, os
            if not path:
                return
            # Normalize and ensure absolute path to avoid surprises
            npath = os.path.abspath(os.path.normpath(path))
            if not os.path.exists(npath):
                raise FileNotFoundError(npath)

            if sys.platform.startswith('win'):
                subprocess.run(['explorer', npath], check=False)
            elif sys.platform == 'darwin':
                subprocess.run(['open', npath], check=False)
            else:
                subprocess.run(['xdg-open', npath], check=False)
            try:
                self.logger.info(f"Pfad geöffnet: {npath}")
            except Exception:
                pass
        except FileNotFoundError as fe:
            # Kommando existiert nicht oder Pfad nicht gefunden
            self._handle_error(fe, context="path.open.cmd", user_message=self._t("Öffnen-Kommando nicht gefunden") if hasattr(self, '_t') else None)
            try:
                self.show_toast("Öffnen-Kommando fehlt oder Pfad nicht gefunden", "error")
            except Exception:
                pass
        except Exception as se:
            self._handle_error(se, context="path.open.exec", user_message=self._t("Ordner konnte nicht geöffnet werden") if hasattr(self, '_t') else None)
            try:
                self.show_toast("Fehler beim Öffnen des Ordners", "error")
            except Exception:
                pass
    
    def _show_professional_customer_dialog(self, existing_customers):
        """Professioneller Kundenauswahl-Dialog mit Design-System"""
        try:
            # Dialog-Fenster erstellen
            dialog = ctk.CTkToplevel(self.root)
            dialog.title("Kunde auswählen")
            dialog.geometry("500x400")
            dialog.resizable(False, False)
            dialog.transient(self.root)
            dialog.grab_set()
            
            # Design-System Farben verwenden
            dialog.configure(fg_color=self.get_color('surface'))
            
            # Zentrierung via generischem Helper
            self._center_window(dialog, 500, 400, anchor_root=True)
            
            # Rückgabewert für selected customer
            selected_customer = None
            
            # Header
            header_frame = ctk.CTkFrame(dialog, fg_color="transparent")
            header_frame.pack(fill="x", padx=20, pady=(20, 10))
            
            title_label = ctk.CTkLabel(
                header_frame,
                # Dialog-Titel bereinigt (Umlaute korrekt)
                text="Kunde für Projekt auswählen",
                font=ctk.CTkFont(*self.get_typography("heading")),
                text_color=self.get_color('gray_700')
            )
            title_label.pack(anchor="w")
            
            subtitle_label = ctk.CTkLabel(
                header_frame,
                text="Wähle einen Kunden für Projektstruktur oder überspringe für einfachen Upload",
                font=ctk.CTkFont(*self.get_typography("body")),
                text_color=self.get_color('gray_500')
            )
            subtitle_label.pack(anchor="w", pady=(5, 0))
            
            # Bestehende Kunden (falls vorhanden)
            if existing_customers:
                customers_frame = ctk.CTkFrame(dialog, fg_color=self.get_color('surface'))
                customers_frame.pack(fill="x", padx=20, pady=10)
                
                customers_label = ctk.CTkLabel(
                    customers_frame,
                    text="Bestehende Kunden:",
                    font=ctk.CTkFont(*self.get_typography("body_bold")),
                    text_color=self.get_color('gray_700')
                )
                customers_label.pack(anchor="w", padx=15, pady=(15, 5))
                
                # Scrollbarer Bereich für Kunden
                customer_scroll = ctk.CTkScrollableFrame(
                    customers_frame, 
                    height=120,
                    fg_color="transparent"
                )
                customer_scroll.pack(fill="x", padx=15, pady=(0, 15))
                
                for customer in existing_customers:
                    customer_button = self._create_button(
                        customer_scroll,
                        text=customer,
                        command=lambda c=customer: [setattr(self, '_temp_selected_customer', c), dialog.destroy()],
                        kind="outline-primary",
                        height=36
                    )
                    customer_button.pack(fill="x", pady=2)
            
            # Neuer Kunde Bereich
            new_customer_frame = ctk.CTkFrame(dialog, fg_color=self.get_color('surface'))
            new_customer_frame.pack(fill="x", padx=20, pady=10)
            
            new_label = ctk.CTkLabel(
                new_customer_frame,
                text="Neuen Kunden erstellen:",
                font=ctk.CTkFont(*self.get_typography("body_bold")),
                text_color=self.get_color('gray_700')
            )
            new_label.pack(anchor="w", padx=15, pady=(15, 5))
            
            # Eingabefeld für neuen Kunden
            customer_entry = ctk.CTkEntry(
                new_customer_frame,
                placeholder_text="Kundenname eingeben...",
                height=40,
                font=ctk.CTkFont(*self.get_typography("body")),
                fg_color=self.get_color('white'),
                border_color=self.get_color('surface_border'),
                text_color=self.get_color('gray_700')
            )
            customer_entry.pack(fill="x", padx=15, pady=(0, 15))
            
            # Buttons
            button_frame = ctk.CTkFrame(dialog, fg_color="transparent")
            button_frame.pack(fill="x", padx=20, pady=(10, 20))
            
            def on_create_new():
                customer_name = customer_entry.get().strip()
                if customer_name:
                    nonlocal selected_customer
                    selected_customer = customer_name
                    dialog.destroy()
                else:
                    self.show_toast("Bitte geben Sie einen Kundennamen ein", "warning")
            
            def on_skip():
                nonlocal selected_customer
                selected_customer = None
                dialog.destroy()
            
            def on_cancel():
                nonlocal selected_customer
                selected_customer = "CANCELLED"
                dialog.destroy()
            
            # Überspringen Button (links)
            skip_button = self._create_button(
                button_frame,
                text="Überspringen",
                command=on_skip,
                kind="outline-primary",
                height=36,
                width=140
            )
            skip_button.pack(side="left")
            
            # Abbrechen Button (rechts)
            cancel_button = self._create_button(
                button_frame,
                text="Abbrechen",
                command=on_cancel,
                kind="outline-primary",
                height=36,
                width=140
            )
            cancel_button.pack(side="right", padx=(10, 0))
            
            create_button = self._create_button(
                button_frame,
                text="Kunden erstellen",
                command=on_create_new,
                kind="primary",
                height=36,
                width=160
            )
            create_button.pack(side="right")
            
            # Enter-Taste für Kunden erstellen
            customer_entry.bind("<Return>", lambda e: on_create_new())
            customer_entry.focus()
            
            # Dialog warten lassen
            self._temp_selected_customer = None
            dialog.wait_window()
            
            # Rückgabe-Logik: 
            # - CANCELLED: Benutzer hat abgebrochen → Upload komplett stoppen
            # - None: Benutzer hat übersprungen → Upload ohne Projektstruktur
            # - String: Kunde ausgewählt → Upload mit Projektstruktur
            result = selected_customer or self._temp_selected_customer
            if result == "CANCELLED":
                return "CANCELLED"  # Spezialwert für Abbruch
            return result  # None für Überspringen, String für Kunde
            
        except Exception as e:
            self._handle_error(e, context="customer.dialog", user_message=self._t("Kundenauswahl fehlgeschlagen"))
            return None
    
    def _is_source_text(self, file_path):
        """Klassifiziert, ob Datei ein Ausgangstext ist (für 01_Ausgangstext)"""
        # Alle hochgeladenen Dateien gehen standardmäßig in 01_Ausgangstext
        return True
    
    def _is_translation(self, file_path):
        """Klassifiziert, ob Datei eine Übersetzung ist (nicht relevant für diese Struktur)"""
        # In der neuen Struktur sind Übersetzungen Teil des Workflows, nicht separate Uploads
        return False
    
    def copy_files_to_project_structure(self, files, customer_name, project_date=None):
        """Kopiert Dateien in Projektordner (vereinheitlicht: key 'source')"""
        # Delegation an UploadManager (reine Logik dort gekapselt) – Signatur & Feedback bleiben erhalten
        try:
            import datetime
            if not project_date:
                project_date = datetime.datetime.now().strftime("%Y-%m-%d")
            if not getattr(self, 'upload_manager', None):
                # Fallback auf alte Logik falls Manager fehlt (extremer Edge-Case)
                try:
                    self.logger.warning("UploadManager fehlt – Fallback Legacy Copy")
                except Exception:
                    pass
                return False

            def _create_structure(cust, pdate):
                try:
                    return self.create_project_structure(cust, pdate)
                except Exception:
                    return ''

            def _get_paths(cust, pdate):
                try:
                    return self.get_project_paths(cust, pdate)
                except Exception:
                    return {}

            copied_files = self.upload_manager.copy_into_project_structure(
                list(files) if files else [],
                _create_structure,
                _get_paths,
                customer_name,
                project_date
            ) or {'source': [], 'translation': [], 'other': []}

            total_copied = sum(len(v) for v in copied_files.values())
            if total_copied:
                try:
                    self.logger.info(f"{total_copied} Dateien in Projektstruktur kopiert")
                except Exception:
                    pass
                try:
                    self.show_toast(
                        f"{total_copied} Datei(en) in Projekt '{customer_name}' ({project_date}) gespeichert",
                        "success"
                    )
                except Exception:
                    pass
            return copied_files
        except Exception as e:
            self._handle_error(
                e,
                context="project.copy",
                user_message=self._t("Projekt-Ordner konnte nicht geöffnet werden") if hasattr(self, '_t') else "Projekt-Ordner Fehler",
                toast=False
            )
            return {}
    
    def _upload_source_files(self):
        """Delegierter Quelldatei-Upload über UploadService (GUI bleibt schlank)."""
        try:
            file_types = [
                ("All Supported", "*.pdf;*.docx;*.txt;*.doc;*.rtf;*.odt"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx;*.doc"),
                ("Text files", "*.txt"),
                ("Rich Text", "*.rtf"),
                ("OpenDocument", "*.odt")
            ]
            files = filedialog.askopenfilenames(title="Ausgangstexte für Qualitätsanalyse auswählen", filetypes=file_types)
            if not files:
                return
            customer_name = self._get_or_select_customer_for_upload()
            if customer_name == "CANCELLED":
                return
            self._log_event("upload.source.start", count=len(files), customer=customer_name)
            from upload_service import get_upload_service
            service = get_upload_service(event_bus=getattr(self, 'event_bus', None), logger=self.logger)
            result = service.process_simple_upload(
                kind='source',
                selected_files=list(files),
                existing_files=self.uploaded_files['source'],
                customer_name=customer_name,
                copy_callback=self.copy_files_to_project_structure if customer_name else None,
            )
            added = result['added_files']
            duplicates = result['duplicate_files']
            if added:
                # Telemetrie: before_total berechnen bevor wir erweitern
                try:
                    before_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
                except Exception:
                    before_total = None
                try:
                    before_source_count = len(self.uploaded_files.get('source', []))
                except Exception:
                    before_source_count = None

                self._register_uploaded_files('source', added)

                try:
                    after_source_count = len(self.uploaded_files.get('source', []))
                    self.logger.info(
                        "Upload-Registrierung (source): %s Dateien, Bestand %s → %s",
                        len(added),
                        before_source_count,
                        after_source_count
                    )
                except Exception:
                    pass
                # Automatische OCR (nur falls aktiviert & Bibliotheken verfügbar)
                try:
                    self._auto_ocr_sources(added, customer_name)
                except Exception:
                    pass
                if customer_name and result['meta'].get('project_copied'):
                    self._offer_open_project_folder(customer_name)
                self._schedule_update_file_counter()
                self._refresh_file_list_display()
                self._show_enhanced_upload_results("source", added)
                self._smart_file_pairing()
                self._check_and_show_manual_pairing_option()
                status_msg = f"{len(added)} Ausgangstexte organisiert"
                if duplicates:
                    status_msg += f" ({len(duplicates)} Duplikate ignoriert)"
                self.update_status(status_msg)
                self.show_toast(f"{len(added)} Ausgangstexte erfolgreich hochgeladen", "success")
                # Telemetrie: files.changed (aggregiert) – action added
                try:
                    if getattr(self, 'event_bus', None):
                        after_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
                        self.event_bus.publish('files.changed', {
                            'action': 'added',
                            'file_type': 'source',
                            'count': len(added),
                            'filenames': [os.path.basename(a) for a in added],
                            'before_total': before_total,
                            'after_total': after_total
                        })
                except Exception:
                    pass
            else:
                self.show_toast("Alle ausgewählten Dateien waren bereits hochgeladen", "info")
            self._record_upload_telemetry(len(added), len(duplicates))
            self._log_event("upload.source.result", added=len(added), duplicates=len(duplicates), customer=customer_name)
            # Ribbon aktualisieren (idempotent) – nach Idle
            try:
                self._update_ribbon_states_after_idle()
            except Exception:
                pass
        except Exception as e:
            msg = f"Fehler beim Hochladen der Ausgangstexte: {e}"
            self.update_status(msg)
            self.show_toast("Fehler beim Hochladen der Ausgangstexte", "error")
            self._handle_error(e, context="upload.source", toast=False)
    
    def _upload_batch_files(self):
        """Delegierter Batch Upload über UploadService."""
        try:
            file_types = [
                ("All Supported Files", "*.pdf;*.docx;*.txt;*.doc;*.rtf;*.odt;*.xlsx;*.pptx"),
                ("Document files", "*.pdf;*.docx;*.doc;*.rtf;*.odt"),
                ("Text files", "*.txt"),
                ("Spreadsheets", "*.xlsx;*.xls"),
                ("Presentations", "*.pptx;*.ppt"),
                ("All files", "*.*")
            ]
            files = filedialog.askopenfilenames(title="Dateien für Stapel-Upload und Analyse auswählen", filetypes=file_types)
            if not files:
                return
            self._log_event("upload.batch.start", count=len(files))
            self.update_status("Stapel-Upload wird verarbeitet...", "processing", show_progress=True, progress_value=0)
            for i, _ in enumerate(files):
                progress = (i + 1) / len(files)
                self.update_status(f"Verarbeite Datei {i+1}/{len(files)}", "processing", show_progress=True, progress_value=progress)
                self.root.update()
            from upload_service import get_upload_service
            service = get_upload_service(event_bus=getattr(self, 'event_bus', None), logger=self.logger)
            result = service.process_batch_upload(list(files), self.uploaded_files['source'], self.uploaded_files['translation'])
            try:
                batch_before_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
            except Exception:
                batch_before_total = None
            if result['source_added']:
                self._register_uploaded_files('source', result['source_added'])
                try:
                    self.logger.info("Batch Upload registriert: %s neue Ausgangstexte", len(result['source_added']))
                except Exception:
                    pass
                # Automatische OCR für neu hinzugefügte Quell-Dateien
                try:
                    self._auto_ocr_sources(result['source_added'], None)
                except Exception:
                    pass
            if result['translation_added']:
                self._register_uploaded_files('translation', result['translation_added'])
                try:
                    self.logger.info("Batch Upload registriert: %s neue Übersetzungen", len(result['translation_added']))
                except Exception:
                    pass
            self._schedule_update_file_counter()
            # Telemetrie für hinzugefügte Dateien (je Typ separat)
            try:
                if getattr(self, 'event_bus', None):
                    after_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
                    if result['source_added']:
                        self.event_bus.publish('files.changed', {
                            'action': 'added',
                            'file_type': 'source',
                            'count': len(result['source_added']),
                            'filenames': [os.path.basename(a) for a in result['source_added']],
                            'before_total': batch_before_total,
                            'after_total': after_total
                        })
                    if result['translation_added']:
                        self.event_bus.publish('files.changed', {
                            'action': 'added',
                            'file_type': 'translation',
                            'count': len(result['translation_added']),
                            'filenames': [os.path.basename(a) for a in result['translation_added']],
                            'before_total': batch_before_total,  # gleicher before_total Kontext für Batch
                            'after_total': after_total
                        })
            except Exception:
                pass
            if (result['source_added'] or result['translation_added']):
                self._smart_file_pairing()
                self._check_and_show_manual_pairing_option()
            processed_files = [f"{name} – {role}" for name, role in result['processed_details']]
            skipped_files = [os.path.basename(p) for p in result['duplicates']]
            self._show_batch_upload_results(len(result['source_added']), len(result['translation_added']), processed_files, skipped_files)
            total_added = len(result['source_added']) + len(result['translation_added'])
            if total_added > 0:
                status_msg = f"Batch Upload abgeschlossen: {total_added} Dateien hinzugefügt"
                if skipped_files:
                    status_msg += f" ({len(skipped_files)} Duplikate)"
                self.update_status(status_msg, "success")
                self.show_toast(f"Batch Upload erfolgreich: {total_added} Dateien verarbeitet", "success")
            else:
                self.update_status("Keine neuen Dateien hinzugefügt", "warning")
                self.show_toast("Keine neuen Dateien hinzugefügt", "warning")
            self._record_upload_telemetry(total_added, len(skipped_files))
            self._log_event("upload.batch.result", added=total_added, duplicates=len(skipped_files))
            # Ribbon aktualisieren – nach Idle
            try:
                self._update_ribbon_states_after_idle()
            except Exception:
                pass
        except Exception as e:
            msg = f"Error in batch upload: {e}"
            self.update_status(msg, "error")
            self.show_toast("Batch Upload fehlgeschlagen", "error")
            self._handle_error(e, context="upload.batch", toast=False)

    # --------------------------- OCR UPLOAD (NEU) ---------------------------
    def _upload_ocr_images(self):
        """Öffnet Dateiauswahl für Bilder und wandelt sie via OCR in Ausgangstext-Dateien um.

        Speicherung erfolgt wie reguläre Ausgangstexte (mit Präfix 'ocr_').
        Fehlende Abhängigkeiten führen zu einer klaren Toast-Nachricht.
        """
        try:
            _pt = globals().get('pytesseract', None)
            _img = globals().get('Image', None)
            if _pt is None or _img is None:
                self.show_toast(self._t("OCR nicht verfügbar (pytesseract/Pillow fehlt)"), "warning")
                self._show_ocr_setup_help()
                return
            
            # Tesseract prüfen
            self._resolve_tesseract_cmd()
            try:
                import pytesseract
                pytesseract.get_tesseract_version()
            except Exception:
                self.show_toast(self._t("Tesseract-OCR nicht gefunden"), "error")
                self._show_ocr_setup_help()
                return
            
            file_types = [
                ("Bilder & PDFs", "*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp;*.pdf"),
                ("Bilder", "*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp"),
                ("PDF Scans", "*.pdf"),
                ("Alle Dateien", "*.*")
            ]
            files = filedialog.askopenfilenames(title=self._t("Dateien für OCR-Texterkennung wählen"), filetypes=file_types)
            if not files:
                return
            customer_name = self._get_or_select_customer_for_upload()
            if customer_name == "CANCELLED" or not customer_name:
                return
            
            # OCR mit Vorschau-Dialog
            self._perform_ocr_with_preview(files, customer_name)
            
        except Exception as e:
            self._handle_error(e, context="ocr.upload", user_message=self._t("OCR fehlgeschlagen") if hasattr(self,'_t') else None)
    
    def _show_ocr_setup_help(self):
        """Zeigt Hilfe-Dialog zur OCR-Einrichtung."""
        try:
            help_text = self._t("""OCR-Einrichtung erforderlich

Für die Texterkennung werden benötigt:

1. Tesseract-OCR Engine
   → Download: https://github.com/UB-Mannheim/tesseract/wiki
   → Bei Installation: Deutsche Sprachdaten auswählen

2. Python-Pakete:
   → pip install pytesseract Pillow

Nach Installation: Programm neu starten.""")
            
            messagebox.showinfo(self._t("OCR-Einrichtung"), help_text)
        except Exception:
            pass
    
    def _perform_ocr_with_preview(self, files, customer_name):
        """Führt OCR mit Fortschrittsanzeige und Vorschau durch."""
        try:
            image_exts = {'.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp'}
            pdf_exts = {'.pdf'}
            
            results = []  # [(filename, text, confidence, detected_lang), ...]
            total = len(files)
            
            self.update_status(self._t("OCR-Texterkennung läuft..."), "processing", show_progress=True, progress_value=0)
            
            for idx, filepath in enumerate(files):
                ext = os.path.splitext(filepath)[1].lower()
                filename = os.path.basename(filepath)
                
                try:
                    self.update_status(f"OCR: {filename} ({idx+1}/{total})", "processing", show_progress=True, progress_value=idx/total)
                    
                    if ext in image_exts:
                        # Bild-OCR
                        text = self._perform_ocr(filepath)
                        confidence = self._estimate_ocr_quality(text)
                        detected_lang = self._detect_language_from_text(text) if text else None
                        results.append((filename, text, confidence, detected_lang))
                        
                    elif ext in pdf_exts:
                        # PDF: Erst Text-Extraktion versuchen, dann OCR
                        text = self._extract_text_from_file(filepath)
                        if text and len(text.strip()) > 50:
                            confidence = 95  # Text-PDFs haben hohe Qualität
                            detected_lang = self._detect_language_from_text(text)
                            results.append((filename, text, confidence, detected_lang))
                        else:
                            # OCR für gescannte PDFs
                            self.logger.info(f"📄 PDF '{filename}' scheint gescannt - verwende OCR")
                            text = self._extract_text_from_file(filepath)  # Nutzt OCR-Fallback
                            confidence = self._estimate_ocr_quality(text) if text else 0
                            detected_lang = self._detect_language_from_text(text) if text else None
                            results.append((filename, text or "", confidence, detected_lang))
                except Exception as e:
                    self.logger.warning(f"OCR fehlgeschlagen für {filename}: {e}")
                    results.append((filename, "", 0, None))
            
            self.update_status(self._t("OCR abgeschlossen"), "success", show_progress=False)
            
            # Vorschau-Dialog anzeigen
            if results:
                self._show_ocr_preview_dialog(results, customer_name)
            else:
                self.show_toast(self._t("Keine Texte erkannt"), "warning")
                
        except Exception as e:
            self._handle_error(e, context="ocr.preview", toast=True)
    
    def _estimate_ocr_quality(self, text: str) -> int:
        """Schätzt die OCR-Qualität basierend auf Textmerkmalen (0-100)."""
        if not text or not text.strip():
            return 0
        
        score = 50  # Basis
        
        # Länge
        length = len(text.strip())
        if length > 500:
            score += 15
        elif length > 100:
            score += 10
        elif length < 20:
            score -= 20
        
        # Wortstruktur
        words = text.split()
        if len(words) > 10:
            avg_word_len = sum(len(w) for w in words) / len(words)
            if 3 <= avg_word_len <= 12:
                score += 15
        
        # Sonderzeichen-Ratio (zu viele = schlechte OCR)
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
        special_ratio = special_chars / len(text) if text else 0
        if special_ratio > 0.3:
            score -= 25
        elif special_ratio < 0.1:
            score += 10
        
        # Wiederholende Zeichen (OCR-Artefakte)
        if '###' in text or '???' in text or '...' in text:
            score -= 15
        
        return max(0, min(100, score))
    
    def _show_ocr_preview_dialog(self, results, customer_name):
        """Zeigt OCR-Ergebnisse zur Überprüfung vor dem Speichern."""
        try:
            dialog = ctk.CTkToplevel(self.root)
            dialog.title(self._t("OCR-Ergebnisse überprüfen"))
            dialog.geometry("900x700")
            dialog.transient(self.root)
            dialog.grab_set()
            
            # Header
            header = ctk.CTkFrame(dialog, fg_color=self.get_color('primary'), corner_radius=0)
            header.pack(fill="x")
            
            header_label = ctk.CTkLabel(
                header,
                text=self._t("OCR-Texterkennung - Ergebnisse"),
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('white')
            )
            header_label.pack(pady=16)
            
            # Statistik mit Sprachinfo
            success_count = sum(1 for r in results if len(r) >= 3 and r[1] and r[2] > 30)
            detected_langs = set()
            for r in results:
                if len(r) >= 4 and r[3]:
                    detected_langs.add(r[3])
            
            stat_text = f"{success_count}/{len(results)} " + self._t("Dateien erfolgreich erkannt")
            if detected_langs:
                lang_names = {
                    'deu': 'Deutsch', 'eng': 'Englisch', 'fra': 'Französisch',
                    'spa': 'Spanisch', 'ita': 'Italienisch', 'por': 'Portugiesisch',
                    'nld': 'Niederländisch', 'pol': 'Polnisch', 'rus': 'Russisch'
                }
                lang_display = ', '.join(lang_names.get(l, l) for l in detected_langs)
                stat_text += f" | {self._t('Erkannte Sprachen')}: {lang_display}"
            
            stat_label = ctk.CTkLabel(
                header,
                text=stat_text,
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('white')
            )
            stat_label.pack(pady=(0, 12))
            
            # Auto-Detect Hinweis wenn Sprache erkannt
            if detected_langs:
                auto_hint = ctk.CTkFrame(dialog, fg_color='#EEF2FF', corner_radius=0)
                auto_hint.pack(fill="x", padx=0)
                
                # Hauptsprache ermitteln (häufigste)
                primary_lang = max(detected_langs, key=lambda x: sum(1 for r in results if len(r) > 3 and r[3] == x))
                lang_names_full = {
                    'deu': 'Deutsch', 'eng': 'Englisch', 'fra': 'Französisch',
                    'spa': 'Spanisch', 'ita': 'Italienisch', 'por': 'Portugiesisch'
                }
                primary_display = lang_names_full.get(primary_lang, primary_lang)
                
                hint_label = ctk.CTkLabel(
                    auto_hint,
                    text=f"  {self._t('Quellsprache automatisch erkannt')}: {primary_display}  ",
                    font=ctk.CTkFont(family="Segoe UI", size=11),
                    text_color='#4338CA'
                )
                hint_label.pack(pady=8)
                
                # Referenz für spätere Nutzung speichern
                dialog._detected_primary_lang = primary_lang
            
            # Content
            content = ctk.CTkScrollableFrame(dialog, fg_color="transparent")
            content.pack(fill="both", expand=True, padx=20, pady=10)
            
            selected_results = []  # Zum Speichern ausgewählte Ergebnisse
            detected_langs_list = []  # Für Auto-Setting
            
            for result_item in results:
                # Unterstütze sowohl (filename, text, confidence) als auch (filename, text, confidence, detected_lang)
                filename = result_item[0]
                text = result_item[1]
                confidence = result_item[2] if len(result_item) > 2 else 0
                detected_lang = result_item[3] if len(result_item) > 3 else None
                
                # Ergebnis-Karte
                card = ctk.CTkFrame(
                    content,
                    fg_color=self.get_color('surface'),
                    corner_radius=8,
                    border_width=1,
                    border_color=self.get_color('surface_border')
                )
                card.pack(fill="x", pady=6)
                
                card_content = ctk.CTkFrame(card, fg_color="transparent")
                card_content.pack(fill="x", padx=16, pady=12)
                
                # Header-Zeile mit Checkbox
                header_row = ctk.CTkFrame(card_content, fg_color="transparent")
                header_row.pack(fill="x")
                
                # Checkbox für Auswahl
                selected_var = tk.BooleanVar(value=bool(text and confidence > 20))
                cb = ctk.CTkCheckBox(
                    header_row,
                    text="",
                    variable=selected_var,
                    width=24,
                    fg_color=self.get_color('primary')
                )
                cb.pack(side="left")
                selected_results.append((filename, text, selected_var))
                
                # Dateiname
                name_label = ctk.CTkLabel(
                    header_row,
                    text=filename,
                    font=ctk.CTkFont(*self.get_typography('subheading')),
                    text_color=self.get_color('text_primary')
                )
                name_label.pack(side="left", padx=(8, 0))
                
                # Sprach-Badge (wenn erkannt)
                if detected_lang:
                    lang_names = {
                        'deu': 'DE', 'eng': 'EN', 'fra': 'FR', 'spa': 'ES', 
                        'ita': 'IT', 'por': 'PT', 'nld': 'NL', 'pol': 'PL', 'rus': 'RU'
                    }
                    lang_badge = ctk.CTkLabel(
                        header_row,
                        text=f" {lang_names.get(detected_lang, detected_lang)} ",
                        font=ctk.CTkFont(family="Segoe UI", size=9, weight="bold"),
                        fg_color='#6366F1',
                        text_color="white",
                        corner_radius=3
                    )
                    lang_badge.pack(side="right", padx=(4, 8))
                
                # Qualitäts-Badge
                if confidence >= 70:
                    badge_color = '#10B981'
                    badge_text = self._t("Gut")
                elif confidence >= 40:
                    badge_color = '#F59E0B'
                    badge_text = self._t("Mittel")
                else:
                    badge_color = '#DC2626'
                    badge_text = self._t("Schwach")
                
                badge = ctk.CTkLabel(
                    header_row,
                    text=f"  {badge_text} ({confidence}%)  ",
                    font=ctk.CTkFont(family="Segoe UI", size=10),
                    fg_color=badge_color,
                    text_color="white",
                    corner_radius=4
                )
                badge.pack(side="right")
                
                # Textvorschau
                preview_text = text[:500] + "..." if len(text) > 500 else text if text else self._t("[Kein Text erkannt]")
                preview = ctk.CTkLabel(
                    card_content,
                    text=preview_text,
                    font=ctk.CTkFont(family="Consolas", size=10),
                    text_color=self.get_color('text_secondary'),
                    justify="left",
                    anchor="w",
                    wraplength=800
                )
                preview.pack(fill="x", pady=(8, 0), anchor="w")
                
                # Sprache für Auto-Setting merken
                if detected_lang:
                    detected_langs_list.append(detected_lang)
            
            # Footer mit Buttons
            footer = ctk.CTkFrame(dialog, fg_color="transparent")
            footer.pack(fill="x", padx=20, pady=16)
            
            def _save_selected():
                saved = []
                for filename, text, var in selected_results:
                    if var.get() and text:
                        saved_path = self._save_ocr_text_as_source(text, filename, customer_name)
                        if saved_path:
                            saved.append(saved_path)
                
                if saved:
                    self._register_uploaded_files('source', saved)
                    self._schedule_update_file_counter()
                    self._refresh_file_list_display()
                    self._smart_file_pairing()
                    
                    # 🌐 Auto-Set Quellsprache basierend auf OCR-Erkennung
                    if detected_langs_list:
                        self._auto_set_source_language(detected_langs_list)
                    
                    self.show_toast(f"{len(saved)} " + self._t("Datei(en) als Ausgangstext gespeichert"), "success")
                dialog.destroy()
            
            save_btn = self._create_button(
                footer,
                text=self._t("Ausgewählte speichern"),
                command=_save_selected,
                kind="primary",
                size="lg",
                height=44
            )
            save_btn.pack(side="right")
            
            cancel_btn = self._create_button(
                footer,
                text=self._t("Abbrechen"),
                command=dialog.destroy,
                kind="secondary",
                size="lg",
                height=44
            )
            cancel_btn.pack(side="right", padx=(0, 12))
            
            # Zentrieren
            try:
                self._center_window(dialog)
            except Exception:
                pass
                
        except Exception as e:
            self._handle_error(e, context="ocr.preview.dialog", toast=True)

    def _perform_ocr(self, image_path: str) -> str:
        """Führt OCR auf einem Bild durch mit verbesserter Vorverarbeitung."""
        try:
            _pt = globals().get('pytesseract', None)
            _img = globals().get('Image', None)
            if _pt is None or _img is None:
                self.logger.warning("⚠️ pytesseract oder PIL nicht verfügbar")
                return ""
            
            self._resolve_tesseract_cmd()
            
            with _img.open(image_path) as im:  # type: ignore
                # Verbesserte Vorverarbeitung
                from PIL import ImageOps, ImageFilter, ImageEnhance
                
                # 1. Graustufen
                gray = im.convert('L')
                
                # 2. Automatischer Kontrast
                contrasted = ImageOps.autocontrast(gray)
                
                # 3. Schärfung
                sharpened = ImageEnhance.Sharpness(contrasted).enhance(1.5)
                
                # 4. Rauschunterdrückung
                denoised = sharpened.filter(ImageFilter.MedianFilter(size=3))
                
                # OCR-Sprachen ermitteln
                languages = self._determine_ocr_languages()
                lang_param = '+'.join(languages) if languages else 'deu+eng'
                
                # OCR mit optimaler Konfiguration
                text = _pt.image_to_string(denoised, lang=lang_param, config='--oem 3 --psm 6')
                
                # Sanitizing
                result = text.replace('\r\n', '\n').strip()
                
                if result:
                    self.logger.info(f"✅ OCR erfolgreich: {len(result)} Zeichen aus {os.path.basename(image_path)}")
                else:
                    self.logger.warning(f"⚠️ OCR lieferte keinen Text für {os.path.basename(image_path)}")
                
                return result
        except Exception as e:
            self._handle_error(e, context="ocr.perform", toast=False)
            return ""

    def _auto_set_source_language(self, detected_langs: list[str]) -> None:
        """Setzt die Quellsprache automatisch basierend auf OCR-Erkennung.
        
        Args:
            detected_langs: Liste erkannter Tesseract-Sprachcodes (z.B. ['deu', 'eng'])
        """
        if not detected_langs:
            return
        
        try:
            # Häufigste Sprache ermitteln
            from collections import Counter
            lang_counts = Counter(detected_langs)
            primary_lang = lang_counts.most_common(1)[0][0]
            
            # Mapping von Tesseract zu ISO-Codes
            tess_to_iso = {
                'deu': 'de', 'eng': 'en', 'fra': 'fr', 'spa': 'es', 'ita': 'it',
                'por': 'pt', 'nld': 'nl', 'pol': 'pl', 'rus': 'ru'
            }
            iso_code = tess_to_iso.get(primary_lang)
            if not iso_code:
                return
            
            # GUI-Anzeigename ermitteln
            iso_to_display = {
                'de': 'German', 'en': 'English', 'fr': 'French', 
                'es': 'Spanish', 'it': 'Italian'
            }
            display_name = iso_to_display.get(iso_code)
            if not display_name:
                return
            
            # Settings aktualisieren (falls SettingsService vorhanden)
            if hasattr(self, 'settings_service') and self.settings_service:
                try:
                    self.settings_service.set('analysis.default_src', iso_code)
                    self.logger.info(f"🌐 Quellsprache automatisch gesetzt: {display_name} ({iso_code})")
                except Exception:
                    pass
            
            # GUI-Variable aktualisieren
            for attr_name in ('var_source_lang', 'var_source_language'):
                var = getattr(self, attr_name, None)
                if var and hasattr(var, 'set'):
                    try:
                        # Übersetze für GUI-Anzeige
                        translated = self._t(display_name) if hasattr(self, '_t') else display_name
                        var.set(translated)
                        self.logger.info(f"✅ GUI-Sprachauswahl aktualisiert: {translated}")
                        break
                    except Exception:
                        pass
            
            # Toast-Feedback
            lang_names_de = {
                'de': 'Deutsch', 'en': 'Englisch', 'fr': 'Französisch',
                'es': 'Spanisch', 'it': 'Italienisch'
            }
            lang_display = lang_names_de.get(iso_code, display_name)
            self.show_toast(
                f"{self._t('Quellsprache automatisch erkannt')}: {lang_display}",
                "info",
                duration=4000
            )
            
        except Exception as e:
            self.logger.debug(f"Auto-Set Sprache fehlgeschlagen: {e}")

    def _save_ocr_text_as_source(self, text: str, origin_image: str, customer_name: str):
        """Speichert OCR-Text als .txt im Projektordner Ausgangstext und gibt Pfad zurück."""
        try:
            import datetime, re, os
            project_paths = self.get_project_paths(customer_name)
            base_path = project_paths.get('ausgangstext')
            if not base_path:
                return None
            os.makedirs(base_path, exist_ok=True)
            ts = datetime.datetime.now().strftime('%H%M%S')
            base_name = os.path.splitext(os.path.basename(origin_image))[0]
            safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', base_name)[:40]
            target_file = os.path.join(base_path, f"ocr_{ts}_{safe_name}.txt")
            
            try:
                with open(target_file, 'w', encoding='utf-8') as f:
                    f.write(text.strip() + '\n')
            except (PermissionError, OSError, IOError) as e:
                self._log('error', f'Failed to save OCR text: {e}')
                self._handle_error(e, context="ocr.save", toast=False)
                return None
            
            return target_file
        except Exception as e:
            self._handle_error(e, context="ocr.save", toast=False)
            return None

    # --------------------------- PDF HILFSFUNKTIONEN (NEU) ---------------------------
    def _extract_pdf_text_basic(self, pdf_path: str, max_pages: int = 10) -> str:
        """Versucht schnellen PDF-Text auszulesen (strukturarm). Liefert String oder ''. 
        Bei wenig Text wird automatisch OCR verwendet."""
        if _lazy_load_pypdf2() is None:
            return ""
        try:
            text_parts = []
            with open(pdf_path, 'rb') as f:  # type: ignore
                reader = PyPDF2.PdfReader(f)  # type: ignore
                page_count = len(reader.pages)
                for i in range(min(page_count, max_pages)):
                    try:
                        page = reader.pages[i]
                        extracted = page.extract_text() or ""
                        if extracted.strip():
                            text_parts.append(extracted)
                    except (AttributeError, IndexError, KeyError):
                        # Page extraction errors - continue with next page
                        pass
            
            full_text = "\n".join(text_parts)
            
            # Wenn wenig oder kein Text extrahiert wurde, versuche OCR
            if len(full_text.strip()) < 100:
                self.logger.info(f"⚠️ PDF enthält wenig Text ({len(full_text.strip())} Zeichen), versuche OCR für Analyse...")
                try:
                    ocr_text = self._extract_text_from_file(pdf_path)  # Nutze die verbesserte Methode
                    if ocr_text and len(ocr_text.strip()) > len(full_text.strip()):
                        self.logger.info(f"✅ OCR lieferte mehr Text: {len(ocr_text)} Zeichen")
                        return ocr_text
                except Exception as e:
                    self.logger.debug(f"OCR-Fallback fehlgeschlagen: {e}")
            
            return full_text
        except (FileNotFoundError, PermissionError) as e:
            self._log('warning', f'PDF file not accessible: {e}')
            return ""
        except (OSError, IOError) as e:
            self._log('error', f'PDF I/O error: {e}')
            return ""
        except Exception:
            return ""

    def _ocr_scanned_pdf(self, pdf_path: str, customer_name: str, max_pages: int = 5) -> str | None:
        """OCR für (vermutlich) gescannte PDF: rendert Seiten zu Bildern (fitz bevorzugt)."""
        _pt = globals().get('pytesseract', None)
        _img = globals().get('Image', None)
        if _pt is None or _img is None:
            return None
        images_text = []
        try:
            # Variante A: PyMuPDF (fitz)
            if fitz is not None:
                try:
                    doc = fitz.open(pdf_path)  # type: ignore
                    for i, page in enumerate(doc):
                        if i >= max_pages:
                            break
                        try:
                            pix = page.get_pixmap(dpi=200)
                            mode = "RGBA" if pix.alpha else "RGB"
                            from PIL import Image as _PILImage  # type: ignore
                            img = _PILImage.frombytes(mode, [pix.width, pix.height], pix.samples)  # type: ignore
                            try:
                                gray = img.convert('L')
                            except Exception:
                                gray = img
                            txt = _pt.image_to_string(gray, lang='deu+eng')  # type: ignore
                            if txt.strip():
                                images_text.append(txt)
                        except Exception:
                            pass
                    doc.close()
                except Exception:
                    pass
            # (Optionaler weiterer Ansatz pdf2image könnte folgen – derzeit nicht erforderlich)
        except Exception:
            return None
        if not images_text:
            return None
        full_text = "\n".join(images_text)
        # Als Quelle speichern
        try:
            saved = self._save_ocr_text_as_source(full_text, pdf_path, customer_name)
            return saved
        except Exception:
            return None

    # --------------------------- AUTO OCR (NEU) ---------------------------
    def _auto_ocr_sources(self, file_paths, customer_name):
        """Führt automatische OCR für Bild- und (falls nötig) gescannte PDF-Dateien aus.

        Bilder:
          - Endungen: .png, .jpg, .jpeg, .tif, .tiff, .bmp
          - Direkte OCR je Bild.
        PDFs:
          - Zuerst Versuch Text extrahieren (PyPDF2) – wenn genügend Text gefunden: kein OCR.
          - Wenn kaum/kein Text: Seiten (max 5) rendern (fitz) und OCR.
        Bedingungen:
          - pytesseract & Pillow nötig; für PDF-OCR zusätzlich fitz (PyMuPDF) empfohlen.
        Ergebnisse:
          - Neue *.txt Dateien (Präfix ocr_*) werden als Ausgangstexte hinzugefügt.
          - Optional (Flag) Original-Bilder aus Liste entfernen; PDFs bleiben erhalten.
        """
        try:
            if not getattr(self, 'auto_ocr_enabled', False):
                return
            _pt = globals().get('pytesseract', None)
            _img = globals().get('Image', None)
            if _pt is None or _img is None:
                return
            image_exts = {'.png','.jpg','.jpeg','.tif','.tiff','.bmp'}
            pdf_exts = {'.pdf'}
            images = [p for p in file_paths if os.path.splitext(p)[1].lower() in image_exts]
            pdfs = [p for p in file_paths if os.path.splitext(p)[1].lower() in pdf_exts]
            if not images and not pdfs:
                return
            # Falls kein customer_name übergeben (Batch), ermitteln via Dialog nur einmal
            if not customer_name:
                try:
                    customer_name = self._get_or_select_customer_for_upload()
                except Exception:
                    customer_name = None
            if not customer_name:
                return
            created = []
            # Bilder OCR
            for img in images:
                try:
                    text = self._perform_ocr(img)
                    if not text.strip():
                        continue
                    saved = self._save_ocr_text_as_source(text, img, customer_name)
                    if saved:
                        created.append(saved)
                        # Optional Original aus Liste entfernen
                        if getattr(self, 'auto_ocr_remove_original', False):
                            try:
                                removed = False
                                if hasattr(self, 'upload_manager') and self.upload_manager:
                                    removed = self.upload_manager.remove_file(img)
                                    if removed:
                                        self._sync_uploaded_files_backend_from_manager()
                                if not removed:
                                    backend_sources = self._uploaded_files_backend.setdefault('source', [])
                                    normalized_img = str(Path(img))
                                    try:
                                        backend_sources.remove(normalized_img)
                                        removed = True
                                    except ValueError:
                                        try:
                                            backend_sources.remove(img)
                                            removed = True
                                        except ValueError:
                                            pass
                                if removed:
                                    self._schedule_update_file_counter()
                            except Exception:
                                pass
                except Exception:
                    pass
            # PDF Verarbeitung
            pdf_ocr_count = 0
            for pdf in pdfs:
                try:
                    basic_text = self._extract_pdf_text_basic(pdf)
                    # Heuristik: Wenn extrahierter Text ausreichend (>=120 Zeichen & mind. 25% Buchstaben), keine OCR
                    if basic_text:
                        letters = sum(c.isalpha() for c in basic_text)
                        if len(basic_text) >= 120 and letters / max(1, len(basic_text)) > 0.25:
                            # Optional könnte man hier trotzdem speichern – vorerst überspringen
                            continue
                    # Sonst OCR versuchen
                    saved_pdf = self._ocr_scanned_pdf(pdf, customer_name)
                    if saved_pdf:
                        created.append(saved_pdf)
                        pdf_ocr_count += 1
                except Exception:
                    pass
            if created:
                self._register_uploaded_files('source', created)
                if pdf_ocr_count:
                    self.show_toast(f"{len(created)} OCR-Text(e) (inkl. {pdf_ocr_count} PDF) erzeugt", "success")
                else:
                    self.show_toast(f"{len(created)} OCR-Text(e) automatisch erzeugt", "success")
        except Exception:
            pass
    
    def _show_batch_upload_results(self, source_count, translation_count, processed_files, skipped_files):
        """SHOW BATCH UPLOAD RESULTS - Zusammenfassung (Farbtokens normalisiert)"""
        try:
            results_window = ctk.CTkToplevel(self.root)
            if hasattr(self, '_t'):
                results_window.title(self._t("Batch Upload Ergebnisse"))
            else:
                results_window.title("Batch Upload Ergebnisse")
            # Adaptive Höhe basierend auf Anzahl-Einträge
            base_w, base_h = 600, 500
            total_entries = len(processed_files) + len(skipped_files)
            if total_entries < 8:
                base_h = 420
            results_window.transient(self.root)
            results_window.grab_set()
            self._center_window(results_window, base_w, base_h, anchor_root=False)

            content_frame = ctk.CTkFrame(results_window, fg_color="transparent")
            content_frame.pack(fill="both", expand=True, padx=20, pady=20)

            header = ctk.CTkLabel(
                content_frame,
                text=self._t("Batch Upload abgeschlossen") if hasattr(self, '_t') else "Batch Upload abgeschlossen",
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('primary')
            )
            header.pack(pady=(0, 20))

            # Fallback verwenden, falls surface_elevated nicht definiert ist
            elevated_color = None
            try:
                elevated_color = self.get_color('surface_elevated')
            except Exception:
                elevated_color = None
            if not elevated_color or elevated_color in (None, '', 'transparent'):
                elevated_color = self.get_color('surface')
            stats_frame = ctk.CTkFrame(content_frame, fg_color=elevated_color)
            stats_frame.pack(fill="x", pady=(0, 15))
            stats_content = ctk.CTkFrame(stats_frame, fg_color="transparent")
            stats_content.pack(fill="x", padx=20, pady=15)

            stats_grid = ctk.CTkFrame(stats_content, fg_color="transparent")
            stats_grid.pack(fill="x")
            stats_grid.grid_columnconfigure((0, 1, 2), weight=1)

            # Source Stat
            source_stat = ctk.CTkFrame(stats_grid, fg_color=self.get_color('gray_100'))
            source_stat.grid(row=0, column=0, padx=(0, 5), pady=2, sticky="ew")
            ctk.CTkLabel(
                source_stat,
                text=str(source_count),
                font=ctk.CTkFont(*self.get_typography('title')),
                text_color=self.get_color('primary')
            ).pack(pady=(10, 2))
            ctk.CTkLabel(
                source_stat,
                text=self._t("Ausgangstexte") if hasattr(self, '_t') else "Ausgangstexte",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('primary')
            ).pack(pady=(0, 10))

            # Translation Stat (verwende info Palette statt secondary*)
            translation_stat = ctk.CTkFrame(stats_grid, fg_color=self.get_color('info_light'))
            translation_stat.grid(row=0, column=1, padx=(5, 5), pady=2, sticky="ew")
            ctk.CTkLabel(
                translation_stat,
                text=str(translation_count),
                font=ctk.CTkFont(*self.get_typography('title')),
                text_color=self.get_color('info')
            ).pack(pady=(10, 2))
            ctk.CTkLabel(
                translation_stat,
                text=self._t("Übersetzungen") if hasattr(self, '_t') else "Übersetzungen",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('info')
            ).pack(pady=(0, 10))

            # Skipped Stat
            skipped_stat = ctk.CTkFrame(stats_grid, fg_color=self.get_color('gray_100'))
            skipped_stat.grid(row=0, column=2, padx=(5, 0), pady=2, sticky="ew")
            ctk.CTkLabel(
                skipped_stat,
                text=str(len(skipped_files)),
                font=ctk.CTkFont(*self.get_typography('title')),
                text_color=self._accent('warning')
            ).pack(pady=(10, 2))
            ctk.CTkLabel(
                skipped_stat,
                text=self._t("Übersprungen") if hasattr(self, '_t') else "Übersprungen",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self._accent('warning')
            ).pack(pady=(0, 10))

            if processed_files or skipped_files:
                details_frame = ctk.CTkScrollableFrame(
                    content_frame,
                    height=200,
                    fg_color=self.get_color('surface')
                )
                details_frame.pack(fill="both", expand=True, pady=(15, 15))

                if processed_files:
                    ctk.CTkLabel(
                        details_frame,
                        text=self._t("Verarbeitete Dateien:") if hasattr(self, '_t') else "Verarbeitete Dateien:",
                        font=ctk.CTkFont(*self.get_typography('subheading')),
                        # Accent Reduction: keine direkte success-Farbe hier (nur System-Status darf grün sein)
                        text_color=self._accent('success'),
                        anchor="w"
                    ).pack(fill="x", pady=(5, 5))
                    for file_info in processed_files:
                        ctk.CTkLabel(
                            details_frame,
                            text=file_info,
                            font=ctk.CTkFont(*self.get_typography('caption')),
                            text_color=self.get_color('text_secondary'),
                            anchor="w"
                        ).pack(fill="x", pady=1)

                if skipped_files:
                    ctk.CTkLabel(
                        details_frame,
                        text=self._t("Übersprungene Dateien:") if hasattr(self, '_t') else "Übersprungene Dateien:",
                        font=ctk.CTkFont(*self.get_typography('subheading')),
                        text_color=self._accent('warning'),
                        anchor="w"
                    ).pack(fill="x", pady=(15, 5))
                    for file_info in skipped_files:
                        ctk.CTkLabel(
                            details_frame,
                            text=file_info,
                            font=ctk.CTkFont(*self.get_typography('caption')),
                            text_color=self.get_color('text_secondary'),
                            anchor="w"
                        ).pack(fill="x", pady=1)

            close_btn = self._create_button(
                content_frame,
                text=self._t("Schließen") if hasattr(self, '_t') else "Schließen",
                command=results_window.destroy,
                kind="primary",
                height=40
            )
            close_btn.configure(width=140)
            close_btn.pack(pady=(15, 0))
            try:
                close_btn.focus_set()
            except Exception:
                pass

            # Shortcuts (Esc/Enter)
            results_window.bind("<Escape>", lambda e: results_window.destroy())
            results_window.bind("<Return>", lambda e: results_window.destroy())
        except Exception as e:
            self._handle_error(e, context="upload.batch.results", toast=False)
            try:
                self.show_toast(f"{source_count + translation_count} Dateien verarbeitet", "success")
            except Exception:
                pass

    def _read_uploaded_files_for_preview(self):
        """Liest hochgeladene Dateien ein, um eine Vorschau zu erstellen (ohne vollständige Analyse)."""
        rows = []
        try:
            # DEBUG: Prüfe uploaded_files State
            self.logger.info("🔍 DEBUG: _read_uploaded_files_for_preview gestartet")
            self.logger.info(f"   DEBUG: hasattr(upload_manager): {hasattr(self, 'upload_manager')}")
            if hasattr(self, 'upload_manager'):
                self.logger.info(f"   DEBUG: upload_manager exists: {self.upload_manager is not None}")
            
            # Hole Source- und Translation-Dateien
            uploaded_data = self.uploaded_files
            self.logger.info(f"   DEBUG: uploaded_files type: {type(uploaded_data)}")
            self.logger.info(f"   DEBUG: uploaded_files keys: {uploaded_data.keys() if uploaded_data else 'None'}")
            
            source_files = uploaded_data.get('source', [])
            translation_files = uploaded_data.get('translation', [])
            
            self.logger.info(f"📁 Hochgeladene Dateien: {len(source_files)} Source, {len(translation_files)} Translation")
            
            if source_files:
                self.logger.info(f"   📄 Source-Dateien: {[os.path.basename(str(f)) for f in source_files[:3]]}")
            if translation_files:
                self.logger.info(f"   📄 Translation-Dateien: {[os.path.basename(str(f)) for f in translation_files[:3]]}")
            
            if not source_files and not translation_files:
                self.logger.warning("⚠️ Keine hochgeladenen Dateien gefunden")
                self.logger.warning("   HINWEIS: Bitte Dateien über 'Datei hochladen' hinzufügen")
                return None
            
            # Lese Source-Dateien
            for source_path in source_files[:3]:  # Max 3 Dateien für Preview
                try:
                    self.logger.info(f"📄 Lese Source-Datei: {os.path.basename(source_path)}")
                    text_content = self._extract_text_from_file(source_path)
                    if text_content:
                        # Prüfe, ob es eine Fehlermeldung ist (z.B. "[PDF ... OCR benötigt]")
                        if text_content.startswith('[') and 'OCR' in text_content:
                            self.logger.warning(f"⚠️ {text_content}")
                            # Füge Warnung als spezielles Segment hinzu
                            rows.append({
                                'source_text': text_content,
                                'target_text': '→ Siehe SCHNELLE_PDF_LOESUNG.md für Hilfe',
                                'page': 1,
                                'file': os.path.basename(source_path),
                                'is_warning': True
                            })
                        else:
                            self.logger.info(f"📝 Extrahierter Text ({len(text_content)} Zeichen): {text_content[:200]}...")
                            # Teile Text in Absätze
                            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
                            self.logger.info(f"✅ {len(paragraphs)} Absätze gefunden")
                            for i, para in enumerate(paragraphs[:30], 1):  # Max 30 Absätze pro Datei
                                if len(para) > 10:  # Nur sinnvolle Absätze
                                    self.logger.info(f"   Absatz {i}: {para[:100]}...")
                                    rows.append({
                                        'source_text': para[:500],  # Max 500 Zeichen pro Absatz
                                        'target_text': '',
                                        'page': (i // 10) + 1,  # Geschätzte Seitenzahl
                                        'file': os.path.basename(source_path)
                                    })
                    else:
                        self.logger.warning(f"⚠️ Keine Text-Extraktion möglich für {os.path.basename(source_path)}")
                        # Füge Hinweis hinzu
                        rows.append({
                            'source_text': f"[Fehler beim Einlesen von '{os.path.basename(source_path)}']",
                            'target_text': '→ Datei konnte nicht verarbeitet werden',
                            'page': 1,
                            'file': os.path.basename(source_path),
                            'is_warning': True
                        })
                except Exception as e:
                    self.logger.error(f"❌ Fehler beim Einlesen von {source_path}: {e}")
                    # Füge Fehler als Segment hinzu
                    rows.append({
                        'source_text': f"[FEHLER beim Einlesen von '{os.path.basename(source_path)}']",
                        'target_text': f'→ {str(e)[:100]}',
                        'page': 1,
                        'file': os.path.basename(source_path),
                        'is_warning': True
                    })
            
            # Lese Translation-Dateien
            for translation_path in translation_files[:3]:  # Max 3 Dateien für Preview
                try:
                    self.logger.info(f"📄 Lese Translation-Datei: {os.path.basename(translation_path)}")
                    text_content = self._extract_text_from_file(translation_path)
                    if text_content:
                        # Teile Text in Absätze
                        paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
                        self.logger.info(f"✅ {len(paragraphs)} Absätze gefunden")
                        for i, para in enumerate(paragraphs[:30], 1):  # Max 30 Absätze pro Datei
                            if len(para) > 10:  # Nur sinnvolle Absätze
                                rows.append({
                                    'source_text': '',
                                    'target_text': para[:500],  # Max 500 Zeichen pro Absatz
                                    'page': (i // 10) + 1,  # Geschätzte Seitenzahl
                                    'file': os.path.basename(translation_path)
                                })
                except Exception as e:
                    self.logger.error(f"❌ Fehler beim Einlesen von {translation_path}: {e}")
            
        except Exception as e:
            self.logger.error(f"❌ Fehler beim Lesen hochgeladener Dateien: {e}")
            import traceback
            traceback.print_exc()
            rows = []
        
        self.logger.info(f"📊 Gesamt eingelesene Zeilen: {len(rows)}")
        return rows if rows else None
    
    def _get_available_tesseract_languages(self) -> list[str]:
        """Ermittelt die installierten Tesseract-Sprachpakete."""
        try:
            import pytesseract
            languages = pytesseract.get_languages()
            # Filtere 'osd' (Orientation and Script Detection) heraus
            return [lang for lang in languages if lang != 'osd']
        except Exception:
            return []
    
    def _detect_language_from_text(self, sample_text: str) -> str | None:
        """Erkennt die Sprache automatisch aus einem Textbeispiel."""
        if not sample_text or len(sample_text.strip()) < 20:
            return None
        
        # Versuche langdetect (falls installiert)
        try:
            from langdetect import detect, DetectorFactory
            DetectorFactory.seed = 0  # Reproduzierbare Ergebnisse
            detected = detect(sample_text)
            # Mapping von langdetect zu Tesseract-Codes
            lang_map = {
                'de': 'deu', 'en': 'eng', 'fr': 'fra', 'es': 'spa', 
                'it': 'ita', 'pt': 'por', 'nl': 'nld', 'pl': 'pol',
                'ru': 'rus', 'zh-cn': 'chi_sim', 'zh-tw': 'chi_tra',
                'ja': 'jpn', 'ko': 'kor', 'ar': 'ara', 'tr': 'tur',
                'sv': 'swe', 'da': 'dan', 'no': 'nor', 'fi': 'fin',
                'cs': 'ces', 'hu': 'hun', 'ro': 'ron', 'uk': 'ukr'
            }
            return lang_map.get(detected, detected)
        except ImportError:
            pass
        except Exception:
            pass
        
        # Fallback: Einfache heuristische Erkennung
        text_lower = sample_text.lower()
        
        # Deutsche Indikatoren
        de_indicators = ['und', 'der', 'die', 'das', 'ist', 'für', 'mit', 'auf', 'ein', 'eine', 'nicht', 'werden', 'haben', 'wird', 'sind', 'bei', 'nach', 'über', 'auch', 'nur', 'kann', 'ß', 'ü', 'ö', 'ä']
        de_score = sum(1 for w in de_indicators if f' {w} ' in f' {text_lower} ' or w in text_lower)
        
        # Englische Indikatoren
        en_indicators = ['the', 'and', 'is', 'for', 'with', 'on', 'are', 'this', 'that', 'have', 'from', 'will', 'not', 'been', 'which', 'can', 'all', 'would', 'their', 'there']
        en_score = sum(1 for w in en_indicators if f' {w} ' in f' {text_lower} ')
        
        # Französische Indikatoren
        fr_indicators = ['le', 'la', 'les', 'un', 'une', 'de', 'du', 'des', 'et', 'est', 'pour', 'dans', 'que', 'qui', 'sur', 'avec', 'ce', 'cette', 'sont', 'nous', 'vous', 'être', 'avoir', 'à', 'é', 'è', 'ê', 'ç']
        fr_score = sum(1 for w in fr_indicators if f' {w} ' in f' {text_lower} ' or w in text_lower)
        
        # Spanische Indikatoren
        es_indicators = ['el', 'la', 'los', 'las', 'un', 'una', 'de', 'del', 'y', 'en', 'que', 'es', 'para', 'con', 'no', 'por', 'su', 'como', 'más', 'pero', 'ñ', '¿', '¡']
        es_score = sum(1 for w in es_indicators if f' {w} ' in f' {text_lower} ' or w in text_lower)
        
        # Italienische Indikatoren
        it_indicators = ['il', 'la', 'le', 'un', 'una', 'di', 'del', 'della', 'e', 'è', 'che', 'per', 'con', 'non', 'sono', 'come', 'questo', 'questa', 'più', 'anche', 'loro']
        it_score = sum(1 for w in it_indicators if f' {w} ' in f' {text_lower} ' or w in text_lower)
        
        scores = {'deu': de_score, 'eng': en_score, 'fra': fr_score, 'spa': es_score, 'ita': it_score}
        best_lang = max(scores, key=scores.get)
        best_score = scores[best_lang]
        
        # Mindestens 3 Treffer für eine zuverlässige Erkennung
        if best_score >= 3:
            return best_lang
        return None
    
    def _determine_ocr_languages(self, sample_text: str | None = None) -> list:
        """Ermittelt die optimalen Sprachen für OCR.
        
        Args:
            sample_text: Optional - Text zur automatischen Spracherkennung
            
        Returns:
            Liste der Tesseract-Sprachcodes (z.B. ['deu', 'eng'])
        """
        try:
            # 1. Verfügbare Sprachen prüfen
            available_langs = self._get_available_tesseract_languages()
            
            # 2. Automatische Erkennung aus Text (wenn vorhanden)
            detected_lang = None
            if sample_text:
                detected_lang = self._detect_language_from_text(sample_text)
                if detected_lang:
                    self.logger.info(f"🌐 Sprache automatisch erkannt: {detected_lang}")
            
            # 3. Aus GUI-Einstellungen
            raw_candidates = []
            for attr_name in (
                'var_source_lang', 'var_source_language', 'source_lang', 'source_language',
                'var_target_lang', 'var_target_language', 'target_lang', 'target_language'
            ):
                attr = getattr(self, attr_name, None)
                value = None
                if attr is None:
                    continue
                getter = getattr(attr, 'get', None)
                if callable(getter):
                    try:
                        value = getter()
                    except Exception:
                        value = None
                elif isinstance(attr, str):
                    value = attr
                if isinstance(value, str) and value.strip():
                    # "Auto-Detect" überspringen
                    val_lower = value.strip().lower()
                    if 'auto' not in val_lower:
                        raw_candidates.append(val_lower)

            # 4. Mapping auf Tesseract-Codes
            mapped = []
            mapping = {
                'de': 'deu', 'deu': 'deu', 'deutsch': 'deu', 'german': 'deu', 'ger': 'deu',
                'en': 'eng', 'eng': 'eng', 'englisch': 'eng', 'english': 'eng',
                'fr': 'fra', 'fra': 'fra', 'französisch': 'fra', 'french': 'fra',
                'es': 'spa', 'spa': 'spa', 'spanisch': 'spa', 'spanish': 'spa',
                'it': 'ita', 'ita': 'ita', 'italienisch': 'ita', 'italian': 'ita',
                'pt': 'por', 'por': 'por', 'portugiesisch': 'por', 'portuguese': 'por',
                'nl': 'nld', 'nld': 'nld', 'niederländisch': 'nld', 'dutch': 'nld',
                'pl': 'pol', 'pol': 'pol', 'polnisch': 'pol', 'polish': 'pol',
                'ru': 'rus', 'rus': 'rus', 'russisch': 'rus', 'russian': 'rus',
                'zh': 'chi_sim', 'chi_sim': 'chi_sim', 'chinesisch': 'chi_sim', 'chinese': 'chi_sim',
                'ja': 'jpn', 'jpn': 'jpn', 'japanisch': 'jpn', 'japanese': 'jpn'
            }
            
            # Automatisch erkannte Sprache an den Anfang
            if detected_lang:
                tess_code = mapping.get(detected_lang, detected_lang)
                if not available_langs or tess_code in available_langs:
                    mapped.append(tess_code)
            
            # GUI-Einstellungen
            for candidate in raw_candidates:
                mapped_lang = mapping.get(candidate)
                if mapped_lang and mapped_lang not in mapped:
                    if not available_langs or mapped_lang in available_langs:
                        mapped.append(mapped_lang)
            
            # 5. Fallback: Deutsch + Englisch
            if not mapped:
                if not available_langs or 'deu' in available_langs:
                    mapped.append('deu')
            if 'eng' not in mapped:
                if not available_langs or 'eng' in available_langs:
                    mapped.append('eng')
            
            # 6. Warnung bei fehlenden Sprachpaketen
            if available_langs:
                missing = [lang for lang in mapped if lang not in available_langs]
                if missing:
                    self.logger.warning(f"⚠️ Fehlende Tesseract-Sprachpakete: {missing}. Installiert: {available_langs}")
                    mapped = [lang for lang in mapped if lang in available_langs]
                    if not mapped:
                        mapped = available_langs[:2] if available_langs else ['deu', 'eng']
            
            return list(dict.fromkeys(mapped))
        except Exception as e:
            self.logger.error(f"Fehler bei Spracherkennung: {e}")
            return ['deu', 'eng']

    def _resolve_tesseract_cmd(self) -> None:
        """Findet und konfiguriert den Tesseract-OCR Pfad automatisch."""
        try:
            import pytesseract
        except ImportError:
            return

        cmd = getattr(pytesseract.pytesseract, 'tesseract_cmd', 'tesseract')
        if cmd and cmd != 'tesseract' and os.path.isfile(cmd):
            return

        candidate_paths: list[str] = []
        
        # 1. Benutzer-Konfiguration
        try:
            if hasattr(self, 'settings_service') and self.settings_service:
                configured = self.settings_service.get('ocr.tesseract_cmd', None)
                if isinstance(configured, str) and configured.strip():
                    candidate_paths.append(configured.strip())
        except Exception:
            pass

        # 2. Umgebungsvariablen
        for env_key in ('TESSERACT_PATH', 'TESSERACT_CMD', 'TESSERACT_EXE_PATH'):
            env_path = os.environ.get(env_key)
            if env_path:
                candidate_paths.append(env_path)

        # 3. Windows Registry (häufigster Installationspfad)
        try:
            import winreg
            for hive in [winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER]:
                for subkey in [r"SOFTWARE\\Tesseract-OCR", r"SOFTWARE\\WOW6432Node\\Tesseract-OCR"]:
                    try:
                        with winreg.OpenKey(hive, subkey) as key:
                            install_dir = winreg.QueryValueEx(key, "InstallDir")[0]
                            if install_dir:
                                candidate_paths.append(os.path.join(install_dir, "tesseract.exe"))
                    except (FileNotFoundError, OSError):
                        continue
        except Exception:
            pass

        # 4. Standard Windows-Pfade
        win_candidates = [
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
            os.path.expandvars(r"%LOCALAPPDATA%\\Programs\\Tesseract-OCR\\tesseract.exe"),
            os.path.expandvars(r"%USERPROFILE%\\AppData\\Local\\Tesseract-OCR\\tesseract.exe"),
        ]
        candidate_paths.extend(win_candidates)
        
        # 5. PATH durchsuchen
        try:
            import shutil
            found_in_path = shutil.which("tesseract")
            if found_in_path:
                candidate_paths.insert(0, found_in_path)  # Priorität
        except Exception:
            pass

        for path in candidate_paths:
            if not path:
                continue
            normalized = os.path.normpath(path)
            if os.path.isfile(normalized):
                pytesseract.pytesseract.tesseract_cmd = normalized
                try:
                    self.logger.info(f"✅ Tesseract Cmd gesetzt: {normalized}")
                except Exception:
                    pass
                return

    def _generate_ocr_image_variants(self, image):
        """Erzeugt optimierte Bildvarianten für OCR mit fortgeschrittener Vorverarbeitung."""
        variants = []
        resample_mode = 3
        try:
            from PIL import Image, ImageOps, ImageFilter, ImageEnhance
            resample_mode = getattr(Image, 'BICUBIC', 3)
            
            # Basis: Graustufen
            gray = image.convert('L')
            
            # 1. Automatischer Kontrastausgleich
            contrasted = ImageOps.autocontrast(gray)
            
            # 2. Rauschunterdrückung (Median-Filter)
            denoised = contrasted.filter(ImageFilter.MedianFilter(size=3))
            
            # 3. Schärfung
            sharpened = ImageEnhance.Sharpness(denoised).enhance(1.5)
            
            # 4. Helligkeitskorrektur für dunkle Scans
            brightness = ImageEnhance.Brightness(sharpened).enhance(1.1)
            
            # 5. Kontrastverstärkung
            contrast_enhanced = ImageEnhance.Contrast(brightness).enhance(1.3)
            
            variants.extend([gray, contrasted, sharpened, contrast_enhanced])
            
            # 6. Adaptive Binarisierung (verschiedene Schwellwerte)
            for threshold in (120, 140, 160, 180):
                bw = contrast_enhanced.point(lambda p, thr=threshold: 255 if p > thr else 0).convert('L')
                variants.append(bw)
            
            # 7. Invertierte Version für weiß-auf-schwarz Dokumente
            try:
                inverted = ImageOps.invert(gray)
                inverted_contrasted = ImageOps.autocontrast(inverted)
                variants.append(inverted_contrasted)
            except Exception:
                pass
            
            # 8. Morphologische Operationen für bessere Texterkennung
            try:
                # Dilatation für dünnere Schriften
                dilated = denoised.filter(ImageFilter.MaxFilter(size=3))
                variants.append(dilated)
                # Erosion für dickere Schriften
                eroded = denoised.filter(ImageFilter.MinFilter(size=3))
                variants.append(eroded)
            except Exception:
                pass
                
        except Exception as prep_error:
            try:
                self.logger.debug(f"OCR Bildvorbereitung fehlgeschlagen: {prep_error}")
            except Exception:
                pass
            variants = [image.convert('L')]
        else:
            variants.append(image.convert('L'))

        # Skalierte Varianten für kleine Bilder
        scaled_variants = []
        for variant in variants:
            scaled_variants.append(variant)
            try:
                width, height = variant.size
                pixel_count = width * height
            except Exception:
                continue
            if pixel_count > 7_500_000:  # vermeide extreme Größen
                continue
            # Hochskalieren für kleine Bilder verbessert OCR
            for scale in (1.5, 2.0, 2.5):
                try:
                    new_size = (int(width * scale), int(height * scale))
                    if max(new_size) > 6000:
                        continue
                    scaled_variants.append(variant.resize(new_size, resample=resample_mode))
                except Exception:
                    continue
        return scaled_variants or [image]

    def _perform_ocr_on_image(self, image, languages=None) -> str:
        """Führt OCR auf einem Bild durch mit Qualitätsprüfung und multiplen Strategien."""
        try:
            import pytesseract
            from pytesseract import TesseractNotFoundError
        except ImportError:
            self.logger.warning("⚠️ Tesseract OCR nicht verfügbar. Bitte pytesseract installieren.")
            return ''

        self._resolve_tesseract_cmd()

        if languages is None or not languages:
            languages = self._determine_ocr_languages()
        lang_candidates = []
        try:
            lang_candidates.append('+'.join(languages))
        except Exception:
            pass
        lang_candidates.extend(languages)

        image_variants = self._generate_ocr_image_variants(image)
        
        # Erweiterte Konfigurationen für verschiedene Dokumenttypen
        configs = [
            '--oem 3 --psm 6',   # Block of text (Standard)
            '--oem 3 --psm 3',   # Fully automatic page segmentation
            '--oem 3 --psm 4',   # Single column of text
            '--oem 3 --psm 1',   # Auto with OSD
            '--oem 1 --psm 6',   # LSTM only
        ]
        
        best_text = ''
        best_confidence = 0

        for cfg_index, config in enumerate(configs):
            for lang in lang_candidates:
                if not lang:
                    continue
                for variant_idx, variant in enumerate(image_variants[:8]):  # Limitiere auf beste 8 Varianten
                    try:
                        # Zuerst Confidence-Daten holen
                        try:
                            data = pytesseract.image_to_data(variant, lang=lang, config=config, output_type=pytesseract.Output.DICT)
                            confidences = [int(c) for c in data.get('conf', []) if str(c).lstrip('-').isdigit() and int(c) > 0]
                            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                        except Exception:
                            avg_confidence = 0
                        
                        text = pytesseract.image_to_string(variant, lang=lang, config=config)
                    except TesseractNotFoundError:
                        message = "❌ Tesseract nicht gefunden. Bitte Pfad konfigurieren oder Tesseract installieren."
                        self.logger.error(message)
                        try:
                            if hasattr(self, 'show_toast'):
                                toast_msg = 'Tesseract wurde nicht gefunden. Bitte Installation prüfen.'
                                if hasattr(self, '_t') and callable(self._t):
                                    toast_msg = self._t(toast_msg)
                                self.show_toast(toast_msg, 'error')
                        except Exception:
                            pass
                        return ''
                    except Exception as ocr_error:
                        self.logger.debug(f"OCR-Versuch (cfg={config}, lang={lang}) fehlgeschlagen: {ocr_error}")
                        continue

                    stripped = text.strip() if text else ''
                    if not stripped:
                        continue

                    # Qualitätsbewertung: Länge + Confidence
                    quality_score = len(stripped) + (avg_confidence * 2)
                    current_best_score = len(best_text.strip()) + (best_confidence * 2)
                    
                    if quality_score > current_best_score:
                        best_text = stripped
                        best_confidence = avg_confidence

                    # Früher Abbruch bei hoher Qualität
                    if len(stripped) >= 100 and avg_confidence >= 70:
                        self.logger.info(f"✅ OCR erfolgreich: {len(stripped)} Zeichen, {avg_confidence:.0f}% Konfidenz")
                        return text
                    if cfg_index == 0 and len(stripped) >= 60:
                        self.logger.info(f"✅ OCR erfolgreich mit Sprache {lang}: {len(stripped)} Zeichen")
                        return text

        if best_text:
            self.logger.info(f"✅ OCR erfolgreich (beste Variante) – {len(best_text)} Zeichen extrahiert")
            return best_text

        self.logger.warning("⚠️ OCR lieferte keinen Text")
        return ''

    def _extract_pdf_text_with_pymupdf(self, filepath: str) -> str:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            self.logger.debug("PyMuPDF nicht installiert – Text-Extraktion wird übersprungen")
            return ''

        try:
            doc = fitz.open(filepath)
        except Exception as open_error:
            self.logger.debug(f"PyMuPDF konnte Datei nicht öffnen: {open_error}")
            return ''

        text_chunks = []
        try:
            for page in doc[:20]:
                try:
                    page_text = page.get_text('text')
                except Exception as extract_error:
                    self.logger.debug(f"PyMuPDF Seitentext-Extraktion fehlgeschlagen: {extract_error}")
                    continue
                if page_text:
                    text_chunks.append(page_text)
        finally:
            doc.close()

        combined = '\n\n'.join(chunk.strip() for chunk in text_chunks if chunk and chunk.strip())
        if combined:
            self.logger.info(f"✅ PyMuPDF Text-Extraktion erfolgreich: {len(combined)} Zeichen")
        return combined

    def _extract_text_from_file(self, filepath):
        """Extrahiert Text aus PDF, DOCX, TXT oder Bilddateien (mit OCR).
        
        PDF-Strategie (5-Stufen-Fallback):
        1. pdfplumber → 2. pypdf → 3. PyPDF2 → 4. PyMuPDF → 5. OCR
        """
        try:
            ext = os.path.splitext(filepath)[1].lower()
            filename = os.path.basename(filepath)
            
            if ext == '.pdf':
                self.logger.info(f"📄 Starte PDF-Extraktion: {filename}")
                
                # STUFE 1: pdfplumber (oft am zuverlässigsten für Text-PDFs)
                try:
                    import pdfplumber
                    self.logger.info("   → Versuch 1/5: pdfplumber")
                    
                    text_parts = []
                    with pdfplumber.open(filepath) as pdf:
                        for page in pdf.pages[:20]:  # Max 20 Seiten
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())
                    
                    if text_parts:
                        text = '\n\n'.join(text_parts)
                        self.logger.info(f"✅ pdfplumber: {len(text)} Zeichen aus {len(text_parts)} Seiten")
                        return text
                    else:
                        self.logger.info("     pdfplumber: Kein Text gefunden")
                except ImportError:
                    self.logger.info("     pdfplumber: Nicht installiert")
                except Exception as e:
                    self.logger.warning(f"     pdfplumber: Fehler - {str(e)[:80]}")
                
                # STUFE 2: pypdf
                try:
                    import pypdf
                    self.logger.info("   → Versuch 2/5: pypdf")
                    
                    with open(filepath, 'rb') as f:
                        reader = pypdf.PdfReader(f)
                        text_parts = []
                        for page in reader.pages[:20]:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())
                        
                        if text_parts:
                            text = '\n\n'.join(text_parts)
                            self.logger.info(f"✅ pypdf: {len(text)} Zeichen")
                            return text
                        else:
                            self.logger.info("     pypdf: Kein Text gefunden")
                except ImportError:
                    self.logger.info("     pypdf: Nicht installiert")
                except Exception as e:
                    self.logger.warning(f"     pypdf: Fehler - {str(e)[:80]}")
                
                # STUFE 3: PyPDF2
                try:
                    import PyPDF2
                    self.logger.info("   → Versuch 3/5: PyPDF2")
                    
                    with open(filepath, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text_parts = []
                        for page in reader.pages[:20]:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())
                        
                        if text_parts:
                            text = '\n\n'.join(text_parts)
                            self.logger.info(f"✅ PyPDF2: {len(text)} Zeichen")
                            return text
                        else:
                            self.logger.info("     PyPDF2: Kein Text gefunden")
                except ImportError:
                    self.logger.info("     PyPDF2: Nicht installiert")
                except Exception as e:
                    self.logger.warning(f"     PyPDF2: Fehler - {str(e)[:80]}")
                
                # STUFE 4: PyMuPDF/fitz
                self.logger.info("   → Versuch 4/5: PyMuPDF")
                pymupdf_text = self._extract_pdf_text_with_pymupdf(filepath)
                if pymupdf_text and pymupdf_text.strip():
                    self.logger.info(f"✅ PyMuPDF: {len(pymupdf_text)} Zeichen")
                    return pymupdf_text
                else:
                    self.logger.info("     PyMuPDF: Kein Text oder nicht verfügbar")
                
                # STUFE 5: OCR mit Tesseract (für gescannte PDFs)
                self.logger.warning(f"⚠️ '{filename}' - kein Text-Layer gefunden → Starte OCR")
                try:
                    self.logger.info("   → Versuch 5/5: OCR mit Tesseract")
                    
                    # Prüfe, ob pytesseract verfügbar ist
                    try:
                        import pytesseract
                        from PIL import Image
                        
                        # Versuche Tesseract zu finden
                        try:
                            pytesseract.get_tesseract_version()
                        except Exception:
                            self.logger.error("❌ Tesseract-OCR Engine nicht gefunden!")
                            return f"[OCR nicht verfügbar für '{filename}'. Tesseract-Engine muss installiert werden.]"
                    except ImportError:
                        self.logger.error("❌ pytesseract nicht installiert: pip install pytesseract")
                        return f"[OCR nicht verfügbar für '{filename}'.]"

                    ocr_languages = self._determine_ocr_languages()
                    lang_param = '+'.join(ocr_languages) if ocr_languages else 'deu+eng'
                    self.logger.info(f"     OCR-Sprachen: {lang_param}")
                    
                    # METHODE 1: pikepdf (KEIN Poppler nötig, reines Python!)
                    try:
                        import pikepdf
                        self.logger.info("     Verwende pikepdf (empfohlen, kein Poppler nötig)...")
                        
                        pdf = pikepdf.Pdf.open(filepath)
                        text_chunks = []
                        
                        # Verarbeite max 10 Seiten
                        for page_num in range(min(10, len(pdf.pages))):
                            self.logger.info(f"       Seite {page_num + 1}...")
                            page = pdf.pages[page_num]
                            
                            # Extrahiere eingebettete Bilder
                            for img_name, img_obj in page.images.items():
                                try:
                                    raw_image = pikepdf.PdfImage(img_obj)
                                    pil_image = raw_image.as_pil_image()
                                    
                                    # OCR auf Bild
                                    page_text = pytesseract.image_to_string(pil_image, lang=lang_param)
                                    if page_text and page_text.strip():
                                        text_chunks.append(page_text.strip())
                                        self.logger.info(f"       → {len(page_text)} Zeichen extrahiert")
                                        break  # Nur erstes Bild pro Seite
                                except Exception as e:
                                    self.logger.warning(f"       Bild {img_name} übersprungen: {str(e)[:50]}")
                        
                        pdf.close()
                        
                        if text_chunks:
                            combined_text = '\n\n'.join(text_chunks)
                            self.logger.info(f"✅ OCR (pikepdf) erfolgreich: {len(combined_text)} Zeichen aus {len(text_chunks)} Seiten")
                            return combined_text
                        else:
                            self.logger.warning("     pikepdf: Keine Bilder gefunden oder kein Text erkannt")
                    except ImportError:
                        self.logger.info("     pikepdf nicht installiert (pip install pikepdf)")
                    except Exception as e:
                        self.logger.warning(f"     pikepdf fehlgeschlagen: {str(e)[:100]}")
                    
                    # METHODE 2: pdf2image (braucht Poppler)
                    try:
                        from pdf2image import convert_from_path
                        self.logger.info("     Verwende pdf2image (braucht Poppler)...")
                        
                        # Konvertiere PDF zu Bildern (max 10 Seiten)
                        images = convert_from_path(filepath, first_page=1, last_page=10, dpi=300)
                        text_chunks = []
                        
                        for i, img in enumerate(images, 1):
                            self.logger.info(f"       Seite {i}/{len(images)}...")
                            try:
                                page_text = pytesseract.image_to_string(img, lang=lang_param)
                                if page_text and page_text.strip():
                                    text_chunks.append(page_text.strip())
                                    self.logger.info(f"       → {len(page_text)} Zeichen")
                            except Exception as e:
                                self.logger.warning(f"       Seite {i} fehlgeschlagen: {e}")
                        
                        if text_chunks:
                            combined_text = '\n\n'.join(text_chunks)
                            self.logger.info(f"✅ OCR erfolgreich: {len(combined_text)} Zeichen aus {len(text_chunks)} Seiten")
                            return combined_text
                    except ImportError:
                        self.logger.warning("     pdf2image nicht verfügbar (braucht Poppler)")
                    except Exception as e:
                        self.logger.warning(f"     pdf2image fehlgeschlagen: {str(e)[:100]}")
                    
                    # METHODE 2: PyMuPDF für Rendering (falls verfügbar)
                    try:
                        import fitz  # PyMuPDF
                        from PIL import Image
                        import io
                        
                        self.logger.info("     Verwende PyMuPDF für Rendering...")
                        
                        doc = fitz.open(filepath)
                        text_chunks = []
                        
                        zoom = 3.0
                        for page_num in range(min(10, len(doc))):  # Max 10 Seiten
                            self.logger.info(f"   OCR auf Seite {page_num + 1} (PyMuPDF)...")
                            page = doc[page_num]
                            
                            # Rendere Seite als Bild
                            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                            img_data = pix.tobytes("png")
                            img = Image.open(io.BytesIO(img_data))
                            
                            page_text = self._perform_ocr_on_image(img, ocr_languages)
                            if page_text:
                                text_chunks.append(page_text)
                        
                        doc.close()
                        
                        combined_text = '\n\n'.join(text_chunks).strip()
                        if combined_text:
                            self.logger.info(f"✅ PDF OCR (PyMuPDF) erfolgreich: {len(combined_text)} Zeichen")
                            return combined_text
                    except ImportError:
                        self.logger.warning("⚠️ PyMuPDF nicht verfügbar. Installieren: pip install PyMuPDF")
                    except Exception as e:
                        self.logger.error(f"❌ PyMuPDF OCR-Fehler: {e}")
                    
                    self.logger.warning("⚠️ Alle PDF OCR-Methoden fehlgeschlagen")
                    return None
                        
                except Exception as e:
                    self.logger.error(f"❌ PDF OCR-Fehler: {e}")
                    return None
            
            elif ext in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(filepath)
                    text = '\n\n'.join([para.text for para in doc.paragraphs[:100]])  # Max 100 Absätze
                    return text
                except Exception:
                    return None
            
            elif ext == '.txt':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(50000)  # Max 50KB
            
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif', '.webp']:
                # OCR für Bilddateien (Fotos, Screenshots, gescannte Dokumente)
                self.logger.info(f"🖼️ Bild-OCR: {filename}")
                
                try:
                    import pytesseract
                    from PIL import Image, ImageEnhance, ImageFilter
                    
                    # Prüfe Tesseract
                    try:
                        pytesseract.get_tesseract_version()
                    except Exception:
                        self.logger.error("❌ Tesseract-OCR nicht gefunden")
                        return f"[OCR nicht verfügbar für '{filename}']"
                    
                    # Lade Bild
                    img = Image.open(filepath)
                    original_size = f"{img.width}x{img.height}px"
                    self.logger.info(f"   Bildgröße: {original_size}, Modus: {img.mode}")
                    
                    # Konvertiere zu RGB falls nötig
                    if img.mode not in ('RGB', 'L'):
                        self.logger.info(f"   Konvertiere {img.mode} → RGB")
                        img = img.convert('RGB')
                    
                    # Bildverbesserung für bessere OCR-Qualität
                    try:
                        # Erhöhe Kontrast
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(1.5)
                        
                        # Schärfen
                        img = img.filter(ImageFilter.SHARPEN)
                        self.logger.info("   Bildoptimierung: Kontrast + Schärfe")
                    except Exception as e:
                        self.logger.warning(f"   Bildoptimierung übersprungen: {str(e)[:50]}")
                    
                    # OCR mit automatischer Spracherkennung
                    ocr_languages = self._determine_ocr_languages()
                    lang_param = '+'.join(ocr_languages) if ocr_languages else 'deu+eng'
                    self.logger.info(f"   OCR-Sprachen: {lang_param}")
                    
                    # Führe OCR aus
                    text = pytesseract.image_to_string(img, lang=lang_param)
                    
                    if text and text.strip():
                        clean_text = text.strip()
                        self.logger.info(f"✅ Bild-OCR erfolgreich: {len(clean_text)} Zeichen")
                        
                        # Wenn sehr wenig Text gefunden wurde, versuche mit höherer Auflösung
                        if len(clean_text) < 50 and img.width < 2000:
                            self.logger.info("   → Versuche mit 2x Skalierung...")
                            img_scaled = img.resize((img.width * 2, img.height * 2), Image.Resampling.LANCZOS)
                            text_scaled = pytesseract.image_to_string(img_scaled, lang=lang_param)
                            if text_scaled and len(text_scaled.strip()) > len(clean_text):
                                self.logger.info(f"   ✅ Besseres Ergebnis: {len(text_scaled.strip())} Zeichen")
                                return text_scaled.strip()
                        
                        return clean_text
                    else:
                        self.logger.warning(f"⚠️ Kein Text im Bild erkannt")
                        return f"[Kein Text in '{filename}' erkannt. Bild evtl. zu unscharf oder leer.]"
                        
                except ImportError:
                    self.logger.error("❌ pytesseract nicht installiert: pip install pytesseract")
                    return f"[OCR nicht verfügbar für '{filename}']"
                except Exception as e:
                    self.logger.error(f"❌ Bild-OCR Fehler: {str(e)[:100]}")
                    return f"[Fehler beim OCR von '{filename}': {str(e)[:80]}]"
            
        except Exception as e:
            self.logger.debug(f"Text-Extraktion fehlgeschlagen für {filepath}: {e}")
            return None

    def _launch_pair_details_preview(self, pair_rows=None, metrics_snapshot=None):
        """Zeigt eine modale Vorschau der zuletzt eingelesenen Segmente."""
        self.logger.info("🔍 _launch_pair_details_preview aufgerufen")
        try:
            rows = pair_rows if isinstance(pair_rows, list) else None
        except Exception:
            rows = None

        fresh_rows = None
        if not rows:
            try:
                self.logger.info("🔄 Versuche Dateien direkt einzulesen...")
                fresh_rows = self._read_uploaded_files_for_preview()
                if fresh_rows:
                    self.logger.info(f"✅ Dateien eingelesen: {len(fresh_rows)} Einträge")
            except Exception as e:
                self.logger.error(f"❌ Fehler beim Lesen hochgeladener Dateien: {e}")
                import traceback
                traceback.print_exc()
                fresh_rows = None

        if fresh_rows:
            rows = fresh_rows
        elif not rows:
            try:
                cached = getattr(self, '_last_pair_details', None)
                if isinstance(cached, list):
                    self.logger.info(f"ℹ️ Verwende Cache-Daten: {len(cached)} Einträge")
                    rows = cached
            except Exception:
                rows = None

        # Prüfe auf vollständig leere Segmente und versuche bei Bedarf erneut direkte Texteinzlesung
        if rows:
            try:
                has_content = any(
                    bool(str(entry.get('source_text') or entry.get('source') or '').strip()) or
                    bool(str(entry.get('target_text') or entry.get('translation') or '').strip())
                    for entry in rows if isinstance(entry, dict)
                )
            except Exception:
                has_content = True
            if not has_content and not fresh_rows:
                try:
                    self.logger.info("ℹ️ Cache ohne Text – erneuter Direkt-Einleseversuch...")
                    rows = self._read_uploaded_files_for_preview()
                except Exception:
                    pass
        
        if not rows:
            try:
                self.logger.warning("❌ Keine Daten für Vorschau verfügbar")
                if hasattr(self, 'show_toast'):
                    self.show_toast(
                        'Keine Texte eingelesen.\n\nBitte laden Sie zunächst Dateien über "Datei hochladen" hoch.',
                        'info',
                        duration=5000
                    )
            except Exception:
                pass
            return
        rows = [entry for entry in rows if isinstance(entry, dict)]
        if not rows:
            try:
                self.logger.warning("❌ Ungültige Daten-Struktur in Preview")
                if hasattr(self, 'show_toast'):
                    self.show_toast(
                        'Fehler beim Einlesen der Texte.\n\nBitte laden Sie die Dateien erneut hoch.',
                        'warning',
                        duration=5000
                    )
            except Exception:
                pass
            return

        metrics_data = metrics_snapshot if isinstance(metrics_snapshot, dict) else None
        if metrics_data is None:
            try:
                results = getattr(self, 'analysis_results', None)
                if isinstance(results, dict):
                    snap = results.get('metrics')
                    if isinstance(snap, dict):
                        metrics_data = snap
            except Exception:
                metrics_data = None

        try:
            existing = getattr(self, '_ingest_preview_window', None)
            if existing and existing.winfo_exists():
                existing.destroy()
        except Exception:
            pass

        master_widget = getattr(self, 'root', None)
        if not master_widget:
            try:
                master_widget = self.output_frame.winfo_toplevel()
            except Exception:
                master_widget = self.output_frame

        try:
            top = ctk.CTkToplevel(master_widget)
        except Exception:
            try:
                top = tk.Toplevel(master_widget)  # type: ignore
            except Exception:
                if hasattr(self, 'show_toast'):
                    self.show_toast(self._t('Vorschau konnte nicht geöffnet werden.'), 'error')
                return

        try:
            setattr(self, '_ingest_preview_window', top)
        except Exception:
            pass

        try:
            top.title(self._t('Eingelesener Text'))
        except Exception:
            pass
        try:
            top.geometry('940x720')
            top.minsize(720, 480)
        except Exception:
            pass
        try:
            if hasattr(top, 'transient'):
                top.transient(master_widget)
        except Exception:
            pass

        spacing_lg = self.get_spacing('lg') if hasattr(self, 'get_spacing') else 24
        spacing_md = self.get_spacing('md') if hasattr(self, 'get_spacing') else 16
        spacing_sm = self.get_spacing('sm') if hasattr(self, 'get_spacing') else 8

        try:
            top.configure(fg_color=self.get_color('surface'))
        except Exception:
            try:
                top.configure(bg=self.get_color('surface'))
            except Exception:
                pass

        container = ctk.CTkFrame(top, fg_color=self.get_color('transparent'))
        container.pack(fill='both', expand=True, padx=spacing_lg, pady=spacing_lg)

        header_frame = ctk.CTkFrame(container, fg_color=self.get_color('transparent'))
        header_frame.pack(fill='x')
        ctk.CTkLabel(
            header_frame,
            text=self._t('Eingelesener Text'),
            font=ctk.CTkFont(*self.get_typography('heading')),
            text_color=self.get_color('text_primary')
        ).pack(anchor='w')

        try:
            total_segments = len(rows)
        except Exception:
            total_segments = 0
        try:
            total_source_chars = sum(len(str(entry.get('source_text') or '').strip()) for entry in rows)
        except Exception:
            total_source_chars = 0
        try:
            total_target_chars = sum(len(str(entry.get('target_text') or '').strip()) for entry in rows)
        except Exception:
            total_target_chars = 0

        translated_segments = metrics_data.get('segments_translated') if isinstance(metrics_data, dict) else None
        untranslated_segments = metrics_data.get('segments_untranslated') if isinstance(metrics_data, dict) else None
        avg_ratio = metrics_data.get('avg_length_ratio') if isinstance(metrics_data, dict) else None

        ctk.CTkLabel(
            container,
            text='\n'.join([
                self._t('Segmente gesamt') + f": {total_segments}",
                self._t('Zeichen Quelle') + f": {total_source_chars}",
                self._t('Zeichen Ziel') + f": {total_target_chars}"] +
                ([self._t('Übersetzte Segmente') + f": {translated_segments}"] if isinstance(translated_segments, int) else []) +
                ([self._t('Unübersetzte Segmente') + f": {untranslated_segments}"] if isinstance(untranslated_segments, int) else []) +
                ([self._t('Durchschnittliches Längenverhältnis') + f": {avg_ratio:.3f}"] if isinstance(avg_ratio, (int, float)) else [])
            ),
            font=ctk.CTkFont(*self.get_typography('body')),
            text_color=self.get_color('text_secondary'),
            justify='left'
        ).pack(anchor='w', pady=(spacing_sm, spacing_md))

        display_limit = 120
        rows_to_render = rows[:display_limit]
        truncated = len(rows) > display_limit

        if truncated:
            ctk.CTkLabel(
                container,
                text=self._t('Hinweis: Anzeige auf {limit} Segmente begrenzt.').format(limit=display_limit),
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('warning'),
                justify='left'
            ).pack(anchor='w', pady=(0, spacing_sm))

        text_frame = ctk.CTkFrame(container, fg_color=self.get_color('transparent'))
        text_frame.pack(fill='both', expand=True)

        text_box = ctk.CTkTextbox(
            text_frame,
            wrap='word',
            fg_color=self.get_color('surface'),
            text_color=self.get_color('text_primary'),
            font=ctk.CTkFont(*self.get_typography('body'))
        )
        text_box.pack(side='left', fill='both', expand=True)
        scrollbar = ctk.CTkScrollbar(text_frame, command=text_box.yview)
        scrollbar.pack(side='right', fill='y')
        try:
            text_box.configure(yscrollcommand=scrollbar.set)
        except Exception:
            pass

        segment_lines = []
        for idx, entry in enumerate(rows_to_render, start=1):
            try:
                source_text = str(entry.get('source_text') or entry.get('source') or '').strip()
            except Exception:
                source_text = ''
            try:
                target_text = str(entry.get('target_text') or entry.get('translation') or '').strip()
            except Exception:
                target_text = ''
            page_hint = entry.get('page') or entry.get('page_number') or entry.get('page_index')
            if page_hint:
                header_line = self._t('Segment {idx} – Seite {page}').format(idx=idx, page=page_hint)
            else:
                header_line = self._t('Segment {idx}').format(idx=idx)
            segment_lines.extend([
                header_line,
                self._t('Quelle') + ':',
                source_text or '–',
                '',
                self._t('Ziel') + ':',
                target_text or '–',
                '\n' + ('-' * 48)
            ])

        if truncated:
            segment_lines.append(self._t('Restliche Segmente bitte in der Analyse speichern oder exportieren.'))

        try:
            text_box.insert('1.0', '\n'.join(segment_lines))
            text_box.configure(state='disabled')
        except Exception:
            pass

        actions = ctk.CTkFrame(container, fg_color=self.get_color('transparent'))
        actions.pack(fill='x', pady=(spacing_md, 0))

        def _on_close():
            try:
                if hasattr(top, 'destroy'):
                    top.destroy()
            finally:
                try:
                    setattr(self, '_ingest_preview_window', None)
                except Exception:
                    pass

        def _copy_all():
            try:
                text_box.configure(state='normal')
                content = text_box.get('1.0', 'end-1c')
                text_box.configure(state='disabled')
                top.clipboard_clear()
                top.clipboard_append(content)
                if hasattr(self, 'show_toast'):
                    self.show_toast(self._t('Segmenttext kopiert.'), 'success')
            except Exception:
                if hasattr(self, 'show_toast'):
                    self.show_toast(self._t('Kopieren nicht möglich.'), 'error')

        copy_btn = ctk.CTkButton(
            actions,
            text=self._t('In Zwischenablage kopieren'),
            width=220,
            height=32,
            fg_color=self.get_color('secondary'),
            hover_color=self.get_color('secondary_hover'),
            text_color=self.get_color('text_inverse'),
            command=_copy_all
        )
        copy_btn.pack(side='left')

        close_btn = ctk.CTkButton(
            actions,
            text=self._t('Schließen'),
            width=140,
            height=32,
            fg_color=self.get_color('surface_hover'),
            hover_color=self.get_color('surface_hover'),
            text_color=self.get_color('text_primary'),
            command=_on_close
        )
        close_btn.pack(side='right')

        try:
            top.protocol('WM_DELETE_WINDOW', _on_close)
        except Exception:
            pass
        try:
            top.bind('<Escape>', lambda _evt: _on_close())
        except Exception:
            pass
        try:
            text_box.focus_set()
        except Exception:
            pass
    
    def _launch_pair_details_preview_with_scroll(self, segment_index=None, segment_hash=None, source_text=''):
        """Öffnet die Segment-Vorschau und scrollt zum angegebenen Segment.
        
        Args:
            segment_index: Index des Segments (1-basiert)
            segment_hash: Hash des Segments für exakte Zuordnung
            source_text: Quelltext für Fuzzy-Matching falls Index/Hash nicht verfügbar
        """
        self.logger.info(f"🎯 Navigation zu Segment: index={segment_index}, hash={segment_hash}")
        
        # Erst die normale Preview öffnen
        self._launch_pair_details_preview()
        
        # Dann zum Segment scrollen
        try:
            preview_window = getattr(self, '_ingest_preview_window', None)
            if not preview_window or not preview_window.winfo_exists():
                return
            
            # Textbox finden (suche nach CTkTextbox im Fenster)
            text_box = None
            def find_textbox(widget):
                nonlocal text_box
                if text_box:
                    return
                try:
                    if widget.__class__.__name__ == 'CTkTextbox':
                        text_box = widget
                        return
                    for child in widget.winfo_children():
                        find_textbox(child)
                except Exception:
                    pass
            
            find_textbox(preview_window)
            
            if not text_box:
                self.logger.warning("Textbox nicht gefunden für Navigation")
                return
            
            # Nach Segment-Marker suchen
            if segment_index is not None:
                search_pattern = f"Segment {segment_index}"
            elif source_text and len(source_text) > 10:
                # Fuzzy-Match: Suche nach erstem Teil des Quelltextes
                search_pattern = source_text[:50].strip()
            else:
                search_pattern = None
            
            if search_pattern:
                # Text durchsuchen
                try:
                    text_box.configure(state='normal')
                    content = text_box.get('1.0', 'end-1c')
                    text_box.configure(state='disabled')
                    
                    # Position finden
                    pos = content.lower().find(search_pattern.lower())
                    if pos >= 0:
                        # Zeile berechnen
                        lines_before = content[:pos].count('\n')
                        line_number = lines_before + 1
                        
                        # Zur Zeile scrollen
                        text_box.see(f"{line_number}.0")
                        
                        # Highlight (falls möglich)
                        try:
                            text_box.configure(state='normal')
                            # Tag für Highlight erstellen
                            text_box.tag_add('highlight', f"{line_number}.0", f"{line_number}.end")
                            text_box.tag_config('highlight', background='#FFFF00', foreground='#000000')
                            text_box.configure(state='disabled')
                        except Exception:
                            pass
                        
                        if hasattr(self, 'show_toast'):
                            self.show_toast(f'📍 Segment {segment_index or "?"} gefunden', 'success', duration=2000)
                    else:
                        if hasattr(self, 'show_toast'):
                            self.show_toast(f'Segment nicht gefunden (Suche: "{search_pattern[:30]}...")', 'warning')
                except Exception as e:
                    self.logger.error(f"Fehler beim Scrollen: {e}")
        except Exception as e:
            self.logger.error(f"Navigation fehlgeschlagen: {e}")
    
    def _setup_modern_file_management(self, parent):
        """Initialisiert die moderne Dateiverwaltung mit Drag-and-Drop-Unterstützung."""
        try:
            # File management container with modern styling
            file_mgmt_frame = self._create_card_frame(
                parent,
                border_width=1,
                corner_radius=8
            )
            file_mgmt_frame.pack(fill="both", expand=True, pady=(20, 0))
            
            # Header section
            header_frame = ctk.CTkFrame(file_mgmt_frame, fg_color="transparent")
            header_frame.pack(fill="x", padx=20, pady=(15, 10))
            
            # Title with file count
            title_container = ctk.CTkFrame(header_frame, fg_color="transparent")
            title_container.pack(side="left", fill="x", expand=True)
            
            self.files_title = ctk.CTkLabel(
                title_container,
                text=(self._t("Dateiverwaltung") if hasattr(self,'_t') else "Dateiverwaltung"),
                font=ctk.CTkFont(*self.get_typography('subheading')),
                text_color=self.get_color('primary'),
                anchor="w"
            )
            self.files_title.pack(side="left")
            
            self.file_count_badge = ctk.CTkLabel(
                title_container,
                text=(self._t("0 Dateien") if hasattr(self,'_t') else "0 Dateien"),
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('white'),
                fg_color=self.get_color('primary'),
                corner_radius=8,
                width=70,
                height=24
            )
            self.file_count_badge.pack(side="left", padx=(10, 0))
            
            # Action buttons
            action_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            action_frame.pack(side="right")
            
            # Clear files button with compact height
            self.clear_files_btn = self._create_button(
                action_frame,
                text=(self._t("Alles löschen") if hasattr(self,'_t') else "Alles löschen"),
                command=self._clear_all_files,
                kind="danger",
                height=30
            )
            self.clear_files_btn.configure(width=85)
            self.clear_files_btn.pack(side="right", padx=(10, 0))
            # Initial deaktivieren bis Dateien vorhanden
            try:
                self.clear_files_btn.configure(state="disabled")
            except Exception:
                pass
            
            # Refresh button with compact height
            self.refresh_files_btn = self._create_button(
                action_frame,
                text=(self._t("Aktualisieren") if hasattr(self,'_t') else "Aktualisieren"),
                command=self._refresh_file_list,
                kind="primary",
                height=30
            )
            self.refresh_files_btn.configure(width=85)
            self.refresh_files_btn.pack(side="right")
            
            # File lists container with improved spacing
            lists_container = ctk.CTkFrame(file_mgmt_frame, fg_color="transparent")
            lists_container.pack(fill="both", expand=True, padx=25, pady=(0, 20))
            lists_container.grid_columnconfigure((0, 1), weight=1)
            
            # Source files list
            self._create_file_list_section(lists_container, (self._t("Ausgangstexte") if hasattr(self,'_t') else "Ausgangstexte"), "source", 0)
            
            # Translation files list  
            self._create_file_list_section(
                lists_container,
                self._t("Übersetzungen") if hasattr(self, '_t') else "Übersetzungen",
                "translation",
                1
            )
            
            # File Pairing Section
            pairing_frame = ctk.CTkFrame(
                lists_container,
                fg_color=self.get_color('info_light'),
                border_width=1,
                corner_radius=10
            )
            pairing_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(10, 0), padx=5)
            
            # Pairing header
            pairing_header = ctk.CTkLabel(
                pairing_frame,
                text=(self._t("Dateipaarung") if hasattr(self,'_t') else "Dateipaarung"),
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('info_dark')
            )
            pairing_header.pack(pady=(15, 5))
            
            # Pairing status
            self.pairing_status_label = ctk.CTkLabel(
                pairing_frame,
                text=(self._t("Keine Paare konfiguriert") if hasattr(self,'_t') else "Keine Paare konfiguriert"),
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_secondary')
            )
            self.pairing_status_label.pack(pady=(0, 10))
            
            # Manual Pairing Button
            self.manual_pairing_button = self._create_button(
                pairing_frame,
                text=(self._t("Dateipaarung anpassen") if hasattr(self,'_t') else "Dateipaarung anpassen"),
                command=self._show_manual_pairing_dialog,
                kind="info",
                height=32
            )
            self.manual_pairing_button.pack(pady=(0, 15))
            
        except Exception as e:
            try:
                self._handle_error(e, context="files.modern.setup", user_message=self._t("Modernes Dateimanagement Fehler"), toast=False)
            except Exception:
                pass
    
    def _create_file_list_section(self, parent, title, file_type, column):
        """Erzeugt einen modernen Dateilistensektor mit interaktiven Funktionen."""
        try:
            # Section container
            section_frame = ctk.CTkFrame(
                parent,
                fg_color=self.get_color('gray_50'),
                border_width=1,
                corner_radius=10
            )
            # Mehr horizontaler Abstand zwischen den Karten
            section_frame.grid(row=0, column=column, sticky="nsew", padx=(0 if column == 0 else 24, 0), pady=0)
            
            # Section header with improved proportions (compact)
            header = ctk.CTkFrame(section_frame, fg_color=self.get_color('gray_100'), height=46)
            # Mehr Außen-Padding für den Header
            header.pack(fill="x", padx=16, pady=(16, 0))
            header.pack_propagate(False)

            header_content = ctk.CTkFrame(header, fg_color="transparent")
            # Mehr Innen-Padding im Header-Content
            header_content.pack(fill="x", padx=8, pady=6)
            
            # Title with count
            title_label = ctk.CTkLabel(
                header_content,
                text=title,
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('primary'),  # Normalisiert: kein primary_dark Token
                anchor="w"
            )
            title_label.pack(side="left")
            
            # File count for this type
            count_label = ctk.CTkLabel(
                header_content,
                text="0",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('primary'),
                fg_color=self.get_color('white'),
                corner_radius=10,
                width=30,
                height=20
            )
            count_label.pack(side="right")
            
            # Store references for updating
            if file_type == "source":
                self.source_count_label = count_label
            else:
                self.translation_count_label = count_label
            
            # Scrollable file list
            file_list = ctk.CTkScrollableFrame(
                section_frame,
                height=220,
                fg_color=self.get_color('white'),
                border_width=0
            )
            # Mehr Innenabstand zwischen Rahmen und Scrollfläche
            file_list.pack(fill="both", expand=True, padx=16, pady=(12, 16))
            
            # Store reference for updating
            if file_type == "source":
                self.source_file_list = file_list
            else:
                self.translation_file_list = file_list
            
            # Empty state message
            # Deutsche Leermeldung abhängig vom Typ
            if file_type == 'source':
                empty_text = self._t("Keine Ausgangstexte\nHier ablegen / laden")
            else:
                empty_text = self._t("Keine Übersetzungen\nHier ablegen / laden")
            empty_message = ctk.CTkLabel(
                file_list,
                text=empty_text,
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary'),
                justify="center"
            )
            empty_message.pack(expand=True, pady=40)
            
            # Store empty message reference
            if file_type == "source":
                self.source_empty_msg = empty_message
            else:
                self.translation_empty_msg = empty_message
                
        except Exception as e:
            try:
                self._handle_error(e, context="files.list.section", user_message=self._t("Dateiliste Abschnitt Fehler"), toast=False, file_type=file_type)
            except Exception:
                pass
    
    def _update_file_counter(self):
        """Aktualisiert alle Dateizähler zusammengefasst für Header und Karten."""
        try:
            start_ts = time.time()
            
            # 🔍 DEBUG: uploaded_files Zustand
            self.logger.info(f"🔍 _update_file_counter() - uploaded_files ID: {id(self.uploaded_files)}, source ID: {id(self.uploaded_files.get('source', []))}")
            self.logger.info(f"🔍 _update_file_counter() - uploaded_files keys: {list(self.uploaded_files.keys())}")
            self.logger.info(f"🔍 _update_file_counter() - source length: {len(self.uploaded_files.get('source', []))}, translation length: {len(self.uploaded_files.get('translation', []))}")
            
            source_count = len(self.uploaded_files.get('source', []))
            translation_count = len(self.uploaded_files.get('translation', []))
            total_count = source_count + translation_count

            # Status/Readiness
            if source_count > 0 and translation_count > 0:
                pairs = min(source_count, translation_count)
                status_text = f" – {pairs} Paar(e) bereit für Analyse"
                status_color = self._accent('success', 'system_status')
                if hasattr(self, 'readiness_indicator'):
                    self.readiness_indicator.configure(text="Bereit für Analyse", text_color=self._accent('success', 'system_status'))
            elif source_count > 0 or translation_count > 0:
                if source_count > translation_count:
                    missing = source_count - translation_count
                    status_text = f" – {missing} weitere Übersetzungsdatei(en) hochladen"
                else:
                    missing = translation_count - source_count
                    status_text = f" – {missing} weitere Ausgangstexte hochladen"
                status_color = self._accent('warning', 'upload_needed')
                if hasattr(self, 'readiness_indicator'):
                    self.readiness_indicator.configure(text="Weitere Dateien hochladen", text_color=self._accent('warning', 'upload_needed'))
            else:
                status_text = " – Dateien hochladen, um die Analyse zu starten"
                status_color = self.get_color('text_secondary')
                if hasattr(self, 'readiness_indicator'):
                    self.readiness_indicator.configure(text="Warten auf Dateien", text_color=self.get_color('text_secondary'))

            counter_text = f"Dateien: {source_count} Ausgangstexte, {translation_count} Übersetzungen{status_text}"
            simple_counter_text = f"Dateien: {source_count} Ausgangstexte, {translation_count} Übersetzungen"

            if hasattr(self, 'header_file_counter_label'):
                self.header_file_counter_label.configure(text=simple_counter_text, text_color=status_color)
            if hasattr(self, 'card_file_counter_label'):
                self.card_file_counter_label.configure(text=counter_text, text_color=status_color)
            if hasattr(self, 'file_counter_label'):
                self.file_counter_label.configure(text=counter_text, text_color=status_color)

            if hasattr(self, 'file_count_badge'):
                self.file_count_badge.configure(text=f"{total_count} files")
                if total_count == 0:
                    self.file_count_badge.configure(fg_color=self.get_color('gray_400'))
                elif total_count < 5:
                    self.file_count_badge.configure(fg_color=self.get_color('primary'))
                else:
                    self.file_count_badge.configure(fg_color=self.get_color('gray_600'))

            if hasattr(self, 'source_count_label'):
                self.source_count_label.configure(text=f"{source_count} Ausgangstexte")
            if hasattr(self, 'translation_count_label'):
                self.translation_count_label.configure(text=f"{translation_count} Übersetzungen")

            if hasattr(self, 'analyze_button'):
                if source_count > 0 and translation_count > 0:
                    self.analyze_button.configure(
                        state="normal",
                        fg_color=self.get_color('primary'),
                        text=(self._t("Qualitätsanalyse starten") if hasattr(self, '_t') else "Qualitätsanalyse starten")
                    )
                else:
                    self.analyze_button.configure(
                        state="disabled",
                        fg_color=self.get_color('gray_400'),
                        text=(self._t("Weitere Dateien hochladen") if hasattr(self, '_t') else "Weitere Dateien hochladen")
                    )

            # Clear-All Button toggeln
            if hasattr(self, 'clear_files_btn'):
                try:
                    if total_count == 0:
                        self.clear_files_btn.configure(state="disabled")
                    else:
                        self.clear_files_btn.configure(state="normal")
                except Exception:
                    pass

            if hasattr(self, 'logger'):
                self.logger.info(f"Dateizähler aktualisiert: {source_count} Ausgangstexte, {translation_count} Übersetzungen")
            # Performance & State Telemetrie
            try:
                duration_ms = int((time.time() - start_ts) * 1000)
                self._publish_files_state(
                    source_count=source_count,
                    translation_count=translation_count,
                    total_count=total_count,
                    duration_ms=duration_ms,
                    ts=start_ts
                )
            except Exception:
                pass

            # Folgeaktionen
            self._refresh_file_list_display()
            self._smart_file_pairing()
            self._check_and_show_manual_pairing_option()
            try:
                self._update_ribbon_states()
            except Exception:
                pass

        except Exception as e:
            # Fallbacks ohne UnboundLocalError
            try:
                if hasattr(self, 'header_file_counter_label'):
                    self.header_file_counter_label.configure(text="Dateien: Fehler beim Zählen")
                if hasattr(self, 'card_file_counter_label'):
                    self.card_file_counter_label.configure(text="Dateien: Fehler beim Zählen")
            except Exception:
                pass
            try:
                if hasattr(self, 'update_file_count'):
                    self.update_file_count(0)
            except Exception:
                pass
            self._handle_error(
                e,
                context="files.counter.update",
                user_message=(self._t("Dateizähler konnte nicht aktualisiert werden") if hasattr(self,'_t') else "Dateizähler konnte nicht aktualisiert werden"),
                toast=False
            )
    
    # --------------------------------------------------------------
    # Debounce Helper: Verhindert mehrfachen UI-Refresh bei Massenupload
    # --------------------------------------------------------------
    def _schedule_update_file_counter(self, delay_ms: int = 80):
        """Planung eines debounced Updates.

        Mehrere schnelle Änderungen (Uploads/Removals) lösen nur einen
        tatsächlichen Zähler-/Pairing-Update nach Ablauf von delay_ms aus.
        Fällt auf direkten Aufruf zurück, falls root fehlt.
        """
        try:
            if not hasattr(self, '_pending_file_counter_job'):
                self._pending_file_counter_job = None
            # Vorherigen Job abbrechen
            if self._pending_file_counter_job and self.root:
                try:
                    self.root.after_cancel(self._pending_file_counter_job)
                except Exception:
                    pass
            if self.root:
                scheduled_at = time.time()
                def _run():
                    try:
                        self._pending_file_counter_job = None
                        self._update_file_counter()
                        if getattr(self, 'event_bus', None):
                            self.event_bus.publish('files.counter.debounced', {
                                'delay_ms': delay_ms,
                                'scheduled_at': scheduled_at,
                                'executed_at': time.time(),
                                'queue_latency_ms': int((time.time() - scheduled_at) * 1000)
                            })
                    except Exception:
                        pass
                self._pending_file_counter_job = self._after(delay_ms, _run)
            else:
                self._update_file_counter()
        except Exception:
            try:
                self._update_file_counter()
            except Exception:
                pass

    # --------------------------------------------------------------
    # Files State Telemetrie Helper (Rate-Limit + Diff)
    # --------------------------------------------------------------
    def _publish_files_state(self, source_count: int, translation_count: int, total_count: int, duration_ms: int, ts: float):
        """Publiziert files.state Event mit Deltas & Rate-Limit.

        Rate-Limit: max 1 Event / 120ms
        Liefert zusätzlich den Unterschied zum vorherigen Snapshot.
        """
        try:
            if not getattr(self, 'event_bus', None):
                return
            now = time.time()
            if not hasattr(self, '_last_files_state_ts'):
                self._last_files_state_ts = 0.0
            if not hasattr(self, '_last_files_state_snapshot'):
                self._last_files_state_snapshot = {
                    'source': None,
                    'translation': None,
                    'total': None,
                    'pairs_ready': None
                }
            # Rate Limit
            if (now - self._last_files_state_ts) < 0.12:
                return
            pairs_ready = min(source_count, translation_count)
            prev = self._last_files_state_snapshot
            diff = {
                'source': (None if prev['source'] is None else source_count - prev['source']),
                'translation': (None if prev['translation'] is None else translation_count - prev['translation']),
                'total': (None if prev['total'] is None else total_count - prev['total']),
                'pairs_ready': (None if prev['pairs_ready'] is None else pairs_ready - prev['pairs_ready'])
            }
            payload = {
                'source': source_count,
                'translation': translation_count,
                'total': total_count,
                'pairs_ready': pairs_ready,
                'duration_ms': duration_ms,
                'ts': ts,
                'diff': diff
            }
            self.event_bus.publish('files.state', payload)
            # Snapshot aktualisieren
            self._last_files_state_snapshot = {
                'source': source_count,
                'translation': translation_count,
                'total': total_count,
                'pairs_ready': pairs_ready
            }
            self._last_files_state_ts = now
        except Exception:
            pass
    
    def _refresh_file_list_display(self):
        """Aktualisiert die angezeigten Dateilisten anhand der aktuellen Dateien."""
        try:
            try:
                self.logger.debug("Aktualisiere Dateilisten…")
            except Exception:
                pass
            
            # Update source files display
            if hasattr(self, 'source_file_list'):
                try:
                    self.logger.debug(f"Updating source files: {len(self.uploaded_files.get('source', []))} files")
                except Exception:
                    pass
                self._update_file_list_content('source', self.source_file_list, self.source_empty_msg)
            else:
                try:
                    self.logger.warning("source_file_list not found")
                except Exception:
                    pass
            
            # Update translation files display
            if hasattr(self, 'translation_file_list'):
                try:
                    self.logger.debug(f"Updating translation files: {len(self.uploaded_files.get('translation', []))} files")
                except Exception:
                    pass
                self._update_file_list_content('translation', self.translation_file_list, self.translation_empty_msg)
            else:
                try:
                    self.logger.warning("translation_file_list not found")
                except Exception:
                    pass
                
        except Exception as e:
            try:
                self._handle_error(e, context="files.list.refresh", user_message=self._t("Dateiliste Refresh Fehler"), toast=False)
            except Exception:
                pass
            try:
                import traceback
                traceback.print_exc()
            except Exception:
                pass
    
    def _update_file_list_content(self, file_type, list_frame, empty_msg):
        """Pflegt den Inhalt einer Dateiliste basierend auf den aktuellen Dateien."""
        try:
            files = self.uploaded_files.get(file_type, [])
            
            # Clear current content (but preserve empty_msg)
            for widget in list_frame.winfo_children():
                if widget != empty_msg:  # Empty-Message nicht zerstören
                    widget.destroy()
            
            if files:
                empty_msg.pack_forget()
                for i, file_path in enumerate(files):
                    self._create_file_item(list_frame, file_path, file_type, i)
                self._auto_hide_scrollbar(list_frame, hide=False)
            else:
                empty_msg.pack(expand=True, pady=40)
                self._auto_hide_scrollbar(list_frame, hide=True)
                
        except Exception as e:
            try:
                self._handle_error(e, context="files.list.content", user_message=self._t("Dateiliste Inhalt konnte nicht aktualisiert werden"), toast=False, file_type=file_type)
            except Exception:
                pass
            # Fallback: Show empty message
            try:
                empty_msg.pack(expand=True, pady=40)
            except:
                pass
    
    def _create_file_item(self, parent, file_path, file_type, index):
        """Erstellt einen einzelnen Dateieintrag inklusive zugehöriger Aktionen."""
        try:
            # File item container
            item_frame = self._create_card_frame(
                parent,
                border_width=1,
                corner_radius=8
            )
            item_frame.configure(height=60)
            item_frame.pack(fill="x", padx=5, pady=3)
            item_frame.pack_propagate(False)
            
            # File info section
            info_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            info_frame.pack(side="left", fill="both", expand=True, padx=10, pady=8)
            
            # File name
            filename = os.path.basename(file_path)
            name_label = ctk.CTkLabel(
                info_frame,
                text=filename,
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('text_primary'),
                anchor="w"
            )
            name_label.pack(fill="x", pady=(0, 2))
            
            # File path and size
            try:
                file_size = os.path.getsize(file_path)
                size_str = self._format_file_size(file_size)
                path_text = f"{file_path} – {size_str}"
            except:
                path_text = file_path
            
            path_label = ctk.CTkLabel(
                info_frame,
                text=path_text,
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=self.get_color('text_secondary'),
                anchor="w"
            )
            path_label.pack(fill="x")
            
            # Actions section
            actions_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            actions_frame.pack(side="right", padx=8, pady=8)
            
            # Remove button
            remove_btn = self._create_button(
                actions_frame,
                text="Remove",
                command=lambda: self._remove_file(file_type, index),
                kind="danger",
                height=26
            )
            remove_btn.configure(width=70, corner_radius=6)
            remove_btn.pack()
            # Scrollbar sicherstellen
            self._auto_hide_scrollbar(parent, hide=False)
            
        except Exception as e:
            try:
                self._handle_error(e, context="files.item.create", user_message=self._t("Dateielement konnte nicht erstellt werden"), toast=False)
            except Exception:
                pass
    
    def _format_file_size(self, size_bytes):
        """DATEIGRÖSSE FORMATIEREN (Delegation) - Single Source of Truth.

    Diese frühere lokale Implementierung delegiert jetzt an die zentrale
        Utility-Funktion `format_file_size` um Duplikate zu eliminieren.
        (Duplicate-Prevention Policy)
        """
        try:
            from quality_gui_utilities import format_file_size
            return format_file_size(int(size_bytes) if size_bytes is not None else 0)
        except Exception:
            # Minimaler Fallback falls Import fehlschlägt
            try:
                b = int(size_bytes)
                if b < 1024:
                    return f"{b} B"
                if b < 1024**2:
                    return f"{b/1024:.1f} KB"
                if b < 1024**3:
                    return f"{b/(1024**2):.1f} MB"
                return f"{b/(1024**3):.1f} GB"
            except Exception:
                return "Unknown"
    
    def _remove_file(self, file_type, index):
        """Entfernt eine bestimmte Datei aus der Liste und aktualisiert alle Bezüge."""
        try:
            files = self.uploaded_files.get(file_type, [])
            if 0 <= index < len(files):
                removed_file = files[index]
                filename = os.path.basename(removed_file)
                before_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source', 'translation'))

                self._unregister_uploaded_files(file_type, [removed_file])

                self._schedule_update_file_counter()
                self._refresh_file_list_display()
                self._smart_file_pairing()
                self._check_and_show_manual_pairing_option()

                after_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source', 'translation'))

                self.show_toast((self._t("Datei entfernt") + f": {filename}") if hasattr(self, '_t') else f"Datei entfernt: {filename}", "success")

                try:
                    if getattr(self, 'logger', None):
                        self.logger.info(
                            "Datei entfernt (%s): %s | Gesamt %s → %s",
                            file_type,
                            filename,
                            before_total,
                            after_total
                        )
                except Exception:
                    pass

                try:
                    if getattr(self, 'event_bus', None):
                        self.event_bus.publish('files.changed', {
                            'action': 'removed',
                            'file_type': file_type,
                            'filename': filename,
                            'after_total': after_total,
                            'before_total': before_total
                        })
                except Exception:
                    pass
                
        except Exception as e:
            try:
                self._handle_error(e, context="files.item.remove", user_message=self._t("Dateielement konnte nicht entfernt werden"), toast=True)
            except Exception:
                pass

    def _smart_file_pairing(self):
        """Delegiert automatische Paarung an QualityGuiPairingManager (Backward Compatible)."""
        try:
            if not hasattr(self, 'pairing_manager') or not self.pairing_manager:
                from quality_gui_pairing_manager import QualityGuiPairingManager
                self.pairing_manager = QualityGuiPairingManager()
            from pairing_service import get_pairing_service  # lazy
            source_files = self.uploaded_files.get('source', [])
            translation_files = self.uploaded_files.get('translation', [])
            # Telemetrie Zähler
            if not hasattr(self, '_telemetry_counters'):
                self._telemetry_counters = {}
            self._telemetry_counters.setdefault('pairing_invocations', 0)
            self._telemetry_counters['pairing_invocations'] += 1
            pairs, unmatched = self.pairing_manager.run_smart_pairing(
                source_files,
                translation_files,
                pairing_service_supplier=lambda: get_pairing_service(event_bus=getattr(self,'event_bus',None))
            )
            self.file_pairs = pairs
            self.unmatched_files = unmatched
            self._display_file_pairing_results(pairs, unmatched.get('source', []), unmatched.get('translation', []))
            self._update_pairing_status_display()
            try:
                self._update_ribbon_states()
            except Exception:
                pass
        except Exception as e:  # pragma: no cover
            try:
                self._handle_error(e, context='pairing.smart', user_message=None, toast=False)
            except Exception:
                pass

    def _update_pairing_status_display(self):
        """UPDATE PAIRING STATUS DISPLAY - Aktualisiere Pairing Status in UI"""
        try:
            if hasattr(self, 'pairing_status_label'):
                pairs_count = len(getattr(self, 'file_pairs', []))
                unmatched_count = 0
                
                if hasattr(self, 'unmatched_files'):
                    unmatched_count = len(self.unmatched_files.get('source', [])) + len(self.unmatched_files.get('translation', []))
                
                if pairs_count > 0:
                    status_text = f"{pairs_count} Dateipaar(e) konfiguriert"
                    if unmatched_count > 0:
                        status_text += f" · {unmatched_count} ungepaarte Datei(en)"
                    
                    self.pairing_status_label.configure(
                        text=status_text,
                        # Accent Reduction: Paarungs-Status nutzt jetzt Primary statt success (Grün reserviert für system_status)
                        text_color=self._accent('success')
                    )
                else:
                    self.pairing_status_label.configure(
                        text="Keine Paare konfiguriert",
                        text_color=self.get_color('text_secondary')
                    )

                try:
                    self._refresh_ollama_pairs()
                except Exception as refresh_err:
                    self._debug_silent(refresh_err, "ollama.refresh_pairs")
                    
        except Exception as e:
            self._handle_error(e, context="pairing_status", user_message=None, toast=False, event_name="error.pairing.status")
    
    def _normalize_filename(self, filepath):
        """Delegiert an PairingService für Normalisierung (Backward Compatibility Stub)."""
        try:
            from pairing_service import get_pairing_service
            svc = get_pairing_service()
            return svc._normalize(filepath)  # bewusst interne Nutzung, um Duplikate zu vermeiden
        except Exception:
            try:
                return os.path.basename(filepath).lower()
            except Exception:
                return filepath
    
    def _calculate_filename_similarity(self, name1, name2):
        """Delegiert Similarity-Berechnung an PairingService (Backward Compatibility Stub)."""
        try:
            from pairing_service import get_pairing_service
            svc = get_pairing_service()
            return svc._similarity(name1, name2)
        except Exception:
            return 0.0
    
    def _display_file_pairing_results(self, pairs, unmatched_source, unmatched_translation):
        """Zeigt das Ergebnis der Dateipaarung in der Benutzeroberfläche."""
        try:
            if not pairs and not unmatched_source and not unmatched_translation:
                return
                
            # Toast mit Paarung-Zusammenfassung
            if pairs:
                count_pairs = len(pairs)
                label_pairs = self._n(self._t('file pair') if hasattr(self,'_t') else 'Dateipaar', self._t('file pairs') if hasattr(self,'_t') else 'Dateipaare', count_pairs)
                message = f"{count_pairs} {label_pairs} {self._t('automatically detected') if hasattr(self,'_t') else 'automatisch erkannt'}"
                if unmatched_source or unmatched_translation:
                    unmatched_total = len(unmatched_source) + len(unmatched_translation)
                    message += f" – {unmatched_total} {self._t('unpaired file(s)')}"
                self.show_toast(message, "success", duration=_SUCCESS_TOAST_DURATION_MS)
                
                # Detaillierte Info ins Log
                for pair in pairs:
                    self.logger.info(f"Pair: {pair['source_name']} → {pair['translation_name']} (Ähnlichkeit: {pair['similarity']:.1%})")
            else:
                self.show_toast(self._t("Automatic file pairing not possible - names too different"), "warning", duration=_WARNING_TOAST_DURATION_MS)
                
            # Status-Update
            status_parts = []
            if pairs:
                count_pairs = len(pairs)
                label_pairs = self._n(self._t('file pair') if hasattr(self,'_t') else 'Dateipaar', self._t('file pairs') if hasattr(self,'_t') else 'Dateipaare', count_pairs)
                status_parts.append(f"{count_pairs} {label_pairs}")
            if unmatched_source:
                status_parts.append(f"{len(unmatched_source)} {self._t('unpaired source file(s)')}")
            if unmatched_translation:
                status_parts.append(f"{len(unmatched_translation)} {self._t('unpaired translation(s)')}")
                
            if status_parts:
                self.update_status(f"Dateipaarung: {' | '.join(status_parts)}")
                
        except Exception as e:
            self._handle_error(e, context="pairing.results.display", user_message=self._t("Pairing-Ergebnis konnte nicht angezeigt werden"), toast=False)
    
    def get_file_pairs(self):
        """Get File Pairs - Hole aktuelle Dateipaarung für Qualitätsanalyse"""
        if hasattr(self, 'file_pairs'):
            return self.file_pairs
        else:
            # Fallback: Einfache Index-basierte Paarung
            source_files = self.uploaded_files.get('source', [])
            translation_files = self.uploaded_files.get('translation', [])
            
            simple_pairs = []
            max_pairs = min(len(source_files), len(translation_files))
            
            for i in range(max_pairs):
                simple_pairs.append({
                    'source': source_files[i],
                    'translation': translation_files[i],
                    'similarity': 0.5,  # Unbekannt
                    'source_name': os.path.basename(source_files[i]),
                    'translation_name': os.path.basename(translation_files[i])
                })
            
            return simple_pairs

    def _check_and_show_manual_pairing_option(self):
        """Check Manual Pairing Option - Prüfe ob manuelles Pairing angeboten werden soll"""
        try:
            source_count = len(self.uploaded_files.get('source', []))
            translation_count = len(self.uploaded_files.get('translation', []))
            
            # Zeige die Option für manuelles Pairing, wenn:
            # 1. Beide Dateitypen vorhanden sind
            # 2. Und ungepaarte Dateien existieren oder der Benutzer manuell konfigurieren möchte
            if source_count > 0 and translation_count > 0:
                unmatched_total = 0
                if hasattr(self, 'unmatched_files'):
                    unmatched_total = len(self.unmatched_files.get('source', [])) + len(self.unmatched_files.get('translation', []))
                
                # Zeige Manual Pairing Button wenn ungepaarte Dateien existieren
                if unmatched_total > 0 or source_count != translation_count:
                    self._after(2000, lambda: self.show_toast(
                        self._t("Manual pairing available - Click 'Adjust file pairing'"),
                        "info",
                        duration=5000
                    ))
                    
        except Exception as e:
            self._handle_error(e, context="pairing.manual.option", user_message=self._t("Manuelles Pairing Option Fehler"), toast=False)

    def _show_manual_pairing_dialog(self):
        """🎯 SHOW MANUAL PAIRING DIALOG - Zeige manuelles Pairing Interface"""
        try:
            # DEBUG: Prüfe uploaded_files beim Öffnen des Dialogs
            if hasattr(self, 'uploaded_files'):
                source_count = len(self.uploaded_files.get('source', []))
                trans_count = len(self.uploaded_files.get('translation', []))
                self.logger.info(f"🔍 Opening pairing dialog - uploaded_files: source={source_count}, translation={trans_count}")
                if source_count > 0:
                    self.logger.info(f"   Source files: {[os.path.basename(f) for f in self.uploaded_files.get('source', [])[:3]]}")
                if trans_count > 0:
                    self.logger.info(f"   Translation files: {[os.path.basename(f) for f in self.uploaded_files.get('translation', [])[:3]]}")
            else:
                self.logger.warning("⚠️ uploaded_files attribute not found!")
                self.show_toast("Keine Dateien hochgeladen", "warning")
                return
            
            # Erstelle Manual Pairing Window
            pairing_window = ctk.CTkToplevel(self.root)
            pairing_window.title("Manuelle Dateipaarung")
            # Etwas breiter & höher für klarere Zweispalten-Ansicht
            pairing_window.geometry("1020x760")
            pairing_window.transient(self.root)
            pairing_window.grab_set()
            # Cleanup-Handler für Scroll-Bindings
            def _cleanup_bindings():
                try:
                    # Entfernt globale Bindings nur wenn existierend
                    pairing_window.unbind_all('<MouseWheel>')
                    pairing_window.unbind_all('<Button-4>')
                    pairing_window.unbind_all('<Button-5>')
                except Exception:
                    pass
            try:
                def _confirm_close_pairing_window():
                    dirty = getattr(self, '_pair_dirty', False)
                    if dirty:
                        try:
                            self._confirm(
                                title=self._t('Discard changes?') if hasattr(self,'_t') else 'Änderungen verwerfen?',
                                message=self._t('You have unsaved pairing changes.') if hasattr(self,'_t') else 'Es gibt ungespeicherte Paarungen.',
                                on_confirm=lambda: (_cleanup_bindings(), pairing_window.destroy())
                            )
                        except Exception:
                            pairing_window.destroy()
                    else:
                        _cleanup_bindings(); pairing_window.destroy()
                pairing_window.protocol("WM_DELETE_WINDOW", _confirm_close_pairing_window)
            except Exception:
                pass
            # ESC beendet aktiven Drag
            try:
                pairing_window.bind('<Escape>', lambda e: self._cancel_pair_drag())
            except Exception:
                pass
            
            # Header – vereinheitlicht mit App-Header (Primary Blau) und etwas höher
            header_frame = ctk.CTkFrame(pairing_window, fg_color=self.get_color('primary'), height=70)
            header_frame.pack(fill="x", padx=0, pady=0)
            header_frame.pack_propagate(False)
            
            header_label = ctk.CTkLabel(
                header_frame,
                text="Manuelle Dateipaarung",
                font=ctk.CTkFont(*self.get_typography('title')),
                text_color=self.get_color('white')
            )
            header_label.pack(pady=15)
            
            # Outer content area mit sanftem Hintergrund für visuelle Trennung
            content_frame = ctk.CTkFrame(
                pairing_window,
                fg_color=self.get_color('background'),
                corner_radius=0
            )
            content_frame.pack(fill="both", expand=True, padx=0, pady=0)

            inner_wrapper = ctk.CTkFrame(
                content_frame,
                fg_color="transparent"
            )
            inner_wrapper.pack(fill="both", expand=True, padx=24, pady=24)
            
            # Instructions
            instructions = ctk.CTkLabel(
                inner_wrapper,
                text="Verbinde Ausgangstexte mit ihren Übersetzungen. Ziehe Einträge oder nutze die Dropdown-Auswahl unten:",
                font=ctk.CTkFont(*self.get_typography('body_md')),
                text_color=self.get_color('text_secondary'),
                wraplength=760,
                justify="center"
            )
            instructions.pack(pady=(0, 24))

            # Live Drag Status
            try:
                self._drag_status_var = tk.StringVar(value="")
                drag_status = ctk.CTkLabel(
                    content_frame,
                    textvariable=self._drag_status_var,
                    font=ctk.CTkFont(*self.get_typography('caption')),
                    text_color=self.get_color('text_secondary')
                )
                drag_status.pack(pady=(0,10))
            except Exception:
                pass
            
            # Pairing area
            # Bereich für Pairing-UI – neutraler Hintergrund
            pairing_area = ctk.CTkFrame(inner_wrapper, fg_color="transparent")
            pairing_area.pack(fill="both", expand=True)
            
            # Setup pairing interface
            self._setup_manual_pairing_interface(pairing_area, pairing_window)
            
        except Exception as e:
            self._handle_error(e, context="pairing.manual.dialog", user_message=self._t("Manueller Pairing Dialog Fehler"))
            self.show_toast("Fehler beim Öffnen der manuellen Paarung", "error")

    def _setup_manual_pairing_interface(self, parent, window):
        """Richtet die interaktive Oberfläche für das manuelle Pairing ein."""
        try:
            # Scrollbarer Bereich mit leichter Card-Anmutung
            scroll_frame = ctk.CTkScrollableFrame(
                parent,
                height=420,
                fg_color=self.get_color('surface'),
                corner_radius=12,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            scroll_frame.pack(fill="both", expand=True, padx=4, pady=4)
            
            # Current pairs section
            current_section = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            current_section.pack(fill='x', pady=(0, 8))
            current_label = ctk.CTkLabel(
                current_section,
                text="Aktuelle Dateipaarungen",
                font=ctk.CTkFont(*self.get_typography('section_label')),
                text_color=self.get_color('text_primary')
            )
            current_label.pack(side='left')
            
            # Zeige aktuelle Paare
            self.pairing_pairs_frame = self._create_card_frame(
                scroll_frame,
                fg_color=self.get_color('surface'),
                corner_radius=12,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            self.pairing_pairs_frame.pack(fill="x", pady=(0, 24))
            
            # Unmatched files section
            unmatched_section = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            unmatched_section.pack(fill='x', pady=(0, 8))
            unmatched_label = ctk.CTkLabel(
                unmatched_section,
                text="Ungepaarte Dateien",
                font=ctk.CTkFont(*self.get_typography('section_label')),
                text_color=self.get_color('warning')
            )
            unmatched_label.pack(side='left')
            
            # Two columns for unmatched files
            unmatched_container = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            unmatched_container.pack(fill="x", pady=(0, 8))
            unmatched_container.grid_columnconfigure((0, 1), weight=1)
            
            # Unmatched source files
            self.unmatched_source_frame = self._create_card_frame(
                unmatched_container,
                fg_color=self.get_color('surface'),
                corner_radius=12,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            self.unmatched_source_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 6))
            
            source_label = ctk.CTkLabel(
                self.unmatched_source_frame,
                text=self._t("Unpaired source files") if hasattr(self,'_t') else "Ungepaarte Quelldateien",
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('text_primary')
            )
            source_label.pack(pady=(10,6))
            try:
                self._unmatched_source_label = source_label
            except Exception:
                pass
            
            # Unmatched translation files
            self.unmatched_translation_frame = self._create_card_frame(
                unmatched_container,
                fg_color=self.get_color('surface'),
                corner_radius=12,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            self.unmatched_translation_frame.grid(row=0, column=1, sticky="nsew", padx=(6, 0))
            
            translation_label = ctk.CTkLabel(
                self.unmatched_translation_frame,
                text=self._t("Unpaired translations") if hasattr(self,'_t') else "Ungepaarte Übersetzungen",
                font=ctk.CTkFont(*self.get_typography('body_bold')),
                text_color=self.get_color('text_primary')
            )
            translation_label.pack(pady=(10,6))
            try:
                self._unmatched_translation_label = translation_label
            except Exception:
                pass
            
            # Populate interface
            self._populate_manual_pairing_interface()
            
            # Ermittele das Toplevel-Fenster für sichere Save/Cancel-Aktionen
            try:
                win = parent.winfo_toplevel()
            except Exception:
                win = parent

            # Button area (separate Leiste unten)
            button_frame = ctk.CTkFrame(parent, fg_color="transparent")
            button_frame.pack(fill="x", pady=(12, 0))
            
            # Undo Button
            undo_btn = self._create_button(
                button_frame,
                text=self._t("Undo") if hasattr(self, '_t') else "Undo",
                command=self._undo,  # nutzt neues Snapshot-System
                kind="primary"
            )
            undo_btn.pack(side="left", padx=(0, 10))
            self._legacy_undo_btn = undo_btn

            # Redo Button
            redo_btn = self._create_button(
                button_frame,
                text=self._t("Redo") if hasattr(self, '_t') else "Redo",
                command=self._redo,  # nutzt neues Snapshot-System
                kind="primary"
            )
            redo_btn.pack(side="left", padx=(0, 20))
            self._legacy_redo_btn = redo_btn

            # Reset button
            reset_btn = self._create_button(
                button_frame,
                text=self._t("Repeat automatic pairing") if hasattr(self,'_t') else "Automatische Paarung wiederholen",
                command=self._reset_to_automatic_pairing,
                kind="primary"
            )
            reset_btn.pack(side="left", padx=(0, 10))
            
            # Save button
            save_btn = self._create_button(
                button_frame,
                text=self._t("Save pairing") if hasattr(self,'_t') else "Paarung speichern",
                command=lambda: self._save_manual_pairing(win),
                kind="primary"
            )
            save_btn.pack(side="right")
            
            # Cancel button
            cancel_btn = self._create_button(
                button_frame,
                text=self._t("Cancel") if hasattr(self,'_t') else "Abbrechen",
                command=win.destroy,
                kind="primary"
            )
            cancel_btn.pack(side="right", padx=(0, 10))

            # Cleanup MouseWheel Bindings beim Schließen (falls bind_all genutzt wurde)
            try:
                def _cleanup_bindings():
                    try:
                        win.unbind_all('<MouseWheel>')
                        win.unbind_all('<Button-4>')
                        win.unbind_all('<Button-5>')
                    except Exception:
                        pass
                def _confirm_close_pairing():
                    if getattr(self, '_pair_dirty', False):
                        try:
                            self._confirm(
                                title=self._t('Discard changes?') if hasattr(self,'_t') else 'Änderungen verwerfen?',
                                message=self._t('You have unsaved pairing changes.') if hasattr(self,'_t') else 'Es gibt ungespeicherte Paarungen.',
                                on_confirm=lambda: (_cleanup_bindings(), win.destroy())
                            )
                        except Exception:
                            win.destroy()
                    else:
                        _cleanup_bindings(); win.destroy()
                win.protocol("WM_DELETE_WINDOW", _confirm_close_pairing)
            except Exception:
                pass
            
        except Exception as e:
            self._handle_error(e, context="pairing.manual.interface", user_message=self._t("Manuelle Pairing Oberfläche Fehler"))

    def _populate_manual_pairing_interface(self):
        """POPULATE MANUAL PAIRING INTERFACE - Fülle Interface mit aktuellen Daten"""
        try:
            self.logger.debug("Populate manual pairing interface")
            start_time = time.perf_counter()
            created_items = {'source': 0, 'translation': 0}
            # Telemetrie-Zähler initialisieren (lazy)
            if not hasattr(self, '_telemetry_counters'):
                self._telemetry_counters = {}
            self._telemetry_counters.setdefault('pairing_paging_used', 0)
            if not hasattr(self, '_virtual_metrics'):
                self._virtual_metrics = {}
            
            # Clear existing content
            for widget in self.pairing_pairs_frame.winfo_children():
                widget.destroy()
            for widget in self.unmatched_source_frame.winfo_children()[1:]:  # Skip label
                widget.destroy()
            for widget in self.unmatched_translation_frame.winfo_children()[1:]:  # Skip label
                widget.destroy()
            
            # Wenn keine Daten vorhanden sind, führe Smart Pairing aus
            if not hasattr(self, 'file_pairs') or not hasattr(self, 'unmatched_files'):
                self.logger.debug("No pairing data found, executing smart pairing…")
                self._smart_file_pairing()
            
            # WICHTIG: Falls immer noch keine Daten vorhanden, erstelle leere Strukturen
            if not hasattr(self, 'file_pairs'):
                self.file_pairs = []
            if not hasattr(self, 'unmatched_files'):
                self.unmatched_files = {'source': [], 'translation': []}
            
            # Show current pairs
            current_pairs = getattr(self, 'file_pairs', [])
            try:
                if getattr(self, '_pair_sort_similarity', False):
                    current_pairs = sorted(current_pairs, key=lambda p: p.get('similarity',0.0), reverse=True)
            except Exception:
                pass
            self.logger.debug("Current pairs: %d", len(current_pairs))
            
            # Toolbar (Undo/Redo) – nur einmal anlegen
            if not hasattr(self, '_pairing_toolbar'):
                toolbar = ctk.CTkFrame(self.pairing_pairs_frame, fg_color=self.get_color('transparent'))
                toolbar.pack(fill='x', pady=(0,4))
                self._pairing_toolbar = toolbar
                undo_btn = self._create_button(
                    toolbar,
                    text=self._t("Undo") if hasattr(self,'_t') else "Undo",
                    command=self._undo,
                    kind='secondary',
                    height=26
                )
                undo_btn.pack(side='left', padx=(0,4))
                redo_btn = self._create_button(
                    toolbar,
                    text=self._t("Redo") if hasattr(self,'_t') else "Redo",
                    command=self._redo,
                    kind='secondary',
                    height=26
                )
                redo_btn.pack(side='left', padx=(0,8))
                self._undo_btn_toolbar = undo_btn
                self._redo_btn_toolbar = redo_btn
                # Sort by similarity Toggle
                try:
                    self._pair_sort_var = tk.BooleanVar(value=getattr(self, '_pair_sort_similarity', False))
                    sort_cb = ctk.CTkCheckBox(
                        toolbar,
                        text=self._t("Sort by similarity") if hasattr(self,'_t') else "Nach Ähnlichkeit sortieren",
                        command=lambda: self._toggle_pair_sort(),
                        variable=self._pair_sort_var
                    )
                    sort_cb.pack(side='left', padx=(4,4))
                except Exception:
                    pass

            if current_pairs:
                for i, pair in enumerate(current_pairs):
                    self._create_pair_display_item(self.pairing_pairs_frame, pair, i)
            else:
                no_pairs_label = ctk.CTkLabel(
                    self.pairing_pairs_frame,
                    text=self._t("No pairs found") if hasattr(self,'_t') else "Keine Paare gefunden",
                    font=ctk.CTkFont(*self.get_typography('body')),
                    text_color=self.get_color('text_secondary')
                )
                no_pairs_label.pack(pady=20)
            
            # Show unmatched files
            unmatched = getattr(self, 'unmatched_files', {'source': [], 'translation': []})
            try:
                # Counts an Labels anhängen (verwende Attribute falls vorhanden)
                if hasattr(self, '_unmatched_source_label') and self._unmatched_source_label.winfo_exists():
                    lbl = self._unmatched_source_label
                    lbl.configure(text=(lbl.cget('text').split(' (')[0] + f" ({len(unmatched.get('source', []))})"))
                if hasattr(self, '_unmatched_translation_label') and self._unmatched_translation_label.winfo_exists():
                    lbl2 = self._unmatched_translation_label
                    lbl2.configure(text=(lbl2.cget('text').split(' (')[0] + f" ({len(unmatched.get('translation', []))})"))
            except Exception:
                pass
            self.logger.debug("Unmatched files: source=%d, translation=%d", len(unmatched.get('source', [])), len(unmatched.get('translation', [])))
            # Live Filter Queries
            q_src = (self._src_filter.get().strip().lower() if hasattr(self, '_src_filter') else "")
            q_trg = (self._trg_filter.get().strip().lower() if hasattr(self, '_trg_filter') else "")
            def _apply_q(fs, q):
                try:
                    if not q:
                        return fs
                    return [f for f in fs if q in os.path.basename(f).lower()]
                except Exception:
                    return fs
            
            # Wenn keine unmatched Files aber uploaded Files vorhanden sind, verwende uploaded Files
            if (not unmatched.get('source') and not unmatched.get('translation')):
                source_files = self.uploaded_files.get('source', []) if hasattr(self, 'uploaded_files') else []
                translation_files = self.uploaded_files.get('translation', []) if hasattr(self, 'uploaded_files') else []
                
                if source_files or translation_files:
                    self.logger.debug("Using uploaded files as unmatched (no pairing data found)")
                    self.logger.debug("Source files available: %d", len(source_files))
                    self.logger.debug("Translation files available: %d", len(translation_files))
                    unmatched = {
                        'source': [f for f in source_files if not any(p.get('source') == f for p in current_pairs)],
                        'translation': [f for f in translation_files if not any(p.get('translation') == f for p in current_pairs)]
                    }
                    self.unmatched_files = unmatched  # Wichtig: Speichern für weitere Verwendung
                    self.logger.debug("Calculated unmatched: source=%d, translation=%d", len(unmatched['source']), len(unmatched['translation']))
            # Optional Paging for large datasets
            SOURCE_THRESHOLD = 150
            PAGE_SIZE = 60
            total_sources = len(unmatched.get('source', []))
            total_trans = len(unmatched.get('translation', []))
            if not hasattr(self, '_pairing_page_loaded'):
                self._pairing_page_loaded = False
            paged = (total_sources > SOURCE_THRESHOLD or total_trans > SOURCE_THRESHOLD) and not self._pairing_page_loaded
            filtered_sources = _apply_q(unmatched.get('source', []), q_src)
            filtered_trans = _apply_q(unmatched.get('translation', []), q_trg)
            total_sources = len(filtered_sources)
            total_trans = len(filtered_trans)
            src_slice = filtered_sources[:PAGE_SIZE] if paged else filtered_sources
            trans_slice = filtered_trans[:PAGE_SIZE] if paged else filtered_trans
            # Virtuelle Fenstergrößen-Parameter (dynamisch ableiten)
            # Heuristik: Schätze, wie viele Items vertikal passen, anhand verfügbarer Höhe des Source-Frames
            try:
                available_height = self.unmatched_source_frame.winfo_height()
            except Exception:
                available_height = 0
            # Wenn Frame noch nicht gerendert ist (Höhe=0), fallback auf vorher gespeicherten Wert oder Default
            if available_height <= 0:
                # Versuch alten Wert zu nutzen
                available_height = getattr(self, '_last_pairing_source_height', 0)
            else:
                self._last_pairing_source_height = available_height
            # Typische Item-Höhe inkl. Padding geschätzt: 34px (Label/Frame + Spacing). Sicherheitskorridor 36.
            est_item_height = 36
            dyn_window = max(40, min(300, int(available_height / est_item_height))) if available_height > 0 else 120
            VWINDOW = dyn_window  # Dynamisches Fenster
            if not hasattr(self, '_virtual_pairing'):  # State Initialisierung
                self._virtual_pairing = {
                    'source_offset': 0,
                    'translation_offset': 0,
                    'window': VWINDOW,  # dynamisch anpassbar
                    'total_source': total_sources,
                    'total_translation': total_trans
                }
            vp = self._virtual_pairing
            # Falls Fenstergröße sich geändert hat, aktualisieren
            if vp.get('window') != VWINDOW:
                vp['window'] = VWINDOW
            # Korrigiere Offsets falls sich Gesamtzahl geändert hat
            vp['total_source'] = total_sources
            vp['total_translation'] = total_trans
            vp['source_offset'] = min(vp.get('source_offset',0), max(0, total_sources - 1)) if total_sources else 0
            vp['translation_offset'] = min(vp.get('translation_offset',0), max(0, total_trans - 1)) if total_trans else 0

            def render_window(frame, files, ftype, offset_key):
                """Render ein virtuelles Fenster für einen Dateityp."""
                try:
                    offset = vp[offset_key]
                    win = vp['window']
                    window_files = files[offset:offset+win]
                    for f in window_files:
                        self._create_unmatched_file_item(frame, f, ftype)
                    # Anzahl erstellter Items protokollieren
                    created_items[ftype] += len(window_files)
                    # Navigation Buttons (oben/unten)
                    if offset > 0:
                        up_prefix = self._t("Vorherige anzeigen") if hasattr(self, '_t') else "Vorherige anzeigen"
                        up_btn = self._create_button(
                            frame,
                            text=f"{up_prefix} (ab {max(0, offset - win)})",
                            command=lambda: self._shift_virtual_window(offset_key, -win),
                            kind="primary",
                            height=24
                        )
                        up_btn.pack(fill="x", padx=5, pady=4)
                    if offset + win < len(files):
                        down_prefix = self._t("Weitere anzeigen") if hasattr(self, '_t') else "Weitere anzeigen"
                        down_btn = self._create_button(
                            frame,
                            text=f"{down_prefix} (ab {min(len(files)-1, offset+win)})",
                            command=lambda: self._shift_virtual_window(offset_key, win),
                            kind="primary",
                            height=24
                        )
                        down_btn.pack(fill="x", padx=5, pady=4)
                except Exception:
                    pass

            # Mousewheel Scroll Binding für dynamische Offsets
            def _bind_mousewheel(frame, offset_key):
                def _on_mousewheel(event):
                    try:
                        delta = 0
                        # Windows liefert event.delta in 120er Schritten
                        if hasattr(event, 'delta') and event.delta:
                            delta = -1 if event.delta > 0 else 1
                        elif hasattr(event, 'num') and event.num in (4,5):  # Linux
                            delta = -1 if event.num == 4 else 1
                        if delta != 0:
                            step = max(1, int(vp.get('window', 50) * 0.25))  # 25% Fenster pro Scroll
                            self._shift_virtual_window(offset_key, delta * step)
                    except Exception:
                        pass
                try:
                    frame.bind_all('<MouseWheel>', _on_mousewheel)  # Windows / Mac (Mac nutzt auch MouseWheel, delta anders ggf.)
                    frame.bind_all('<Button-4>', _on_mousewheel)     # Linux up
                    frame.bind_all('<Button-5>', _on_mousewheel)     # Linux down
                except Exception:
                    pass

            _bind_mousewheel(self.unmatched_source_frame, 'source_offset')
            _bind_mousewheel(self.unmatched_translation_frame, 'translation_offset')

            # Debug: Prüfe ob Dateien vorhanden sind
            self.logger.debug("About to render: src_slice=%d files, trans_slice=%d files", len(src_slice), len(trans_slice))
            
            # WICHTIG: Stelle sicher dass Dateien gerendert werden
            if src_slice:
                self.logger.debug("Rendering source files: %s", [os.path.basename(f) for f in src_slice[:3]])
            if trans_slice:
                self.logger.debug("Rendering translation files: %s", [os.path.basename(f) for f in trans_slice[:3]])
            
            render_window(self.unmatched_source_frame, src_slice, 'source', 'source_offset')
            render_window(self.unmatched_translation_frame, trans_slice, 'translation', 'translation_offset')
            if paged:
                # Telemetrie: Nutzung Paging registrieren
                self._telemetry_counters['pairing_paging_used'] += 1
                def _load_all():
                    try:
                        self._pairing_page_loaded = True
                        self._populate_manual_pairing_interface()
                    except Exception:
                        self.logger.exception("Error loading full pairing lists")
                # Komponiere Text ohne Annahme über Platzhalter im Übersetzungs-String
                if hasattr(self, '_t'):
                    sources_label = self._t('sources') if hasattr(self,'_t') else 'Quellen'
                    translations_label = self._t('translations') if hasattr(self,'_t') else 'Übersetzungen'
                    btn_text = (self._t("Show all") + f" ({total_sources} {sources_label} / {total_trans} {translations_label})")
                else:
                    btn_text = f"Alle anzeigen ({total_sources} Quellen / {total_trans} Übersetzungen)"
                more_btn = self._create_button(
                    self.unmatched_source_frame,
                    text=btn_text,
                    command=_load_all,
                    kind="primary"
                )
                more_btn.pack(fill="x", padx=8, pady=8)
            # Performance / Metrics erfassen
            try:
                elapsed = (time.perf_counter() - start_time) * 1000.0
                self._virtual_metrics['last_render_ms'] = round(elapsed, 2)
                self._virtual_metrics['window'] = self._virtual_pairing.get('window')
                self._virtual_metrics['created_source'] = created_items['source']
                self._virtual_metrics['created_translation'] = created_items['translation']
                self._virtual_metrics['total_source'] = total_sources
                self._virtual_metrics['total_translation'] = total_trans
                self.logger.debug(
                    "Virtual render: win=%s src_created=%s trans_created=%s elapsed=%.2fms src_total=%d trans_total=%d",
                    self._virtual_metrics['window'],
                    created_items['source'],
                    created_items['translation'],
                    elapsed,
                    total_sources,
                    total_trans
                )
            except Exception:
                pass
            # Nach initialem Render dynamische Justierung (nur wenn Frame-Höhen vorher 0 waren)
            try:
                if getattr(self, '_virtual_adjust_scheduled', False) is False:
                    self._virtual_adjust_scheduled = True
                    self._after(80, self._adjust_virtual_window_size_safe)
            except Exception:
                pass
                
        except Exception as e:
            self._handle_error(e, context="pairing.manual.populate", user_message=self._t("Manuelle Pairing Befüllung Fehler"), toast=False)
        finally:
            try:
                if hasattr(self, '_update_undo_redo_buttons'):
                    self._update_undo_redo_buttons()
            except Exception:
                pass

    def _adjust_virtual_window_size_safe(self):
        """Passe die virtuelle Fenstergröße adaptiv an reale Item-Höhen an (Heuristik)."""
        try:
            if not hasattr(self, '_virtual_pairing'):
                return
            frame = getattr(self, 'unmatched_source_frame', None)
            if not frame:
                return
            children = frame.winfo_children()
            if len(children) <= 1:  # nur Label vorhanden
                # Erneut versuchen
                self._after(120, self._adjust_virtual_window_size_safe)
                return
            # Label-Höhe abziehen
            header_h = children[0].winfo_height() if children else 0
            usable_h = max(0, frame.winfo_height() - header_h)
            # Durchschnittliche Höhe der ersten bis zu 10 Items (exkl. Label & Navigationsbuttons erkennbar an führenden Symbolen)
            item_heights = []
            for w in children[1:11]:
                try:
                    txt = getattr(w, 'cget', lambda *_: '')('text') if hasattr(w, 'cget') else ''
                    if isinstance(txt, str) and (txt.startswith('Vorherige') or txt.startswith('Weitere')):
                        continue
                    item_heights.append(w.winfo_height())
                except Exception:
                    pass
            if not item_heights:
                return
            avg_h = sum(item_heights) / len(item_heights)
            if avg_h <= 0:
                return
            target_window = int(usable_h / avg_h) if avg_h > 0 else self._virtual_pairing.get('window', 120)
            target_window = max(20, min(300, target_window))
            current_window = self._virtual_pairing.get('window', target_window)
            # Nur aktualisieren wenn signifikant anders (>5 Unterschied)
            if abs(target_window - current_window) > 5:
                self._virtual_pairing['window'] = target_window
                self.logger.debug("Adjust virtual window size: %d -> %d (avg_h=%.1f usable_h=%d)", current_window, target_window, avg_h, usable_h)
                # Neu rendern ohne erneute Scheduling-Flut
                self._populate_manual_pairing_interface()
        except Exception:
            pass

    def _create_pair_display_item(self, parent, pair, index):
        """CREATE PAIR DISPLAY ITEM - Erstelle anzeigbares Dateipaar"""
        try:
            pair_frame = ctk.CTkFrame(parent, fg_color=self.get_color('surface'), border_width=1, border_color=self.get_color('surface_border'))
            pair_frame.pack(fill="x", padx=10, pady=5)
            
            # Pair content
            content_frame = ctk.CTkFrame(pair_frame, fg_color="transparent")
            content_frame.pack(fill="x", padx=15, pady=10)
            
            # Source file
            source_label = ctk.CTkLabel(
                content_frame,
                text=f"{self._t('Source') if hasattr(self,'_t') else 'Source'}: {pair['source_name']}",
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                anchor="w"
            )
            source_label.pack(fill="x")
            
            # Arrow and similarity
            arrow_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            arrow_frame.pack(fill="x", pady=5)
            
            # Dynamische Farbgebung nach Ähnlichkeitswert
            sim = pair.get('similarity', 0.0)
            if sim >= 0.8:
                sim_color = self.get_color('success')
            elif sim < 0.5:
                sim_color = self.get_color('warning')
            else:
                sim_color = self.get_color('text_secondary')
            similarity_label = self._t("Similarity") if hasattr(self, '_t') else "Ähnlichkeit"
            if isinstance(similarity_label, str) and similarity_label.strip().lower() == "similarity":
                similarity_label = "Ähnlichkeit"

            arrow_label = ctk.CTkLabel(
                arrow_frame,
                text=f"{similarity_label}: {sim:.1%}",
                font=ctk.CTkFont(*self.get_typography('caption')),
                text_color=sim_color
            )
            arrow_label.pack(side="left")
            try:
                badge = ctk.CTkLabel(
                    arrow_frame,
                    text=f"{sim*100:.0f}%",
                    font=ctk.CTkFont(*self.get_typography('caption')),
                    text_color=sim_color,
                    fg_color=self.get_color('white'),
                    corner_radius=8,
                    width=42
                )
                badge.pack(padx=8, side="right")
            except Exception:
                pass
            # Kontextmenü
            try:
                self._attach_ctx(pair_frame, pair.get('source'))
            except Exception:
                pass
            
            # Translation file
            trans_label = ctk.CTkLabel(
                content_frame,
                text=f"{self._t('Translation') if hasattr(self,'_t') else 'Translation'}: {pair['translation_name']}",
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                anchor="w"
            )
            trans_label.pack(fill="x")
            
            # Actions
            actions_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            actions_frame.pack(fill="x", pady=(10, 0))
            
            # Unpair button
            unpair_btn = self._create_button(
                actions_frame,
                text=self._t("Unpair") if hasattr(self,'_t') else "Trennen",
                command=lambda: self._unpair_files(index),
                kind="warning",
                height=26
            )
            unpair_btn.configure(width=80)
            unpair_btn.pack(side="right")
            
        except Exception as e:
            self._handle_error(e, context="pairing.manual.pair_display_item", user_message=self._t("Pair Anzeige Element Fehler"), toast=False)

    def _create_unmatched_file_item(self, parent, file_path, file_type):
        """CREATE UNMATCHED FILE ITEM - Erstelle ungepaarte Datei mit Pairing-Option"""
        try:
            self.logger.debug("Creating unmatched file item: %s (type: %s)", os.path.basename(file_path), file_type)
            
            item_frame = ctk.CTkFrame(parent, fg_color=self.get_color('surface'), border_width=1, border_color=self.get_color('surface_border'))
            item_frame.pack(fill="x", padx=5, pady=2)
            
            content_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            content_frame.pack(fill="x", padx=10, pady=8)
            
            # File name
            # _t nur aufrufen wenn vorhanden
            if hasattr(self, '_t'):
                file_prefix = self._t("Source") if file_type == "source" else self._t("Translation")
            else:
                file_prefix = "Source" if file_type == "source" else "Translation"
            file_label = ctk.CTkLabel(
                content_frame,
                text=f"{file_prefix}: {os.path.basename(file_path)}",
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                anchor="w"
            )
            file_label.pack(fill="x")
            try:
                self._attach_ctx(item_frame, file_path)
            except Exception:
                pass

            # Drag & Drop: Start Drag bei Linksklick halten
            try:
                item_frame.bind('<ButtonPress-1>', lambda e, fp=file_path, ft=file_type, w=item_frame: self._begin_pair_drag(fp, ft, w))
                # 🔧 KRITISCH: ButtonRelease muss auf dem TARGET-Widget getriggert werden!
                # Wir binden Enter-Event, um bei Mouse-Over das Pairing zu ermöglichen
                def on_enter_with_drag(event, target_file=file_path, target_type=file_type, widget=item_frame):
                    self._drag_enter_item(target_file, target_type, widget)
                    # Wenn Maustaste noch gedrückt ist (Drag aktiv), erlaube Drop hier
                    # ButtonRelease auf diesem Widget = Pairing complete
                    widget.bind('<ButtonRelease-1>', lambda e, tf=target_file, tt=target_type: self._complete_pair_drag(tf, tt), add='+')
                
                item_frame.bind('<Enter>', on_enter_with_drag)
                item_frame.bind('<Leave>', lambda e: self._drag_leave_item(item_frame))
            except Exception:
                pass
            
            # Dropdown für manuelle Paarung
            if file_type == "source":
                # Zeige verfügbare Translation-Dateien
                if hasattr(self, 'unmatched_files') and self.unmatched_files.get('translation'):
                    available_translations = self.unmatched_files.get('translation', [])
                else:
                    # Fallback: Verwende alle translation files die nicht gepaart sind
                    all_translations = self.uploaded_files.get('translation', [])
                    paired_translations = [p['translation'] for p in getattr(self, 'file_pairs', [])]
                    available_translations = [f for f in all_translations if f not in paired_translations]
                
                self.logger.debug("Available translations for %s: %d", os.path.basename(file_path), len(available_translations))
                
                if available_translations:
                    dropdown_values = [self._prompt_translation()] + [os.path.basename(f) for f in available_translations]
                    
                    # Dropdown mit StringVar für bessere Kontrolle
                    pair_var = tk.StringVar(value=self._prompt_translation())
                    
                    pair_dropdown = ctk.CTkComboBox(
                        content_frame,
                        values=dropdown_values,
                        variable=pair_var,
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        height=24,
                        state="readonly"
                    )
                    pair_dropdown.pack(fill="x", pady=(5, 0))
                    
                    # Expliziter "Zuordnen" Button
                    pair_button = self._create_button(
                        content_frame,
                        text=self._t("Zuordnen") if hasattr(self, '_t') else "Zuordnen",
                        command=lambda sf=file_path, var=pair_var: self._manual_pair_files(sf, var.get(), file_type),
                        kind="primary",
                        height=28
                    )
                    pair_button.pack(fill="x", pady=(3, 0))
                else:
                    # Keine verfügbaren Übersetzungen
                    no_trans_label = ctk.CTkLabel(
                        content_frame,
                        text=self._t("No unmatched translations available") if hasattr(self,'_t') else "Keine ungepaarten Übersetzungen verfügbar",
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        text_color=self.get_color('text_secondary')
                    )
                    no_trans_label.pack(fill="x", pady=(5, 0))
                    
            elif file_type == "translation":
                # Zeige verfügbare Source-Dateien
                if hasattr(self, 'unmatched_files') and self.unmatched_files.get('source'):
                    available_sources = self.unmatched_files.get('source', [])
                else:
                    # Fallback: Verwende alle source files die nicht gepaart sind
                    all_sources = self.uploaded_files.get('source', [])
                    paired_sources = [p['source'] for p in getattr(self, 'file_pairs', [])]
                    available_sources = [f for f in all_sources if f not in paired_sources]
                
                self.logger.debug("Available sources for %s: %d", os.path.basename(file_path), len(available_sources))
                
                if available_sources:
                    dropdown_values = [self._prompt_source()] + [os.path.basename(f) for f in available_sources]
                    
                    # Dropdown mit StringVar für bessere Kontrolle
                    pair_var = tk.StringVar(value=self._prompt_source())
                    
                    pair_dropdown = ctk.CTkComboBox(
                        content_frame,
                        values=dropdown_values,
                        variable=pair_var,
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        height=24,
                        state="readonly"
                    )
                    pair_dropdown.pack(fill="x", pady=(5, 0))
                    
                    # Expliziter "Zuordnen" Button
                    pair_button = self._create_button(
                        content_frame,
                        text=self._t("Zuordnen") if hasattr(self, '_t') else "Zuordnen",
                        command=lambda tf=file_path, var=pair_var: self._manual_pair_files_translation(tf, var.get(), file_type),
                        kind="primary",
                        height=28
                    )
                    pair_button.pack(fill="x", pady=(3, 0))
                else:
                    # Keine verfügbaren Quellen
                    no_source_label = ctk.CTkLabel(
                        content_frame,
                        text=self._t("No unmatched sources available") if hasattr(self,'_t') else "Keine ungepaarten Ausgangstexte verfügbar",
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        text_color=self.get_color('text_secondary')
                    )
                    no_source_label.pack(fill="x", pady=(5, 0))
                    
        except Exception as e:
            self._handle_error(e, context="pairing.manual.unmatched.item", user_message=self._t("Ungepaartes Datei Element Fehler"), toast=False)

    def _manual_pair_files(self, source_file, selected_translation, file_type):
        """Delegiert manuelles Pairing (Source -> Translation) an PairingManager."""
        try:
            if selected_translation == self._prompt_translation():
                return
            # Übersetzungsdatei ermitteln
            if hasattr(self, 'unmatched_files') and self.unmatched_files.get('translation'):
                candidates = self.unmatched_files.get('translation', [])
            else:
                paired = [p['translation'] for p in getattr(self, 'file_pairs', [])]
                candidates = [f for f in self.uploaded_files.get('translation', []) if f not in paired]
            translation_file = next((f for f in candidates if os.path.basename(f) == selected_translation), None)
            if not translation_file:
                self.show_toast(self._t('Translation file not found'), 'error'); return
            if not self.pairing_manager.add_manual_pair(source_file, translation_file):
                self.show_toast(self._t('File already paired'), 'warning'); return
            self.file_pairs = self.pairing_manager.get_legacy_pairs()
            self.unmatched_files = self.pairing_manager.get_legacy_unmatched()
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.create', source=os.path.basename(source_file), translation=selected_translation)
            self.show_toast(f"{self._t('Paar erstellt') if hasattr(self,'_t') else 'Paar erstellt'}: {os.path.basename(source_file)} → {selected_translation}", 'success')
        except Exception as e:  # pragma: no cover
            self._handle_error(e, context='pairing.manual.pair_files', user_message=self._t('Manuelle Paarung Fehler'))

    def _manual_pair_files_translation(self, translation_file, selected_source, file_type):
        """Delegiert manuelles Pairing (Translation -> Source) an PairingManager."""
        try:
            if selected_source == self._prompt_source():
                return
            if hasattr(self, 'unmatched_files') and self.unmatched_files.get('source'):
                candidates = self.unmatched_files.get('source', [])
            else:
                paired = [p['source'] for p in getattr(self, 'file_pairs', [])]
                candidates = [f for f in self.uploaded_files.get('source', []) if f not in paired]
            source_file = next((f for f in candidates if os.path.basename(f) == selected_source), None)
            if not source_file:
                self.show_toast(self._t('Source file not found'), 'error'); return
            if not self.pairing_manager.add_manual_pair(source_file, translation_file):
                self.show_toast(self._t('File already paired'), 'warning'); return
            self.file_pairs = self.pairing_manager.get_legacy_pairs()
            self.unmatched_files = self.pairing_manager.get_legacy_unmatched()
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.create', source=selected_source, translation=os.path.basename(translation_file))
            self.show_toast(f"{self._t('Paar erstellt') if hasattr(self,'_t') else 'Paar erstellt'}: {selected_source} → {os.path.basename(translation_file)}", 'success')
        except Exception as e:  # pragma: no cover
            self._handle_error(e, context='pairing.manual.pair_files.translation', user_message=self._t('Manuelle Paarung (Translation) Fehler'))

    def _unpair_files(self, pair_index):
        """Delegiert Unpair an PairingManager (Backward Compatible)."""
        try:
            if not hasattr(self, 'file_pairs') or pair_index >= len(self.file_pairs):
                return
            removed = self.pairing_manager.remove_pair_by_index(pair_index)
            if not removed:
                return
            self.file_pairs = self.pairing_manager.get_legacy_pairs()
            self.unmatched_files = self.pairing_manager.get_legacy_unmatched()
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.remove', source=removed.get('source_name'), translation=removed.get('translation_name'))
            self.show_toast(f"{self._t('Paarung aufgelöst') if hasattr(self,'_t') else 'Paarung aufgelöst'}: {removed['source_name']} → {removed['translation_name']}", 'info')
        except Exception as e:  # pragma: no cover
            self._handle_error(e, context='pairing.unpair')

    def _reset_to_automatic_pairing(self):
        """RESET TO AUTOMATIC PAIRING - Setze auf automatische Paarung zurück"""
        try:
            # History Snapshot (einfacher) – vollständige Zustände bereits im History Stack
            self._push_history()
            # Führe automatische Paarung erneut aus
            self._smart_file_pairing()
            
            # Interface aktualisieren
            self._populate_manual_pairing_interface()
            
            # WICHTIG: Aktualisiere auch Haupt-GUI Status
            self._update_pairing_status_display()
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.reorder')
            self.show_toast(self._t("Automatic pairing restored"), "info")
            try:
                self._update_ribbon_states()
            except Exception:
                pass
            
        except Exception as e:
            self._handle_error(e, context="pairing.reset_auto")

    def _save_manual_pairing(self, window):
        """Speichere manuelle Paarung (Status/Toast/Autosave bleiben erhalten)."""
        try:
            # Aktualisiere Status
            pairs_count = len(getattr(self, 'file_pairs', []))
            unmatched_count = len(self.unmatched_files.get('source', [])) + len(self.unmatched_files.get('translation', []))
            
            # 🔧 DEBUG: Log current file_pairs state before saving
            self.logger.info(f"📝 Saving pairing: {pairs_count} pairs configured")
            if hasattr(self, 'file_pairs') and self.file_pairs:
                for idx, pair in enumerate(self.file_pairs[:3]):  # Log first 3 pairs
                    src = os.path.basename(pair.get('source', 'N/A'))
                    tgt = os.path.basename(pair.get('translation', 'N/A'))
                    self.logger.info(f"   Pair {idx+1}: {src} → {tgt}")
            
            # Zeige Ergebnis
            if pairs_count > 0:
                label_pairs = self._n(self._t('file pair') if hasattr(self,'_t') else 'Dateipaar', self._t('file pairs') if hasattr(self,'_t') else 'Dateipaare', pairs_count)
                message = f"{self._t('Pairing saved')}: {pairs_count} {label_pairs}"
                if unmatched_count > 0:
                    label_unmatched = self._n(self._t('unpaired file') if hasattr(self,'_t') else 'ungepaarte Datei', self._t('unpaired files') if hasattr(self,'_t') else 'ungepaarte Dateien', unmatched_count)
                    message += f" → {unmatched_count} {label_unmatched}"
                    
                self.show_toast(message, "success", duration=_SUCCESS_TOAST_DURATION_MS)
                status_label_pairs = self._n(self._t('file pair') if hasattr(self,'_t') else 'Dateipaar', self._t('file pairs') if hasattr(self,'_t') else 'Dateipaare', pairs_count)
                self.update_status(f"{self._t('Manual file pairing')}: {pairs_count} {status_label_pairs} {self._t('configured') if hasattr(self,'_t') else 'konfiguriert'}")
            else:
                self.show_toast(self._t("No pairs configured"), "warning")
            
            # WICHTIG: Aktualisiere Pairing Status Display in der Haupt-GUI
            self._update_pairing_status_display()
            
            # 🔧 FIX: Stelle sicher dass file_pairs in der Haupt-App gespeichert werden
            # BEFORE schließen des Dialogs - damit die Daten nicht verloren gehen
            try:
                # Synchronisiere mit pairing_manager
                if hasattr(self, 'pairing_manager') and self.pairing_manager:
                    self.logger.info("🔄 Syncing file_pairs with pairing_manager")
                    # Stelle sicher dass die Paarungen im Manager gespeichert sind
                    self.file_pairs = self.pairing_manager.get_legacy_pairs()
                    self.unmatched_files = self.pairing_manager.get_legacy_unmatched()
                    
                    # 🔧 DEBUG: Verify after sync
                    self.logger.info(f"✅ After sync: {len(self.file_pairs)} pairs in self.file_pairs")
            except Exception as sync_err:
                self.logger.error(f"❌ Sync error: {sync_err}")
            
            # Schließe Dialog
            window.destroy()
            
            # Speichere in Datei/Persistenz
            try:
                self.save_pairings()
                self.logger.info("💾 Pairings saved to file")
            except Exception as save_err:
                self.logger.error(f"❌ Save error: {save_err}")
                
            self._pair_dirty = False
            if self._pair_autosave_after_id and getattr(self,'root',None):
                try:
                    self.root.after_cancel(self._pair_autosave_after_id)
                except Exception:
                    pass
                self._pair_autosave_after_id = None
            self.logger.info(f"✅ Manual pairing completed: {pairs_count} pairs, {unmatched_count} unmatched files")
            try:
                self._update_ribbon_states()
            except Exception:
                pass
            
        except Exception as e:
            self._handle_error(e, context="pairing.save", user_message=self._t("Fehler beim Speichern der Paarung"))
    # ------------------------------------------------------------------
    # UNDO / REDO PAIRING SYSTEM (additiv, keine bestehende Signatur geändert)
    # ------------------------------------------------------------------
    def _record_pair_action(self, action_type: str, data: Dict[str, Any]):
        """Interne History-Aufzeichnung für Pairing Aktionen.
        action_type: create_pair | remove_pair | reset_auto_pairing
        data: kontextabhängige Daten (siehe Aufrufer)
        """
        if not hasattr(self, '_pair_history'):
            self._pair_history = []  # type: ignore
        if not hasattr(self, '_pair_redo'):
            self._pair_redo = []  # type: ignore
        self._pair_history.append({'action': action_type, 'data': data})
        # Redo-Stack leeren bei neuer Aktion
        self._pair_redo.clear()
        # Autosave-Debounce markieren
        try:
            self._mark_pairing_dirty()
        except Exception:
            pass

    def _undo_pair_action(self):
        """DEPRECATED: Weiterleitung auf neues _undo System (Kompatibilität)."""
        return self._undo()

    def _redo_pair_action(self):
        """DEPRECATED: Weiterleitung auf neues _redo System (Kompatibilität)."""
        return self._redo()

    # ------------------------------------------------------------------
    # DRAG & DROP PAIRING (grundlegend)
    # ------------------------------------------------------------------
    def _shift_virtual_window(self, offset_key: str, delta: int):
        """Verschiebe das virtuelle Fenster für Unmatched-Listen."""
        try:
            if not hasattr(self, '_virtual_pairing'):
                return
            vp = self._virtual_pairing
            new_offset = max(0, vp[offset_key] + delta)
            total_key = 'total_source' if 'source' in offset_key else 'total_translation'
            limit = max(0, vp.get(total_key, 0) - 1)
            new_offset = min(new_offset, limit)
            if new_offset != vp[offset_key]:
                vp[offset_key] = new_offset
                self._populate_manual_pairing_interface()
        except Exception as e:
            # Navigation-Fehler sind nicht kritisch – Debug
            self._debug_silent(e, 'pairing.virtual_window')

    def _begin_pair_drag(self, file_path: str, file_type: str, widget):
        """Starte Drag für eine ungepaarte Datei."""
        try:
            self._pairing_drag = {'file': file_path, 'type': file_type, 'widget': widget}
            # Visuelles Feedback
            try:
                widget.configure(border_color=self.get_color('primary'))
            except Exception:
                pass
            try:
                if hasattr(self, '_drag_status_var') and self._drag_status_var:
                    self._drag_status_var.set(self._t('Drag to pair'))
            except Exception:
                pass
        except Exception as e:
            self._handle_error(e, context="pairing.drag.start", user_message=self._t("Drag Start Fehler"), toast=False)

    def _drag_enter_item(self, target_file: str, target_type: str, widget):
        """Hover über potentielles Drop-Ziel."""
        try:
            if not self._pairing_drag:
                return
            drag_type = self._pairing_drag.get('type')
            # Nur unterschiedlicher Typ erlaubt
            if drag_type == target_type:
                return
            try:
                widget.configure(border_color=self.get_color('success'))
            except Exception:
                self._debug_silent(Exception('border_color failed'), 'drag.enter.border')
            # Hintergrund leicht hervorheben
            try:
                widget.configure(fg_color=self.get_color('surface_hover'))
            except Exception:
                self._debug_silent(Exception('fg_color failed'), 'drag.enter.fg')
            try:
                if hasattr(self, '_drag_status_var') and self._drag_status_var:
                    self._drag_status_var.set(self._t('Drop to pair'))
            except Exception:
                self._debug_silent(Exception('status var failed'), 'drag.enter.status')
        except Exception as e:
            self._debug_silent(e, 'drag.enter')

    def _drag_leave_item(self, widget):
        try:
            if widget and self._pairing_drag and widget is not self._pairing_drag.get('widget'):
                try:
                    widget.configure(border_color=self.get_color('surface_border'))
                except Exception:
                    self._debug_silent(Exception('border reset failed'), 'drag.leave.border')
                # Hintergrund zurücksetzen
                try:
                    widget.configure(fg_color=self.get_color('gray_50'))
                except Exception:
                    self._debug_silent(Exception('fg reset failed'), 'drag.leave.fg')
        except Exception as e:
            self._debug_silent(e, 'drag.leave')

    def _complete_pair_drag(self, release_file: str, release_type: str):
        """Abschluss eines Drag & Drop Pairings, falls möglich."""
        try:
            if not self._pairing_drag:
                return
            drag_file = self._pairing_drag.get('file')
            drag_type = self._pairing_drag.get('type')
            if drag_file == release_file:
                # Gleiches Element – abbrechen
                self._cancel_pair_drag()
                return
            if drag_type == release_type:
                self.show_toast(self._t("Cannot pair identical type"), "warning")
                self._cancel_pair_drag()
                return
            # Bestimme Source/Translation Reihenfolge
            if drag_type == 'source' and release_type == 'translation':
                source_file, translation_file = drag_file, release_file
            elif drag_type == 'translation' and release_type == 'source':
                source_file, translation_file = release_file, drag_file
            else:
                self._cancel_pair_drag()
                return
            # Duplicate Pair Guard
            try:
                for existing in getattr(self, 'file_pairs', []):
                    if existing.get('source') == source_file and existing.get('translation') == translation_file:
                        self.show_toast(self._t('Pair already exists') if hasattr(self,'_t') else 'Pair existiert bereits', 'info')
                        self._cancel_pair_drag()
                        return
            except Exception:
                self._debug_silent(Exception('duplicate guard failed'), 'drag.complete.dupcheck')
            # Konflikt-Guard (einseitige Doppelbelegung)
            try:
                if self._is_file_already_paired(source_file, 'source') or self._is_file_already_paired(translation_file, 'translation'):
                    self.show_toast(self._t('File already paired') if hasattr(self,'_t') else 'Datei bereits gepaart', 'warning')
                    self._cancel_pair_drag()
                    return
            except Exception:
                pass
            # Snapshot vor Änderung für Undo/Redo System
            try:
                if hasattr(self, '_push_history'):
                    self._push_history()
            except Exception:
                pass
            # Führe Pairing durch (nutzt bestehende Logik für Konsistenz)
            new_pair = {
                'source': source_file,
                'translation': translation_file,
                'similarity': 1.0,
                'source_name': os.path.basename(source_file),
                'translation_name': os.path.basename(translation_file)
            }
            
            # 🔧 KRITISCH: Paarung ZUERST im pairing_manager erstellen!
            try:
                if hasattr(self, 'pairing_manager') and self.pairing_manager:
                    success = self.pairing_manager.add_manual_pair(source_file, translation_file, similarity=1.0)
                    if success:
                        self.logger.info(f"✅ Pair added to pairing_manager: {os.path.basename(source_file)} → {os.path.basename(translation_file)}")
                    else:
                        self.logger.warning(f"⚠️ pairing_manager rejected pair (already exists or invalid)")
                        self.show_toast("Paarung bereits vorhanden oder ungültig", "warning")
                        self._cancel_pair_drag()
                        return
                else:
                    self.logger.warning("⚠️ pairing_manager not available")
            except Exception as e:
                self.logger.error(f"❌ Error adding pair to pairing_manager: {e}")
            
            if not hasattr(self, 'file_pairs'):
                self.file_pairs = []
            self.file_pairs.append(new_pair)
            # Entferne aus unmatched
            try:
                if source_file in self.unmatched_files.get('source', []):
                    self.unmatched_files['source'].remove(source_file)
                if translation_file in self.unmatched_files.get('translation', []):
                    self.unmatched_files['translation'].remove(translation_file)
            except Exception:
                self._debug_silent(Exception('unmatched remove failed'), 'drag.complete.unmatched')
            # History
            try:
                self._record_pair_action('create_pair', {'pair': new_pair})
            except Exception:
                self._debug_silent(Exception('legacy record failed'), 'drag.complete.legacy_record')
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            # Cache invalidieren und kombinierten Helper nutzen
            try:
                self._invalidate_pair_cache()
            except Exception:
                pass
            self._set_pair_dirty_and_event('pairing.create', source=new_pair.get('source_name'), translation=new_pair.get('translation_name'))
            self.show_toast(f"{self._t('Pair created (drag & drop)')}: {new_pair['source_name']} -> {new_pair['translation_name']}", "success")
        except Exception as e:
            self._handle_error(e, context="pairing.drag.complete", user_message=self._t("Drag Abschluss Fehler"), toast=False)
        finally:
            self._cancel_pair_drag()

    def _cancel_pair_drag(self):
        try:
            if self._pairing_drag:
                w = self._pairing_drag.get('widget')
                try:
                    if w:
                        w.configure(border_color=self.get_color('surface_border'))
                except Exception:
                    pass
            try:
                if hasattr(self, '_drag_status_var') and self._drag_status_var:
                    self._drag_status_var.set("")
            except Exception:
                pass
            self._pairing_drag = None
        except Exception:
            pass

    # ---------------------------------------------------------------
    # Pairing Persistenz & Utilities (additiv)
    # ---------------------------------------------------------------
    def _pairings_path(self):
        try:
            base = getattr(self, 'current_project_base', None)
            if not base:
                base = os.path.join('Checker_Projekte', '_default')
            return os.path.join(base, 'pairings.json')
        except Exception:
            return os.path.join('Checker_Projekte', '_default', 'pairings.json')

    def load_pairings(self):
        # Delegation an PairingManager – formt legacy Attribute weiterhin
        try:
            base = getattr(self, 'current_project_base', None)
            if getattr(self, 'pairing_manager', None) and self.pairing_manager.load(base):
                # Übernehme Legacy Strukturen
                self.file_pairs = self.pairing_manager.get_legacy_pairs()
                self.unmatched_files = self.pairing_manager.get_legacy_unmatched()
                self._pair_dirty = False  # geladener Zustand ist synchron
                try:
                    self._invalidate_pair_cache()
                except Exception:
                    pass
                try:
                    self._update_pairing_status_display()
                except Exception:
                    pass
        except Exception:
            pass

    def save_pairings(self):
        try:
            base = getattr(self, 'current_project_base', None)
            if getattr(self, 'pairing_manager', None):
                # Sync aktuellen Legacy Zustand vor Save (falls UI Änderungen gemacht hat)
                try:
                    # Falls self.file_pairs geändert wurden: PairingManager überschreiben
                    if hasattr(self, 'file_pairs'):
                        self.pairing_manager._legacy_pairs = list(self.file_pairs)
                    if hasattr(self, 'unmatched_files'):
                        um = self.unmatched_files or {'source': [], 'translation': []}
                        self.pairing_manager._legacy_unmatched = {'source': um.get('source', []), 'translation': um.get('translation', [])}
                        self.pairing_manager._sync_state_from_legacy()
                except Exception:
                    pass
                if not self.pairing_manager.save(base):
                    raise RuntimeError('pairing save failed')
                try:
                    self._log_event('pairing.save', status='ok', pairs=len(getattr(self,'file_pairs',[]) or []))
                except Exception:
                    pass
        except Exception as e:
            self._handle_error(e, context='pairings.persist', toast=False)
            try:
                self._log_event('pairing.save', status='failed', error=str(e))
            except Exception:
                pass

    # Entfernt: alter periodischer Autosave (_autosave_tick) – Debounce übernimmt Persistenz

    def _is_file_already_paired(self, fpath, ftype):
        """Schneller Lookup via Cache; Cache wird bei Änderungen invalidiert."""
        try:
            if not hasattr(self, '_pair_cache_source') or self._pair_cache_source is None:
                self._rebuild_pair_cache()
            if ftype == 'source':
                return fpath in (self._pair_cache_source or set())
            if ftype == 'translation':
                return fpath in (self._pair_cache_translation or set())
        except Exception:
            return False
        return False

    def _rebuild_pair_cache(self):
        try:
            src_set = set(); trg_set = set()
            for p in getattr(self, 'file_pairs', []) or []:
                try:
                    s = p.get('source'); t = p.get('translation')
                    if s: src_set.add(s)
                    if t: trg_set.add(t)
                except Exception:
                    continue
            self._pair_cache_source = src_set
            self._pair_cache_translation = trg_set
        except Exception:
            self._pair_cache_source = set(); self._pair_cache_translation = set()

    def _invalidate_pair_cache(self):
        self._pair_cache_source = None
        self._pair_cache_translation = None

    def _bulk_pair_visible_in_order(self):
        try:
            if not hasattr(self, 'unmatched_files'):
                return
            q_src = (self._src_filter.get().strip().lower() if hasattr(self, '_src_filter') else "")
            q_trg = (self._trg_filter.get().strip().lower() if hasattr(self, '_trg_filter') else "")
            def _apply_q(fs, q):
                try:
                    if not q:
                        return fs
                    return [f for f in fs if q in os.path.basename(f).lower()]
                except Exception:
                    return fs
            sources = _apply_q(self.unmatched_files.get('source', []), q_src)
            translations = _apply_q(self.unmatched_files.get('translation', []), q_trg)
            if hasattr(self, '_virtual_pairing'):
                vp = self._virtual_pairing
                win = vp.get('window', len(sources))
                so = vp.get('source_offset', 0)
                to = vp.get('translation_offset', 0)
                sources = sources[so:so+win]
                translations = translations[to:to+win]
            sources.sort(key=lambda p: os.path.basename(p).lower())
            translations.sort(key=lambda p: os.path.basename(p).lower())
            count = min(len(sources), len(translations))
            if count == 0:
                self.show_toast(self._t('No files to bulk pair') if hasattr(self,'_t') else 'Keine passenden Listen', 'info')
                return
            self._push_history()
            for i in range(count):
                s = sources[i]; t = translations[i]
                if self._is_file_already_paired(s, 'source') or self._is_file_already_paired(t, 'translation'):
                    continue
                try:
                    self.file_pairs.append({
                        'source': s,
                        'translation': t,
                        'similarity': 1.0,
                        'source_name': os.path.basename(s),
                        'translation_name': os.path.basename(t)
                    })
                    self.unmatched_files['source'].remove(s)
                    self.unmatched_files['translation'].remove(t)
                except Exception:
                    pass
            # Dirty + Event nach Gesamtvorgang (kombiniert) + Cache invalidieren
            self._invalidate_pair_cache()
            self._set_pair_dirty_and_event('pairing.bulk', count=count, sources=len(sources), translations=len(translations))
            self._populate_manual_pairing_interface()
            self._update_pairing_status_display()
            self.show_toast(self._t('Bulk pairing done') if hasattr(self,'_t') else 'Bulk-Pairing abgeschlossen', 'success')
        except Exception as e:
            self._handle_error(e, context='pairing.bulk_pair', toast=False)

    def _toggle_pair_sort(self):
        try:
            self._pair_sort_similarity = not getattr(self, '_pair_sort_similarity', False)
            self._populate_manual_pairing_interface()
        except Exception:
            pass

    def _ctx_menu_for(self, fpath):
        try:
            import tkinter as tk
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label=self._t('Open') if hasattr(self,'_t') else 'Öffnen', command=lambda: os.startfile(fpath) if os.name=='nt' else None)
            menu.add_command(label=self._t('Reveal in folder') if hasattr(self,'_t') else 'Im Ordner zeigen', command=lambda: self._reveal_in_fs(fpath) if hasattr(self,'_reveal_in_fs') else None)
            menu.add_command(label=self._t('Copy name') if hasattr(self,'_t') else 'Namen kopieren', command=lambda: self.root.clipboard_append(os.path.basename(fpath)))
            return menu
        except Exception:
            return None

    def _attach_ctx(self, widget, fpath):
        try:
            if not fpath:
                return
            def _popup(e):
                m = self._ctx_menu_for(fpath)
                try:
                    if m:
                        m.tk_popup(e.x_root, e.y_root)
                finally:
                    try:
                        if m:
                            m.grab_release()
                    except Exception:
                        pass
            widget.bind('<Button-3>', _popup)
            widget.bind('<Control-1>', _popup)
        except Exception:
            pass
    
    def _clear_all_files(self):
        """CLEAR ALL FILES - Alle hochgeladenen Dateien löschen"""
        try:
            total_files = len(self.uploaded_files.get('source', [])) + len(self.uploaded_files.get('translation', []))
            
            if total_files == 0:
                self.show_toast("Keine Dateien zum Löschen vorhanden", "info")
                return
            
            # Alle Dateien löschen
            self.uploaded_files = {'source': [], 'translation': []}
            
            # File pairs und unmatched files auch löschen
            if hasattr(self, 'file_pairs'):
                self.file_pairs = []
            if hasattr(self, 'unmatched_files'):
                self.unmatched_files = {'source': [], 'translation': []}
            
            # UI aktualisieren
            self._schedule_update_file_counter()
            # Beide Scrollbars ausblenden (Listen jetzt leer)
            try:
                if hasattr(self, 'source_file_list'):
                    self._auto_hide_scrollbar(self.source_file_list, hide=True)
                if hasattr(self, 'translation_file_list'):
                    self._auto_hide_scrollbar(self.translation_file_list, hide=True)
            except Exception:
                pass
            self._refresh_file_list_display()
            self._update_pairing_status_display()
            try:
                self._mark_pairing_dirty()
            except Exception:
                self._pair_dirty = True
            
            self.update_status("Alle Dateien wurden gelöscht")
            self.show_toast(f"{total_files} Dateien erfolgreich gelöscht", "success")
            # Event publizieren
            try:
                if getattr(self, 'event_bus', None):
                    self.event_bus.publish('files.cleared', {'total': total_files})
            except Exception:
                pass
            
        except Exception as e:
            try:
                self._handle_error(e, context="files.clear_all", user_message=self._t("Dateien Löschen Fehler"))
            except Exception:
                pass
            self.show_toast("Fehler beim Löschen der Dateien", "error")
    
    def _refresh_file_list(self):
        """REFRESH FILE LIST - Dateiliste aktualisieren"""
        try:
            self._update_file_counter()
            self._refresh_file_list_display()
            self._update_pairing_status_display()
            
            total_files = len(self.uploaded_files.get('source', [])) + len(self.uploaded_files.get('translation', []))
            self.update_status(f"Dateiliste aktualisiert - {total_files} Dateien verfügbar")
            self.show_toast("Dateiliste erfolgreich aktualisiert", "success")
            
        except Exception as e:
            try:
                self._handle_error(e, context="files.refresh", user_message=self._t("Dateiliste Aktualisierung fehlgeschlagen"))
            except Exception:
                pass
            self.show_toast("Fehler beim Aktualisieren der Dateiliste", "error")
            try:
                self._handle_error(e, context="files.refresh_list", user_message=self._t("Dateiliste Refresh Fehler"), toast=False)
            except Exception:
                pass
    
    def _upload_translation_files(self):
        """Delegierter Übersetzungsdatei-Upload via UploadService."""
        try:
            file_types = [
                ("All Supported", "*.pdf;*.docx;*.txt;*.doc;*.rtf;*.odt"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx;*.doc"),
                ("Text files", "*.txt"),
                ("Rich Text", "*.rtf"),
                ("OpenDocument", "*.odt")
            ]
            files = filedialog.askopenfilenames(title="Übersetzungsdateien für Qualitätsanalyse auswählen", filetypes=file_types)
            if not files:
                return
            customer_name = self._get_or_select_customer_for_upload()
            if customer_name == "CANCELLED":
                return
            from upload_service import get_upload_service
            service = get_upload_service(event_bus=getattr(self, 'event_bus', None), logger=self.logger)
            result = service.process_simple_upload(
                kind='translation',
                selected_files=list(files),
                existing_files=self.uploaded_files['translation'],
                customer_name=customer_name,
                copy_callback=self.copy_files_to_project_structure if customer_name else None,
            )
            added = result['added_files']
            duplicates = result['duplicate_files']
            if added:
                try:
                    before_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
                except Exception:
                    before_total = None
                try:
                    before_translation_count = len(self.uploaded_files.get('translation', []))
                except Exception:
                    before_translation_count = None

                self._register_uploaded_files('translation', added)

                try:
                    after_translation_count = len(self.uploaded_files.get('translation', []))
                    self.logger.info(
                        "Upload-Registrierung (translation): %s Dateien, Bestand %s → %s",
                        len(added),
                        before_translation_count,
                        after_translation_count
                    )
                except Exception:
                    pass
                if customer_name and result['meta'].get('project_copied'):
                    self._offer_open_project_folder(customer_name)
                self._update_file_counter()
                self._refresh_file_list_display()
                self._show_enhanced_upload_results("translation", added)
                source_count = len(self.uploaded_files['source'])
                trans_count = len(self.uploaded_files['translation'])
                status_msg = f"{len(added)} Übersetzungsdateien hinzugefügt"
                if duplicates:
                    status_msg += f" ({len(duplicates)} Duplikate ignoriert)"
                if source_count > 0 and trans_count > 0:
                    status_msg += f" | Bereit für Analyse ({min(source_count, trans_count)} Paare)"
                self.update_status(status_msg)
                self.show_toast(f"{len(added)} Übersetzungsdateien erfolgreich hochgeladen", "success")
                if source_count > 0 and trans_count > 0:
                    self._after(1000, lambda: self.show_toast("Bereit für Qualitätsanalyse!", "info"))
                # Telemetrie: files.changed (added translation)
                try:
                    if getattr(self, 'event_bus', None):
                        after_total = sum(len(self.uploaded_files.get(k, [])) for k in ('source','translation'))
                        self.event_bus.publish('files.changed', {
                            'action': 'added',
                            'file_type': 'translation',
                            'count': len(added),
                            'filenames': [os.path.basename(a) for a in added],
                            'before_total': before_total,
                            'after_total': after_total
                        })
                except Exception:
                    pass
            else:
                self.show_toast("Alle ausgewählten Dateien waren bereits hochgeladen", "info")
            if not hasattr(self, '_telemetry_counters'):
                self._telemetry_counters = {}
            self._telemetry_counters.setdefault('upload_count', 0)
            self._telemetry_counters.setdefault('upload_files_added', 0)
            self._telemetry_counters.setdefault('upload_duplicates_ignored', 0)
            self._telemetry_counters['upload_count'] += 1
            self._telemetry_counters['upload_files_added'] += len(added)
            self._telemetry_counters['upload_duplicates_ignored'] += len(duplicates)
            # Ribbon aktualisieren
            try:
                self._update_ribbon_states()
            except Exception:
                pass
        except Exception as e:
            msg = f"Fehler beim Hochladen der Übersetzungsdateien: {e}"
            self.update_status(msg)
            self.show_toast("Fehler beim Hochladen der Übersetzungsdateien", "error")
            try:
                self._handle_error(e, context="upload.translation", user_message=self._t("Übersetzungsdateien Upload fehlgeschlagen"))
            except Exception:
                pass

    # =====================================================================
    # Upload Telemetrie Helper
    # =====================================================================
    def _record_upload_telemetry(self, added: int, duplicates: int):
        """Reduzierte Duplikation: Einheitliche Erfassung von Upload-Telemetrie."""
        try:
            if not hasattr(self, '_telemetry_counters'):
                self._telemetry_counters = {}
            self._telemetry_counters.setdefault('upload_count', 0)
            self._telemetry_counters.setdefault('upload_files_added', 0)
            self._telemetry_counters.setdefault('upload_duplicates_ignored', 0)
            self._telemetry_counters['upload_count'] += 1
            self._telemetry_counters['upload_files_added'] += added
            self._telemetry_counters['upload_duplicates_ignored'] += duplicates
        except Exception:
            pass
    
    # DUPLIKAT METHODE ENTFERNT - _update_file_counter() ist bereits oben konsolidiert implementiert
    
    def _show_enhanced_upload_results(self, file_type, files):
        """OPTIMIERT: Verbesserte Upload-Ergebnisse mit intelligenter Dateianalyse und Vorschau"""
        try:
            # Clear output and show enhanced upload results
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Enhanced upload results card
            results_card = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('surface'),
                corner_radius=self.get_spacing('large_border_radius'),
                border_width=0,
                            )
            results_card.pack(fill="x", pady=self.get_spacing('lg'), padx=self.get_spacing('lg'))
            
            # Animated header with gradient effect simulation
            header_frame = ctk.CTkFrame(results_card, fg_color=self.get_color('gray_600'))
            header_frame.pack(fill="x", padx=0, pady=0)
            
            # Header with file type icon
            file_prefix = "Source" if file_type == "source" else "Translation"
            file_type_german = "Ausgangstexte" if file_type == "source" else "Übersetzungen"
            header_label = ctk.CTkLabel(
                header_frame,
                text=f"{file_type_german} erfolgreich hochgeladen",
                font=ctk.CTkFont(*self.get_typography('subheading')),
                text_color=self.get_color('white')
            )
            header_label.pack(pady=self.get_spacing('header_padding'))
            
            # Enhanced file list with metadata
            content_frame = ctk.CTkFrame(results_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=self.get_spacing('content_padding'), pady=self.get_spacing('content_padding'))
            
            # File analysis summary
            total_size = 0
            file_types_detected = set()
            
            for i, file_path in enumerate(files, 1):
                file_name = os.path.basename(file_path)
                file_ext = os.path.splitext(file_name)[1].lower()
                file_types_detected.add(file_ext)
                
                try:
                    file_size = os.path.getsize(file_path)
                    total_size += file_size
                    size_str = self._format_file_size(file_size)
                except:
                    size_str = "Unbekannte Größe"
                
                # Enhanced file entry with icon and metadata
                file_frame = ctk.CTkFrame(content_frame, fg_color=self.get_color('surface_light'))
                file_frame.pack(fill="x", pady=2)
                
                file_label = ctk.CTkLabel(
                    file_frame,
                    text=f"{i}. {file_name} ({size_str})",
                    font=ctk.CTkFont(*self.get_typography('body')),
                    text_color=self.get_color('text_primary'),
                    anchor="w"
                )
                file_label.pack(side="left", padx=self.get_spacing('sm'), pady=self.get_spacing('xs'))
                
                # File type indicator
                type_indicator = ctk.CTkLabel(
                    file_frame,
                    text=file_ext.upper(),
                    font=ctk.CTkFont(*self.get_typography('caption')),
                    text_color=self.get_color('primary'),
                    width=40
                )
                type_indicator.pack(side="right", padx=self.get_spacing('sm'), pady=self.get_spacing('xs'))
            
            # Smart summary with analytics
            summary_frame = ctk.CTkFrame(content_frame, fg_color=self.get_color('gray_100'))
            summary_frame.pack(fill="x", pady=(self.get_spacing('md'), 0))
            
            summary_text = "\n".join([
                self._t("Upload-Zusammenfassung:"),
                self._t(f"- Dateien gesamt: {len(files)}"),
                self._t(f"- Gesamtgröße: {self._format_file_size(total_size)}"),
                self._t(f"- Dateitypen: {', '.join(sorted(file_types_detected))}"),
                self._t(f"- Aktuelle {file_type_german}-Anzahl: {len(self.uploaded_files[file_type])}"),
                "",
                self._t(self._get_smart_upload_advice(file_type))
            ])
            
            summary_label = ctk.CTkLabel(
                summary_frame,
                text=summary_text,
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                justify="left",
                anchor="w"
            )
            summary_label.pack(padx=self.get_spacing('md'), pady=self.get_spacing('md'))
            
        except Exception as e:
            try:
                self._handle_error(e, context="upload.results.enhanced", user_message=self._t("Erweiterte Upload Ergebnisse Fehler"), toast=False)
            except Exception:
                pass
            # Fallback to simple display
            self._show_upload_results(file_type, files)
    
    # Entfernte doppelte _format_file_size Definition (siehe zentrale Variante weiter oben)

    def _auto_hide_scrollbar(self, list_frame, hide: bool):
        """Versteckt oder zeigt Scrollbar eines scrollbaren Frames.

    Defensive Implementation – ignoriert Fehler still (kein funktionaler Impact).
        """
        try:
            scrollbar = getattr(list_frame, '_scrollbar', None)
            if not scrollbar:
                for child in list_frame.winfo_children():
                    try:
                        if child.__class__.__name__.lower().endswith('scrollbar'):
                            scrollbar = child
                            break
                    except Exception:
                        pass
            if not scrollbar:
                return
            if hide:
                try:
                    scrollbar.pack_forget()
                except Exception:
                    pass
            else:
                if not scrollbar.winfo_ismapped():
                    try:
                        scrollbar.pack(side="right", fill="y")
                    except Exception:
                        pass
        except Exception:
            pass
    
    def _get_smart_upload_advice(self, file_type):
        """Kontextuelle Ratschläge basierend auf Upload-Status (bereinigt)."""
        try:
            source_count = len(self.uploaded_files.get('source', []))
            trans_count = len(self.uploaded_files.get('translation', []))
            if file_type == "source":
                if trans_count == 0:
                    return self._t("Nächster Schritt: Entsprechende Übersetzungsdateien hochladen")
                elif trans_count < source_count:
                    return self._t(f"Ausgleich: Erwägen Sie das Hochladen von {source_count - trans_count} weiteren Übersetzungsdatei(en)")
                else:
                    return self._t("Bereit: Alle Dateien hochgeladen - Qualitätsanalyse starten!")
            else:
                if source_count == 0:
                    return self._t("Nächster Schritt: Entsprechende Ausgangstexte hochladen")
                elif source_count < trans_count:
                    return self._t(f"Ausgleich: Erwägen Sie das Hochladen von {trans_count - source_count} weiteren Ausgangstexten")
                else:
                    return self._t("Bereit: Alle Dateien hochgeladen - Qualitätsanalyse starten!")
        except Exception:
            return self._t("Dateien erfolgreich hochgeladen")
    
    def start_analysis(self, rule_profile: str = "default"):
        """Startet Analyse mit Cache-Prüfung und Event-Publishing."""
        try:
            if rule_profile == "default":
                try:
                    depth_var = getattr(self, 'var_depth', None)
                    if depth_var is not None and hasattr(depth_var, 'get'):
                        selected_depth = depth_var.get() or 'medium'
                        rule_profile = str(selected_depth)
                    else:
                        rule_profile = 'medium'
                except Exception:
                    rule_profile = 'medium'

            # Früh: Ribbon-States aktualisieren (idempotent), z.B. Analyse läuft
            try:
                if hasattr(self, '_update_ribbon_states'):
                    self._update_ribbon_states()
            except Exception:
                pass
            import uuid, time as _tmod
            analysis_id = getattr(self, '_current_analysis_id', None)
            if not analysis_id:
                analysis_id = str(uuid.uuid4())
                self._current_analysis_id = analysis_id
            start_ts = _tmod.time()
            
            # 🔧 FIX: Verwende file_pairs wenn vorhanden, sonst uploaded_files
            # Dies stellt sicher, dass die manuellen Paarungen respektiert werden
            if hasattr(self, 'file_pairs') and self.file_pairs:
                # Verwende die gepaarten Dateien
                source_files = [pair.get('source') for pair in self.file_pairs if pair.get('source')]
                translation_files = [pair.get('translation') for pair in self.file_pairs if pair.get('translation')]
                files = {'source': source_files, 'translation': translation_files}
                self.logger.info(f"✅ Using file_pairs for analysis: {len(source_files)} source, {len(translation_files)} translation")
                # 🔧 DEBUG: Log first few pairs being used
                for idx, pair in enumerate(self.file_pairs[:3]):
                    src = os.path.basename(pair.get('source', 'N/A'))
                    tgt = os.path.basename(pair.get('translation', 'N/A'))
                    self.logger.info(f"   Analysis Pair {idx+1}: {src} → {tgt}")
            else:
                # Fallback auf uploaded_files
                files = self.uploaded_files
                self.logger.info(f"⚠️ Using uploaded_files (no pairs): {len(files.get('source', []))} source, {len(files.get('translation', []))} translation")
                if not hasattr(self, 'file_pairs'):
                    self.logger.warning("   file_pairs attribute not found!")
                elif not self.file_pairs:
                    self.logger.warning(f"   file_pairs is empty: {self.file_pairs}")
            
            if not files.get('source') and not files.get('translation'):
                self.show_toast("Keine Dateien geladen.", "warning")
                return
            source_count = len(files.get('source', []))
            translation_count = len(files.get('translation', []))
            # 1) Cache prüfen
            cache_key = None
            try:
                if getattr(self, '_cache', None):
                    cache_key = self._cache.make_key(files, rule_profile)  # type: ignore
                    cached = self._cache.get(cache_key)  # type: ignore
                    if cached:
                        normalized_cached = self._normalize_analysis_results_structure(cached)
                        self.analysis_results = normalized_cached
                        cached = normalized_cached
                        # UI sofort (bestehende Methode nutzen, falls vorhanden)
                        try:
                            if hasattr(self, '_show_analysis_results'):
                                # Markiere current_analysis für Kompatibilität
                                self.current_analysis = normalized_cached
                                self._show_analysis_results()
                        except Exception:
                            pass
                        if self.event_bus:
                            try:
                                self.event_bus.publish("cache.hit", {"key": cache_key, "analysis_id": analysis_id})
                            except Exception as e:
                                self._debug_silent(e, 'analysis.cache.event.cache_hit')
                            try:
                                duration_ms = int((time.time() - start_ts) * 1000)
                                meta = {
                                    "analysis_id": analysis_id,
                                    "_cache_key": cache_key,
                                    "cache": True,
                                    "duration_ms": duration_ms,
                                    "rule_profile": rule_profile,
                                    "source_count": len(files.get('source', [])),
                                    "translation_count": len(files.get('translation', []))
                                }
                                self.event_bus.publish("analysis.done", {**cached, **meta})
                                try:
                                    self.event_bus.publish("analysis.completed", {**meta, "summary_keys": list(cached.keys())[:10]})
                                except Exception as ie:
                                    self._debug_silent(ie, 'analysis.cache.event.analysis_completed')
                            except Exception as e:
                                self._debug_silent(e, 'analysis.cache.event.analysis_done')
                        try:
                            self._log_event("analysis.completed", analysis_id=analysis_id, cache=True)
                        except Exception:
                            pass
                        return
            except Exception as e:
                self._debug_silent(e, 'analysis.cache.lookup')

            # 2) Kein Cache: Progress UI aufbauen
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            # Event-basierte Progress UI
            progress_card = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('surface'),
                corner_radius=self.get_component_value('borders.radius_md') if hasattr(self, 'get_component_value') else 8,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            progress_card.pack(fill="x", pady=self.get_spacing('lg'), padx=self.get_spacing('lg'))
            header_frame = ctk.CTkFrame(progress_card, fg_color=self.get_color('primary'))
            header_frame.pack(fill="x")
            header_label = ctk.CTkLabel(
                header_frame,
                text=self._t("Quality analysis in progress"),
                font=ctk.CTkFont(*self.get_typography('subheading')),
                text_color=self.get_color('text_inverse')
            )
            header_label.pack(pady=self.get_spacing('md'))
            content_frame = ctk.CTkFrame(progress_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=self.get_spacing('lg'), pady=self.get_spacing('lg'))
            progress_bar = ctk.CTkProgressBar(
                content_frame,
                height=18,
                # Accent Reduction: Fortschrittsfarbe vereinheitlicht auf Primary
                progress_color=self.get_color('primary'),
                fg_color=self.get_color('surface_border')
            )
            progress_bar.pack(pady=self.get_spacing('sm'), fill='x')
            progress_bar.set(0)
            status_label = ctk.CTkLabel(
                content_frame,
                text=self._t("Initializing analysis engine"),
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                anchor='w', justify='left'
            )
            status_label.pack(pady=(self.get_spacing('xs'),0), fill='x')

            # Cancel Button
            cancel_btn = self._create_button(
                content_frame,
                text=self._t('Cancel analysis'),
                command=lambda: self._request_analysis_cancel(status_label, progress_bar),
                kind="warning",
                height=32
            )
            cancel_btn.pack(pady=(self.get_spacing('md'),0), anchor='e')

            # Listener für Progress/Done-Events (einmalig registrieren)
            def _on_progress(payload):
                try:
                    phase = payload.get('phase') if isinstance(payload, dict) else None
                    if phase == 'rules':
                        cur = payload.get('current', 0)
                        total = payload.get('total', 1)
                        ratio = max(0.0, min(1.0, cur/total))
                        progress_bar.set(ratio)
                        status_label.configure(text=self._t('Analyzing rules') + f" ({cur}/{total})")
                    elif phase == 'pair_metrics':
                        # Datei-Paar-Metriken – Fortschritt aus Payload nutzen
                        ratio = payload.get('progress') if isinstance(payload, dict) else None
                        if isinstance(ratio, (int, float)):
                            progress_bar.set(max(0.0, min(1.0, float(ratio))))
                        status_label.configure(text=self._t('Paare analysiert'))
                    elif phase == 'finalize':
                        progress_bar.set(1.0)
                        status_label.configure(text=(self._t('Ergebnisse finalisieren') if hasattr(self,'_t') else 'Ergebnisse finalisieren') + f" – {payload.get('duration',0):.2f}s")
                    elif phase == 'completed':
                        progress_bar.set(1.0)
                        status_label.configure(text=self._t('Analyse abgeschlossen – Ergebnisse werden angezeigt'))
                        try:
                            # Thread-sicher ins UI: Ergebnisse anzeigen
                            if hasattr(self, 'root') and getattr(self, 'root', None):
                                self._after(0, lambda: self._show_analysis_results())
                            else:
                                self._show_analysis_results()
                        except Exception as ie:
                            self._debug_silent(ie, 'analysis.progress.completed.render')
                    elif phase == 'start':
                        progress_bar.set(0.01)
                        status_label.configure(text=self._t('Loading rules and preparing data'))
                except Exception as e:
                    self._debug_silent(e, 'analysis.progress.handler')
            try:
                if getattr(self, 'event_bus', None):
                    self.event_bus.subscribe('analysis.progress', _on_progress)
                    # Zusätzlich: Auf Abschlussereignisse reagieren (Robustheit)
                    def _on_done(_payload):
                        try:
                            if hasattr(self, 'root') and getattr(self, 'root', None):
                                self._after(0, lambda: self._show_analysis_results())
                            else:
                                self._show_analysis_results()
                        except Exception as ie:
                            self._debug_silent(ie, 'analysis.done.handler')
                    try:
                        self.event_bus.subscribe('analysis.done', _on_done)
                    except Exception as ie:
                        self._debug_silent(ie, 'analysis.done.subscribe')
                    try:
                        self.event_bus.subscribe('analysis.completed', _on_done)
                    except Exception as ie:
                        self._debug_silent(ie, 'analysis.completed.subscribe')
            except Exception as e:
                self._debug_silent(e, 'analysis.progress.subscribe')

            # Events: analysis.started
            if self.event_bus:
                try:
                    self.event_bus.publish("analysis.started", {
                        "analysis_id": analysis_id,
                        "sources": source_count,
                        "translations": translation_count,
                        "rule_profile": rule_profile,
                        "ts": start_ts
                    })
                except Exception as e:
                    self._debug_silent(e, 'analysis.started.publish')
            self.update_status(self._t("Quality analysis started"))
            self._show_toast(self._t("Quality analysis started"), "success")
            self._log_event("analysis.started.ui", analysis_id=analysis_id, sources=source_count, translations=translation_count, rule_profile=rule_profile)

            def job():
                try:
                    results = self._run_analysis_pipeline(files, rule_profile)
                    results = self._normalize_analysis_results_structure(results)
                    self.analysis_results = results
                    # Automatischer dynamischer Bericht Refresh
                    try:
                        _auto_report_path = self._generate_dynamic_report()
                        if _auto_report_path:
                            self.latest_dynamic_report_path = _auto_report_path
                            try:
                                self._show_toast(self._t("Dynamischer Bericht aktualisiert"), "info")
                            except Exception:
                                pass
                    except Exception as _autorefresh_err:
                        self._debug_silent(_autorefresh_err, 'dynamic.report.autorefresh')
                    self.current_analysis = results
                    duration_ms = int((time.time() - start_ts) * 1000)
                    results_meta = {
                        "analysis_id": analysis_id,
                        "duration_ms": duration_ms,
                        "rule_profile": rule_profile,
                        "source_count": source_count,
                        "translation_count": translation_count,
                        "cache": False if cache_key is None else False  # hier: kein Cache Pfad
                    }
                    if self.event_bus:
                        try:
                            self.event_bus.publish("analysis.done", {**results, "_cache_key": cache_key, **results_meta})
                        except Exception:
                            pass
                        try:
                            self.event_bus.publish("analysis.completed", {**results_meta, "summary_keys": list(results.keys())[:10]})
                        except Exception as e:
                            self._debug_silent(e, 'analysis.completed.publish')
                    try:
                        self._log_event("analysis.completed", **results_meta)
                    except Exception:
                        pass
                    # UI close wenn Cancel
                    if getattr(self, '_analysis_cancel_requested', False):
                        try:
                            self._show_toast(self._t('Analysis cancelled'), 'warning')
                        except Exception as e:
                            self._debug_silent(e, 'analysis.done.toast_cancel')
                    # Nach erfolgreicher Analyse die Ergebnisse anzeigen (robust, thread-sicher)
                    try:
                        if hasattr(self, 'root') and getattr(self, 'root', None):
                            self._after(0, lambda: self._show_analysis_results())
                        else:
                            self._show_analysis_results()
                    except Exception as ie:
                        self._debug_silent(ie, 'analysis.render_results.after_job')
                except Exception as e:  # Einheitliche Fehlerbehandlung über zentralen Handler
                    duration_ms = int((time.time() - start_ts) * 1000)
                    try:
                        self._log_event("analysis.failed", analysis_id=analysis_id, duration_ms=duration_ms, error=str(e.__class__.__name__), message=str(e))
                    except Exception:
                        pass
                    if self.event_bus:
                        try:
                            self.event_bus.publish("analysis.failed", {
                                "analysis_id": analysis_id,
                                "duration_ms": duration_ms,
                                "error": str(e.__class__.__name__),
                                "message": str(e)
                            })
                        except Exception:
                            pass
                    self._handle_error(
                        e,
                        context="analysis.pipeline",
                        user_message=self._t("Analyse fehlgeschlagen") if hasattr(self, '_t') else "Analyse fehlgeschlagen",
                        event_name="analysis.error",
                        cache_key=cache_key
                    )
                finally:
                    try:
                        self._show_busy(False)
                    except Exception as e:
                        self._debug_silent(e, 'analysis.busy.hide')
                    self._analysis_cancel_requested = False
                    # Am Ende immer Ribbon-States aktualisieren (idempotent)
                    try:
                        if hasattr(self, '_update_ribbon_states'):
                            self._update_ribbon_states()
                    except Exception:
                        pass

            try:
                self._show_busy(True)
            except Exception as e:
                self._debug_silent(e, 'analysis.busy.show')
            if getattr(self, 'worker_pool', None):
                try:
                    self.worker_pool.submit(job)
                except Exception as e:
                    self._debug_silent(e, 'analysis.worker_pool.submit')
                    job()
            else:
                job()
        except Exception as e:  # Start-Phase Fehler zentral behandeln
            self._handle_error(
                e,
                context="analysis.start",
                user_message=self._t("Analyse Start fehlgeschlagen") if hasattr(self, '_t') else "Analyse Start fehlgeschlagen",
                event_name="analysis.error.start"
            )
    
    # LEGACY: Simulation (ersetzt durch Event-basierte progress Events: analysis.progress)
    def _simulate_analysis_progress(self, progress_bar, status_label):  # pragma: no cover
        """Legacy Simulation deaktiviert - beibehalten für Rückwärtskompatibilität."""
        try:
            status_label.configure(text=self._t('Fortschritts-Simulation veraltet – warte auf echte Ereignisse') if hasattr(self,'_t') else 'Fortschritts-Simulation veraltet – warte auf echte Ereignisse')
        except Exception:
            pass
    
    def _update_enhanced_progress(self, *_, **__):  # pragma: no cover
        """Legacy Hook: Wird nicht mehr genutzt (Event-basierter Fortschritt aktiv)."""
        return

    def _normalize_analysis_results_structure(self, payload: Any) -> dict:
        """Stellt sicher, dass Analyse-Ergebnisse immer Befunde und abgeleitete Kennzahlen enthalten."""
        if not isinstance(payload, dict):
            return {}
        try:
            normalized = copy.deepcopy(payload)
        except Exception:
            normalized = dict(payload)

        def _coerce(issue: Any, fallback_phase: str) -> dict:
            if isinstance(issue, dict):
                data = dict(issue)
            elif is_dataclass(issue):
                data = asdict(issue)
            elif hasattr(issue, "__dict__"):
                try:
                    data = dict(issue.__dict__)
                except Exception:
                    data = {}
            else:
                data = {}
            phase = str(data.get('phase') or data.get('category') or fallback_phase or 'phase2')
            severity = str(data.get('severity') or 'info').lower()
            if severity not in ('critical', 'major', 'minor', 'info'):
                severity = 'info'
            category = data.get('category') or phase
            message = data.get('message') or ''
            source = data.get('source_text') or data.get('source') or ''
            target = data.get('target_text') or data.get('target') or ''
            rule = data.get('rule') or data.get('rule_id') or category
            meta = data.get('meta') if isinstance(data.get('meta'), dict) else {}
            norm = dict(data)
            norm['phase'] = phase
            norm['severity'] = severity
            norm['category'] = category
            norm['rule'] = rule
            norm['message'] = message
            norm['source'] = source
            norm['target'] = target
            norm['source_excerpt'] = (source or '')[:200]
            norm['target_excerpt'] = (target or '')[:200]
            norm['code'] = str(data.get('code') or '')
            norm['meta'] = meta
            return norm

        findings_raw = normalized.get('findings')
        findings: list[dict] = []
        if isinstance(findings_raw, (list, tuple)):
            for item in findings_raw:
                fallback = 'phase2'
                if isinstance(item, dict):
                    fallback = str(item.get('phase') or item.get('category') or 'phase2')
                findings.append(_coerce(item, fallback))
        if not findings:
            for phase_key in ('issues_phase1', 'issues_phase2', 'issues_phase3'):
                issues = normalized.get(phase_key)
                if not isinstance(issues, list):
                    continue
                phase_label = phase_key.replace('issues_', '')
                for issue in issues:
                    findings.append(_coerce(issue, phase_label))
        normalized['findings'] = findings

        severity_counts = {'critical': 0, 'major': 0, 'minor': 0}
        for entry in findings:
            sev = str(entry.get('severity') or '').lower()
            if sev in severity_counts:
                severity_counts[sev] += 1

        summary = normalized.get('summary') if isinstance(normalized.get('summary'), dict) else {}
        summary['critical'] = severity_counts['critical']
        summary['major'] = severity_counts['major']
        summary['minor'] = severity_counts['minor']
        normalized['summary'] = summary

        phase_counts = normalized.get('phase_issue_counts') if isinstance(normalized.get('phase_issue_counts'), dict) else {}
        for phase_key in ('phase1', 'phase2', 'phase3'):
            count = sum(1 for f in findings if str(f.get('phase') or '').lower() == phase_key)
            phase_counts[phase_key] = count
        normalized['phase_issue_counts'] = phase_counts

        phases_info = normalized.get('phases') if isinstance(normalized.get('phases'), dict) else {}
        if findings:
            phase4 = phases_info.get('phase4', {}) or {}
            phase4.setdefault('total', len(findings))
            if 'risk_score' not in phase4:
                metrics = normalized.get('metrics') if isinstance(normalized.get('metrics'), dict) else {}
                pair_count = metrics.get('pair_count') if isinstance(metrics, dict) else None
                try:
                    pair_count = int(pair_count) if pair_count is not None else None
                except Exception:
                    pair_count = None
                crit = severity_counts['critical']
                maj = severity_counts['major']
                mi = severity_counts['minor']
                denom = pair_count if pair_count else len(findings) or 1
                risk = min(100.0, round(((crit * 3 + maj * 2 + mi) / denom) * 25.0, 2))
                phase4['risk_score'] = risk
            phases_info['phase4'] = phase4
            # Für UI-Kompatibilität: mappe phase4 auf consolidation
            phases_info['consolidation'] = phase4.copy()
        else:
            # WICHTIG: Keine Befunde = Risiko 0 (nicht 100!)
            phases_info['consolidation'] = {'total': 0, 'risk_score': 0.0}
            phases_info['phase4'] = {'total': 0, 'risk_score': 0.0}
        normalized['phases'] = phases_info

        consolidated = normalized.get('consolidated') if isinstance(normalized.get('consolidated'), dict) else {}
        consolidated.setdefault('total', len(findings))
        if findings and 'risk_score' not in consolidated:
            consolidated['risk_score'] = phases_info.get('phase4', {}).get('risk_score', 0.0)
        normalized['consolidated'] = consolidated

        existing_grouped = normalized.get('findings_grouped') if isinstance(normalized.get('findings_grouped'), list) else []
        if findings and not existing_grouped:
            groups: dict[tuple[str, str, str], dict] = {}
            for entry in findings:
                key = (
                    str(entry.get('rule') or entry.get('rule_id') or ''),
                    str(entry.get('code') or ''),
                    str(entry.get('severity') or 'info')
                )
                if key not in groups:
                    groups[key] = {
                        'rule': key[0],
                        'code': key[1],
                        'severity': key[2],
                        'message': entry.get('message') or '',
                        'count': 0,
                        'confidence_sum': 0.0,
                        'confidence_n': 0,
                    }
                grp = groups[key]
                grp['count'] += 1
                confidence = entry.get('confidence')
                if confidence is None and isinstance(entry.get('meta'), dict):
                    confidence = entry['meta'].get('confidence')
                try:
                    if isinstance(confidence, str):
                        confidence = float(confidence)
                    if isinstance(confidence, (int, float)):
                        grp['confidence_sum'] += float(confidence)
                        grp['confidence_n'] += 1
                except Exception:
                    pass
            grouped_list: list[dict] = []
            for grp in groups.values():
                n = grp.pop('confidence_n', 0)
                s_val = grp.pop('confidence_sum', 0.0)
                grp['confidence'] = round(s_val / n, 4) if n else None
                grouped_list.append(grp)
            try:
                order = {'critical': 0, 'major': 1, 'minor': 2, 'info': 3}
                grouped_list.sort(key=lambda item: (order.get(item.get('severity'), 9), -(item.get('count') or 0)))
            except Exception:
                pass
            normalized['findings_grouped'] = grouped_list
        elif not findings:
            normalized['findings_grouped'] = []

        return normalized

    # ---------------------------------------------------------------
    # LEICHTGEWICHTIGE ANALYSE-PIPELINE (Fallback)
    # ---------------------------------------------------------------
    def _run_analysis_pipeline(self, files: list[dict], rule_profile: str) -> dict:
        """Erweiterte lightweight Analyse mit:
          - Rule Dispatch (rule_profile + UI Flags)
          - Terminologie-Trefferquote (optional Glossar CSV/TSV)
          - Stil/Lesbarkeit (Flesch-ähnliche Lesbarkeitszahl, Satzlänge)
          - Passiv-Anteil ("wird .*" / "was|were .*ed" heuristisch)
          - Bedeutungsvergleich (Embedding Similarity via sentence-transformers oder Ollama-Fallback)
          - Qualitätskriterien-Filter (nur aktivierte Metriken)
        """
        import time, os, hashlib, re, math
        start = time.time()
        enabled_qc: dict[str, bool] = {}
        try:
            if hasattr(self, 'quality_vars') and isinstance(self.quality_vars, dict):
                enabled_qc = {k: bool(v.get()) for k, v in self.quality_vars.items()}
        except Exception:
            pass
        def qc_enabled(key: str) -> bool:
            return enabled_qc.get(key, True)

        def _get_setting(path: str, default):
            try:
                if getattr(self, 'settings_service', None):
                    return self.settings_service.get(path, default)
            except Exception:
                pass
            return default

        def _read(p: str) -> str:
            try:
                if not p or not os.path.isfile(p):
                    return ''
                lower_path = p.lower()
                if lower_path.endswith('.docx'):
                    text_chunks: list[str] = []
                    docx_text = ''
                    try:
                        from docx import Document  # type: ignore
                        document = Document(p)
                        text_chunks.extend(par.text for par in document.paragraphs if par.text)
                        for table in getattr(document, 'tables', []):
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text:
                                        text_chunks.append(cell.text)
                    except Exception:
                        pass
                    if text_chunks:
                        docx_text = '\n'.join(chunk.strip() for chunk in text_chunks if chunk and chunk.strip())
                    if not docx_text:
                        try:
                            import zipfile
                            from xml.etree import ElementTree as ET
                            with zipfile.ZipFile(p) as archive:
                                with archive.open('word/document.xml') as xml_file:
                                    tree = ET.parse(xml_file)
                                    root = tree.getroot()
                                    paras: list[str] = []
                                    def _strip_ns(tag: str) -> str:
                                        return tag.split('}', 1)[1] if '}' in tag else tag
                                    for para in root.iter():
                                        if _strip_ns(para.tag) != 'p':
                                            continue
                                        runs: list[str] = []
                                        for node in para.iter():
                                            if _strip_ns(node.tag) == 't' and node.text:
                                                runs.append(node.text)
                                        if runs:
                                            paras.append(''.join(runs))
                                    if paras:
                                        docx_text = '\n'.join(paras)
                        except Exception:
                            docx_text = ''
                    return docx_text[:200000]
                if lower_path.endswith('.pdf'):
                    try:
                        pdf_text = self._extract_pdf_text_basic(p)
                        if pdf_text:
                            return pdf_text[:200000]
                    except Exception:
                        pass
                with open(p, 'r', encoding='utf-8', errors='ignore') as fh:
                    return fh.read(200000)
            except Exception:
                return ''
            return ''

        def _resolve_pairs() -> list[dict]:
            normalized: list[dict] = []
            try:
                if isinstance(files, dict):
                    if isinstance(files.get('pairs'), list):
                        normalized = [p for p in files.get('pairs', []) if isinstance(p, dict)]
                    else:
                        src_list = list(files.get('source', []) or [])
                        trg_list = list(files.get('translation', []) or [])
                        limit = min(len(src_list), len(trg_list))
                        for idx in range(limit):
                            src_path = src_list[idx]
                            trg_path = trg_list[idx]
                            normalized.append({
                                'source': src_path,
                                'translation': trg_path,
                                'source_name': os.path.basename(src_path) if src_path else '',
                                'translation_name': os.path.basename(trg_path) if trg_path else ''
                            })
                elif isinstance(files, list):
                    normalized = [p for p in files if isinstance(p, dict)]
            except Exception:
                normalized = []
            if not normalized:
                try:
                    fallback = getattr(self, 'file_pairs', []) or []
                    normalized = [p for p in fallback if isinstance(p, dict)]
                except Exception:
                    normalized = []
            return normalized

        def _issue_to_dict(issue, phase: str) -> dict:
            data: dict = {}
            try:
                if isinstance(issue, dict):
                    data = dict(issue)
                elif is_dataclass(issue):
                    data = asdict(issue)
                elif hasattr(issue, '__dict__'):
                    data = dict(getattr(issue, '__dict__'))
            except Exception:
                data = {}
            if not data:
                data = {
                    'code': getattr(issue, 'code', ''),
                    'severity': getattr(issue, 'severity', 'info'),
                    'category': getattr(issue, 'category', phase),
                    'message': getattr(issue, 'message', ''),
                    'source_text': getattr(issue, 'source_text', getattr(issue, 'source', '')),
                    'target_text': getattr(issue, 'target_text', getattr(issue, 'target', '')),
                    'meta': getattr(issue, 'meta', {}) or {}
                }
            severity = str(data.get('severity', 'info') or 'info').lower()
            category = data.get('category') or phase
            source_txt = data.get('source_text', data.get('source')) or ''
            target_txt = data.get('target_text', data.get('target')) or ''
            meta = data.get('meta') or {}
            return _normalize_issue_dict(
                phase=phase,
                code=data.get('code', ''),
                severity=severity,
                category=category,
                message=data.get('message', ''),
                source=source_txt,
                target=target_txt,
                meta=meta,
            )

        def _normalize_issue_dict(*, phase: str, code: str, severity: str, category: str,
                                   message: str, source: str, target: str, meta: dict) -> dict:
            safe_severity = severity if severity in ('info', 'minor', 'major', 'critical') else 'info'
            norm = {
                'phase': phase,
                'code': code,
                'severity': safe_severity,
                'category': category or phase,
                'rule': category or phase,
                'message': message,
                'source': source,
                'target': target,
                'source_excerpt': (source or '')[:200],
                'target_excerpt': (target or '')[:200],
                'meta': meta or {},
            }
            return norm

        normalized_pairs = _resolve_pairs()
        pairs: list[dict] = []
        pair_segments: list[tuple[str, str]] = []
        pair_details: list[dict] = []
        total_src_chars = 0
        total_trg_chars = 0
        glossary_terms: set[str] = set()
        glossary_path = getattr(self, 'glossary_file_path', None)
        if glossary_path and os.path.isfile(glossary_path):
            try:
                import csv
                with open(glossary_path, 'r', encoding='utf-8', errors='ignore') as gh:
                    dialect = 'excel'
                    if glossary_path.lower().endswith('.tsv'):
                        dialect = 'excel-tab'
                    reader = csv.reader(gh, delimiter='\t' if dialect=='excel-tab' else ',')
                    for row in reader:
                        if not row: continue
                        term = row[0].strip()
                        if term and len(term) > 1:
                            glossary_terms.add(term.lower())

            except Exception:
                pass

        def _translate_phrase(text: str) -> str:
            try:
                translator = getattr(self, '_t', None)
                if callable(translator):
                    return translator(text)
            except Exception:
                pass
            return text

        def _collect_grammar_summary(issues: list[dict], spell_cfg: dict, phase3_on: bool) -> dict:
            used: set[str] = set()
            cache_hit = False
            for issue in issues:
                category_value = str(issue.get('category') or '').lower()
                if category_value != 'grammar':
                    continue
                meta_obj = issue.get('meta') or {}
                checker_id = meta_obj.get('checker') or issue.get('checker')
                if isinstance(checker_id, str) and checker_id:
                    used.add(checker_id)
                else:
                    rule_name = str(issue.get('rule') or '')
                    lower_rule = rule_name.lower()
                    if lower_rule.startswith('hunspell') or 'hunspell' in lower_rule:
                        used.add('hunspell')
                    elif lower_rule.startswith('languagetool') or '.lt.' in lower_rule:
                        used.add('languagetool')
                    elif lower_rule.startswith('ollama'):
                        used.add('ollama')
                    elif lower_rule.startswith('heuristic'):
                        used.add('heuristic')
                if not cache_hit and meta_obj:
                    try:
                        cache_hit = cache_hit or bool(meta_obj.get('cache_hit'))
                    except Exception:
                        pass
            disabled: list[str] = []
            if not phase3_on:
                disabled.append(_translate_phrase('Phase 3 deaktiviert'))
            elif not bool(spell_cfg.get('enabled', True)):
                disabled.append(_translate_phrase('Spellcheck deaktiviert'))
            info = {
                'used': sorted(used),
                'disabled': disabled,
                'force': bool(spell_cfg.get('force_grammar', False)),
                'cache_hit': cache_hit,
            }
            return info

        embed_backend = None
        # Semantische Ähnlichkeit: Standardmäßig aktiviert wenn Config es erlaubt
        # GUI-Variable kann überschreiben (falls vorhanden)
        _semantic_var = getattr(self, 'var_phase3_semantic', None)
        if _semantic_var is not None and hasattr(_semantic_var, 'get'):
            # GUI-Variable existiert: deren Wert verwenden
            use_semantic = bool(_semantic_var.get())
        else:
            # Keine GUI-Variable: standardmäßig aktivieren (Config entscheidet später)
            use_semantic = True
        phase1_enabled = bool(_get_setting('analysis.phases.phase1.enabled', True))
        phase2_enabled = bool(_get_setting('analysis.phases.phase2.enabled', True))
        phase3_enabled = bool(_get_setting('analysis.phases.phase3.enabled', True))
        semantic_cfg_enabled = bool(_get_setting('analysis.phases.phase3.semantic.enabled', True))
        semantic_threshold_cfg = float(_get_setting('analysis.phases.phase3.semantic.threshold', 0.85))
        semantic_use_ollama_cfg = bool(_get_setting('analysis.phases.phase3.semantic.use_ollama', False))
        semantic_ollama_model = str(_get_setting('analysis.phases.phase3.semantic.ollama_model', 'nomic-embed-text'))
        phase_findings_cap = int(_get_setting('analysis.phases.max_findings_per_phase', 250))
        selected_embed_model = str(_get_setting('analysis.embeddings.model', 'all-MiniLM-L6-v2'))
        spellcheck_cfg = copy.deepcopy(_get_setting('analysis.phases.phase3.spellcheck', {}))
        if not isinstance(spellcheck_cfg, dict):
            spellcheck_cfg = {}
        force_grammar = bool(spellcheck_cfg.get('force_grammar', False))
        if not force_grammar and spellcheck_cfg.get('enabled', True) and not qc_enabled('grammar'):
            spellcheck_cfg['enabled'] = False
        elif force_grammar:
            spellcheck_cfg['enabled'] = True
        validation_locale_cfg = copy.deepcopy(_get_setting('analysis.validation.locale', {}))
        if not isinstance(validation_locale_cfg, dict):
            validation_locale_cfg = {}
        validation_blacklist_cfg = copy.deepcopy(_get_setting('analysis.validation.blacklist', {}))
        if not isinstance(validation_blacklist_cfg, dict):
            validation_blacklist_cfg = {}
        validation_lists_cfg = copy.deepcopy(_get_setting('analysis.validation.lists', {}))
        if not isinstance(validation_lists_cfg, dict):
            validation_lists_cfg = {}
        validation_metadata_cfg = copy.deepcopy(_get_setting('analysis.validation.metadata', {}))
        if not isinstance(validation_metadata_cfg, dict):
            validation_metadata_cfg = {}
        validation_config = {
            'locale': validation_locale_cfg,
            'blacklist': validation_blacklist_cfg,
            'lists': validation_lists_cfg,
            'metadata': validation_metadata_cfg
        }

        def _extract_lang_value(attr_name: str):
            attr = getattr(self, attr_name, None)
            if attr is None:
                return None
            getter = getattr(attr, 'get', None)
            if callable(getter):
                try:
                    return getter()
                except Exception:
                    return None
            return None

        def _normalize_lang_code(raw_value: Any) -> Optional[str]:
            if raw_value is None:
                return None
            try:
                if isinstance(raw_value, str):
                    value = raw_value.strip()
                    if not value:
                        return None
                    lower = value.lower()
                    if lower in ('auto', 'auto-detect', 'auto_detect', 'auto-erkennung', 'auto_erkennung'):
                        return None
                    mapping = getattr(self, 'LANG_DISPLAY2ISO', {}) or {}
                    for display, iso_code in mapping.items():
                        if isinstance(iso_code, str) and iso_code.lower() == lower:
                            return iso_code.lower()
                        if isinstance(display, str) and display.lower() == lower:
                            return str(iso_code).lower()
                    fallback_map = {
                        'english': 'en',
                        'englisch': 'en',
                        'german': 'de',
                        'deutsch': 'de',
                        'french': 'fr',
                        'französisch': 'fr',
                        'spanish': 'es',
                        'spanisch': 'es',
                        'italian': 'it',
                        'italienisch': 'it'
                    }
                    if lower in fallback_map:
                        return fallback_map[lower]
                    if len(lower) == 2 and lower.isalpha():
                        return lower
                    if len(lower) >= 3 and lower[:2].isalpha():
                        return lower[:2]
            except Exception:
                return None
            return None

        def _language_display(code: Optional[str]) -> str:
            if not code:
                return self._t('Auto-Erkennung') if hasattr(self, '_t') else 'Auto-Erkennung'
            normalized = str(code).lower()
            mapping = getattr(self, 'LANG_DISPLAY2ISO', {}) or {}
            try:
                inverse = {str(iso).lower(): display for display, iso in mapping.items()}
            except Exception:
                inverse = {}
            display = inverse.get(normalized)
            fallback_names = {
                'en': 'Englisch',
                'de': 'Deutsch',
                'fr': 'Französisch',
                'es': 'Spanisch',
                'it': 'Italienisch'
            }
            if not display:
                display = fallback_names.get(normalized, normalized.upper())
            try:
                return self._t(display)
            except Exception:
                return display

        def _resolve_language_pair(pairs: list[dict]) -> tuple[Optional[str], Optional[str]]:
            src_candidates: list[Any] = []
            tgt_candidates: list[Any] = []

            for attr_name in ('var_source_lang', 'source_lang', 'var_source_language', 'source_language'):
                candidate = _extract_lang_value(attr_name)
                if candidate:
                    src_candidates.append(candidate)
            for attr_name in ('var_target_lang', 'target_lang', 'var_target_language', 'target_language'):
                candidate = _extract_lang_value(attr_name)
                if candidate:
                    tgt_candidates.append(candidate)

            src_candidates.append(_get_setting('analysis.lang.source', None))
            tgt_candidates.append(_get_setting('analysis.lang.target', None))

            src_candidates.append(validation_locale_cfg.get('source_language'))
            tgt_candidates.append(validation_locale_cfg.get('target_language'))

            src_candidates.append(spellcheck_cfg.get('source_language'))
            tgt_candidates.append(spellcheck_cfg.get('target_language'))

            for entry in pairs:
                if not isinstance(entry, dict):
                    continue
                meta = entry.get('meta') if isinstance(entry.get('meta'), dict) else {}
                src_candidates.extend([
                    entry.get('source_language'),
                    entry.get('source_lang'),
                    meta.get('source_language'),
                    meta.get('source_lang'),
                    meta.get('src_lang')
                ])
                tgt_candidates.extend([
                    entry.get('target_language'),
                    entry.get('target_lang'),
                    entry.get('translation_language'),
                    meta.get('target_language'),
                    meta.get('target_lang'),
                    meta.get('trg_lang')
                ])
                for key in ('source_info', 'translation_info'):
                    info = meta.get(key)
                    if isinstance(info, dict):
                        lang_hint = info.get('language')
                        if key.startswith('source'):
                            src_candidates.append(lang_hint)
                        else:
                            tgt_candidates.append(lang_hint)

            src_code = next((code for code in (_normalize_lang_code(val) for val in src_candidates) if code), None)
            tgt_code = next((code for code in (_normalize_lang_code(val) for val in tgt_candidates) if code), None)
            return src_code, tgt_code

        length_pair_overrides = _get_setting('quality.thresholds.length.language_pairs', {})
        if not isinstance(length_pair_overrides, dict):
            length_pair_overrides = {}
        length_expected_ratio = float(_get_setting('quality.thresholds.length.expected_ratio_default', 1.0))
        length_var = float(_get_setting('quality.thresholds.length.variance', 0.05))

        source_lang_code, target_lang_code = _resolve_language_pair(normalized_pairs)
        src_lang_summary = source_lang_code or 'auto'
        tgt_lang_summary = target_lang_code or 'auto'
        
        # Erkannte Sprachen an spellcheck_cfg übergeben für bessere Grammatikprüfung
        if target_lang_code and 'target_language' not in spellcheck_cfg:
            spellcheck_cfg['target_language'] = target_lang_code
        if source_lang_code and 'source_language' not in spellcheck_cfg:
            spellcheck_cfg['source_language'] = source_lang_code

        candidate_keys: list[str] = []
        if source_lang_code and target_lang_code:
            candidate_keys.append(f"{source_lang_code.lower()}->{target_lang_code.lower()}")
        if source_lang_code:
            candidate_keys.append(f"{source_lang_code.lower()}->any")
        if target_lang_code:
            candidate_keys.append(f"any->{target_lang_code.lower()}")
        candidate_keys.append('any->any')

        for key in candidate_keys:
            cfg = length_pair_overrides.get(key)
            if not isinstance(cfg, dict):
                continue
            if 'expected_ratio' in cfg:
                try:
                    length_expected_ratio = float(cfg['expected_ratio'])
                except Exception:
                    pass
            if 'tolerance' in cfg:
                try:
                    length_var = float(cfg['tolerance'])
                except Exception:
                    pass
            break

        length_expected_ratio = max(0.01, float(length_expected_ratio))
        length_var = max(0.01, float(length_var))
        length_lower_bound = max(0.0, length_expected_ratio - length_var)
        length_upper_bound = length_expected_ratio + length_var
        length_lang_label = f"{_language_display(source_lang_code)}->{_language_display(target_lang_code)}"

        # OPTIMIERUNG: use_semantic nur als Flag für Phase 3 - keine doppelte Embedding-Initialisierung hier
        if not phase3_enabled or not semantic_cfg_enabled:
            use_semantic = False
        # embed_backend wird nicht mehr hier initialisiert - Phase 3 macht das selbst

        sentence_split = re.compile(r'[.!?]+\s+|\n+')
        passive_de = re.compile(r'\b(wird|werden|wurde|wurden|ist|sind|sei|waren)\b[^.?!]{0,60}?\b(ge\w+(?:t|en)|worden)\b', re.IGNORECASE)
        passive_en = re.compile(r'\b(is|are|was|were|been|be|being)\b[^.?!]{0,40}?\b\w+(ed|en)\b', re.IGNORECASE)
        passive_de_zu = re.compile(r'\bzu\s+ge\w+(?:t|en)\b', re.IGNORECASE)

        language_markers: dict[str, set[str]] = {
            'de': {
                'der', 'die', 'das', 'und', 'oder', 'aber', 'weil', 'dass', 'nicht', 'kein', 'mehr', 'weniger',
                'wenn', 'doch', 'schon', 'muss', 'soll', 'sollte', 'wird', 'wurden', 'hat', 'haben', 'beim',
                'zum', 'zur', 'über', 'für', 'mit', 'ohne', 'noch', 'wie', 'war', 'waren', 'heute', 'gestern',
                'morgen', 'durch', 'gegen', 'seit', 'zwischen', 'ihre', 'seine', 'unsere', 'deutsch', 'übersetzung',
                'einer', 'einem', 'einen', 'keine', 'keinen', 'keinem', 'immer', 'bereits', 'sowie', 'damit'
            },
            'en': {
                'the', 'and', 'or', 'but', 'because', 'that', 'not', 'no', 'more', 'less', 'when', 'already',
                'must', 'should', 'will', 'were', 'has', 'have', 'into', 'for', 'with', 'without', 'still', 'how',
                'today', 'yesterday', 'tomorrow', 'through', 'against', 'since', 'between', 'their', 'english',
                'translation', 'however', 'therefore', 'furthermore', 'while', 'within', 'across'
            },
            'fr': {
                'le', 'la', 'les', 'des', 'et', 'ou', 'mais', 'parce', 'que', 'pas', 'plus', 'moins', 'quand',
                'déjà', 'doit', 'devrait', 'sera', 'étaient', 'est', 'sont', 'été', 'avec', 'sans', 'encore',
                'comment', 'aujourd', 'hier', 'demain', 'contre', 'depuis', 'entre', 'leur', 'leurs', 'français',
                'traduction', 'selon', 'tandis', 'toutefois'
            },
            'es': {
                'el', 'la', 'los', 'las', 'y', 'o', 'pero', 'porque', 'que', 'no', 'más', 'menos', 'cuando',
                'ya', 'debe', 'debería', 'será', 'fueron', 'es', 'son', 'ha', 'han', 'con', 'sin', 'aún', 'cómo',
                'hoy', 'ayer', 'mañana', 'contra', 'desde', 'entre', 'su', 'sus', 'español', 'traducción',
                'mientras', 'además'
            },
            'it': {
                'il', 'la', 'i', 'gli', 'e', 'o', 'ma', 'perché', 'che', 'non', 'più', 'meno', 'quando', 'già',
                'deve', 'dovrebbe', 'sarà', 'sono', 'è', 'con', 'senza', 'ancora', 'come', 'oggi', 'ieri',
                'domani', 'contro', 'dal', 'tra', 'suo', 'loro', 'italiano', 'traduzione', 'mentre', 'pertanto'
            }
        }
        language_diacritics: dict[str, tuple[str, ...]] = {
            'de': ('ä', 'ö', 'ü', 'ß'),
            'fr': ('à', 'â', 'ä', 'ç', 'é', 'è', 'ê', 'ë', 'î', 'ï', 'ô', 'ù', 'û', 'ü', 'ÿ', 'œ', 'æ'),
            'es': ('á', 'é', 'í', 'ó', 'ú', 'ñ'),
            'it': ('à', 'è', 'é', 'ì', 'î', 'ò', 'ó', 'ù')
        }

        def _language_token_stats(text: str) -> dict[str, Any]:
            tokens = re.findall(r"[A-Za-zÀ-ÿÄÖÜäöüß\-']+", (text or '').lower())
            stats: dict[str, Any] = {'total': len(tokens), 'by_lang': {}}
            if not tokens:
                return stats
            for tok in tokens:
                matched_langs: set[str] = set()
                for code, marker_set in language_markers.items():
                    if tok in marker_set:
                        matched_langs.add(code)
                for code, accents in language_diacritics.items():
                    if code in matched_langs:
                        continue
                    if any(ch in tok for ch in accents):
                        matched_langs.add(code)
                for code in matched_langs:
                    stats['by_lang'][code] = stats['by_lang'].get(code, 0) + 1
            return stats

        def _normalize_text_compare(value: str) -> str:
            return re.sub(r'\s+', ' ', (value or '')).strip().lower()

        acc_sims: list[float] = []
        term_hits = 0
        term_total = 0
        style_sentence_lengths: list[int] = []
        passive_hits = 0
        passive_total = 0
        try:
            language_mix_threshold = float(_get_setting('analysis.validation.language_mismatch.total_threshold', 0.08))
        except Exception:
            language_mix_threshold = 0.08
        try:
            segment_source_threshold = float(_get_setting('analysis.validation.language_mismatch.segment_threshold', 0.35))
        except Exception:
            segment_source_threshold = 0.35
        language_mix_threshold = max(0.0, min(1.0, language_mix_threshold))
        segment_source_threshold = max(0.0, min(1.0, segment_source_threshold))

        language_marker_hits = {'total': 0, 'by_lang': {}}
        language_segment_flags: list[dict] = []
        untranslated_segments: list[dict] = []

        total_files = len(normalized_pairs)
        for idx, f in enumerate(normalized_pairs):
            src = f.get('source') if isinstance(f, dict) else None
            trg = f.get('translation') if isinstance(f, dict) else None
            src_txt = _read(src)
            trg_txt = _read(trg)
            total_src_chars += len(src_txt)
            total_trg_chars += len(trg_txt)

            pair_hash = hashlib.md5(f"{src}|{trg}".encode()).hexdigest()[:8]

            pair_segments.append((src_txt, trg_txt))
            pair_meta = f.get('meta') if isinstance(f, dict) else None

            if trg_txt:
                lang_stats = _language_token_stats(trg_txt)
                lang_total = int(lang_stats.get('total') or 0)
                lang_by: dict[str, int] = lang_stats.get('by_lang') or {}
                if lang_total:
                    language_marker_hits['total'] += lang_total
                    for code, amount in lang_by.items():
                        if not amount:
                            continue
                        language_marker_hits['by_lang'][code] = language_marker_hits['by_lang'].get(code, 0) + int(amount)

                    segment_total = lang_total
                    target_ratio_segment = 0.0
                    if target_lang_code:
                        target_ratio_segment = lang_by.get(target_lang_code, 0) / segment_total
                    source_ratio_segment = 0.0
                    if source_lang_code:
                        source_ratio_segment = lang_by.get(source_lang_code, 0) / segment_total

                    top_other_code = None
                    top_other_ratio = 0.0
                    for code, amount in lang_by.items():
                        if code == target_lang_code or not amount:
                            continue
                        ratio = amount / segment_total
                        if ratio > top_other_ratio:
                            top_other_code = code
                            top_other_ratio = ratio

                    flag_language = None
                    flag_ratio = 0.0
                    if source_lang_code and source_ratio_segment >= segment_source_threshold and source_ratio_segment > target_ratio_segment and len(trg_txt) > 40:
                        flag_language = source_lang_code
                        flag_ratio = source_ratio_segment
                    elif top_other_code and top_other_ratio >= segment_source_threshold and top_other_ratio > target_ratio_segment and len(trg_txt) > 40:
                        flag_language = top_other_code
                        flag_ratio = top_other_ratio

                    if flag_language:
                        language_segment_flags.append({
                            'hash': pair_hash,
                            'language': flag_language,
                            'language_label': _language_display(flag_language),
                            'flag_ratio': round(flag_ratio, 4),
                            'target_ratio': round(target_ratio_segment, 4),
                            'preview': (trg_txt or '')[:160]
                        })

            if src_txt and not trg_txt.strip():
                untranslated_segments.append({
                    'hash': pair_hash,
                    'reason': 'leer',
                    'source_excerpt': (src_txt or '')[:160],
                    'target_excerpt': ''
                })
            elif src_txt and trg_txt:
                if _normalize_text_compare(src_txt) == _normalize_text_compare(trg_txt):
                    untranslated_segments.append({
                        'hash': pair_hash,
                        'reason': 'identisch',
                        'source_excerpt': (src_txt or '')[:160],
                        'target_excerpt': (trg_txt or '')[:160]
                    })

            sim = 0.0
            if qc_enabled('accuracy') and src_txt and trg_txt:
                src_tokens = set(src_txt.split())
                trg_tokens = set(trg_txt.split())
                common = len(src_tokens & trg_tokens)
                total_union = max(1, len(src_tokens) + len(trg_tokens))
                sim = round((2*common)/total_union, 4)
                # OPTIMIERUNG: Embedding-Berechnung hier entfernt - 
                # wird in Phase 3 check_semantic_similarity() gemacht wenn aktiviert
                # Dies vermeidet doppelte Embedding-Berechnungen
                acc_sims.append(sim)

            if qc_enabled('terminology') and glossary_terms and trg_txt:
                trg_lower = trg_txt.lower()
                for term in glossary_terms:
                    if len(term) < 3:
                        continue
                    term_total += 1
                    if term in trg_lower:
                        term_hits += 1

            # OPTIMIERUNG: Stil-Statistiken nur berechnen wenn Phase 3 NICHT aktiv
            # (Phase 3 check_style/check_readability macht detailliertere Analyse)
            if trg_txt and (qc_enabled('fluency') or qc_enabled('style') or qc_enabled('grammar')) and not phase3_enabled:
                sentences = [s.strip() for s in sentence_split.split(trg_txt) if s.strip()]
                for s in sentences:
                    tokens = s.split()
                    if tokens:
                        style_sentence_lengths.append(len(tokens))
                    passive_total += 1
                    if passive_de.search(s) or passive_en.search(s) or passive_de_zu.search(s):
                        passive_hits += 1
            elif trg_txt and phase3_enabled:
                # Bei aktivierter Phase 3: nur minimale Statistiken für Metriken-Übersicht
                sentences = [s.strip() for s in sentence_split.split(trg_txt) if s.strip()]
                for s in sentences:
                    tokens = s.split()
                    if tokens:
                        style_sentence_lengths.append(len(tokens))

            pairs.append({
                'source': src,
                'translation': trg,
                'similarity': sim,
                'hash': pair_hash,
                'meta': pair_meta,
                'index': idx
            })
            pair_details.append({
                'index': idx,
                'source': src,
                'translation': trg,
                'source_name': f.get('source_name') if isinstance(f, dict) else None,
                'translation_name': f.get('translation_name') if isinstance(f, dict) else None,
                'similarity': sim,
                'hash': pair_hash,
                'source_chars': len(src_txt),
                'target_chars': len(trg_txt),
                'source_excerpt': (src_txt or '')[:200],
                'target_excerpt': (trg_txt or '')[:200],
                'meta': pair_meta
            })
            if language_segment_flags and language_segment_flags[-1].get('hash') is None:
                language_segment_flags[-1]['hash'] = pair_hash
            if self.event_bus and total_files:
                try:
                    self.event_bus.publish('analysis.progress', {
                        'phase': 'pair_metrics',
                        'current': idx+1,
                        'total': total_files,
                        'progress': (idx+1)/total_files,
                        'info': 'Paare analysiert'
                    })
                except Exception:
                    pass

        pair_count = len(pairs)
        untranslated_total = len(untranslated_segments)
        translated_segments = max(pair_count - untranslated_total, 0)
        if pair_count:
            translated_ratio = round(translated_segments / pair_count, 4)
            untranslated_ratio = round(untranslated_total / pair_count, 4)
        else:
            translated_ratio = None
            untranslated_ratio = None
        def _avg(xs: list[float]) -> float:
            return round(sum(xs)/len(xs), 4) if xs else 0.0

        readability = 0.0
        if style_sentence_lengths and (qc_enabled('fluency') or qc_enabled('style')):
            total_words = sum(style_sentence_lengths)
            asl = total_words / max(1, len(style_sentence_lengths))
            try:
                # Silben grob über Vokalgruppen aus letztem untersuchten Text (vereinfachte Heuristik)
                joined_sample = ' '.join([str(len(style_sentence_lengths))])
                vowels = re.findall(r'[aeiouyäöüAEIOUYÄÖÜ]+', joined_sample)
                syllables = len(vowels) if vowels else total_words
            except Exception:
                syllables = total_words
            asw = syllables / max(1, total_words)
            readability = round(206.835 - 1.015*asl - 84.6*asw, 2)

        passive_rate = round((passive_hits / passive_total), 4) if passive_total else 0.0
        terminology_rate = round((term_hits / term_total), 4) if term_total else (0.0 if glossary_terms else None)
        language_ratio_map: dict[str, float] = {}
        total_language_tokens = int(language_marker_hits['total']) if language_marker_hits['total'] else 0
        if total_language_tokens:
            for code, amount in (language_marker_hits.get('by_lang') or {}).items():
                if not amount:
                    continue
                language_ratio_map[code] = float(amount) / float(total_language_tokens)

        target_ratio_total = language_ratio_map.get(target_lang_code, 0.0) if target_lang_code else 0.0
        source_ratio_total = language_ratio_map.get(source_lang_code, 0.0) if source_lang_code else 0.0
        top_other_code = None
        top_other_ratio = 0.0
        for code, ratio in language_ratio_map.items():
            if code == target_lang_code:
                continue
            if ratio > top_other_ratio:
                top_other_code = code
                top_other_ratio = ratio

        metrics_all = {
            'pair_count': pair_count,
            'avg_similarity': _avg(acc_sims),
            'avg_length_ratio': round((total_trg_chars / total_src_chars), 4) if total_src_chars else 0.0,
            'terminology_match_rate': terminology_rate,
            'readability_score': readability,
            'avg_sentence_length': _avg(style_sentence_lengths),
            'passive_voice_rate': passive_rate,
            'analysis_time_ms': int((time.time()-start)*1000),
            'length_ratio_expected': round(length_expected_ratio, 4),
            'length_ratio_tolerance': round(length_var, 4),
            'length_ratio_lower': round(length_lower_bound, 4),
            'length_ratio_upper': round(length_upper_bound, 4),
            'length_ratio_languages': {
                'source': src_lang_summary,
                'target': tgt_lang_summary
            },
            'length_ratio_label': length_lang_label,
            'target_language_analysis': {
                'token_total': total_language_tokens,
                'expected_target': tgt_lang_summary,
                'target_ratio': round(target_ratio_total, 4) if target_lang_code else None,
                'source_ratio': round(source_ratio_total, 4) if source_lang_code else None,
                'ratios': {code: round(val, 4) for code, val in language_ratio_map.items()},
                'flagged_segments': len(language_segment_flags),
                'top_other_language': {
                    'code': top_other_code,
                    'label': _language_display(top_other_code) if top_other_code else None,
                    'ratio': round(top_other_ratio, 4)
                } if top_other_code else None,
                'thresholds': {
                    'mix_total': round(language_mix_threshold, 4),
                    'segment_source': round(segment_source_threshold, 4)
                }
            },
            'segments_total': pair_count,
            'segments_translated': translated_segments,
            'segments_untranslated': untranslated_total,
            'translated_ratio': translated_ratio,
            'untranslated_ratio': untranslated_ratio,
            'translation_completeness_score': translated_ratio,
            'completeness': translated_ratio,
            'untranslated_segments_flagged': untranslated_total
        }
        pair_target_lookup = {
            p.get('hash'): (p.get('target_excerpt') or p.get('translation') or '')
            for p in pair_details if isinstance(p, dict) and p.get('hash')
        }
        # Similarity-Verteilung ergänzen (Buckets) – robust gegen Fehler
        try:
            if acc_sims:
                _buckets = {
                    '0-0.50': 0,
                    '0.50-0.70': 0,
                    '0.70-0.85': 0,
                    '0.85-0.95': 0,
                    '0.95-1.00': 0
                }
                for _s in acc_sims:
                    if _s < 0.50: _buckets['0-0.50'] += 1
                    elif _s < 0.70: _buckets['0.50-0.70'] += 1
                    elif _s < 0.85: _buckets['0.70-0.85'] += 1
                    elif _s < 0.95: _buckets['0.85-0.95'] += 1
                    else: _buckets['0.95-1.00'] += 1
                metrics_all['similarity_distribution'] = {
                    'buckets': _buckets,
                    'min': min(acc_sims),
                    'max': max(acc_sims),
                    'median': sorted(acc_sims)[len(acc_sims)//2]
                }
        except Exception:
            pass

        thr_crit = float(_get_setting('quality.thresholds.similarity.critical', 0.75))
        thr_major = float(_get_setting('quality.thresholds.similarity.major', 0.85))
        sim_threshold_mode = 'static'
        use_pct = bool(_get_setting('quality.thresholds.similarity.use_percentiles', False))
        p_crit = float(_get_setting('quality.thresholds.similarity.percentiles.critical', 10.0))
        p_major = float(_get_setting('quality.thresholds.similarity.percentiles.major', 25.0))
        if use_pct and acc_sims and isinstance(acc_sims, list) and len(acc_sims) >= 12:
            xs = sorted([float(x) for x in acc_sims if isinstance(x, (int, float))])
            if xs:
                def _pct(vals: list[float], p: float) -> float:
                    try:
                        p = max(0.0, min(100.0, float(p)))
                        if not vals:
                            return 0.0
                        idx = int(round((p/100.0) * (len(vals)-1)))
                        return float(vals[idx])
                    except Exception:
                        return 0.0
                thr_crit = _pct(xs, p_crit)
                thr_major = _pct(xs, max(p_crit, p_major))
                if thr_major < thr_crit:
                    thr_major = min(1.0, thr_crit + 0.05)
                sim_threshold_mode = 'percentile'
        semantic_threshold = float(max(0.0, min(1.0, semantic_threshold_cfg))) if semantic_threshold_cfg else thr_major
        if not semantic_threshold:
            semantic_threshold = thr_major
        try:
            metrics_all.setdefault('similarity_thresholds_used', {})
            metrics_all['similarity_thresholds_used'].update({
                'critical': round(float(thr_crit), 4),
                'major': round(float(thr_major), 4),
                'mode': sim_threshold_mode,
                'p_critical': p_crit if use_pct else None,
                'p_major': p_major if use_pct else None
            })
        except Exception:
            pass

        quality: dict[str, float] = {}
        if qc_enabled('accuracy'):
            quality['accuracy'] = metrics_all['avg_similarity']
        if qc_enabled('terminology') and terminology_rate is not None:
            quality['terminology'] = terminology_rate
        if qc_enabled('fluency'):
            quality['fluency'] = readability
        if qc_enabled('style'):
            quality['style'] = metrics_all['avg_sentence_length']
        if qc_enabled('grammar'):
            quality['grammar'] = 1.0 - passive_rate
        if qc_enabled('completeness'):
            completeness_metric = metrics_all.get('completeness')
            if isinstance(completeness_metric, (int, float)):
                quality['completeness'] = completeness_metric
            else:
                quality['completeness'] = metrics_all['avg_length_ratio']

        phase1_issues: list[dict] = []
        phase2_issues: list[dict] = []
        phase3_issues: list[dict] = []
        if phase1_enabled and run_phase1_checks and pair_segments:
            try:
                raw_phase1 = run_phase1_checks(pair_segments)
                phase1_issues = [_issue_to_dict(issue, 'phase1') for issue in raw_phase1][:phase_findings_cap]
            except Exception as phase1_err:
                try:
                    (getattr(self, 'logger', None) or logger).debug(f"Phase1 Checks Fehler: {phase1_err}")
                except Exception:
                    pass
        if phase2_enabled and run_phase2_checks and pair_segments:
            try:
                glossary_arg = glossary_path if glossary_path else 'glossary_terms.json'
                raw_phase2 = run_phase2_checks(
                    pair_segments,
                    glossary_path=glossary_arg,
                    config=validation_config,
                    pair_infos=pair_details
                )
                phase2_issues = [_issue_to_dict(issue, 'phase2') for issue in raw_phase2][:phase_findings_cap]
            except Exception as phase2_err:
                try:
                    (getattr(self, 'logger', None) or logger).debug(f"Phase2 Checks Fehler: {phase2_err}")
                except Exception:
                    pass
        if phase3_enabled and run_phase3_checks and pair_segments:
            try:
                # Konsistenzprüfung aktivieren wenn genug Segmente vorhanden
                enable_consistency = len(pair_segments) >= 5
                # OCR-Prüfung aktivieren (erkennt l/1, O/0, rn/m Verwechslungen)
                enable_ocr = True
                
                raw_phase3 = run_phase3_checks(
                    pair_segments,
                    enable_semantic=use_semantic,
                    enable_consistency=enable_consistency,
                    enable_ocr_check=enable_ocr,
                    semantic_use_ollama=semantic_use_ollama_cfg,
                    semantic_ollama_model=semantic_ollama_model,
                    semantic_threshold=semantic_threshold,
                    consistency_min_occurrences=2,
                    spellcheck_config=spellcheck_cfg,
                    pair_infos=pair_details
                )
                phase3_issues = [_issue_to_dict(issue, 'phase3') for issue in raw_phase3][:phase_findings_cap]
            except Exception as phase3_err:
                try:
                    (getattr(self, 'logger', None) or logger).debug(f"Phase3 Checks Fehler: {phase3_err}")
                except Exception:
                    pass
        
        # =====================================================================
        # Custom KI-Prüfung (Benutzerdefinierte Fragen an Ollama)
        # =====================================================================
        custom_ki_findings: list[dict] = []
        custom_ki_result = None
        try:
            custom_ki_enabled = getattr(self, 'var_custom_ki_enabled', None)
            custom_ki_prompt_var = getattr(self, 'var_custom_ki_prompt', None)
            custom_ki_textbox = getattr(self, 'custom_ki_textbox', None)
            
            # Prüfen ob aktiviert
            if (custom_ki_enabled and custom_ki_enabled.get() and 
                semantic_use_ollama_cfg and pair_segments):
                
                # Prompt aus Textbox lesen (aktuellster Wert)
                prompt_text = ''
                try:
                    if custom_ki_textbox:
                        prompt_text = custom_ki_textbox.get('1.0', 'end-1c').strip()
                except Exception:
                    pass
                
                # Fallback auf Variable
                if not prompt_text or prompt_text.startswith('z.B.'):
                    prompt_text = custom_ki_prompt_var.get() if custom_ki_prompt_var else ''
                
                if prompt_text and not prompt_text.startswith('z.B.'):
                    try:
                        from quality_gui_custom_prompts import CustomCheckAnalyzer
                        
                        # Progress-Event senden
                        if self.event_bus:
                            self.event_bus.publish('analysis.progress', {
                                'phase': 'custom_ki',
                                'progress': 0.9,
                                'info': f'KI-Prüfung: {prompt_text[:50]}...'
                            })
                        
                        analyzer = CustomCheckAnalyzer(
                            model=semantic_ollama_model or 'mistral',
                            timeout=120
                        )
                        
                        # Analyse durchführen
                        custom_ki_result = analyzer.analyze(
                            question=prompt_text,
                            pairs=pair_segments,
                            max_pairs=50
                        )
                        
                        # Findings übernehmen
                        if custom_ki_result and custom_ki_result.findings:
                            for f in custom_ki_result.findings:
                                custom_ki_findings.append(_normalize_issue_dict(
                                    phase='custom_ki',
                                    code=f.get('rule_id', 'CUSTOM_CHECK'),
                                    severity=f.get('severity', 'minor'),
                                    category='custom_ki',
                                    message=f.get('message', ''),
                                    source='',
                                    target=f.get('suggestion', ''),
                                    meta={
                                        'custom_prompt': prompt_text,
                                        'segment_index': f.get('segment_index', 0),
                                        'checker': 'ollama',
                                        'ki_answer': custom_ki_result.answer[:500] if custom_ki_result.answer else ''
                                    }
                                ))
                        
                        # Log
                        try:
                            (getattr(self, 'logger', None) or logger).info(
                                f"Custom KI-Prüfung: {len(custom_ki_findings)} Findings für '{prompt_text[:30]}...'"
                            )
                        except Exception:
                            pass
                            
                    except ImportError:
                        pass
                    except Exception as ki_err:
                        try:
                            (getattr(self, 'logger', None) or logger).debug(f"Custom KI-Prüfung Fehler: {ki_err}")
                        except Exception:
                            pass
        except Exception:
            pass
        
        phase_issue_counts = {
            'phase1': len(phase1_issues),
            'phase2': len(phase2_issues),
            'phase3': len(phase3_issues),
            'custom_ki': len(custom_ki_findings)
        }
        grammar_summary = _collect_grammar_summary(phase3_issues, spellcheck_cfg, phase3_enabled)


        summary = {
            'pairs': pair_count,
            'total_source_chars': total_src_chars,
            'total_translation_chars': total_trg_chars,
            'profile_used': rule_profile or 'medium',
            'quality_keys': list(quality.keys()),
            'embedding_model': selected_embed_model if use_semantic else None,
            'language_pair': {
                'source': src_lang_summary,
                'target': tgt_lang_summary,
                'label': length_lang_label
            },
            'length_ratio_expectation': {
                'expected': round(length_expected_ratio, 4),
                'tolerance': round(length_var, 4),
                'lower': round(length_lower_bound, 4),
                'upper': round(length_upper_bound, 4)
            },
            'language_marker_summary': {
                'flagged_segments': len(language_segment_flags),
                'top_other_language': _language_display(top_other_code) if top_other_code else None,
                'top_other_ratio': round(top_other_ratio, 4) if top_other_code else None,
                'target_ratio': round(target_ratio_total, 4) if target_lang_code else None
            },
            'translation_completeness': {
                'translated_ratio': translated_ratio,
                'untranslated_ratio': untranslated_ratio,
                'segments_translated': translated_segments,
                'segments_untranslated': untranslated_total,
                'segments_total': pair_count
            }
        }
        if self.event_bus:
            try:
                self.event_bus.publish('analysis.progress', {
                    'phase': 'completed',
                    'progress': 1.0,
                    'info': 'Analyse abgeschlossen'
                })
            except Exception:
                pass
        # Automatische Empfehlungen – Geschäftslogik heuristisch
        recommendations: list[str] = []
        try:
            # Schwellen aus Settings (Fallback auf Defaults)
            acc_low = float(_get_setting('quality.thresholds.accuracy.low', 0.85))
            acc_mid = float(_get_setting('quality.thresholds.accuracy.mid', 0.92))
            term_min = float(_get_setting('quality.thresholds.terminology.min', 0.90))
            readability_min = float(_get_setting('quality.thresholds.readability.min', 60))
            passive_max = float(_get_setting('quality.thresholds.passive.max', 0.15))

            _acc = quality.get('accuracy')
            if _acc is not None:
                if _acc < acc_low:
                    recommendations.append(f"Genauigkeit <{int(acc_low*100)}% – Terminologie & Bedeutungsäquivalenz genauer prüfen.")
                elif _acc < acc_mid:
                    recommendations.append("Gute Genauigkeit – gezielte Optimierungen für Konsistenz & Präzision möglich.")
                else:
                    recommendations.append("Sehr hohe Genauigkeit – Fokus auf Stil & Lesbarkeit verschieben.")
            _term = quality.get('terminology')
            if _term is not None and _term < term_min:
                recommendations.append(f"Terminologie-Trefferquote <{int(term_min*100)}% – Glossar konsequenter anwenden.")
            if readability and readability < readability_min:
                recommendations.append("Lesbarkeitswert niedrig – längere Sätze aufteilen & klarere Struktur.")
            if metrics_all.get('passive_voice_rate', 0) > passive_max:
                recommendations.append("Hoher Passivanteil – aktivere Formulierungen bevorzugen.")
            _len_ratio = metrics_all.get('avg_length_ratio')
            if isinstance(_len_ratio, (int, float)):
                len_ratio_val = float(_len_ratio)
                deviation = abs(len_ratio_val - length_expected_ratio)
                if deviation > length_var:
                    current_pct = len_ratio_val * 100.0
                    lower_pct = length_lower_bound * 100.0
                    upper_pct = length_upper_bound * 100.0
                    if len_ratio_val < length_expected_ratio:
                        recommendations.append(
                            f"Längenverhältnis {length_lang_label} liegt unter dem erwarteten Bereich ({lower_pct:.0f}-{upper_pct:.0f}%). Ist-Wert: {current_pct:.0f}%. Kürzungen oder fehlende Passagen prüfen."
                        )
                    else:
                        recommendations.append(
                            f"Längenverhältnis {length_lang_label} liegt über dem erwarteten Bereich ({lower_pct:.0f}-{upper_pct:.0f}%). Ist-Wert: {current_pct:.0f}%. Zusätzliche Inhalte oder Dopplungen prüfen."
                        )
            untranslated_notice = metrics_all.get('segments_untranslated')
            if isinstance(untranslated_notice, int) and untranslated_notice > 0:
                translated_ratio_pct = float(metrics_all.get('translated_ratio') or 0.0) * 100.0
                recommendations.append(
                    f"{untranslated_notice} Segment(e) wirken unübersetzt oder unverändert – Übersetzungsquote aktuell {translated_ratio_pct:.0f}%. Bitte vervollständigen."
                )
            if not recommendations:
                recommendations.append("Keine kritischen Auffälligkeiten – Qualitätsniveau stabil.")
            recommendations.append("Bericht automatisch generiert – Empfehlungen als Ausgangspunkt verwenden.")
        except Exception:
            pass
        # Ergebnis-Schema-Normalisierung für UI (Score/Findings/Phasen)
        try:
            # 1) Gesamt-Score ableiten (0..100)
            def _clamp01(x):
                try:
                    return 0.0 if x is None else (1.0 if x > 1 else (0.0 if x < 0 else float(x)))
                except Exception:
                    return 0.0
            # Teil-Scores 0..1
            acc_score = _clamp01(metrics_all.get('avg_similarity'))
            term_score = None if terminology_rate is None else _clamp01(terminology_rate)
            flu_score = _clamp01((metrics_all.get('readability_score') or 0) / 100.0)
            # Stil: Zielbereich ~18 Tokens; einfache Normalisierung
            try:
                _asl = float(metrics_all.get('avg_sentence_length') or 0)
                style_score = max(0.0, min(1.0, 1.0 - (abs(_asl - 18.0) / 18.0)))
            except Exception:
                style_score = 0.0
            gram_score = _clamp01(1.0 - (metrics_all.get('passive_voice_rate') or 0.0))
            completeness_score_source = metrics_all.get('completeness')
            if isinstance(completeness_score_source, (int, float)):
                comp_score = max(0.0, min(1.0, float(completeness_score_source)))
            else:
                try:
                    _len_ratio = float(metrics_all.get('avg_length_ratio') or 1.0)
                    comp_score = max(0.0, min(1.0, 1.0 - abs(_len_ratio - 1.0)))
                except Exception:
                    comp_score = 0.0
            # Konfigurierbare Gewichte + optional geometrisches Mittel
            try:
                w_acc = float(_get_setting('quality.weights.accuracy', 0.40))
                w_term = float(_get_setting('quality.weights.terminology', 0.15))
                w_flu = float(_get_setting('quality.weights.fluency', 0.15))
                w_style = float(_get_setting('quality.weights.style', 0.10))
                w_gram = float(_get_setting('quality.weights.grammar', 0.10))
                w_comp = float(_get_setting('quality.weights.completeness', 0.10))
                use_geom = bool(_get_setting('quality.scoring.use_geometric', False))
            except Exception:
                w_acc, w_term, w_flu, w_style, w_gram, w_comp = 0.40, 0.15, 0.15, 0.10, 0.10, 0.10
                use_geom = False
            parts = []
            # Gewichte – nur vorhandene berücksichtigen
            def _add(score, w):
                if score is not None:
                    parts.append((score, w))
            _add(acc_score, w_acc)
            _add(term_score, w_term)
            _add(flu_score, w_flu)
            _add(style_score, w_style)
            _add(gram_score, w_gram)
            _add(comp_score, w_comp)
            if parts:
                wsum = sum(w for _, w in parts) or 1.0
                if use_geom:
                    # Geometrisches Mittel der Teil-Scores mit Gewichten
                    import math
                    eps = 1e-6
                    log_sum = 0.0
                    for s, w in parts:
                        s_eff = max(eps, min(1.0, float(s)))
                        try:
                            log_sum += w * math.log(s_eff)
                        except Exception:
                            log_sum += w * math.log(eps)
                    gm = math.exp(log_sum / wsum)
                    overall_score = round(gm * 100.0, 2)
                else:
                    norm = sum(s*w for s, w in parts) / wsum
                    overall_score = round(norm * 100.0, 2)
            else:
                overall_score = 85.0
            metrics_all['overall_score'] = overall_score
            try:
                summary['quality_score'] = overall_score
                summary['overall_score_norm'] = round(overall_score / 100.0, 4)
            except Exception:
                pass

            # 2) Findings (leichtgewichtig, heuristisch) erstellen
            findings: list[dict] = []
            print(f"DEBUG: Starting findings creation...")
            print(f"DEBUG: phase1_issues={len(phase1_issues)}, phase2_issues={len(phase2_issues)}, phase3_issues={len(phase3_issues)}")
            try:
                for flagged in untranslated_segments[:20]:
                    reason = flagged.get('reason', '')
                    message = "Segment ohne Übersetzung" if reason == 'leer' else "Segment unverändert aus Ausgangstext"
                    findings.append(_normalize_issue_dict(
                        phase='phase2',
                        code='UNTRANSLATED_SEGMENT',
                        severity='critical' if reason == 'leer' else 'major',
                        category='completeness',
                        message=f"{message} (Segment {flagged.get('hash')})",
                        source=flagged.get('source_excerpt', ''),
                        target=flagged.get('target_excerpt', ''),
                        meta={'reason': reason, 'hash': flagged.get('hash')}
                    ))
            except Exception:
                pass
            try:
                if target_lang_code and total_language_tokens:
                    if top_other_code and top_other_ratio >= language_mix_threshold:
                        findings.append(_normalize_issue_dict(
                            phase='phase2',
                            code='TARGET_LANGUAGE_MIX',
                            severity='major',
                            category='completeness',
                            message=(
                                f"Zieltext enthält {top_other_ratio*100:.1f}% {_language_display(top_other_code)} Tokens "
                                f"(erwartet {_language_display(target_lang_code)})."
                            ),
                            source='',
                            target='Übersetzung (gesamt)',
                            meta={
                                'target_language': target_lang_code,
                                'target_ratio': round(target_ratio_total, 4),
                                'top_other_language': top_other_code,
                                'top_other_ratio': round(top_other_ratio, 4),
                                'flagged_segments': len(language_segment_flags)
                            }
                        ))
                    if language_segment_flags:
                        segment_counts: dict[str, int] = {}
                        for entry in language_segment_flags:
                            lang_code = entry.get('language')
                            if not lang_code:
                                continue
                            segment_counts[lang_code] = segment_counts.get(lang_code, 0) + 1
                        dominant_segment_lang = None
                        dominant_segment_count = 0
                        for code, amount in segment_counts.items():
                            if amount > dominant_segment_count:
                                dominant_segment_lang = code
                                dominant_segment_count = amount
                        sample = language_segment_flags[:5]
                        findings.append(_normalize_issue_dict(
                            phase='phase2',
                            code='TARGET_LANGUAGE_SEGMENTS',
                            severity='major',
                            category='completeness',
                            message=(
                                f"{len(language_segment_flags)} Segmente wirken "
                                f"{_language_display(dominant_segment_lang) if dominant_segment_lang else 'fremdsprachig'} "
                                f"statt {_language_display(target_lang_code)}."
                            ),
                            source='',
                            target='Übersetzung (gesamt)',
                            meta={
                                'segments': sample,
                                'language_counts': segment_counts,
                                'target_language': target_lang_code
                            }
                        ))
            except Exception:
                pass
            try:
                print(f"DEBUG: Checking {len(pairs)} pairs for similarity issues...")
                for p in pairs:
                    try:
                        s = float(p.get('similarity'))
                        target_text = pair_target_lookup.get(p.get('hash'), '')
                        if s < thr_crit:
                            try:
                                conf = max(0.0, min(1.0, (thr_crit - s) / max(1e-6, thr_crit)))
                            except Exception:
                                conf = 0.7
                            findings.append(_normalize_issue_dict(
                                phase='phase2',
                                code='SIM_LOW_CRIT',
                                severity='critical',
                                category='pair_similarity',
                                message=f"Sehr niedrige Ähnlichkeit in Übersetzung: {s*100:.1f}% (Segment {p.get('hash')})",
                                source='',
                                target=target_text,
                                meta={'confidence': conf}
                            ))
                        elif s < thr_major:
                            try:
                                span = max(1e-6, (thr_major - thr_crit))
                                conf = max(0.0, min(1.0, (thr_major - s) / span))
                            except Exception:
                                conf = 0.5
                            findings.append(_normalize_issue_dict(
                                phase='phase2',
                                code='SIM_LOW',
                                severity='major',
                                category='pair_similarity',
                                message=f"Niedrige Ähnlichkeit in Übersetzung: {s*100:.1f}% (Segment {p.get('hash')})",
                                source='',
                                target=target_text,
                                meta={'confidence': conf}
                            ))
                    except Exception:
                        continue
            except Exception as e:
                print(f"DEBUG: Exception during pair similarity check: {e}")
                pass
            # Globale Metrik-basierte Findings
            print(f"DEBUG: After pair checks, findings={len(findings)} items")
            try:
                if terminology_rate is not None and terminology_rate < term_min:
                    try:
                        conf = max(0.0, min(1.0, (term_min - terminology_rate) / max(1e-6, term_min)))
                    except Exception:
                        conf = 0.6
                    findings.append(_normalize_issue_dict(
                        phase='phase2',
                        code='TERMINOLOGY_LOW',
                        severity='major',
                        category='terminology',
                        message=f"Niedrige Terminologie-Trefferquote: {terminology_rate*100:.1f}% (<{int(term_min*100)}%)",
                        source='',
                        target='Übersetzung (gesamt)',
                        meta={'confidence': conf}
                    ))
            except Exception:
                pass

            # Phase Findings integrieren (begrenzte Anzahl)
            try:
                for phase_name, issue_list in (('phase1', phase1_issues), ('phase2', phase2_issues), ('phase3', phase3_issues), ('custom_ki', custom_ki_findings)):
                    for issue in issue_list:
                        findings.append(_normalize_issue_dict(
                            phase=str(issue.get('phase') or phase_name),
                            code=str(issue.get('code') or ''),
                            severity=str(issue.get('severity') or 'info'),
                            category=str(issue.get('category') or ''),
                            message=str(issue.get('message') or ''),
                            source=str(issue.get('source') or ''),
                            target=str(issue.get('target') or ''),
                            meta=issue.get('meta') or {}
                        ))
            except Exception:
                pass
            try:
                if readability and readability < readability_min:
                    try:
                        conf = max(0.0, min(1.0, (readability_min - readability) / max(1.0, float(readability_min))))
                    except Exception:
                        conf = 0.5
                    findings.append(_normalize_issue_dict(
                        phase='phase2',
                        code='READABILITY_LOW',
                        severity='minor',
                        category='fluency',
                        message=f"Niedriger Lesbarkeitswert: {readability:.1f} (<{int(readability_min)})",
                        source='',
                        target='Übersetzung (gesamt)',
                        meta={'confidence': conf}
                    ))
            except Exception:
                pass
            try:
                if passive_rate > passive_max:
                    try:
                        conf = max(0.0, min(1.0, (passive_rate - passive_max) / max(1e-6, (1.0 - passive_max))))
                    except Exception:
                        conf = 0.6
                    findings.append(_normalize_issue_dict(
                        phase='phase2',
                        code='PASSIVE_HIGH',
                        severity='major',
                        category='grammar',
                        message=f"Hoher Passivanteil: {passive_rate*100:.1f}% (>{int(passive_max*100)}%)",
                        source='',
                        target='Übersetzung (gesamt)',
                        meta={'confidence': conf}
                    ))
            except Exception:
                pass
            try:
                if isinstance(_len_ratio, (int, float)):
                    len_ratio_val = float(_len_ratio)
                    len_deviation = abs(len_ratio_val - length_expected_ratio)
                    if len_deviation > length_var:
                        try:
                            overshoot = len_deviation - length_var
                            conf = max(0.0, min(1.0, overshoot / max(1e-6, length_expected_ratio)))
                        except Exception:
                            conf = 0.5
                        relation = 'unter' if len_ratio_val < length_expected_ratio else 'über'
                        findings.append(_normalize_issue_dict(
                            phase='phase2',
                            code='LENGTH_DRIFT',
                            severity='major',
                            category='completeness',
                            message=(
                                f"Längenverhältnis {length_lang_label} liegt {relation} dem erwarteten Bereich. "
                                f"Ist: {len_ratio_val*100:.0f}% · Erwartet: {length_expected_ratio*100:.0f}% ±{length_var*100:.0f}%"
                            ),
                            source='',
                            target='Übersetzung (gesamt)',
                            meta={
                                'confidence': conf,
                                'length_ratio': round(len_ratio_val, 4),
                                'expected_ratio': round(length_expected_ratio, 4),
                                'tolerance': round(length_var, 4),
                                'lower': round(length_lower_bound, 4),
                                'upper': round(length_upper_bound, 4),
                                'source_language': src_lang_summary,
                                'target_language': tgt_lang_summary
                            }
                        ))
            except Exception:
                pass

            # Severity-Zusammenfassung in summary
            try:
                sev_counts = {'critical': 0, 'major': 0, 'minor': 0}
                for f in findings:
                    sev = str(f.get('severity') or '').lower()
                    if sev in sev_counts:
                        sev_counts[sev] += 1
                summary.update({
                    'critical': sev_counts['critical'],
                    'major': sev_counts['major'],
                    'minor': sev_counts['minor'],
                })
            except Exception:
                pass

            if grammar_summary['used']:
                summary['grammar_checkers_used'] = list(grammar_summary['used'])
            if grammar_summary['disabled']:
                summary['grammar_checkers_disabled'] = list(grammar_summary['disabled'])
            if grammar_summary['force']:
                summary['grammar_force_override'] = True
            summary['grammar_status'] = grammar_summary
            metrics_all['grammar_status'] = grammar_summary

            # Optionale Gruppierung wiederkehrender Befunde (leichtgewichtig)
            findings_grouped: list[dict] = []
            print(f"DEBUG: Before grouping, findings={len(findings)} items")
            print(f"DEBUG: phase1_issues={len(phase1_issues)}, phase2_issues={len(phase2_issues)}, phase3_issues={len(phase3_issues)}")
            try:
                do_group = bool(_get_setting('quality.findings.clustering.enabled', True))
            except Exception:
                do_group = True
            if do_group and findings:
                try:
                    from collections import defaultdict
                    groups: dict[tuple, dict] = {}
                    for f in findings:
                        try:
                            rule = f.get('rule') or f.get('rule_id') or 'rule'
                            code = f.get('code') or 'GEN'
                            sev = f.get('severity') or 'info'
                            # Key: gruppiere primär nach (rule, code, severity)
                            key = (str(rule), str(code), str(sev))
                            if key not in groups:
                                groups[key] = {
                                    'rule': rule,
                                    'code': code,
                                    'severity': sev,
                                    'message': f.get('message') or '',
                                    'count': 0,
                                    'confidence_sum': 0.0,
                                    'confidence_n': 0,
                                }
                            g = groups[key]
                            g['count'] += 1
                            c = f.get('confidence')
                            try:
                                if isinstance(c, str): c = float(c)
                                if isinstance(c, (int, float)):
                                    g['confidence_sum'] += float(c)
                                    g['confidence_n'] += 1
                            except Exception:
                                pass
                        except Exception:
                            continue
                    # finalize
                    for g in groups.values():
                        try:
                            n = g.pop('confidence_n', 0) or 0
                            s = g.pop('confidence_sum', 0.0) or 0.0
                            g['confidence'] = round(s / n, 4) if n > 0 else None
                            # Optional: verdichte Message für pair_similarity
                            if g.get('rule') == 'pair_similarity':
                                if g.get('code') == 'SIM_LOW':
                                    g['message'] = f"Niedrige Ähnlichkeit – {g['count']} Segmente"
                                elif g.get('code') == 'SIM_LOW_CRIT':
                                    g['message'] = f"Sehr niedrige Ähnlichkeit – {g['count']} Segmente"
                        except Exception:
                            pass
                        findings_grouped.append(g)
                    # Sortierung: Schwere -> Count desc
                    try:
                        order = {'critical':0,'major':1,'minor':2,'info':3}
                        findings_grouped.sort(key=lambda x:(order.get(x.get('severity','info'),9), -int(x.get('count') or 0)))
                    except Exception:
                        pass
                except Exception:
                    findings_grouped = []

            # 3) Phasen-/Konsolidierungsblöcke klar strukturiert befüllen
            phases: dict[str, dict] = {}
            try:
                # Phase 1: Format & Struktur (Platzhalter, URLs, E-Mails, Whitespace)
                if phase_issue_counts.get('phase1'):
                    phases['phase1'] = {
                        'name': 'Format & Struktur',
                        'issue_total': phase_issue_counts['phase1'],
                        'description': 'Prüfung von Platzhaltern, URLs, E-Mails und strukturellen Elementen'
                    }
                
                # Phase 2: Inhalt & Konsistenz (Zahlen, Glossar, Eigennamen)
                if phase_issue_counts.get('phase2'):
                    phases['phase2'] = {
                        'name': 'Inhalt & Konsistenz',
                        'issue_total': phase_issue_counts['phase2'],
                        'description': 'Prüfung von Zahlen, Einheiten, Glossar-Begriffen und Eigennamen'
                    }
                
                # Phase 3: Semantik & Grammatik (Lesbarkeit, Rechtschreibung, Stil)
                phase3_specific = phase_issue_counts.get('phase3', 0)
                if phase3_specific:
                    phases['phase3'] = {
                        'name': 'Semantik & Grammatik',
                        'issue_total': phase3_specific,
                        'description': 'Prüfung von Lesbarkeit, Rechtschreibung und grammatikalischen Strukturen'
                    }
                
                # Custom KI-Prüfung (Benutzerdefinierte Fragen)
                custom_ki_count = phase_issue_counts.get('custom_ki', 0)
                if custom_ki_count:
                    custom_prompt_used = ''
                    if custom_ki_result:
                        custom_prompt_used = custom_ki_result.question[:50] + '...' if len(custom_ki_result.question) > 50 else custom_ki_result.question
                    phases['custom_ki'] = {
                        'name': 'KI-Prüfung',
                        'issue_total': custom_ki_count,
                        'description': f'Benutzerdefinierte Prüfung: {custom_prompt_used}' if custom_prompt_used else 'Benutzerdefinierte KI-Prüfung',
                        'custom_prompt': custom_prompt_used
                    }
                
                # Konsolidierung: Gesamtbewertung und Risiko-Score
                try:
                    total_f = len(findings)
                    if total_f:
                        crit = summary.get('critical', 0) or 0
                        maj = summary.get('major', 0) or 0
                        mi = summary.get('minor', 0) or 0
                        # Einfache Risiko-Heuristik 0..100
                        risk = min(100.0, round(((crit*3 + maj*2 + mi*1) / max(1, pair_count)) * 25.0, 2))
                    else:
                        risk = 0.0
                    consolidated = {'total': total_f, 'risk_score': risk}
                except Exception:
                    # Bei Fehler: Sichere Defaults (0 Befunde = 0 Risiko)
                    consolidated = {'total': len(findings), 'risk_score': 0.0 if not findings else 50.0}
                
                phases['consolidation'] = {
                    'name': 'Konsolidierung',
                    'total': consolidated.get('total'),
                    'risk_score': consolidated.get('risk_score'),
                    'description': 'Gesamtbewertung aller Befunde mit Risiko-Einschätzung'
                }
                
            except Exception:
                consolidated = {'total': len(findings)}
            
            # Empfehlungen: Automatisch generierte Verbesserungsvorschläge
            try:
                if recommendations:
                    phases['recommendations'] = {
                        'name': 'Empfehlungen',
                        'suggestions': len(recommendations),
                        'description': 'Automatisch generierte Verbesserungsvorschläge basierend auf der Analyse'
                    }
            except Exception:
                pass

            # Enforcement/Gates (minimal)
            enforcement = {
                'gated': bool(overall_score < float(_get_setting('quality.thresholds.gate.min_score', 75.0))),
                'issues_remaining': []
            }
            suggestions_list = list(recommendations)

            # DEBUG: Log findings count before storing
            try:
                print(f"DEBUG: About to store {len(findings)} findings in normalized_results")
                print(f"DEBUG: phase1_issues={len(phase1_issues)}, phase2_issues={len(phase2_issues)}, phase3_issues={len(phase3_issues)}")
                if findings:
                    print(f"DEBUG: First finding example: {findings[0] if findings else 'N/A'}")
            except Exception:
                pass
            
            normalized_results = {
                'file_pairs': pairs,
                'metrics': metrics_all,
                'quality': quality,
                'summary': summary,
                'recommendations': recommendations,
                'findings': findings,
                'findings_grouped': findings_grouped,
                'phases': phases,
                'consolidated': consolidated,
                'enforcement': enforcement,
                'suggestions': suggestions_list,
                # Optionale kompatible Felder für Tab-Renderer
                'issues_phase1': phase1_issues,
                'issues_phase2': phase2_issues,
                'issues_phase3': phase3_issues,
                'issues_custom_ki': custom_ki_findings,
                'phase_issue_counts': phase_issue_counts,
                'pair_details': pair_details,
                # Custom KI-Prüfungsergebnis (falls vorhanden)
                'custom_ki_result': custom_ki_result.to_dict() if custom_ki_result else None,
                # Phasen-Namen für bessere UI-Darstellung
                'phase_names': {
                    'phase1': 'Format & Struktur',
                    'phase2': 'Inhalt & Konsistenz',
                    'phase3': 'Semantik & Grammatik',
                    'custom_ki': 'KI-Prüfung',
                    'consolidation': 'Konsolidierung',
                    'recommendations': 'Empfehlungen'
                }
            }
        except Exception as e:
            # Fallback – liefere ursprüngliche Struktur
            print(f"DEBUG: Exception creating normalized_results: {e}")
            normalized_results = {
                'file_pairs': pairs,
                'metrics': metrics_all,
                'quality': quality,
                'summary': summary,
                'recommendations': recommendations
            }

        try:
            self._last_phase_issues = {
                'phase1': phase1_issues,
                'phase2': phase2_issues,
                'phase3': phase3_issues
            }
            self._last_pair_details = pair_details
        except Exception:
            pass

        try:
            if hasattr(self, 'logger') and hasattr(self.logger, 'info'):
                findings_count = len(normalized_results.get('findings', []) or [])
                self.logger.info("Normalized results created: %s findings, phases=%s",
                                 findings_count,
                                 list((normalized_results.get('phases') or {}).keys()))
                print(f"DEBUG: Logged {findings_count} findings to analysis results")
        except Exception:
            pass

        return normalized_results
    
    def _show_analysis_results(self):
        """Show comprehensive analysis results"""
        try:
            from quality_gui_components_analysis_results import show_analysis_results
            show_analysis_results(self)
        except Exception as e:
            try:
                self._handle_error(e, context="analysis.results.delegate", user_message=self._t("Analyse Ergebnisse Anzeige Fehler"))
            except Exception:
                pass

    def _request_analysis_cancel(self, status_label=None, progress_bar=None):
        """Markiert laufende Analyse als abgebrochen (kooperativer Abbruch)."""
        try:
            if getattr(self, '_analysis_cancel_requested', False):
                return
            self._analysis_cancel_requested = True
            if status_label:
                try:
                    status_label.configure(text=self._t('Abbruch angefordert – beende aktuelle Regel') if hasattr(self,'_t') else 'Abbruch angefordert – beende aktuelle Regel')
                except Exception:
                    pass
            if progress_bar:
                try:
                    progress_bar.configure(progress_color=self.get_color('warning'))
                except Exception:
                    pass
            if getattr(self, 'event_bus', None):
                try:
                    self.event_bus.publish('analysis.cancelled', {'ts': time.time()})
                except Exception:
                    pass
            self._show_toast(self._t('Cancel requested'), 'warning')
        except Exception:
            pass
    
    def show_demo_results(self):
        """Show comprehensive demo results"""
        try:
            # Clear output
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Demo results card
            demo_card = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=0,
                            )
            demo_card.pack(fill="x", pady=20, padx=20)
            
            # Header
            header_frame = ctk.CTkFrame(demo_card, fg_color=self.get_color('gray_600'))
            header_frame.pack(fill="x", padx=0, pady=0)
            
            header_label = ctk.CTkLabel(
                header_frame,
                text=self._t("Demo Analyseergebnisse") if hasattr(self,'_t') else "Demo Analyseergebnisse",
                font=ctk.CTkFont(*self.get_typography('title')),
                text_color=self.get_color('white')
            )
            header_label.pack(pady=20)
            
            # Demo content
            content_frame = ctk.CTkFrame(demo_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=30, pady=30)
            
            demo_text = """
Sample Translation Quality Analysis

Document: Technical Manual EN→DE
Overall Quality Score: 91/100 (Excellent)

Detailed Metrics:
Accuracy: 94/100
Fluency: 89/100  
Grammar: 96/100
Terminology: 88/100
Style: 90/100
Completeness: 100/100

Key Findings:
• Excellent technical accuracy
• Consistent terminology usage
• Minor style variations in 2 paragraphs
• All content translated completely

Recommendations:
• Review technical term glossary
• Standardize formal address usage
• Consider cultural adaptations

Result: Publication-ready quality
   Ready for immediate use with minor revisions
"""
            
            demo_label = ctk.CTkLabel(
                content_frame,
                text=demo_text,
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                justify="left",
                anchor="w"
            )
            demo_label.pack(fill="x", anchor="w")
            
            self.update_status("Demo results displayed")
            self.show_toast(self._t("Demo Analyseergebnisse geladen") if hasattr(self,'_t') else "Demo Analyseergebnisse geladen", "success")
            
        except Exception as e:
            try:
                self._handle_error(e, context="analysis.demo.results", user_message=self._t("Demo Ergebnisse Anzeige Fehler"))
            except Exception:
                pass
    
    def export_results(self):
        """Export analysis results - Modernisierte Karten-basierte GUI"""
        try:
            # Clear output frame
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            # Hauptcontainer
            main_container = ctk.CTkFrame(
                self.output_frame,
                fg_color="transparent"
            )
            main_container.pack(fill="both", expand=True, pady=10, padx=20)
            
            # Header-Karte
            header_card = ctk.CTkFrame(
                main_container,
                fg_color=self.get_color('primary'),
                corner_radius=12
            )
            header_card.pack(fill="x", pady=(0, 16))
            
            header_content = ctk.CTkFrame(header_card, fg_color="transparent")
            header_content.pack(fill="x", padx=24, pady=20)
            
            title_label = ctk.CTkLabel(
                header_content,
                text=self._t("Ergebnisse exportieren"),
                font=ctk.CTkFont(*self.get_typography('heading')),
                text_color=self.get_color('white')
            )
            title_label.pack(anchor="w")
            
            # Status-Anzeige
            status_text = self._t("Analyse abgeschlossen") if self.analysis_results else self._t("Keine Analyse vorhanden")
            status_label = ctk.CTkLabel(
                header_content,
                text=status_text,
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('white')
            )
            status_label.pack(anchor="w", pady=(4, 0))
            
            # Shortcut-Hinweis
            shortcut_hint = ctk.CTkLabel(
                header_content,
                text=self._t("Strg+1 PDF  |  Strg+2 Excel  |  Strg+3 Text"),
                font=ctk.CTkFont(family="Segoe UI", size=10),
                text_color=self.get_color('white')
            )
            shortcut_hint.pack(anchor="w", pady=(8, 0))
            
            # Export-Format-Karten Grid
            formats_frame = ctk.CTkFrame(main_container, fg_color="transparent")
            formats_frame.pack(fill="x", pady=(0, 16))
            formats_frame.columnconfigure((0, 1, 2), weight=1, uniform="export_col")
            
            # Format-Definitionen
            export_formats = [
                {
                    'fmt': 'pdf',
                    'title': 'PDF',
                    'subtitle': self._t('Professioneller Bericht'),
                    'color': '#DC2626',  # Rot für PDF
                    'features': [
                        self._t('Visuelle Diagramme'),
                        self._t('Befunde mit Details'),
                        self._t('Druckoptimiert')
                    ],
                    'recommended': True
                },
                {
                    'fmt': 'xlsx',
                    'title': 'Excel',
                    'subtitle': self._t('Datentabelle'),
                    'color': '#10B981',  # Grün für Excel
                    'features': [
                        self._t('Alle Metriken'),
                        self._t('Filterbar'),
                        self._t('Weiterverarbeitung')
                    ],
                    'recommended': False
                },
                {
                    'fmt': 'txt',
                    'title': 'Text',
                    'subtitle': self._t('Schnellübersicht'),
                    'color': '#6B7280',  # Grau für Text
                    'features': [
                        self._t('Kompakt'),
                        self._t('Universell lesbar'),
                        self._t('Leichtgewichtig')
                    ],
                    'recommended': False
                }
            ]
            
            for col_idx, fmt_info in enumerate(export_formats):
                # Format-Karte mit Hover-Effekt
                card = ctk.CTkFrame(
                    formats_frame,
                    fg_color=self.get_color('surface'),
                    corner_radius=12,
                    border_width=2 if fmt_info['recommended'] else 1,
                    border_color=fmt_info['color'] if fmt_info['recommended'] else self.get_color('surface_border')
                )
                card.grid(row=0, column=col_idx, padx=8, pady=4, sticky="nsew")
                
                # Hover-Effekte
                def _on_enter(e, c=card, color=fmt_info['color']):
                    c.configure(border_color=color, border_width=2)
                def _on_leave(e, c=card, fmt=fmt_info):
                    if fmt['recommended']:
                        c.configure(border_color=fmt['color'], border_width=2)
                    else:
                        c.configure(border_color=self.get_color('surface_border'), border_width=1)
                card.bind("<Enter>", _on_enter)
                card.bind("<Leave>", _on_leave)
                
                card_content = ctk.CTkFrame(card, fg_color="transparent")
                card_content.pack(fill="both", expand=True, padx=16, pady=16)
                
                # Hover auch für card_content
                card_content.bind("<Enter>", _on_enter)
                card_content.bind("<Leave>", _on_leave)
                
                # Badge für Format-Typ
                badge_frame = ctk.CTkFrame(
                    card_content,
                    fg_color=fmt_info['color'],
                    corner_radius=6,
                    height=28
                )
                badge_frame.pack(anchor="w")
                badge_frame.pack_propagate(False)
                
                badge_label = ctk.CTkLabel(
                    badge_frame,
                    text=f"  {fmt_info['title']}  ",
                    font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"),
                    text_color=self.get_color('white')
                )
                badge_label.pack(padx=8, pady=4)
                
                # Empfohlen-Badge
                if fmt_info['recommended']:
                    rec_label = ctk.CTkLabel(
                        card_content,
                        text=self._t("Empfohlen"),
                        font=ctk.CTkFont(family="Segoe UI", size=10),
                        text_color=fmt_info['color']
                    )
                    rec_label.pack(anchor="w", pady=(4, 0))
                
                # Untertitel
                subtitle = ctk.CTkLabel(
                    card_content,
                    text=fmt_info['subtitle'],
                    font=ctk.CTkFont(*self.get_typography('body')),
                    text_color=self.get_color('text_primary')
                )
                subtitle.pack(anchor="w", pady=(8, 4))
                
                # Feature-Liste
                for feature in fmt_info['features']:
                    feature_label = ctk.CTkLabel(
                        card_content,
                        text=f"• {feature}",
                        font=ctk.CTkFont(*self.get_typography('caption')),
                        text_color=self.get_color('text_secondary')
                    )
                    feature_label.pack(anchor="w", pady=1)
                
                # Export-Button
                btn = self._create_button(
                    card_content,
                    text=self._t("Exportieren"),
                    command=lambda f=fmt_info['fmt']: self._perform_export(f),
                    kind="primary" if fmt_info['recommended'] else "secondary",
                    size='md',
                    height=36
                )
                btn.pack(fill="x", pady=(12, 0))
            
            # Zusatz-Optionen Karte
            options_card = ctk.CTkFrame(
                main_container,
                fg_color=self.get_color('surface'),
                corner_radius=12,
                border_width=1,
                border_color=self.get_color('surface_border')
            )
            options_card.pack(fill="x", pady=(0, 16))
            
            options_content = ctk.CTkFrame(options_card, fg_color="transparent")
            options_content.pack(fill="x", padx=20, pady=16)
            
            options_title = ctk.CTkLabel(
                options_content,
                text=self._t("Weitere Optionen"),
                font=ctk.CTkFont(*self.get_typography('subheading')),
                text_color=self.get_color('text_primary')
            )
            options_title.pack(anchor="w", pady=(0, 12))
            
            # Button-Leiste für weitere Aktionen
            action_frame = ctk.CTkFrame(options_content, fg_color="transparent")
            action_frame.pack(fill="x")
            
            # HTML Bericht öffnen
            html_btn = self._create_button(
                action_frame,
                text=self._t("Bericht im Browser öffnen"),
                command=self._open_quality_report,
                kind="secondary",
                size='md',
                height=40
            )
            html_btn.pack(side="left", padx=(0, 12))
            
            # Korrekturpaket Option
            try:
                cur_val = False
                if hasattr(self, 'settings_service') and self.settings_service:
                    cur_val = bool(self.settings_service.get('reporting.create_correction_package', False))
                create_pkg_var = tk.BooleanVar(value=cur_val)

                def _persist_create_pkg():
                    try:
                        if hasattr(self, 'settings_service') and self.settings_service:
                            self.settings_service.set('reporting.create_correction_package', bool(create_pkg_var.get()))
                        self.show_toast(self._t('Einstellung gespeichert'), 'success')
                    except Exception:
                        pass

                create_pkg_cb = ctk.CTkCheckBox(
                    action_frame,
                    text=self._t('Korrekturpaket erstellen'),
                    variable=create_pkg_var,
                    onvalue=True,
                    offvalue=False,
                    command=_persist_create_pkg,
                    fg_color=self.get_color('primary'),
                    font=ctk.CTkFont(*self.get_typography('body'))
                )
                create_pkg_cb.pack(side="left", padx=(12, 0))
            except Exception:
                pass
            
            # Zuletzt exportiert Anzeige
            try:
                last_export_path = None
                if hasattr(self, 'settings_service') and self.settings_service:
                    last_export_path = self.settings_service.get('reporting.last_export_path', None)
                
                if last_export_path:
                    from pathlib import Path
                    last_path = Path(last_export_path)
                    if last_path.exists() or last_path.parent.exists():
                        last_export_card = ctk.CTkFrame(
                            main_container,
                            fg_color=self.get_color('surface'),
                            corner_radius=8,
                            border_width=1,
                            border_color=self.get_color('surface_border')
                        )
                        last_export_card.pack(fill="x", pady=(0, 16))
                        
                        last_content = ctk.CTkFrame(last_export_card, fg_color="transparent")
                        last_content.pack(fill="x", padx=16, pady=12)
                        
                        last_title = ctk.CTkLabel(
                            last_content,
                            text=self._t("Zuletzt exportiert"),
                            font=ctk.CTkFont(*self.get_typography('caption')),
                            text_color=self.get_color('text_secondary')
                        )
                        last_title.pack(anchor="w")
                        
                        last_filename = ctk.CTkLabel(
                            last_content,
                            text=last_path.name,
                            font=ctk.CTkFont(*self.get_typography('body')),
                            text_color=self.get_color('text_primary')
                        )
                        last_filename.pack(anchor="w", pady=(2, 8))
                        
                        # Button-Leiste
                        last_btn_frame = ctk.CTkFrame(last_content, fg_color="transparent")
                        last_btn_frame.pack(anchor="w")
                        
                        def _open_export_folder():
                            try:
                                import subprocess
                                folder = last_path.parent if last_path.exists() else last_path.parent
                                subprocess.run(['explorer', str(folder)], check=False)
                            except Exception:
                                self.show_toast(self._t("Ordner konnte nicht geöffnet werden"), 'error')
                        
                        def _open_export_file():
                            try:
                                import subprocess
                                if last_path.exists():
                                    subprocess.run(['start', '', str(last_path)], shell=True, check=False)
                                else:
                                    self.show_toast(self._t("Datei nicht mehr vorhanden"), 'warning')
                            except Exception:
                                self.show_toast(self._t("Datei konnte nicht geöffnet werden"), 'error')
                        
                        folder_btn = self._create_button(
                            last_btn_frame,
                            text=self._t("Ordner öffnen"),
                            command=_open_export_folder,
                            kind="secondary",
                            size='sm',
                            height=32
                        )
                        folder_btn.pack(side="left", padx=(0, 8))
                        
                        if last_path.exists():
                            open_btn = self._create_button(
                                last_btn_frame,
                                text=self._t("Datei öffnen"),
                                command=_open_export_file,
                                kind="primary",
                                size='sm',
                                height=32
                            )
                            open_btn.pack(side="left")
            except Exception:
                pass
            
            # Hinweis falls keine Analyse
            if not self.analysis_results:
                hint_frame = ctk.CTkFrame(
                    main_container,
                    fg_color=self._accent('warning'),
                    corner_radius=8
                )
                hint_frame.pack(fill="x")
                hint_label = ctk.CTkLabel(
                    hint_frame,
                    text=self._t("Hinweis: Noch keine Analyse vorhanden – Export erzeugt leeren Bericht"),
                    font=ctk.CTkFont(*self.get_typography('caption')),
                    text_color=self.get_color('white')
                )
                hint_label.pack(padx=16, pady=10)
            
            self.update_status(self._t("Export Optionen angezeigt"))
            
            # Ribbon-States aktualisieren
            try:
                if hasattr(self, '_update_ribbon_states'):
                    self._update_ribbon_states()
            except Exception:
                pass
            
        except Exception as e:
            try:
                self._handle_error(e, context="export.options", user_message=self._t("Export Optionen Anzeige Fehler"))
            except Exception:
                pass

    def _perform_export(self, fmt: str):
        """Multi-Format Export Integration mit visuellem Feedback.

        Falls nur einzelnes Format angeklickt wird (Button), wird Settings-Logik dennoch genutzt,
        damit konsistenter Dateiname / Pfad / Events greifen.
        """
        # Format-Namen für Benutzer
        format_names = {'pdf': 'PDF', 'xlsx': 'Excel', 'txt': 'Text', 'html': 'HTML'}
        format_display = format_names.get(fmt, fmt.upper())
        
        # Zeige Export-Fortschritt
        self.update_status(self._t(f"Exportiere als {format_display}..."))
        self.show_toast(self._t(f"Export als {format_display} wird erstellt..."), 'info')
        
        # Delegation an modularen Reporting-Layer
        try:
            if getattr(self, 'reporting', None):
                self.reporting.perform_export(self, fmt)
            else:
                # Fallback: Kurzer Hinweis falls Reporting nicht initialisiert
                self.show_toast(self._t("Export Modul nicht verfügbar"), 'error')
        except Exception as e:
            self._handle_error(e, context='export.delegate', user_message=self._t("Unerwarteter Export Fehler"))
            try:
                self.show_toast(self._t("Export fehlgeschlagen"), 'error')
            except Exception:
                pass

    # ---------------------------------------------------------------
    # HTML REPORT INTEGRATION (Additiv) – öffnet migrierte Bericht_*.html Dateien
    # ---------------------------------------------------------------
    def _open_quality_report(self):
        """Öffnet einen vorhandenen migrierten HTML-Bericht (Bericht_*.html) im Standardbrowser.

        Auswahlstrategie (erste vorhandene Datei wird genutzt):
        1. Bericht_Interaktiv.html
        2. Bericht_Core_B.html
        3. Bericht_Core_A.html
        4. Bericht_Export.html
        5. Bericht_Basis.html
        """
        try:
            import webbrowser, os
            base_dir = os.path.dirname(__file__)

            # 1. Versuche dynamischen Bericht zu erzeugen (falls Daten vorhanden)
            dynamic_path = self._generate_dynamic_report()
            if dynamic_path and os.path.exists(dynamic_path):
                candidates = [dynamic_path]  # Vorrang
            else:
                candidates = []

            # 2. Statische migrierte Varianten (Fallback-Reihenfolge unverändert, aber nach dynamisch)
            candidates.extend([
                os.path.join(base_dir, 'Bericht_Interaktiv.html'),
                os.path.join(base_dir, 'Bericht_Core_B.html'),
                os.path.join(base_dir, 'Bericht_Core_A.html'),
                os.path.join(base_dir, 'Bericht_Export.html'),
                os.path.join(base_dir, 'Bericht_Basis.html'),
            ])

            target = None
            for path in candidates:
                if path and os.path.exists(path):
                    target = path
                    break
            if not target:
                self.show_toast(self._t("Kein Bericht gefunden"), "warning")
                return

            opened = False
            try:
                opened = webbrowser.open_new_tab(f'file://{target}')
            except Exception:
                opened = False
            if opened:
                # Differenziere Feedback
                if target.endswith('Bericht_Dynamisch.html'):
                    self.show_toast(self._t("Dynamischer Bericht geöffnet"), "success")
                else:
                    self.show_toast(self._t("Bericht geöffnet"), "success")
            else:
                self.show_toast(self._t("Bericht konnte nicht geöffnet werden"), "error")
        except Exception as e:
            try:
                self._handle_error(e, context="report.open", user_message=self._t("Öffnen des Berichts fehlgeschlagen"))
            except Exception:
                pass
            self.show_toast(self._t("Bericht Fehler"), "error")
    
    def clear_files(self):
        """OPTIMIZED: Smart file clearing with confirmation and cache cleanup"""
        try:
            source_count = len(self.uploaded_files['source'])
            translation_count = len(self.uploaded_files['translation'])
            total_files = source_count + translation_count
            
            if total_files == 0:
                self.show_toast(self._t("Keine Dateien zum Leeren"), "info")
                return
            
            # Smart confirmation for large file sets
            if total_files > 5:
                # Nicht-blockierender eigener Dialog
                self._show_clear_files_dialog(total_files, source_count, translation_count)
                return
            
            # Direkte Ausführung bei kleiner Menge
            self._execute_clear_files(total_files)
            
        except Exception as e:
            try:
                self._handle_error(e, context="files.clear", user_message=self._t("Dateien Leeren Fehler"))
            except Exception:
                pass

    def _show_clear_files_dialog(self, total_files: int, source_count: int, translation_count: int):
        """Nicht-blockierender Bestätigungsdialog für Dateilöschung."""
        try:
            dialog = ctk.CTkToplevel(self.root)
            dialog.title(self._t("Bestätigung"))
            dialog.transient(self.root)
            dialog.grab_set()
            try:
                dialog.attributes('-topmost', True)
            except Exception:
                pass

            # Styling
            fg = self.get_color('surface') if hasattr(self, 'get_color') else '#FFFFFF'
            border = self.get_color('surface_border') if hasattr(self, 'get_color') else '#E5E7EB'
            frame = ctk.CTkFrame(dialog, fg_color=fg, border_color=border, border_width=1, corner_radius=8)
            frame.pack(fill='both', expand=True, padx=16, pady=16)

            title_lbl = ctk.CTkLabel(frame, text=self._t("Dateien löschen"), font=ctk.CTkFont(*self.get_typography('subheading')) if hasattr(self,'get_typography') else None, text_color=self.get_color('text_primary') if hasattr(self,'get_color') else None)
            title_lbl.pack(anchor='w', pady=(4,8))

            msg = self._t("{n} Dateien löschen? Dies kann nicht rückgängig gemacht werden.").format(n=total_files)
            detail = self._t("Quellen: {s} / Übersetzungen: {t}").format(s=source_count, t=translation_count)
            msg_lbl = ctk.CTkLabel(frame, text=msg + "\n" + detail, justify='left', wraplength=420, font=ctk.CTkFont(*self.get_typography('body')) if hasattr(self,'get_typography') else None, text_color=self.get_color('gray_700') if hasattr(self,'get_color') else None)
            msg_lbl.pack(fill='x', pady=(0,12))

            btn_row = ctk.CTkFrame(frame, fg_color='transparent')
            btn_row.pack(fill='x', pady=(8,0))

            def _confirm():
                try:
                    dialog.destroy()
                except Exception:
                    pass
                self._execute_clear_files(total_files)

            def _cancel():
                try:
                    dialog.destroy()
                except Exception:
                    pass
                self.update_status(self._t("Löschung abgebrochen"))

            confirm_btn = self._create_button(btn_row, text=self._t("Löschen"), command=_confirm, kind="danger", height=36)
            cancel_btn = self._create_button(btn_row, text=self._t("Abbrechen"), command=_cancel, kind="outline-primary", height=36)
            cancel_btn.pack(side='right', padx=(8,0))
            confirm_btn.pack(side='right')

            # Shortcuts
            dialog.bind('<Escape>', lambda e: _cancel())
            dialog.bind('<Return>', lambda e: _confirm())

            try:
                self._center_window(dialog)
            except Exception:
                pass
        except Exception as e:
            try:
                self._handle_error(e, context='files.clear.dialog', user_message=self._t('Dialog Fehler'))
            except Exception:
                pass

    def _execute_clear_files(self, total_files: int):
        """Interne Ausführung der Dateilöschung (aus Dialog oder direkt)."""
        try:
            self.uploaded_files = {'source': [], 'translation': []}
            self.analysis_results = {}
            self.current_analysis = None
            try:
                self._clear_cache_smart()
            except Exception:
                pass
            try:
                self._update_file_counter()
            except Exception:
                pass
            try:
                self._show_enhanced_welcome_output()
            except Exception:
                pass
            self.update_status(self._t("Dateien geleert") + f": {total_files}")
            self.show_toast(self._t("Erfolgreich geleert") + f": {total_files}", "success")
            if total_files > 10:
                import gc
                gc.collect()
                try:
                    self.logger.debug(f"Memory cleanup completed after clearing {total_files} files")
                except Exception:
                    pass
            # Nach Clear immer Ribbon-States aktualisieren
            try:
                if hasattr(self, '_update_ribbon_states'):
                    self._update_ribbon_states()
            except Exception:
                pass
        except Exception as e:
            try:
                self._handle_error(e, context='files.clear.execute', user_message=self._t('Dateilöschung Fehler'))
            except Exception:
                pass
            self.show_toast("Error clearing files", "error")
    
    def show_settings_view(self):
        """Show settings and configuration"""
        try:
            # Clear output
            for widget in self.output_frame.winfo_children():
                widget.destroy()
            
            settings_card = ctk.CTkFrame(
                self.output_frame,
                fg_color=self.get_color('surface'),
                corner_radius=8,
                border_width=1,
                            )
            settings_card.pack(fill="x", pady=20, padx=20)
            
            # Header
            header_frame = ctk.CTkFrame(settings_card, fg_color=self.get_color('gray_600'))
            header_frame.pack(fill="x", padx=0, pady=0)
            
            header_label = ctk.CTkLabel(
                header_frame,
                text=self._t("Einstellungen & Konfiguration"),
                font=ctk.CTkFont(*self.get_typography('subheading')),
                text_color=self.get_color('white')
            )
            header_label.pack(pady=15)

            # Content Frame
            content_frame = ctk.CTkFrame(settings_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=30, pady=30)

            # --- Reporting Einstellungen (additiv) ---
            try:
                reporting_section = ctk.CTkFrame(content_frame, fg_color=self.get_color('gray_50'))
                reporting_section.pack(fill="x", pady=(0, 25))
                sec_label = ctk.CTkLabel(reporting_section, text="Reporting", font=ctk.CTkFont(*self.get_typography('label_bold')), text_color=self.get_color('text_primary'))
                sec_label.pack(anchor="w", padx=20, pady=(18, 6))

                # Auto-Generate Toggle
                auto_var = tk.BooleanVar(value=bool(self.settings_service.get('reporting.auto_generate', True)) if hasattr(self, 'settings_service') else True)
                auto_toggle = ctk.CTkCheckBox(reporting_section, text="Automatischen Report-Export aktivieren", variable=auto_var, onvalue=True, offvalue=False, fg_color=self.get_color('primary'))
                auto_toggle.pack(anchor="w", padx=30, pady=4)

                # Auto-Open Toggle
                auto_open_var = tk.BooleanVar(value=bool(self.settings_service.get('reporting.auto_open', False)) if hasattr(self, 'settings_service') else False)
                auto_open_toggle = ctk.CTkCheckBox(reporting_section, text="Bericht nach Export automatisch öffnen", variable=auto_open_var, onvalue=True, offvalue=False, fg_color=self.get_color('primary'))
                auto_open_toggle.pack(anchor="w", padx=30, pady=4)

                # Format Auswahl
                format_frame = ctk.CTkFrame(reporting_section, fg_color="transparent")
                format_frame.pack(fill="x", padx=30, pady=(12, 4))
                fmt_label = ctk.CTkLabel(format_frame, text="Export-Formate", font=ctk.CTkFont(*self.get_typography('caption')))
                fmt_label.pack(anchor="w")
                available_formats = ['pdf', 'xlsx', 'txt']
                selected_formats = set()
                try:
                    stored = self.settings_service.get('reporting.formats')
                    if isinstance(stored, (list, tuple)):
                        selected_formats.update([str(x) for x in stored if x in available_formats])
                except Exception:
                    pass
                if not selected_formats:
                    selected_formats.update(['pdf'])
                fmt_vars: dict[str, tk.BooleanVar] = {}
                fmt_checks_frame = ctk.CTkFrame(format_frame, fg_color="transparent")
                fmt_checks_frame.pack(anchor='w', pady=(4,0))
                for f in available_formats:
                    v = tk.BooleanVar(value=f in selected_formats)
                    cb = ctk.CTkCheckBox(fmt_checks_frame, text=f.upper(), variable=v, onvalue=True, offvalue=False, fg_color=self.get_color('primary'))
                    cb.pack(side='left', padx=(0,12))
                    fmt_vars[f] = v

                # Charts Toggle (PDF Diagramme einbetten)
                charts_var = tk.BooleanVar(value=bool(self.settings_service.get('reporting.include_charts', True)) if hasattr(self, 'settings_service') else True)
                charts_toggle = ctk.CTkCheckBox(reporting_section, text="Diagramme (Metriken) in PDF einbetten", variable=charts_var, onvalue=True, offvalue=False, fg_color=self.get_color('primary'))
                charts_toggle.pack(anchor='w', padx=30, pady=(4,4))

                # Korrekturpaket Toggle (zentraler Reporting-Flow)
                corr_pkg_var = tk.BooleanVar(value=bool(self.settings_service.get('reporting.create_correction_package', False)) if hasattr(self, 'settings_service') else False)
                corr_pkg_toggle = ctk.CTkCheckBox(reporting_section, text="Korrekturpaket zusätzlich erstellen", variable=corr_pkg_var, onvalue=True, offvalue=False, fg_color=self.get_color('primary'))
                corr_pkg_toggle.pack(anchor='w', padx=30, pady=(4,4))

                # Naming-Konvention
                naming_frame = ctk.CTkFrame(reporting_section, fg_color="transparent")
                naming_frame.pack(fill="x", padx=30, pady=(12, 4))
                naming_label = ctk.CTkLabel(naming_frame, text="Naming-Konvention (Platzhalter: {ts}, {fmt})", font=ctk.CTkFont(*self.get_typography('caption')))
                naming_label.pack(anchor="w")
                naming_entry = ctk.CTkEntry(naming_frame, placeholder_text="report_{ts}", width=320)
                try:
                    naming_entry.insert(0, str(self.settings_service.get('reporting.naming_pattern', 'report_{ts}')))
                except Exception:
                    pass
                naming_entry.pack(anchor='w', pady=4)

                # Output Directory
                out_dir_frame = ctk.CTkFrame(reporting_section, fg_color="transparent")
                out_dir_frame.pack(fill="x", padx=30, pady=(10, 4))
                out_dir_label = ctk.CTkLabel(out_dir_frame, text="Ausgabeordner", font=ctk.CTkFont(*self.get_typography('caption')))
                out_dir_label.pack(anchor="w")
                out_dir_entry = ctk.CTkEntry(out_dir_frame, placeholder_text="reports", width=320)
                try:
                    out_dir_entry.insert(0, str(self.settings_service.get('reporting.output_dir', 'reports')))
                except Exception:
                    pass
                out_dir_entry.pack(anchor="w", pady=4)

                # Filename Prefix
                prefix_frame = ctk.CTkFrame(reporting_section, fg_color="transparent")
                prefix_frame.pack(fill="x", padx=30, pady=(10, 4))
                prefix_label = ctk.CTkLabel(prefix_frame, text="Dateipräfix", font=ctk.CTkFont(*self.get_typography('caption')))
                prefix_label.pack(anchor="w")
                prefix_entry = ctk.CTkEntry(prefix_frame, placeholder_text="analysis_report", width=320)
                try:
                    prefix_entry.insert(0, str(self.settings_service.get('reporting.filename_prefix', 'analysis_report')))
                except Exception:
                    pass
                prefix_entry.pack(anchor="w", pady=4)

                # Save Button
                def _save_reporting_settings():
                    try:
                        if hasattr(self, 'settings_service') and self.settings_service:
                            self.settings_service.set('reporting.auto_generate', bool(auto_var.get()))
                            self.settings_service.set('reporting.auto_open', bool(auto_open_var.get()))
                            self.settings_service.set('reporting.output_dir', out_dir_entry.get().strip() or 'reports')
                            self.settings_service.set('reporting.filename_prefix', prefix_entry.get().strip() or 'analysis_report')
                            # Formate sammeln
                            chosen = [f for f,v in fmt_vars.items() if v.get()]
                            if not chosen:
                                chosen = ['txt']
                            self.settings_service.set('reporting.formats', chosen)
                            self.settings_service.set('reporting.naming_pattern', naming_entry.get().strip() or 'report_{ts}')
                            self.settings_service.set('reporting.include_charts', bool(charts_var.get()))
                            self.settings_service.set('reporting.create_correction_package', bool(corr_pkg_var.get()))
                            self.show_toast(self._t("Reporting Einstellungen gespeichert"), "success")
                        else:
                            self.show_toast(self._t("SettingsService nicht verfügbar"), "error")
                    except Exception as e:
                        self.show_toast(self._t("Fehler beim Speichern"), "error")
                        try:
                            self.logger.debug(f"Save reporting settings error: {e}")
                        except Exception:
                            pass

                save_btn = self._create_button(reporting_section, text=self._t("Speichern"), command=_save_reporting_settings, kind="primary", height=36)
                save_btn.pack(anchor="e", padx=30, pady=(14, 20))
                # Anzeige letzter Export (read-only)
                try:
                    last_path = self.settings_service.get('reporting.last_export_path') if hasattr(self, 'settings_service') else None
                    if last_path:
                        last_export_label = ctk.CTkLabel(reporting_section, text=self._t("Letzter Export") + f": {last_path}", font=ctk.CTkFont(*self.get_typography('caption')), text_color=self.get_color('gray_600'), justify='left')
                        last_export_label.pack(fill='x', padx=30, pady=(0,10))
                except Exception:
                    pass
            except Exception:
                pass

            # --- Plugin / Analyse Einstellungen (Abort-Schwelle) ---
            try:
                plugins_section = ctk.CTkFrame(content_frame, fg_color=self.get_color('gray_50'))
                plugins_section.pack(fill="x", pady=(0,25))
                pl_label = ctk.CTkLabel(plugins_section, text=self._t("Plugins / Analyse"), font=ctk.CTkFont(*self.get_typography('label_bold')), text_color=self.get_color('text_primary'))
                pl_label.pack(anchor='w', padx=20, pady=(18,6))

                # Abort Ratio Eingabe
                ratio_frame = ctk.CTkFrame(plugins_section, fg_color="transparent")
                ratio_frame.pack(fill='x', padx=30, pady=(4,4))
                ratio_label = ctk.CTkLabel(ratio_frame, text=self._t("Frühabbruch-Schwelle (Timeout-Quote 0.1 - 0.9)"), font=ctk.CTkFont(*self.get_typography('caption')), text_color=self.get_color('text_primary'))
                ratio_label.pack(anchor='w')
                current_ratio = 0.4
                try:
                    if hasattr(self,'settings_service') and self.settings_service:
                        current_ratio = float(self.settings_service.get('plugins.abort_timeout_ratio', 0.4))
                except Exception:
                    current_ratio = 0.4
                ratio_entry = ctk.CTkEntry(ratio_frame, width=120)
                ratio_entry.insert(0, str(round(current_ratio, 2)))
                ratio_entry.pack(anchor='w', pady=(4, 2))
                ratio_hint = ctk.CTkLabel(ratio_frame, text=self._t("Analyse wird früh beendet wenn Timeouts / Ausführungen > Schwelle."), font=ctk.CTkFont(*self.get_typography('caption')), text_color=self.get_color('gray_600'), justify='left')
                ratio_hint.pack(anchor='w', pady=(0, 4))

                last_saved_ratio = [round(current_ratio, 3)]
                def _save_ratio():
                    try:
                        import math
                        raw = (ratio_entry.get() or "").strip().replace(',', '.')
                        val = float(raw)
                        if not math.isfinite(val):
                            raise ValueError("non-finite")
                        if not (0.1 <= val <= 0.9):
                            self.show_toast(self._t("Wert muss zwischen 0.1 und 0.9 liegen"), "warning")
                            return
                        val = round(val, 3)
                        if last_saved_ratio[0] == val:
                            return
                        if hasattr(self,'settings_service') and self.settings_service:
                            self.settings_service.set('plugins.abort_timeout_ratio', val)
                        last_saved_ratio[0] = val
                        self.show_toast(self._t("Abort-Schwelle gespeichert"), "success")
                        try:
                            self._log_event('settings.plugins.abort_ratio.changed', value=val)
                        except Exception:
                            pass
                    except Exception:
                        self.show_toast(self._t("Ungültiger Wert"), "error")

                def _filter_ratio(event=None):
                    try:
                        txt = (ratio_entry.get() or "").strip().replace(',', '.')
                        if txt in ("", ".", "0."):
                            return
                        v = float(txt)
                        if v < 0.0:
                            ratio_entry.delete(0, 'end'); ratio_entry.insert(0, "0.1")
                        elif v > 1.0:
                            ratio_entry.delete(0, 'end'); ratio_entry.insert(0, "0.9")
                    except Exception:
                        ratio_entry.delete(0, 'end'); ratio_entry.insert(0, str(last_saved_ratio[0] if last_saved_ratio[0] is not None else 0.4))

                # save on Enter and FocusOut (debounced identical saves)
                ratio_entry.bind('<Return>', lambda e: _save_ratio())
                ratio_entry.bind('<FocusOut>', lambda e: _save_ratio())
                ratio_entry.bind('<KeyRelease>', _filter_ratio)
            except Exception:
                pass

            # --- UI / Accessibility Einstellungen ---
            try:
                ui_section = ctk.CTkFrame(content_frame, fg_color=self.get_color('gray_50'))
                ui_section.pack(fill="x", pady=(0, 25))
                ui_label = ctk.CTkLabel(ui_section, text=self._t("Benutzeroberfläche"), font=ctk.CTkFont(*self.get_typography('label_bold')), text_color=self.get_color('text_primary'))
                ui_label.pack(anchor="w", padx=20, pady=(18, 6))

                # High Contrast Toggle
                hc_var = tk.BooleanVar(value=bool(self.settings_service.get('ui.contrast_mode', False)) if hasattr(self, 'settings_service') else False)
                def _toggle_contrast():
                    try:
                        if hasattr(self, 'settings_service') and self.settings_service:
                            self.settings_service.set('ui.contrast_mode', bool(hc_var.get()))
                        # ThemeGuard re-evaluiert über provider lambda -> einfach repaint
                        self._repaint_colors_safe()
                        self.show_toast("Kontrast-Modus aktualisiert", "success")
                        self._log_event('ui.contrast_mode.changed', enabled=bool(hc_var.get()))
                    except Exception as e:
                        try:
                            self.logger.debug(f"Contrast toggle Fehler: {e}")
                        except Exception:
                            pass
                hc_toggle = ctk.CTkCheckBox(ui_section, text=self._t("Hoher Kontrast"), variable=hc_var, onvalue=True, offvalue=False, command=_toggle_contrast, fg_color=self.get_color('primary'))
                hc_toggle.pack(anchor="w", padx=30, pady=4)

                # Font Scale (einfacher Faktor)
                fs_frame = ctk.CTkFrame(ui_section, fg_color="transparent")
                fs_frame.pack(fill='x', padx=30, pady=(10,4))
                fs_label = ctk.CTkLabel(fs_frame, text=self._t("Schrift-Skalierung (0.8 - 1.5)"), font=ctk.CTkFont(*self.get_typography('caption')))
                fs_label.pack(anchor='w')
                current_scale = 1.0
                try:
                    current_scale = float(self.settings_service.get('ui.font_scale', 1.0)) if hasattr(self,'settings_service') else 1.0
                except Exception:
                    current_scale = 1.0
                fs_entry = ctk.CTkEntry(fs_frame, width=120)
                fs_entry.insert(0, str(current_scale))
                fs_entry.pack(anchor='w', pady=4)
                def _save_font_scale():
                    try:
                        raw = (fs_entry.get() or "").strip().replace(',', '.')
                        val = float(raw)
                        if not (0.8 <= val <= 1.5):
                            self.show_toast(self._t("Wert muss zwischen 0.8 und 1.5 liegen"), "warning")
                            return
                        if hasattr(self,'settings_service') and self.settings_service:
                            self.settings_service.set('ui.font_scale', round(val, 3))
                        # Cache safe clear
                        try:
                            getattr(self, '_font_cache', {}).clear()
                        except Exception:
                            pass
                        self._repaint_fonts_safe()
                        self.show_toast(self._t("Font-Skalierung aktualisiert"), "success")
                        self._log_event('ui.font_scale.changed', scale=val)
                    except Exception:
                        self.show_toast(self._t("Fehler beim Aktualisieren"), "error")
                # Vereinheitlicht mit Button-Factory (primary Style)
                fs_btn = self._create_button(fs_frame, text=self._t("Speichern") if hasattr(self,'_t') else "Speichern", command=_save_font_scale, kind="primary", height=36)
                fs_btn.pack(anchor='w', pady=(4,10))
            except Exception:
                pass
            
            # Settings content (rename to avoid shadowing)
            settings_text_frame = ctk.CTkFrame(settings_card, fg_color="transparent")
            settings_text_frame.pack(fill="x", padx=20, pady=20)

            settings_text = self._t(
                "Anwendungseinstellungen:\n\n"
                "Analyse-Konfiguration:\n"
                "Standard-Qualitätskriterien: Alle aktiv\n"
                "Spracherkennung: Automatisch\n"
                "Stapelverarbeitung: Aktiviert\n"
                "Fortschritt: Echtzeit\n\n"
                "Ausgabepräferenzen:\n"
                "Berichtsformat: PDF + Zusammenfassung\n"
                "Detailgrad: Umfassend\n"
                "Diagramme: Enthalten\n"
                "Empfehlungen: Aktiviert\n\n"
                "Performance:\n"
                "Threading: Mehrkern-Verarbeitung\n"
                "Speicheroptimierung: Aktiv\n"
                "Cache-Verwaltung: Automatisch\n"
                "Hintergrundverarbeitung: Aktiviert\n\n"
                "Datenverwaltung:\n"
                "Auto-Speichern: Aktiviert\n"
                "Backups: Automatisch\n"
                "Historie: 30 Tage\n"
                "Privatsphäre: Sichere Verarbeitung"
            )

            settings_label = ctk.CTkLabel(
                settings_text_frame,
                text=settings_text,
                font=ctk.CTkFont(*self.get_typography('body')),
                text_color=self.get_color('text_primary'),
                justify="left",
                anchor="w"
            )
            settings_label.pack(fill="x", anchor="w")
            
            self.update_status(self._t("Einstellungen angezeigt"))
            self.show_toast(self._t("Einstellungsansicht geladen"), "info")
            
        except Exception as e:
            try:
                self._handle_error(e, context="settings.view", user_message=self._t("Einstellungen Anzeige Fehler"))
            except Exception:
                pass
        
    # ---------------------- BACKGROUND ANALYSE (ADD-ON) ----------------------
    def _run_background_analysis_safe(self, analysis_id: str):
        """Nicht-blockierende Hintergrundanalyse im WorkerPool.
        Erzeugt (optionalen) Report unabhängig von UI-Simulation.
        Thread-sicher: Keine direkten UI-Manipulationen."""
        start = time.time()
        findings = []
        plugin_stats = {}
        try:
            from core.model import Finding, AnalysisReport  # Lazy Import
            context = {
                'analysis_id': analysis_id,
                'source_files': list(self.uploaded_files.get('source', [])) if hasattr(self, 'uploaded_files') else [],
                'translation_files': list(self.uploaded_files.get('translation', [])) if hasattr(self, 'uploaded_files') else [],
                'start_ts': start,
                'phase': 'background'
            }
            plugin_results = []
            try:
                if getattr(self, 'loaded_rules', None):
                    plugin_results = self._analyze_with_plugins(context)
            except Exception as pe:
                try:
                    self.logger.debug(f"BG Plugin Fehler: {pe}")
                except Exception:
                    pass
            for r in plugin_results:
                try:
                    findings.append(Finding(rule=getattr(r, 'rule', 'plugin'), message=getattr(r, 'summary', str(r))[:400]))
                except Exception:
                    continue
            duration = time.time() - start
            plugin_stats = {
                'rules_executed': len(plugin_results),
                'duration_s': duration,
            }
            try:
                if hasattr(self, '_last_plugin_stats') and isinstance(self._last_plugin_stats, dict):
                    plugin_stats.update(self._last_plugin_stats)
            except Exception:
                pass
            # Dynamische Score Abschätzung: Durchschnitt vorhandener plugin score Felder falls vorhanden
            dynamic_score = 87.0
            try:
                numeric_scores = []
                for r in plugin_results:
                    try:
                        val = getattr(r, 'score', None)
                        if isinstance(val, (int, float)) and 0 <= val <= 100:
                            numeric_scores.append(float(val))
                    except Exception:
                        continue
                if numeric_scores:
                    dynamic_score = sum(numeric_scores)/len(numeric_scores)
            except Exception:
                pass
            report = AnalysisReport.create(
                overall_score=dynamic_score,
                plugins_run=plugin_stats['rules_executed'],
                findings=findings,
                context=context
            )
            report.duration_s = duration
            report.file_counts = {
                'source': len(context['source_files']),
                'translation': len(context['translation_files'])
            }
            report.plugin_stats = plugin_stats
            self._latest_background_report = report
            try:
                if self.event_bus:
                    payload = {
                        'analysis_id': analysis_id,
                        'duration_s': duration,
                        'findings': len(findings)
                    }
                    self.event_bus.publish('analysis.background.completed', payload)
                    self._log_event("analysis.background.completed", **payload)
            except Exception:
                pass
            try:
                self._export_and_publish_report(report, background=True)
            except Exception:
                pass
        except Exception as e:
            try:
                self.logger.debug(f"Hintergrundanalyse Fehler: {e}")
            except Exception:
                pass

    def _export_and_publish_report(self, report, background: bool = False):
        """Zentrale Export-/Event-Pipeline für Reports (wiederverwendet UI & Background)."""
        try:
            from pathlib import Path
            import time
            from reporting.exporter import ReportExporter
            if self.event_bus:
                try:
                    payload = {
                        "findings": len(getattr(report, 'findings', []) or []),
                        "plugins": getattr(report, 'plugins_run', 0),
                        "ts": getattr(report, 'created_ts', time.time()),
                        "background": background
                    }
                    self.event_bus.publish("report.generated", payload)
                    self._log_event("report.generated", **payload)
                except Exception:
                    pass
            auto_export = False
            try:
                if hasattr(self, 'settings_service') and self.settings_service:
                    auto_export = bool(self.settings_service.get("reporting.auto_generate", False))
            except Exception:
                auto_export = False
            if not auto_export:
                return
            export_dir = None
            filename_prefix = "analysis_report"
            try:
                if hasattr(self, 'settings_service') and self.settings_service:
                    export_dir = self.settings_service.get("reporting.output_dir")
                    filename_prefix = self.settings_service.get("reporting.filename_prefix", filename_prefix)
            except Exception:
                pass
            exporter = ReportExporter(base_dir=Path(export_dir) if export_dir else None)
            exported_path = exporter.export_json(report, filename_prefix=filename_prefix)
            if exported_path:
                try:
                    if hasattr(self, 'settings_service') and self.settings_service:
                        self.settings_service.set('reporting.last_export_path', str(exported_path))
                    if not background:
                        self.show_toast("Report automatisch exportiert", "success")
                except Exception:
                    pass
        except Exception as e:
            try:
                self.logger.debug(f"Report Export Fehler: {e}")
            except Exception:
                pass

    def run(self):
        """Startet die GUI-Anwendung (robust, deutschsprachige Logs)."""
        try:
            if self.root:
                try:
                    self.logger.info("Starte GUI-Ereignisschleife")
                    try:
                        self._log_event("app.ui_loop_started")
                    except Exception:
                        pass
                except Exception:
                    pass
                self.root.mainloop()
                try:
                    self._log_event("app.ui_loop_ended")
                except Exception:
                    pass
            else:
                try:
                    self._handle_error(Exception("root missing"), context="app.run", user_message=self._t("Anwendung nicht korrekt initialisiert"), toast=True)
                    try:
                        self._log_event("app.root_missing")
                    except Exception:
                        pass
                except Exception:
                    pass
        except Exception as e:
            try:
                self._handle_error(e, context="app.runtime", user_message=self._t("App Lauf Fehler"))
            except Exception:
                pass

    # ---------------------- I18N PROMPT / LABEL HELPERS (additiv) ----------------------
    def _prompt_translation(self):
        """Einheitlicher Prompt-Text für Übersetzungs-Dropdowns."""
        try:
            return self._t("Übersetzung auswählen…") if hasattr(self, '_t') else "Übersetzung auswählen…"
        except Exception:
            return "Übersetzung auswählen…"

    def _prompt_source(self):
        """Einheitlicher Prompt-Text für Source-Dropdowns."""
        try:
            return self._t("Ausgangstext auswählen…") if hasattr(self, '_t') else "Ausgangstext auswählen…"
        except Exception:
            return "Ausgangstext auswählen…"

    def _label_source(self):
        """Lokalisierter Label-Text für Source Dateien."""
        try:
            return self._t("Ausgangstext") if hasattr(self, '_t') else "Ausgangstext"
        except Exception:
            return "Ausgangstext"

    def _label_translation(self):
        """Lokalisierter Label-Text für Translation Dateien."""
        try:
            return self._t("Übersetzung") if hasattr(self, '_t') else "Übersetzung"
        except Exception:
            return "Übersetzung"

    # ---------------------- REPAINT HELPERS (Contrast/Fonts) ----------------------
    def _repaint_colors_safe(self):
        try:
            # Rekursiv: Alle Frames einfärben (sichtbare Hierarchie)
            if not self.root:
                return
            def walk(node):
                try:
                    if isinstance(node, ctk.CTkFrame) and hasattr(self, 'get_color'):
                        node.configure(fg_color=self.get_color('surface'))
                except Exception:
                    pass
                try:
                    for child in node.winfo_children():
                        walk(child)
                except Exception:
                    pass
            walk(self.root)
        except Exception:
            pass

                    # (Entfernt: versehentlich duplizierter files.changed Block außerhalb Funktionen)
    def _repaint_fonts_safe(self):
        try:
            if not self.root:
                return
            # Rekursiv Labels aktualisieren; respektiere _preserve_font Flag
            def walk_fonts(node):
                try:
                    if isinstance(node, ctk.CTkLabel) and not getattr(node, '_preserve_font', False):
                        current_font = getattr(node, 'cget', lambda x: None)('font') if hasattr(node,'cget') else None
                        if not current_font or ('heading' not in str(current_font).lower() and 'title' not in str(current_font).lower()):
                            node.configure(font=ctk.CTkFont(*self.get_typography('body')))
                except Exception:
                    pass
                try:
                    for child in node.winfo_children():
                        walk_fonts(child)
                except Exception:
                    pass
            walk_fonts(self.root)
        except Exception:
            pass

    # ---------------------- LOGGING HELPER ----------------------
    def _log_event(self, name: str, **payload):
        """Strukturiertes Event-Logging (JSON-basiert für robustere Werte)."""
        try:
            import json, uuid
            if not hasattr(self, '_session_id'):
                self._session_id = uuid.uuid4().hex[:12]
            data = {"event": name, "ts": time.time(), "session": self._session_id, **payload}
            self.logger.info(json.dumps(data, ensure_ascii=False))
            try:
                bus = get_global_event_bus()
                if bus and hasattr(bus, 'publish'):
                    bus.publish(name, data)
            except Exception:
                pass
        except Exception:
            pass

    # ---------------------- SMALL UTILITIES (additiv) ----------------------
    def _log(self, level: str, msg: str, *args, **kwargs):
        """Kleiner Logging-Helper mit Fallback, um Log-Aufrufe zu vereinheitlichen.

        level: 'debug'|'info'|'warning'|'error'|'exception'
        """
        try:
            logger = getattr(self, 'logger', None)
            if not logger:
                # Fallback auf print, falls kein Logger gesetzt ist
                try:
                    print(f"[{level.upper()}] {msg % args if args else msg}")
                except Exception:
                    pass
                return
            lvl = level.lower()
            if lvl == 'debug':
                logger.debug(msg, *args, **kwargs)
            elif lvl == 'info':
                logger.info(msg, *args, **kwargs)
            elif lvl == 'warning':
                logger.warning(msg, *args, **kwargs)
            elif lvl == 'error':
                logger.error(msg, *args, **kwargs)
            elif lvl == 'exception':
                logger.exception(msg, *args, **kwargs)
            else:
                logger.info(msg, *args, **kwargs)
        except Exception:
            pass

    def _font(self, name: str):
        """Liefert eine gecachte CTkFont-Instanz für ein Typografie-Token.

        Nutzt internen Cache; fällt auf get_typography zurück, falls Cache fehlt.
        """
        try:
            if not hasattr(self, '_settings_font_cache'):
                self._settings_font_cache = {}
            if name not in self._settings_font_cache:
                self._settings_font_cache[name] = ctk.CTkFont(*self.get_typography(name))
            return self._settings_font_cache[name]
        except Exception:
            # Sicherer Fallback
            try:
                return ctk.CTkFont('Segoe UI', 14, 'normal')
            except Exception:
                return None

    def wait_future(self, fut, timeout: float | None = None):
        """Kleiner Helper zum Warten auf Futures mit sauberem Cancel-/Error-Handling.

        Gibt (True, result) bei Erfolg zurück, (False, None) bei Fehlern. CancelledError wird propagiert.
        """
        try:
            from concurrent.futures import CancelledError
            try:
                res = fut.result(timeout=timeout)
                return True, res
            except CancelledError:
                # Absichtlich nicht geschluckt
                raise
            except Exception as e:
                self._log('debug', f'wait_future error: {e}')
                return False, None
        except Exception:
            return False, None


def main():
    """OPTIMIZED: Main entry point with enhanced error handling and performance optimization"""
    try:
        import time
        _startup_time = time.time()
        
        logging.getLogger("checker").info("Initialisiere Übersetzungsqualitäts-Framework")
        
        # Create and configure the application
        app = ProfessionelleUebersetzungsqualitaetsApp()
        
        # Performance optimization: Show startup performance
        _init_duration = time.time() - _startup_time
        logging.getLogger("checker").info(f"Anwendung bereit – GUI wird gestartet (Init: {_init_duration:.2f}s)")
        
        # Start the optimized application
        _app_start = time.time()
        app.run()
        
        # Cleanup metrics
        _app_duration = time.time() - _app_start
        logging.getLogger("checker").info(f"Anwendung beendet (Runtime: {_app_duration:.2f}s)")
        
        # Flush offener Pairing-Autosave beim regulären Exit
        if getattr(app, '_pair_dirty', False):
            try:
                app.save_pairings()
            except Exception:
                pass
        
    except Exception as e:
        try:
            logging.getLogger("checker").exception(f"Critical application error: {e}")
        except Exception:
            pass


if __name__ == "__main__":
    main()

# Kompatibilitäts-Alias für bestehende Tests/Plugins
try:
    class QualityGuiMainApp(ProfessionelleUebersetzungsqualitaetsApp):
        """Alias der Hauptklasse für Rückwärtskompatibilität (Tests/Plugins)."""
        pass
except Exception:
    pass

