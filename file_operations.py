import os
import json
import re
from datetime import datetime
import sys # For resource_path
from docx import Document
from PyPDF2 import PdfReader, errors as PyPDF2Errors
import chardet
import logging

# --- OCR Imports ---
OCR_ENABLED = False
PDF2IMAGE_ENABLED = False
PYTESSERACT_ENABLED = False
MIN_TEXT_LENGTH_FOR_PDF_OCR_THRESHOLD = 100 # If PDF text < this, try OCR

# Flag to avoid repeated warnings
_WARNINGS_SHOWN = False

try:
    from PIL import Image
    OCR_ENABLED = True # PIL is a base requirement for pytesseract and pdf2image
except ImportError:
    if not _WARNINGS_SHOWN:
        print("WARNUNG: Pillow (PIL) nicht gefunden. OCR-Funktionalität ist eingeschränkt.")
        _WARNINGS_SHOWN = True
    Image = None

if OCR_ENABLED:
    try:
        import pytesseract
        # Attempt to get Tesseract version to check if it's installed and accessible
        # pytesseract.get_tesseract_version() # This can be slow if tesseract is not in PATH
        PYTESSERACT_ENABLED = True
        if not _WARNINGS_SHOWN:
            print("INFO: Pytesseract erfolgreich importiert.")
    except ImportError:
        if not _WARNINGS_SHOWN:
            print("WARNUNG: Pytesseract nicht gefunden. OCR für Bilder und gescannte PDFs ist nicht verfügbar.")
    except Exception as e: # Catching generic exception for Tesseract not found/configured
        if not _WARNINGS_SHOWN:
            print(f"WARNUNG: Tesseract OCR nicht korrekt konfiguriert oder nicht gefunden: {e}. OCR ist nicht verfügbar.")
            print("           Stellen Sie sicher, dass Tesseract installiert und im System-PATH ist oder pytesseract.tesseract_cmd konfiguriert wurde.")
        PYTESSERACT_ENABLED = False

if OCR_ENABLED:
    try:
        from pdf2image import convert_from_path
        from poppler_config import POPPLER_CONFIG, get_poppler_path_for_pdf2image
        
        # Verwende die automatische Poppler-Konfiguration
        if POPPLER_CONFIG and POPPLER_CONFIG.is_configured:
            PDF2IMAGE_ENABLED = True
            if not _WARNINGS_SHOWN:
                print("INFO: pdf2image mit automatischer Poppler-Konfiguration erfolgreich importiert.")
        else:
            # pdf2image ist installiert, aber Poppler fehlt
            if not _WARNINGS_SHOWN:
                print("INFO: pdf2image verfügbar, Poppler-Konfiguration empfohlen für vollständige PDF-Unterstützung.")
            # Aktiviere trotzdem grundlegende Funktionalität
            PDF2IMAGE_ENABLED = True
            
    except ImportError:
        if not _WARNINGS_SHOWN:
            print("WARNUNG: pdf2image nicht gefunden. OCR für PDF-Dateien ist nicht verfügbar.")
            print("         Nur relevant für PDF-OCR; mit 'pip install pdf2image' beheben.")
        PDF2IMAGE_ENABLED = False
    except Exception as e: # Catch other pdf2image related errors
        if not _WARNINGS_SHOWN:
            print(f"WARNUNG: Fehler beim Initialisieren von pdf2image: {e}. OCR für PDFs könnte eingeschränkt sein.")
        PDF2IMAGE_ENABLED = False
else:
    PDF2IMAGE_ENABLED = False

# Mark warnings as shown
_WARNINGS_SHOWN = True


# Mapping from common language codes to Tesseract language codes
# (Tesseract uses 3-letter ISO 639-2 codes typically)
LANG_CODE_TO_TESSERACT = {
    "de-DE": "deu", "de": "deu",
    "en-US": "eng", "en-GB": "eng", "en": "eng",
    "fr-FR": "fra", "fr": "fra",
    "es-ES": "spa", "es": "spa",
    "it-IT": "ita", "it": "ita",
    # Add more mappings as needed
}
DEFAULT_TESSERACT_LANG = "eng" # Fallback Tesseract language

# --- Helper function for OCR on a single image object ---
def _ocr_image_to_text(image_obj, lang_code_tesseract, status_update_func=None):
    if not PYTESSERACT_ENABLED or not OCR_ENABLED or image_obj is None:
        return ""
    try:
        if status_update_func:
            status_update_func(f"OCR wird für Bild durchgeführt (Sprache: {lang_code_tesseract})...")
        text = pytesseract.image_to_string(image_obj, lang=lang_code_tesseract)
        if status_update_func:
            status_update_func("OCR für Bild abgeschlossen.")
        return text
    except Exception as e:
        error_msg = f"Fehler bei OCR mit Tesseract: {e}"
        if "Tesseract is not installed or it's not in your PATH" in str(e) or \
           "tesseract is not installed or not in your path" in str(e).lower():
            error_msg += "\nBitte Tesseract OCR installieren und zum System-PATH hinzufügen."
        print(f"FEHLER: {error_msg}")
        if status_update_func:
            status_update_func(f"OCR fehlgeschlagen: {e}")
        return ""

# --- OCR Datei Funktion ---
def ocr_datei(dateipfad, lang_code_tesseract, status_update_func=None):
    """OCR-Funktion mit automatischer Poppler-Konfiguration"""
    if not PYTESSERACT_ENABLED or not OCR_ENABLED:
        if status_update_func:
            status_update_func("OCR nicht verfügbar (pytesseract oder Pillow fehlt).")
        return ""

    _, ext = os.path.splitext(dateipfad.lower())
    extracted_text = ""

    if ext == ".pdf":
        if not PDF2IMAGE_ENABLED:
            if status_update_func:
                status_update_func("OCR für PDF nicht verfügbar (pdf2image fehlt).")
            return ""
            
        try:
            if status_update_func:
                status_update_func(f"Konvertiere PDF zu Bildern für OCR: {os.path.basename(dateipfad)}...")
            
            # Verwende die automatische Poppler-Konfiguration
            poppler_path_env = get_poppler_path_for_pdf2image()

            images = convert_from_path(dateipfad, poppler_path=poppler_path_env)
            if status_update_func:
                status_update_func(f"PDF konvertiert, {len(images)} Seiten. Starte OCR...")
            for i, image_obj in enumerate(images):
                if status_update_func:
                    status_update_func(f"OCR für Seite {i+1}/{len(images)}...")
                extracted_text += _ocr_image_to_text(image_obj, lang_code_tesseract, status_update_func) + "\n\n"
            if status_update_func:
                status_update_func(f"OCR für PDF '{os.path.basename(dateipfad)}' abgeschlossen.")
        except Exception as e:
            print(f"FEHLER bei PDF zu Bild Konvertierung oder PDF OCR: {e}")
            if "PDFInfoNotInstalledError" in str(e) or "PDFPageCountError" in str(e) or "PDFSyntaxError" in str(e):
                 print("       Stellen Sie sicher, dass Poppler Utilities installiert sind.")
                 print("       Verwenden Sie 'python poppler_config.py' für Installationshilfe.")
            if status_update_func:
                status_update_func(f"PDF OCR fehlgeschlagen: {e}")
            return "" # Return empty if PDF conversion/OCR fails
    elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
        try:
            if Image is None: raise ImportError("Pillow (PIL) Image module not loaded.")
            img = Image.open(dateipfad)
            extracted_text = _ocr_image_to_text(img, lang_code_tesseract, status_update_func)
        except ImportError: # Should be caught by global OCR_ENABLED check
             if status_update_func: status_update_func("OCR für Bilder nicht möglich: Pillow fehlt.")
        except Exception as e:
            print(f"FEHLER bei Bild-OCR für {dateipfad}: {e}")
            if status_update_func:
                status_update_func(f"Bild OCR fehlgeschlagen: {e}")
            return ""
    else:
        if status_update_func:
            status_update_func(f"Dateityp {ext} wird für OCR nicht direkt unterstützt.")
        return "" # Unsupported file type for direct OCR by this function

    if status_update_func:
        status_update_func(f"OCR für {ext.upper()} abgeschlossen.")
    return extracted_text.strip()


def lese_datei(dateipfad, language_hint="en-US", status_update_func=None):
    """
    Liest den Inhalt einer Datei. Unterstützt .txt, .docx, .pdf.
    Versucht OCR für PDFs mit wenig Text oder für Bilddateien, falls OCR-Bibliotheken verfügbar sind.
    language_hint: z.B. 'de-DE', 'en-US' für die OCR-Sprache.
    status_update_func: Callback-Funktion zur Statusaktualisierung.
    """
    if not dateipfad or not os.path.exists(dateipfad):
        if status_update_func:
            status_update_func(f"Datei nicht gefunden: {dateipfad}")
        return ""
        
    _, ext = os.path.splitext(dateipfad.lower())
    text_inhalt = ""

    if status_update_func:
        status_update_func(f"Lese Datei: {os.path.basename(dateipfad)}...")

    try:
        if ext == ".txt":
            with open(dateipfad, "r", encoding="utf-8") as f:
                text_inhalt = f.read()
        elif ext == ".docx":
            doc = Document(dateipfad)
            text_inhalt = "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".pdf":
            try:
                with open(dateipfad, "rb") as f:
                    reader = PdfReader(f)
                    if reader.is_encrypted:
                        # Try to decrypt with an empty password, common for some PDFs
                        try:
                            reader.decrypt("")
                        except Exception as e_decrypt:
                            print(f"INFO: PDF {os.path.basename(dateipfad)} ist verschlüsselt und konnte nicht entschlüsselt werden: {e_decrypt}")
                            # Fall through to potentially try OCR if it's a scanned encrypted PDF.
                            # Or return an error message. For now, let it try OCR.
                            pass # Allow to fall through to OCR attempt

                    for page in reader.pages:
                        text_inhalt += page.extract_text() or "" # Ensure None is handled
                
                # If PDF text extraction yields very little, and OCR is available, try OCR
                if PYTESSERACT_ENABLED and PDF2IMAGE_ENABLED and OCR_ENABLED and \
                   len(text_inhalt.strip()) < MIN_TEXT_LENGTH_FOR_PDF_OCR_THRESHOLD:
                    if status_update_func:
                        status_update_func(f"Wenig Text aus PDF '{os.path.basename(dateipfad)}' extrahiert. Versuche OCR...")
                    tess_lang = LANG_CODE_TO_TESSERACT.get(language_hint.split('-')[0], DEFAULT_TESSERACT_LANG)
                    ocr_text = ocr_datei(dateipfad, tess_lang, status_update_func)
                    if ocr_text.strip(): # Use OCR text if it's substantial
                        text_inhalt = ocr_text
                    elif status_update_func:
                        status_update_func(f"OCR für PDF '{os.path.basename(dateipfad)}' lieferte keinen zusätzlichen Text.")

            except PyPDF2Errors.PdfReadError as e:
                print(f"PyPDF2 Lesefehler für {dateipfad}: {e}. Versuche OCR, falls Bild-PDF.")
                if PYTESSERACT_ENABLED and PDF2IMAGE_ENABLED and OCR_ENABLED:
                    tess_lang = LANG_CODE_TO_TESSERACT.get(language_hint.split('-')[0], DEFAULT_TESSERACT_LANG)
                    text_inhalt = ocr_datei(dateipfad, tess_lang, status_update_func)
            except Exception as e_pdf: # Catch other PDF related errors
                print(f"Allgemeiner Fehler beim Lesen von PDF {dateipfad}: {e_pdf}. Versuche OCR.")
                if PYTESSERACT_ENABLED and PDF2IMAGE_ENABLED and OCR_ENABLED:
                    tess_lang = LANG_CODE_TO_TESSERACT.get(language_hint.split('-')[0], DEFAULT_TESSERACT_LANG)
                    text_inhalt = ocr_datei(dateipfad, tess_lang, status_update_func)

        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            if PYTESSERACT_ENABLED and OCR_ENABLED:
                if status_update_func:
                    status_update_func(f"Bilddatei '{os.path.basename(dateipfad)}' erkannt. Versuche OCR...")
                tess_lang = LANG_CODE_TO_TESSERACT.get(language_hint.split('-')[0], DEFAULT_TESSERACT_LANG)
                text_inhalt = ocr_datei(dateipfad, tess_lang, status_update_func)
            else:
                message = f"Bilddatei {os.path.basename(dateipfad)} kann nicht ohne OCR-Funktionalität gelesen werden."
                if status_update_func: status_update_func(message)
                print(f"WARNUNG: {message}")
                text_inhalt = "" # Cannot read image without OCR
        else:
            if status_update_func:
                status_update_func(f"Nicht unterstützter Dateityp: {ext} für {os.path.basename(dateipfad)}")
            text_inhalt = "" # Fallback for unsupported types
            
    except Exception as e:
        error_msg = f"Fehler beim Lesen der Datei {dateipfad}: {e}"
        print(f"FEHLER: {error_msg}")
        if status_update_func:
            status_update_func(error_msg)
        return "" # Return empty string on error

    if status_update_func:
        status_update_func(f"Datei '{os.path.basename(dateipfad)}' erfolgreich gelesen.")
    return text_inhalt.strip()


def read_file_content(file_path):
    """
    Reads the content of a file, automatically detecting encoding for text files
    and handling .docx files.

    Args:
        file_path (str): The path to the file.

    Returns:
        str: The content of the file, or an empty string if reading fails.
    """
    if not file_path or not isinstance(file_path, str):
        logging.warning("read_file_content received an invalid file_path.")
        return ""

    try:
        if file_path.lower().endswith('.docx'):
            try:
                doc = Document(file_path)
                full_text = [para.text for para in doc.paragraphs]
                return "\n".join(full_text)
            except Exception as e:
                logging.error(f"Error reading .docx file {file_path}: {e}")
                return ""
        else:
            try:
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                
                # Detect encoding
                result = chardet.detect(raw_data)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'
                logging.info(f"Detected encoding for {file_path}: {encoding} with confidence {result['confidence']}")

                # Handle potential BOM
                if raw_data.startswith(b'\xef\xbb\xbf'):
                    content = raw_data[3:].decode('utf-8')
                else:
                    content = raw_data.decode(encoding, errors='replace')
                
                return content

            except FileNotFoundError:
                logging.error(f"File not found: {file_path}")
                return ""
            except Exception as e:
                logging.error(f"Error reading text file {file_path}: {e}")
                # Fallback to a more robust read
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        return f.read()
                except Exception as final_e:
                    logging.error(f"Final attempt to read {file_path} failed: {final_e}")
                    return ""

    except Exception as e:
        logging.error(f"An unexpected error occurred in read_file_content for {file_path}: {e}")
        return ""

# --- Placeholder for other functions that might be in file_operations.py ---

def _sanitize_foldername(name):
    """Converts a string to a safe folder name."""
    if not name:
        return "Unbenannter_Kunde"
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', '_', name)
    name = re.sub(r'__+', '_', name)
    name = re.sub(r'--+', '-', name)
    name = name.strip('_-.' + ' ')
    if not name:
        return "Bereinigt_Kunde"
    return name

def _extract_version_tag_from_filename(filename):
    """Extracts a version tag like v1, r2, rev3 from a filename."""
    match = re.search(r'[._-](v|r|rev)(\d+)', filename, re.IGNORECASE)
    if match:
        prefix = match.group(1).lower()
        number = match.group(2)
        return f"v{number}" # Normalize rX and revX to vX
    return "v_unk"

def _extract_language_pair_from_filename(filename):
    """Extracts a language pair like DE_EN, EN-FR from a filename."""
    match = re.search(r'([a-zA-Z]{2})[_-]([a-zA-Z]{2})', filename, re.IGNORECASE)
    if match:
        return match.group(1).upper(), match.group(2).upper()
    return None, None

# --- Data management paths and functions ---

_PROJECT_DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "project_data")
aktives_glossar_file_ops = [] # Internal to file_operations

def lade_fachglossar(fachgebiet, status_update_func=None):
    """Lädt ein Fachglossar. `status_update_func` ist eine optionale Callback-Funktion für UI-Updates."""
    global aktives_glossar_file_ops
    dateiname_fachgebiet = fachgebiet.lower().replace(" ", "_").replace("/", "_").replace("\\", "_")
    glossar_dir = resource_path("glossaries") # resource_path defined in this module
    if not os.path.exists(glossar_dir):
        os.makedirs(glossar_dir, exist_ok=True)
        if status_update_func: status_update_func(f"Glossar-Verzeichnis '{glossar_dir}' erstellt.")
        else: print(f"ℹ️ Glossar-Verzeichnis '{glossar_dir}' erstellt.")

    pfad = os.path.join(glossar_dir, f"{dateiname_fachgebiet}.json")
    
    try:
        with open(pfad, "r", encoding="utf-8") as f:
            neues_glossar = json.load(f)
        if isinstance(neues_glossar, list) and all(isinstance(term, str) for term in neues_glossar):
            aktives_glossar_file_ops = neues_glossar
            msg = f"Fachglossar '{fachgebiet}' ({len(aktives_glossar_file_ops)} Begriffe) geladen."
            if status_update_func: status_update_func(msg)
            else: print(f"📚 {msg}")
        else:
            aktives_glossar_file_ops = []
            msg = f"Warnung: Glossar '{fachgebiet}' fehlerhaft. Wird als leer behandelt."
            if status_update_func: status_update_func(msg)
            else: print(f"⚠️ {msg}")
    except FileNotFoundError:
        aktives_glossar_file_ops = []
        msg = f"Kein spezifisches Glossar für '{fachgebiet}' gefunden. Standard (leer) verwendet."
        if status_update_func: status_update_func(msg)
        else: print(f"ℹ️ {msg}")
    except json.JSONDecodeError:
        aktives_glossar_file_ops = []
        msg = f"Fehler beim Lesen des Glossars '{fachgebiet}'. Formatfehler?"
        if status_update_func: status_update_func(msg)
        else: print(f"⚠️ {msg}")
    except Exception as e:
        aktives_glossar_file_ops = []
        msg = f"Unerwarteter Fehler beim Laden des Glossars '{fachgebiet}': {e}"
        if status_update_func: status_update_func(msg)
        else: print(f"⚠️ {msg}")
    return aktives_glossar_file_ops


def get_active_glossar():
    """Gibt das aktuell geladene Fachglossar zurück."""
    return aktives_glossar_file_ops

# Path for last inputs, relative to where project_data_manager stores its data
# This needs to be robust. If project_data_manager is not used or its path is unknown,
# default to a path relative to this file_operations.py script.
try:
    from project_data_manager import DATA_FILE as PDM_DATA_FILE
    PROJECT_DATA_DIR_FO = os.path.dirname(PDM_DATA_FILE) if PDM_DATA_FILE else os.path.abspath(os.path.dirname(__file__))
    if not PROJECT_DATA_DIR_FO:
         PROJECT_DATA_DIR_FO = os.path.abspath(os.path.dirname(__file__))
except (ImportError, AttributeError):
    PROJECT_DATA_DIR_FO = os.path.abspath(os.path.dirname(__file__))
    # Silently use script directory for last_inputs.json (project_data_manager is optional)

LAST_INPUTS_FILE_PATH_FO = os.path.join(PROJECT_DATA_DIR_FO, "last_inputs.json")

def _save_last_inputs(kunden_name, kundenbetreuer, auftragsnummer, zielsprache):
    """Saves the last entered project details to a JSON file."""
    inputs_to_save = {
        "kunden_name": kunden_name,
        "kundenbetreuer_name": kundenbetreuer,
        "auftragsnummer": auftragsnummer,
        "zielsprache": zielsprache
    }
    try:
        os.makedirs(os.path.dirname(LAST_INPUTS_FILE_PATH_FO), exist_ok=True)
        with open(LAST_INPUTS_FILE_PATH_FO, "w", encoding="utf-8") as f:
            json.dump(inputs_to_save, f, indent=4)
        print(f"ℹ️ Letzte Projekteingaben gespeichert in: {LAST_INPUTS_FILE_PATH_FO}")
    except Exception as e:
        print(f"⚠️ Fehler beim Speichern der letzten Projekteingaben: {e}")

def _load_last_inputs():
    """Loads the last entered project details from a JSON file."""
    defaults = {
        "kunden_name": "",
        "kundenbetreuer_name": "",
        "auftragsnummer": "",
        "zielsprache": "de-DE" # Default language
    }
    try:
        if os.path.exists(LAST_INPUTS_FILE_PATH_FO):
            with open(LAST_INPUTS_FILE_PATH_FO, "r", encoding="utf-8") as f:
                loaded_inputs = json.load(f)
            # Merge loaded inputs with defaults to ensure all keys are present
            defaults.update(loaded_inputs)
            print(f"ℹ️ Letzte Projekteingaben geladen von: {LAST_INPUTS_FILE_PATH_FO}")
        else:
            print(f"ℹ️ Keine Datei für letzte Projekteingaben gefunden ({LAST_INPUTS_FILE_PATH_FO}). Verwende Standardwerte.")
    except json.JSONDecodeError:
        print(f"⚠️ Fehler beim Lesen der JSON für letzte Projekteingaben. Verwende Standardwerte.")
    except Exception as e:
        print(f"⚠️ Unerwarteter Fehler beim Laden der letzten Projekteingaben: {e}. Verwende Standardwerte.")
    return defaults

def resource_path(relative_path):
    """Gibt den absoluten Pfad zu einer Ressource zurück, funktioniert auch bei gebündelten Anwendungen."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)

if __name__ == '__main__':
    print("OCR ist nicht aktiviert. Bitte überprüfen Sie die Installation von Tesseract und die Konfiguration der Umgebungsvariablen.")

# --- FileOperations class wrapper ---
class FileOperations:
    """
    Class wrapper for file operations to provide object-oriented interface.
    Wraps the existing functions in this module for compatibility.
    """
    
    def __init__(self):
        """Initialize FileOperations instance."""
        pass
    
    def lese_datei(self, dateipfad, language_hint="en-US", status_update_func=None):
        """Read file content - wrapper for lese_datei function."""
        return lese_datei(dateipfad, language_hint, status_update_func)
    
    def ocr_datei(self, dateipfad, lang_code_tesseract, status_update_func=None):
        """OCR file content - wrapper for ocr_datei function."""
        return ocr_datei(dateipfad, lang_code_tesseract, status_update_func)
    
    def lade_fachglossar(self, fachgebiet, status_update_func=None):
        """Load domain glossary - wrapper for lade_fachglossar function."""
        return lade_fachglossar(fachgebiet, status_update_func)
    
    def get_active_glossar(self):
        """Get active glossary - wrapper for get_active_glossar function."""
        return get_active_glossar()
    
    def resource_path(self, relative_path):
        """Get resource path - wrapper for resource_path function."""
        return resource_path(relative_path)

