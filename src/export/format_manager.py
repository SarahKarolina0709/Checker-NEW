"""
Format Export Manager
====================

Handles multi-format export capabilities including PDF, DOCX, HTML, XML, JSON, and CSV.
Provides unified interface for exporting data in various formats with quality preservation.
"""

from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple, Union
import json
import logging
import os

import csv

class FormatExportManager:
    """Manager for multi-format export operations."""

    def __init__(self, app_instance=None):
        """
        Initialize the Format Export Manager.

        Args:
            app_instance: Reference to the main application instance
        """
        self.app = app_instance
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['PDF', 'JSON', 'CSV', 'HTML', 'TXT']
        self.format_handlers = {}

        # Initialize format handlers
        self._init_format_handlers()

    def _init_format_handlers(self):
        """Initialize handlers for different export formats."""
        try:
            # JSON handler
            self.format_handlers['JSON'] = self._export_json
            self.format_handlers['json'] = self._export_json

            # CSV handler
            self.format_handlers['CSV'] = self._export_csv
            self.format_handlers['csv'] = self._export_csv

            # HTML handler
            self.format_handlers['HTML'] = self._export_html
            self.format_handlers['html'] = self._export_html

            # TXT handler
            self.format_handlers['TXT'] = self._export_txt
            self.format_handlers['txt'] = self._export_txt

            # Try to add DOCX support
            try:
                from docx import Document
                self.format_handlers['DOCX'] = self._export_docx
                self.format_handlers['docx'] = self._export_docx
                self.supported_formats.append('DOCX')
                self.logger.info("[EXPORT] DOCX export available")
            except ImportError:
                self.logger.warning("[EXPORT] python-docx not available - DOCX export disabled")

            # Try to add XML support
            try:
                import xml.etree.ElementTree as ET
                self.format_handlers['XML'] = self._export_xml
                self.format_handlers['xml'] = self._export_xml
                self.supported_formats.append('XML')
                self.logger.info("[EXPORT] XML export available")
            except ImportError:
                self.logger.warning("[EXPORT] XML export not available")

            self.logger.info(f"[EXPORT] Format handlers initialized: {', '.join(self.supported_formats)}")

        except Exception as e:
            self.logger.error(f"[EXPORT] Error initializing format handlers: {e}")

    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats."""
        return self.supported_formats.copy()

    def export_data(self, data: Dict[str, Any], format_type: str, output_path: str = None) -> Tuple[bool, str]:
        """
        Export data in the specified format.

        Args:
            data: Data to export
            format_type: Target format (PDF, JSON, CSV, HTML, etc.)
            output_path: Optional output path

        Returns:
            Tuple of (success: bool, result_path_or_error: str)
        """
        try:
            format_upper = format_type.upper()

            if format_upper not in [f.upper() for f in self.supported_formats]:
                return False, f"Format {format_type} not supported"

            # Get appropriate handler
            handler = self.format_handlers.get(format_type, self.format_handlers.get(format_upper))

            if not handler:
                return False, f"No handler available for format {format_type}"

            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                extension = format_type.lower()
                filename = f"export_{timestamp}.{extension}"
                output_path = self._get_default_export_path(filename)

            # Execute export
            success = handler(data, output_path)

            if success:
                self.logger.info(f"[EXPORT] Data exported to {format_type}: {output_path}")

                # Show notification if possible
                if (self.app and hasattr(self.app, 'notification_center') and
                    self.app.notification_center):
                    self.app.notification_center.show_notification(
                        f"Data exported to {os.path.basename(output_path)}", "success"
                    )

                return True, output_path
            else:
                return False, f"Export to {format_type} failed"

        except Exception as e:
            error_msg = f"Error exporting to {format_type}: {e}"
            self.logger.error(f"[EXPORT] {error_msg}")
            return False, error_msg

    def _export_json(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as JSON."""
        try:
            # Ensure data is JSON serializable
            serializable_data = self._make_json_serializable(data)

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(serializable_data, f, indent=2, ensure_ascii=False)

            return True

        except Exception as e:
            self.logger.error(f"[EXPORT] JSON export error: {e}")
            return False

    def _export_csv(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as CSV."""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)

                # Handle different data structures
                if isinstance(data, dict):
                    # If it's a flat dictionary, write key-value pairs
                    if all(not isinstance(v, (dict, list)) for v in data.values()):
                        writer.writerow(['Key', 'Value'])
                        for key, value in data.items():
                            writer.writerow([key, str(value)])
                    else:
                        # Handle nested data by flattening
                        flattened = self._flatten_dict(data)
                        writer.writerow(['Key', 'Value'])
                        for key, value in flattened.items():
                            writer.writerow([key, str(value)])

                elif isinstance(data, list):
                    # If it's a list of dictionaries (table-like data)
                    if data and isinstance(data[0], dict):
                        headers = list(data[0].keys())
                        writer.writerow(headers)
                        for row in data:
                            writer.writerow([str(row.get(h, '')) for h in headers])
                    else:
                        # Simple list
                        writer.writerow(['Value'])
                        for item in data:
                            writer.writerow([str(item)])

            return True

        except Exception as e:
            self.logger.error(f"[EXPORT] CSV export error: {e}")
            return False

    def _export_html(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as HTML."""
        try:
            html_content = self._generate_html(data)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            return True

        except Exception as e:
            self.logger.error(f"[EXPORT] HTML export error: {e}")
            return False

    def _export_txt(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as plain text."""
        try:
            text_content = self._generate_text(data)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)

            return True

        except Exception as e:
            self.logger.error(f"[EXPORT] TXT export error: {e}")
            return False

    def _export_docx(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as DOCX (if python-docx is available)."""
        try:
            from docx import Document

            doc = Document()

            # Add title
            title = doc.add_heading('Export Report', 0)

            # Add data
            self._add_data_to_docx(doc, data)

            doc.save(output_path)
            return True

        except ImportError:
            self.logger.error("[EXPORT] python-docx not available for DOCX export")
            return False
        except Exception as e:
            self.logger.error(f"[EXPORT] DOCX export error: {e}")
            return False

    def _export_xml(self, data: Dict[str, Any], output_path: str) -> bool:
        """Export data as XML."""
        try:
            import xml.etree.ElementTree as ET

            root = ET.Element("export")
            root.set("timestamp", datetime.now().isoformat())

            self._dict_to_xml(data, root)

            tree = ET.ElementTree(root)
            tree.write(output_path, encoding='utf-8', xml_declaration=True)

            return True

        except Exception as e:
            self.logger.error(f"[EXPORT] XML export error: {e}")
            return False

    def _make_json_serializable(self, obj: Any) -> Any:
        """Make object JSON serializable."""
        if isinstance(obj, dict):
            return {k: self._make_json_serializable(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [self._make_json_serializable(item) for item in obj]
        elif isinstance(obj, (datetime,)):
            return obj.isoformat()
        elif hasattr(obj, '__dict__'):
            return self._make_json_serializable(obj.__dict__)
        else:
            try:
                json.dumps(obj)  # Test if it's already serializable
                return obj
            except (TypeError, ValueError):
                return str(obj)

    def _flatten_dict(self, d: Dict[str, Any], parent_key: str = '', sep: str = '.') -> Dict[str, Any]:
        """Flatten nested dictionary."""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key, sep=sep).items())
            elif isinstance(v, list):
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        items.extend(self._flatten_dict(item, f"{new_key}[{i}]", sep=sep).items())
                    else:
                        items.append((f"{new_key}[{i}]", str(item)))
            else:
                items.append((new_key, v))
        return dict(items)

    def _generate_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML content from data."""
        html = """
<!DOCTYPE html>
<html>
<head>
    <title>Export Report</title>
    <meta charset="UTF-8">
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        .header { color: #333; border-bottom: 2px solid #ccc; padding-bottom: 10px; }
        .section { margin: 20px 0; }
        .key { font-weight: bold; color: #666; }
        .value { margin-left: 20px; }
        table { border-collapse: collapse; width: 100%; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
    </style>
</head>
<body>
    <h1 class="header">Export Report</h1>
    <p><strong>Generated:</strong> """ + datetime.now().strftime("%Y-%m-%d %H:%M:%S") + """</p>
"""

        # Optionaler Abschnitt: Verwendete Ähnlichkeitsschwellen (falls im Datensatz vorhanden)
        try:
            if isinstance(data, dict):
                metrics = data.get('metrics') if isinstance(data.get('metrics'), dict) else None
                thr = metrics.get('similarity_thresholds_used') if metrics else None
                if isinstance(thr, dict) and (thr.get('critical') is not None or thr.get('major') is not None):
                    def _fmt(v):
                        try:
                            return f"{float(v):.0%}" if v is not None else '–'
                        except Exception:
                            return '–'
                    c_txt = _fmt(thr.get('critical'))
                    m_txt = _fmt(thr.get('major'))
                    html += f"<div class='section'><h2 class='key'>Verwendete Ähnlichkeitsschwellen</h2><div>Kritisch ≥ {c_txt}, Wesentlich ≥ {m_txt}</div></div>\n"
        except Exception:
            pass

        # Domain-spezifische Sonderbehandlung: Findings als Tabelle mit Confidence
        try:
            if isinstance(data, dict) and isinstance(data.get('findings'), list) and data.get('findings'):
                findings = data.get('findings')
                html += "<div class='section'>\n<h2 class='key'>Befunde</h2>\n"
                # Prüfe auf optionales 'count' Feld für gruppierte Exporte
                has_count = False
                try:
                    has_count = any(isinstance(f, dict) and 'count' in f for f in findings)
                except Exception:
                    has_count = False
                # Deutsche Spaltenbezeichnungen vereinheitlichen (Lokalisierung)
                cols = ["Schweregrad", "Regel", "Nachricht"]
                if has_count:
                    cols.append("Anzahl")
                # Checker / Sicherheitswert (Confidence)
                cols.extend(["Prüfer", "Sicherheitswert"])  # Prüfer bleibt leer bei Gruppierung
                html += "<table><thead><tr>" + ''.join(f"<th>{c}</th>" for c in cols) + "</tr></thead><tbody>"
                from html import escape as _esc
                for f in findings:
                    if not isinstance(f, dict):
                        continue
                    # Severity lokalisieren (falls noch englisch intern)
                    sev_code = str(f.get('severity', ''))
                    sev_map = { 'critical': 'kritisch', 'major': 'wesentlich', 'minor': 'gering' }
                    sev = _esc(sev_map.get(sev_code, sev_code))
                    rule = _esc(str(f.get('rule_id') or f.get('rule') or ''))
                    msg = _esc(str(f.get('message', '')))
                    # Checker evtl. nicht vorhanden in Gruppierung
                    chk = _esc(str(f.get('checker', '')))
                    # Confidence-Fallbacks (avg_confidence/avg_conf)
                    conf = f.get('confidence')
                    if conf is None:
                        conf = f.get('avg_confidence', f.get('avg_conf'))
                    try:
                        if isinstance(conf, str):
                            conf = float(conf)
                    except Exception:
                        pass
                    conf_txt = f"{conf:.4f}" if isinstance(conf, (int, float)) else ''
                    tds = [sev, rule, msg]
                    if has_count:
                        tds.append(str(f.get('count') or ''))
                    tds.extend([chk, conf_txt])
                    html += "<tr>" + ''.join(f"<td>{_esc(str(x))}</td>" for x in tds) + "</tr>"
                html += "</tbody></table></div>\n"
                # Entferne findings aus generischem Renderer, um Duplikate zu vermeiden
                rest = {k: v for k, v in data.items() if k != 'findings'}
                html += self._dict_to_html(rest)
            else:
                html += self._dict_to_html(data)
        except Exception:
            html += self._dict_to_html(data)
        html += "\n</body>\n</html>"

        return html

    def _dict_to_html(self, data: Dict[str, Any], level: int = 1) -> str:
        """Convert dictionary to HTML representation."""
        html = ""

        for key, value in data.items():
            if isinstance(value, dict):
                html += f"<div class='section'><h{min(level+1, 6)} class='key'>{key}</h{min(level+1, 6)}>\n"
                html += self._dict_to_html(value, level + 1)
                html += "</div>\n"
            elif isinstance(value, list):
                html += f"<div class='section'><span class='key'>{key}:</span>\n<ul>\n"
                for item in value:
                    if isinstance(item, dict):
                        html += "<li>\n" + self._dict_to_html(item, level + 1) + "</li>\n"
                    else:
                        html += f"<li>{str(item)}</li>\n"
                html += "</ul></div>\n"
            else:
                html += f"<div class='section'><span class='key'>{key}:</span> <span class='value'>{str(value)}</span></div>\n"

        return html

    def _generate_text(self, data: Dict[str, Any]) -> str:
        """Generate plain text content from data."""
        text = f"Export Report\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        text += "=" * 50 + "\n\n"

        text += self._dict_to_text(data)

        return text

    def _dict_to_text(self, data: Dict[str, Any], indent: int = 0) -> str:
        """Convert dictionary to text representation."""
        text = ""
        prefix = "  " * indent

        for key, value in data.items():
            if isinstance(value, dict):
                text += f"{prefix}{key}:\n"
                text += self._dict_to_text(value, indent + 1)
            elif isinstance(value, list):
                text += f"{prefix}{key}:\n"
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        text += f"{prefix}  [{i}]:\n"
                        text += self._dict_to_text(item, indent + 2)
                    else:
                        text += f"{prefix}  - {str(item)}\n"
            else:
                text += f"{prefix}{key}: {str(value)}\n"

        text += "\n"
        return text

    def _add_data_to_docx(self, doc, data: Dict[str, Any], level: int = 1):
        """Add data to DOCX document."""
        for key, value in data.items():
            if isinstance(value, dict):
                heading = doc.add_heading(str(key), level=min(level, 9))
                self._add_data_to_docx(doc, value, level + 1)
            elif isinstance(value, list):
                doc.add_heading(str(key), level=min(level, 9))
                for item in value:
                    if isinstance(item, dict):
                        self._add_data_to_docx(doc, item, level + 1)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"• {str(item)}")
            else:
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))

    def _dict_to_xml(self, data: Dict[str, Any], parent_element):
        """Convert dictionary to XML elements."""
        import xml.etree.ElementTree as ET

        for key, value in data.items():
            # Clean key name for XML
            clean_key = str(key).replace(' ', '_').replace('-', '_')
            clean_key = ''.join(c for c in clean_key if c.isalnum() or c == '_')

            if isinstance(value, dict):
                element = ET.SubElement(parent_element, clean_key)
                self._dict_to_xml(value, element)
            elif isinstance(value, list):
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        element = ET.SubElement(parent_element, f"{clean_key}_item")
                        element.set("index", str(i))
                        self._dict_to_xml(item, element)
                    else:
                        element = ET.SubElement(parent_element, f"{clean_key}_item")
                        element.set("index", str(i))
                        element.text = str(item)
            else:
                element = ET.SubElement(parent_element, clean_key)
                element.text = str(value)

    def _get_default_export_path(self, filename: str = None) -> str:
        """Get default export path."""
        try:
            # Try to use app's customer manager for context
            if (self.app and hasattr(self.app, 'kunden_manager') and
                self.app.kunden_manager):
                base_path = os.path.join(self.app.kunden_manager.base_dir, "exports")
            else:
                base_path = os.path.join(os.getcwd(), "exports")

            os.makedirs(base_path, exist_ok=True)

            if filename:
                return os.path.join(base_path, filename)
            else:
                return base_path

        except Exception as e:
            self.logger.error(f"[EXPORT] Error getting default export path: {e}")
            return os.path.join(os.getcwd(), "exports")

    def batch_export(self, data_list: List[Dict[str, Any]], format_type: str, output_dir: str = None) -> Tuple[int, List[str]]:
        """
        Export multiple data sets in batch mode.

        Args:
            data_list: List of data dictionaries to export
            format_type: Target format
            output_dir: Optional output directory

        Returns:
            Tuple of (successful_exports: int, exported_files: List[str])
        """
        try:
            if not output_dir:
                output_dir = self._get_default_export_path()

            os.makedirs(output_dir, exist_ok=True)

            successful_exports = 0
            exported_files = []

            for i, data in enumerate(data_list):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                extension = format_type.lower()
                filename = f"batch_export_{i+1}_{timestamp}.{extension}"
                output_path = os.path.join(output_dir, filename)

                success, result = self.export_data(data, format_type, output_path)
                if success:
                    successful_exports += 1
                    exported_files.append(result)
                else:
                    self.logger.warning(f"[EXPORT] Failed to export item {i+1}: {result}")

            self.logger.info(f"[EXPORT] Batch export completed: {successful_exports}/{len(data_list)} successful")

            return successful_exports, exported_files

        except Exception as e:
            self.logger.error(f"[EXPORT] Error in batch export: {e}")
            return 0, []