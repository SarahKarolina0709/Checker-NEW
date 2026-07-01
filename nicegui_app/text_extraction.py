# -*- coding: utf-8 -*-
"""Text-Extraktion fuer alle unterstuetzten Dateiformate.

Kapselt die zuvor in main.py verstreute Logik:
- TXT/MD/PY (UTF-8 → CP1252 → Latin-1 Fallback-Kette)
- DOCX inkl. Tabellenzellen, Header/Footer, TextBoxen
- PDF (pdfplumber -> PyMuPDF -> OCR-Fallback)
- Bilder (OCR via Tesseract, optional)

Optional-Dependencies werden lazy importiert; fehlt eine, gibt's eine
verstaendliche Meldung statt Crash.
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Cache fuer extrahierten Text (Key: Pfad+mtime+Size, Value: Text)
# Begrenzt durch _EXTRACT_CACHE_MAX (FIFO-Approximation)
# ---------------------------------------------------------------------------
_EXTRACT_CACHE: Dict[Tuple[str, int, int], str] = {}
_EXTRACT_CACHE_MAX = 64

# ---------------------------------------------------------------------------
# Optional OCR
# ---------------------------------------------------------------------------
try:
    from PIL import Image  # type: ignore
except ImportError:
    Image = None  # type: ignore

try:
    import pytesseract  # type: ignore
    _HAS_OCR = True
except ImportError:
    pytesseract = None  # type: ignore
    _HAS_OCR = False

# Gebuendelte Tools (Repo-Root) bevorzugen, sonst System-PATH
try:
    from nicegui_app.ocr_paths import bundled_poppler_bin, configure_pytesseract
except ImportError:
    _POPPLER_BIN = None
else:
    configure_pytesseract(pytesseract)
    _POPPLER_BIN = bundled_poppler_bin()

# Encoding-Fallback-Kette für TXT-Dateien (Windows-Dateien oft cp1252)
_TXT_ENCODINGS = ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1')


def _normalize_text(text: str) -> str:
    """Bereinigt unsichtbare Sonderzeichen die Wort-Matching stören.

    - Soft-Hyphens (\\xad) entfernen
    - Zero-Width-Spaces/Joiners entfernen
    - Non-breaking spaces → normales Leerzeichen
    - Mehrfache Leerzeilen → eine Leerzeile
    """
    # Unsichtbare Zeichen entfernen
    text = text.replace('\xad', '')          # soft hyphen
    text = text.replace('\u200b', '')        # zero-width space
    text = text.replace('\u200c', '')        # zero-width non-joiner
    text = text.replace('\u200d', '')        # zero-width joiner
    text = text.replace('\ufeff', '')        # BOM (falls noch vorhanden)
    text = text.replace('\u00a0', ' ')       # non-breaking space → space
    text = text.replace('\u202f', ' ')       # narrow no-break space → space
    # Mehrfache Leerzeilen auf maximal 2 reduzieren
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _ocr_image(path: str) -> str:
    if not _HAS_OCR or Image is None:
        return f'[OCR nicht verfuegbar: {Path(path).name}]'
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, lang='deu+eng') or ''
    except Exception as exc:
        return f'[OCR-Fehler {Path(path).name}: {exc}]'


def _ocr_pdf_page(path: str) -> str:
    """OCR ueber alle Seiten einer PDF (Fallback wenn Text-Extraktion leer).

    Bevorzugt `pdf2image` (300 DPI, beste Qualitaet); faellt auf PyMuPDF
    bei niedrigerer Aufloesung zurueck wenn pdf2image nicht installiert ist.
    """
    if not _HAS_OCR or Image is None:
        return ''
    # Primaer: pdf2image mit 300 DPI (Qualitaet)
    try:
        from pdf2image import convert_from_path  # type: ignore
        images = []
        try:
            images = convert_from_path(path, dpi=300, poppler_path=_POPPLER_BIN)
            texts = [pytesseract.image_to_string(img, lang='deu+eng') or '' for img in images]
            if any(t.strip() for t in texts):
                return '\n'.join(texts)
        finally:
            for img in images:
                try:
                    img.close()
                except Exception:
                    pass
    except ImportError:
        pass
    except Exception:
        pass
    # Fallback: PyMuPDF
    try:
        import fitz  # type: ignore
    except ImportError:
        return ''
    out: List[str] = []
    try:
        doc = fitz.open(path)
        try:
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                from io import BytesIO
                with Image.open(BytesIO(pix.tobytes('png'))) as img:
                    out.append(pytesseract.image_to_string(img, lang='deu+eng') or '')
        finally:
            doc.close()
    except Exception:
        return ''
    return '\n'.join(out)


def extract_text(path: str) -> str:
    """Extrahiert Plaintext aus einer Datei.

    Unterstuetzte Formate: .txt .md .py .docx .pdf .png .jpg .jpeg .tiff .tif

    Rueckgabe: Plaintext oder eine `[...]`-Fehlermeldung. Wirft nie eine
    Exception nach aussen.

    Cached intern auf (Pfad, mtime, Groesse) - Re-Aufrufe fuer dieselbe
    unveraenderte Datei sind quasi gratis.
    """
    if not path or not os.path.isfile(path):
        return f'[Datei nicht gefunden: {Path(path).name if path else ""}]'
    try:
        st = os.stat(path)
        cache_key = (path, st.st_mtime_ns, st.st_size)
    except OSError:
        cache_key = None
    if cache_key is not None:
        cached = _EXTRACT_CACHE.get(cache_key)
        if cached is not None:
            return cached
    result = _extract_text_impl(path)
    if cache_key is not None and result and not result.startswith('['):
        # Cache nur erfolgreiche Extraktionen (keine Fehlermeldungen)
        if len(_EXTRACT_CACHE) >= _EXTRACT_CACHE_MAX:
            # Einfache LRU-Approximation: aeltesten Eintrag werfen
            try:
                _EXTRACT_CACHE.pop(next(iter(_EXTRACT_CACHE)))
            except StopIteration:
                pass
        _EXTRACT_CACHE[cache_key] = result
    return result


def _extract_text_impl(path: str) -> str:
    """Tatsaechliche Extraktion - ohne Cache."""
    ext = Path(path).suffix.lower()
    if ext in ('.txt', '.py', '.md'):
        # Encoding-Fallback-Kette (UTF-8-BOM → UTF-8 → CP1252 → Latin-1)
        for enc in _TXT_ENCODINGS:
            try:
                text = Path(path).read_text(encoding=enc)
                return _normalize_text(text)
            except (UnicodeDecodeError, UnicodeError):
                continue
        # Letzter Ausweg: UTF-8 mit Ersatz-Zeichen
        try:
            return _normalize_text(Path(path).read_text(encoding='utf-8', errors='replace'))
        except Exception as exc:
            return f'[Lesefehler {Path(path).name}: {exc}]'
    if ext == '.docx':
        try:
            from docx import Document  # type: ignore
            doc = Document(path)
            parts: List[str] = []
            # Header/Footer aller Sektionen
            for section in doc.sections:
                for hf in (section.header, section.footer):
                    try:
                        for p in hf.paragraphs:
                            t = (p.text or '').strip()
                            if t:
                                parts.append(t)
                    except Exception:
                        pass
            # Absätze
            parts.extend(p.text for p in doc.paragraphs if p.text and p.text.strip())
            # Tabellen — merged cells deduplizieren (python-docx gibt sie mehrfach zurück)
            for table in doc.tables:
                for row in table.rows:
                    seen_ids: set = set()
                    for cell in row.cells:
                        if id(cell) in seen_ids:
                            continue
                        seen_ids.add(id(cell))
                        ct = (cell.text or '').strip()
                        if ct:
                            parts.append(ct)
            # TextBoxen / Shapes (werden von python-docx nicht direkt exponiert)
            try:
                from docx.oxml.ns import qn  # type: ignore
                for shape in doc.element.body.iter(qn('wps:txbx')):
                    for p_elem in shape.iter(qn('w:p')):
                        runs = ''.join(r.text or '' for r in p_elem.iter(qn('w:t')))
                        if runs.strip():
                            parts.append(runs.strip())
            except Exception:
                pass
            return _normalize_text('\n'.join(parts))
        except ImportError:
            return f'[python-docx nicht installiert: {Path(path).name}]'
        except Exception as exc:
            return f'[DOCX-Lesefehler {Path(path).name}: {exc}]'
    if ext in ('.png', '.jpg', '.jpeg', '.tiff', '.tif'):
        return _ocr_image(path)
    if ext == '.pdf':
        text = ''
        try:
            import pdfplumber  # type: ignore
            pages_text: List[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_parts: List[str] = []
                    # Fließtext
                    pt = page.extract_text() or ''
                    if pt.strip():
                        page_parts.append(pt)
                    # Tabellen separat — vermeidet zusammengepresste Zellen
                    try:
                        for tbl in page.extract_tables() or []:
                            for row in tbl:
                                cells = [str(c or '').strip() for c in row if c and str(c).strip()]
                                if cells:
                                    page_parts.append(' | '.join(cells))
                    except Exception:
                        pass
                    pages_text.append('\n'.join(page_parts))
            text = '\n'.join(pages_text)
        except Exception:
            try:
                import fitz  # type: ignore
                doc = fitz.open(path)
                try:
                    text = '\n'.join(p.get_text() for p in doc)
                finally:
                    doc.close()
            except Exception:
                pass
        # Scanned PDFs liefern oft wenige Wörter (<5) trotz Seite — OCR-Fallback
        if not text.strip() or len(text.split()) < 5:
            ocr_text = _ocr_pdf_page(path)
            if ocr_text.strip():
                return _normalize_text(ocr_text)
            if not text.strip():
                return f'[PDF-Extraktion fehlgeschlagen: {Path(path).name}]'
        return _normalize_text(text)
    if ext == '.doc':
        return f'[.doc-Format nur mit externem Konverter: {Path(path).name}]'
    return ''


def get_text_stats(fp: str, chars_per_norm_line: int = 36) -> dict:
    """Liefert Zeichen-/Wort-/Normzeilen-Statistik fuer eine Datei.

    Liefert leeres Dict bei Fehler oder leerem Text.
    """
    try:
        text = extract_text(fp)
        if not text:
            return {}
        chars = len(text)
        words = len(text.split())
        return {
            'chars': chars,
            'words': words,
            'norm_lines': round(chars / chars_per_norm_line, 1),
            'cpl': chars_per_norm_line,
        }
    except Exception:
        return {}


def finding_to_dict(f) -> dict:
    """Serialisiert ein QAIssue/Finding-Objekt zu einem JSON-tauglichen Dict."""
    return {
        'severity': getattr(f, 'severity', 'info'),
        'code': getattr(f, 'code', ''),
        'message': getattr(f, 'message', ''),
        'category': getattr(f, 'category', ''),
        'source_text': getattr(f, 'source_text', '') or '',
        'target_text': getattr(f, 'target_text', '') or '',
        'segment_index': getattr(f, 'segment_index', -1),
        'source_file': getattr(f, 'source_file', '') or '',
        'target_file': getattr(f, 'target_file', '') or '',
        'meta': getattr(f, 'meta', {}) or {},
    }


def dict_to_finding(d: dict, qa_issue_cls):
    """Deserialisiert ein Dict zurueck zu einem QAIssue.

    `qa_issue_cls` wird per DI uebergeben um zirkulaere Imports zu vermeiden.
    None-Schutz fuer ALLE Felder (Legacy-Sessions koennen None enthalten).
    """
    seg = d.get('segment_index', -1)
    if seg is None:
        seg = -1
    kwargs = dict(
        severity=d.get('severity') or 'info',
        code=d.get('code') or '',
        message=d.get('message') or '',
        category=d.get('category') or '',
        source_text=d.get('source_text') or '',
        target_text=d.get('target_text') or '',
        segment_index=seg,
        meta=d.get('meta') or {},
    )
    # Per-File-Attribution: nur wenn die Klasse die Felder unterstützt
    try:
        import dataclasses as _dc
        if _dc.is_dataclass(qa_issue_cls):
            field_names = {f.name for f in _dc.fields(qa_issue_cls)}
        else:
            field_names = set(getattr(qa_issue_cls, '__annotations__', {}).keys())
        if 'source_file' in field_names:
            kwargs['source_file'] = d.get('source_file') or ''
        if 'target_file' in field_names:
            kwargs['target_file'] = d.get('target_file') or ''
    except Exception:
        pass
    return qa_issue_cls(**kwargs)
